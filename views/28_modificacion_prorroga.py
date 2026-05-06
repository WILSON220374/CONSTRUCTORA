from io import BytesIO
from datetime import date, datetime, timedelta
import re

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, Inches

from supabase_state import cargar_estado
from supabase_state import guardar_estado as guardar_estado_bd


CLAVE_GUARDADO = "modificacion_prorroga"


# ==========================================================
# Persistencia
# ==========================================================
def guardar_estado(clave, datos):
    def serializar(obj):
        if isinstance(obj, dict):
            return {k: serializar(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [serializar(x) for x in obj]
        if isinstance(obj, (date, datetime)):
            return obj.isoformat()
        return obj

    guardar_estado_bd(clave, serializar(datos))


# ==========================================================
# Helpers generales
# ==========================================================
def _texto(valor):
    if valor is None:
        return ""
    return str(valor).strip()


def _safe_float(valor, default=0.0):
    try:
        if valor is None or valor == "":
            return float(default)
        if isinstance(valor, (int, float)):
            return float(valor)
        txt = str(valor).strip().replace("$", "").replace(" ", "")
        txt = txt.replace(".", "").replace(",", ".")
        return float(txt)
    except Exception:
        return float(default)


def _parse_fecha(valor):
    if isinstance(valor, date):
        return valor
    if isinstance(valor, datetime):
        return valor.date()
    if isinstance(valor, str) and valor.strip():
        txt = valor.strip()
        try:
            return datetime.fromisoformat(txt).date()
        except Exception:
            pass
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
            try:
                return datetime.strptime(txt, fmt).date()
            except Exception:
                continue
    return None


def _fecha_texto(valor):
    fecha = _parse_fecha(valor)
    if not fecha:
        return ""
    return fecha.strftime("%d/%m/%Y")


def _fecha_input(valor):
    return _parse_fecha(valor) or date.today()


def _moneda(valor):
    return f"$ {_safe_float(valor, 0.0):,.2f}"


def _leer_estado(clave):
    datos = cargar_estado(clave) or {}
    return datos if isinstance(datos, dict) else {}


def _primero_no_vacio(*valores):
    for valor in valores:
        txt = _texto(valor)
        if txt:
            return txt
    return ""


def _extraer_dias_plazo(valor):
    txt = _texto(valor).lower()
    if not txt:
        return 0
    nums = re.findall(r"\d+", txt)
    if not nums:
        return 0
    n = int(nums[0])
    if "mes" in txt:
        return n * 30
    if "semana" in txt:
        return n * 7
    return n


def _sumar_valores(rows, campo):
    total = 0.0
    for fila in rows or []:
        if isinstance(fila, dict):
            total += _safe_float(fila.get(campo), 0.0)
    return round(total, 2)


def _df(rows, columns):
    registros = rows if isinstance(rows, list) else []
    return pd.DataFrame(registros, columns=columns)


# ==========================================================
# Lectura de fuentes
# ==========================================================
def _valor_contrato(acta_inicio, contrato_obra):
    for valor in [
        acta_inicio.get("valor_total_contrato_obra"),
        acta_inicio.get("valor_contrato"),
        acta_inicio.get("valor"),
        contrato_obra.get("valor_total_numeros"),
        contrato_obra.get("valor_contrato"),
        contrato_obra.get("valor"),
    ]:
        numero = _safe_float(valor, 0.0)
        if numero > 0:
            return numero
    return 0.0


def _valor_anticipo(control_obra, plan_anticipo):
    anticipo_rows = control_obra.get("anticipo_rows", []) or []
    if isinstance(anticipo_rows, list) and anticipo_rows:
        valor = _safe_float(anticipo_rows[0].get("VALOR INICIAL"), 0.0) if isinstance(anticipo_rows[0], dict) else 0.0
        if valor > 0:
            return valor

    for clave in ["valor_anticipo", "anticipo_valor", "valor_total_anticipo", "anticipo"]:
        valor = _safe_float(plan_anticipo.get(clave), 0.0)
        if valor > 0:
            return valor
    return 0.0


def _fecha_inicio(acta_inicio):
    return _parse_fecha(
        _primero_no_vacio(
            acta_inicio.get("fecha_inicio"),
            acta_inicio.get("fecha_presente_acta"),
        )
    )


def _plazo_texto(acta_inicio, contrato_obra):
    return _primero_no_vacio(
        acta_inicio.get("plazo_ejecucion"),
        acta_inicio.get("plazo"),
        contrato_obra.get("plazo_ejecucion"),
        contrato_obra.get("plazo"),
    )


def _fecha_vencimiento_inicial(acta_inicio, contrato_obra):
    fecha_directa = _primero_no_vacio(
        acta_inicio.get("fecha_terminacion"),
        acta_inicio.get("fecha_terminacion_contrato"),
    )
    if fecha_directa:
        return _parse_fecha(fecha_directa)

    inicio = _fecha_inicio(acta_inicio)
    if not inicio:
        return None

    plazo_dias = int(acta_inicio.get("plazo_ejecucion_dias", 0) or 0)
    if plazo_dias <= 0:
        plazo_dias = int(contrato_obra.get("plazo_ejecucion_dias", 0) or 0)
    if plazo_dias <= 0:
        plazo_dias = _extraer_dias_plazo(_plazo_texto(acta_inicio, contrato_obra))
    if plazo_dias <= 0:
        return None
    return inicio + timedelta(days=plazo_dias)


def _fecha_vencimiento_actual(fecha_inicial, prorrogas, suspensiones):
    fecha_actual = fecha_inicial

    for fila in suspensiones or []:
        if isinstance(fila, dict):
            fecha = _parse_fecha(fila.get("NUEVA FECHA DE FINALIZACIÓN"))
            if fecha and (not fecha_actual or fecha > fecha_actual):
                fecha_actual = fecha

    for fila in prorrogas or []:
        if isinstance(fila, dict):
            fecha = _parse_fecha(fila.get("NUEVA FECHA DE TERMINACIÓN"))
            if fecha and (not fecha_actual or fecha > fecha_actual):
                fecha_actual = fecha

    return fecha_actual


def _datos_generales(acta_inicio, contrato_obra, contrato_interventoria, control_obra):
    adiciones = control_obra.get("adiciones_rows", []) or []
    prorrogas = control_obra.get("prorrogas_rows", []) or []
    suspensiones = control_obra.get("suspensiones_rows", []) or []

    valor_inicial = _valor_contrato(acta_inicio, contrato_obra)
    valor_adiciones = _sumar_valores(adiciones, "VALOR")
    fecha_inicial = _fecha_vencimiento_inicial(acta_inicio, contrato_obra)

    salario_minimo = _safe_float(control_obra.get("salario_minimo_anio_contrato"), 0.0)
    valor_contrato_smmlv = round(valor_inicial / salario_minimo, 4) if salario_minimo > 0 else 0.0
    cincuenta_smmlv = round(valor_contrato_smmlv * 0.5, 4)

    return {
        "numero_contrato": _primero_no_vacio(acta_inicio.get("numero_contrato"), contrato_obra.get("numero_contrato")),
        "contratista": _primero_no_vacio(acta_inicio.get("nombre_firma_contratista"), contrato_obra.get("nombre_contratista")),
        "contratante": _primero_no_vacio(contrato_obra.get("nombre_entidad"), acta_inicio.get("contratante")),
        "objeto": _primero_no_vacio(acta_inicio.get("objeto_contrato"), contrato_obra.get("objeto_general"), contrato_obra.get("objeto_contrato"), contrato_obra.get("objeto")),
        "contrato_interventoria": _primero_no_vacio(contrato_interventoria.get("numero_proceso_contratacion")),
        "interventor": _primero_no_vacio(acta_inicio.get("nombre_firma_interventor"), contrato_interventoria.get("nombre_interventor"), contrato_obra.get("nombre_interventor"), contrato_obra.get("nombre_supervisor")),
        "supervisor": _primero_no_vacio(acta_inicio.get("nombre_firma_supervisor"), contrato_obra.get("nombre_supervisor")),
        "plazo_inicial": _plazo_texto(acta_inicio, contrato_obra),
        "fecha_inicio": _fecha_inicio(acta_inicio),
        "fecha_vencimiento_inicial": fecha_inicial,
        "fecha_vencimiento_actual": _fecha_vencimiento_actual(fecha_inicial, prorrogas, suspensiones),
        "valor_inicial": valor_inicial,
        "valor_adiciones": valor_adiciones,
        "valor_acumulado": round(valor_inicial + valor_adiciones, 2),
        "salario_minimo": salario_minimo,
        "valor_contrato_smmlv": valor_contrato_smmlv,
        "cincuenta_smmlv": cincuenta_smmlv,
    }


# ==========================================================
# Tablas desde hoja 24
# ==========================================================
def _tabla_prorrogas(control_obra):
    rows = control_obra.get("prorrogas_rows", []) or []
    filas = []
    for fila in rows:
        if not isinstance(fila, dict):
            continue
        if not (_texto(fila.get("PRÓRROGA No.")) or _parse_fecha(fila.get("FECHA"))):
            continue
        filas.append(
            {
                "PRÓRROGA No.": _texto(fila.get("PRÓRROGA No.")),
                "FECHA": _fecha_texto(fila.get("FECHA")),
                "DESDE": _fecha_texto(fila.get("DESDE")),
                "HASTA": _fecha_texto(fila.get("HASTA")),
                "NUEVA DURACIÓN": _texto(fila.get("NUEVA DURACIÓN")),
                "NUEVA FECHA DE TERMINACIÓN": _fecha_texto(fila.get("NUEVA FECHA DE TERMINACIÓN")),
            }
        )
    return filas


def _tabla_suspensiones(control_obra):
    rows = control_obra.get("suspensiones_rows", []) or []
    filas = []
    for fila in rows:
        if not isinstance(fila, dict):
            continue
        if not (_texto(fila.get("ACTA DE SUSPENSIÓN No.")) or _texto(fila.get("ACTA DE AMPLIACIÓN SUSPENSIÓN No.")) or _parse_fecha(fila.get("FECHA DEL ACTA"))):
            continue
        filas.append(
            {
                "ACTA DE SUSPENSIÓN No.": _texto(fila.get("ACTA DE SUSPENSIÓN No.")),
                "ACTA DE AMPLIACIÓN SUSPENSIÓN No.": _texto(fila.get("ACTA DE AMPLIACIÓN SUSPENSIÓN No.")),
                "FECHA DEL ACTA": _fecha_texto(fila.get("FECHA DEL ACTA")),
                "DESDE": _fecha_texto(fila.get("DESDE")),
                "HASTA": _fecha_texto(fila.get("HASTA")),
                "PERIODO DE SUSPENSIÓN": _texto(fila.get("PERIODO DE SUSPENSIÓN")),
                "NUEVA FECHA DE FINALIZACIÓN": _fecha_texto(fila.get("NUEVA FECHA DE FINALIZACIÓN")),
            }
        )
    return filas


def _tabla_adiciones(control_obra):
    rows = control_obra.get("adiciones_rows", []) or []
    filas = []
    for fila in rows:
        if not isinstance(fila, dict):
            continue
        if not (_texto(fila.get("ADICIONAL No.")) or _safe_float(fila.get("VALOR"), 0.0) > 0):
            continue
        filas.append(
            {
                "ADICIONAL No.": _texto(fila.get("ADICIONAL No.")),
                "FECHA": _fecha_texto(fila.get("FECHA")),
                "VALOR": _safe_float(fila.get("VALOR"), 0.0),
                "SMMLV DEL AÑO DE LA ADICIÓN": _safe_float(fila.get("SMMLV DEL AÑO DE LA ADICIÓN"), 0.0),
                "ADICIÓN EN SALARIOS MÍNIMOS": _safe_float(fila.get("ADICIÓN EN SALARIOS MÍNIMOS"), 0.0),
                "VALOR ACUMULADO DEL CONTRATO": _safe_float(fila.get("VALOR ACUMULADO DEL CONTRATO"), 0.0),
            }
        )
    return filas


# ==========================================================
# Información derivada de seguimiento físico y financiero
# ==========================================================
def _ultimo_corte_seguimiento(seguimiento_fisico):
    seguimientos = seguimiento_fisico.get("seguimientos_fisicos", {}) or {}
    if not isinstance(seguimientos, dict) or not seguimientos:
        return {}
    claves = sorted(seguimientos.keys())
    corte = seguimientos.get(claves[-1], {})
    return corte if isinstance(corte, dict) else {}


def _avance_fisico(seguimiento_fisico):
    corte = _ultimo_corte_seguimiento(seguimiento_fisico)
    avance_general = corte.get("avance_general", []) or []
    fila = avance_general[0] if isinstance(avance_general, list) and avance_general and isinstance(avance_general[0], dict) else {}

    return {
        "FECHA DE CORTE": _fecha_texto(corte.get("fecha_corte")),
        "% PROGRAMADO": _safe_float(fila.get("% PROGRAMADO"), 0.0),
        "% EJECUTADO": _safe_float(fila.get("% EJECUTADO"), 0.0),
        "$ PROGRAMADO": _safe_float(fila.get("$ PROGRAMADO"), 0.0),
        "$ EJECUTADO": _safe_float(fila.get("$ EJECUTADO"), 0.0),
    }


def _avance_inversion_obra(seguimiento_fisico):
    avance = _avance_fisico(seguimiento_fisico)
    return {
        "INVERSIÓN PROGRAMADA ACUMULADA": avance["$ PROGRAMADO"],
        "INVERSIÓN EJECUTADA ACUMULADA": avance["$ EJECUTADO"],
    }


def _resumen_financiero(control_obra, plan_anticipo):
    anticipo_rows = control_obra.get("anticipo_rows", []) or []
    pagos = control_obra.get("pagos_rows", []) or []

    valor_anticipo = _valor_anticipo(control_obra, plan_anticipo)
    valor_amortizado = 0.0
    saldo_amortizar = valor_anticipo

    if isinstance(anticipo_rows, list):
        for fila in anticipo_rows:
            if isinstance(fila, dict):
                valor_amortizado += _safe_float(fila.get("VALOR AMORTIZADO"), 0.0)
                saldo_amortizar = _safe_float(fila.get("SALDO"), saldo_amortizar)

    valor_facturado = _sumar_valores(pagos, "VALOR FACTURADO")
    ultimo_mes = ""
    saldo_por_ejecutar = 0.0
    for fila in pagos:
        if isinstance(fila, dict):
            if _parse_fecha(fila.get("FECHA")):
                ultimo_mes = _fecha_texto(fila.get("FECHA"))
            saldo_por_ejecutar = _safe_float(fila.get("PENDIENTE POR FACTURAR"), saldo_por_ejecutar)

    return {
        "VALOR ANTICIPO OTORGADO": valor_anticipo,
        "VALOR AMORTIZADO": round(valor_amortizado, 2),
        "SALDO POR AMORTIZAR": round(saldo_amortizar, 2),
        "VALOR FACTURADO": round(valor_facturado, 2),
        "ÚLTIMO MES FACTURADO": ultimo_mes,
        "SALDO POR EJECUTAR": round(saldo_por_ejecutar, 2),
    }


# ==========================================================
# Estado editable de esta hoja
# ==========================================================
def _filas_modificaciones(rows):
    filas = []
    for fila in rows or []:
        base = {"MODIFICACIÓN No.": "", "FECHA": None, "BREVE DESCRIPCIÓN": ""}
        if isinstance(fila, dict):
            base["MODIFICACIÓN No."] = _texto(fila.get("MODIFICACIÓN No."))
            base["FECHA"] = _parse_fecha(fila.get("FECHA"))
            base["BREVE DESCRIPCIÓN"] = _texto(fila.get("BREVE DESCRIPCIÓN"))
        filas.append(base)
    if not filas:
        filas.append({"MODIFICACIÓN No.": "", "FECHA": None, "BREVE DESCRIPCIÓN": ""})
    return filas


def _filas_discriminacion(rows):
    filas = []
    for fila in rows or []:
        base = {
            "DESCRIPCIÓN": "",
            "VALOR INICIAL": 0.0,
            "VALOR ACTUALIZADO": 0.0,
            "VALOR SOLICITADO": 0.0,
            "VALOR ACTUALIZADO PRESENTE SOLICITUD": 0.0,
        }
        if isinstance(fila, dict):
            base["DESCRIPCIÓN"] = _texto(fila.get("DESCRIPCIÓN"))
            base["VALOR INICIAL"] = _safe_float(fila.get("VALOR INICIAL"), 0.0)
            base["VALOR ACTUALIZADO"] = _safe_float(fila.get("VALOR ACTUALIZADO"), 0.0)
            base["VALOR SOLICITADO"] = _safe_float(fila.get("VALOR SOLICITADO"), 0.0)
            base["VALOR ACTUALIZADO PRESENTE SOLICITUD"] = round(base["VALOR ACTUALIZADO"] + base["VALOR SOLICITADO"], 2)
        filas.append(base)

    if not filas:
        filas.append({
            "DESCRIPCIÓN": "",
            "VALOR INICIAL": 0.0,
            "VALOR ACTUALIZADO": 0.0,
            "VALOR SOLICITADO": 0.0,
            "VALOR ACTUALIZADO PRESENTE SOLICITUD": 0.0,
        })
    return filas


def _estado_vacio():
    return {
        "fecha_solicitud": date.today().isoformat(),
        "tipo_solicitud": ["ADICIÓN"],
        "modificaciones_rows": [],
        "fecha_vencimiento_actual_manual": "",
        "valor_acumulado_manual": "",
        "comunicacion_no": "",
        "comunicacion_fecha": date.today().isoformat(),
        "objeto_solicitud": "",
        "alcance_actividades": "",
        "discriminacion_rows": [],
        "valor_amortizado_manual": "",
        "valor_facturado_fecha": date.today().isoformat(),
        "valor_facturado_manual": "",
        "ultimo_mes_facturado_manual": date.today().isoformat(),
        "procesos_multas_sanciones": "",
        "justificacion_tecnica": "",
        "justificacion_juridica": "",
        "justificacion_financiera": "",
        "justificacion_presupuestal": "",
        "justificacion_ambiental_social_predial": "",
        "otras_justificaciones": "",
        "otros_documentos": "",
    }


def _inicializar_estado():
    group_id_actual = _texto(st.session_state.get("group_id"))
    cache_group = _texto(st.session_state.get("_modificacion_prorroga_group"))

    if cache_group != group_id_actual or "modificacion_prorroga_datos" not in st.session_state:
        cargado = cargar_estado(CLAVE_GUARDADO) or {}
        if not isinstance(cargado, dict):
            cargado = {}
        base = _estado_vacio()
        base.update(cargado)
        base["modificaciones_rows"] = _filas_modificaciones(base.get("modificaciones_rows", []))
        base["discriminacion_rows"] = _filas_discriminacion(base.get("discriminacion_rows", []))
        st.session_state["modificacion_prorroga_datos"] = base
        st.session_state["_modificacion_prorroga_group"] = group_id_actual


def _guardar():
    guardar_estado(CLAVE_GUARDADO, st.session_state["modificacion_prorroga_datos"])


# ==========================================================
# Word helpers
# ==========================================================
def _set_doc_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(8)


def _p(doc, texto="", bold=False, align=None, size=8):
    par = doc.add_paragraph()
    run = par.add_run(str(texto))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if align:
        par.alignment = align
    return par


def _cell_text(cell, texto, bold=False, align=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    par = cell.paragraphs[0]
    run = par.add_run(str(texto))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(8)
    if align:
        par.alignment = align


def _valor_word(valor, col=""):
    if isinstance(valor, (date, datetime)):
        return _fecha_texto(valor)
    if isinstance(valor, (int, float)) and ("VALOR" in col or "$" in col or "SALDO" in col or "INVERSIÓN" in col or "ANTICIPO" in col):
        return _moneda(valor)
    return _texto(valor)


def _tabla_simple(doc, filas):
    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    for etiqueta, valor in filas:
        cells = tabla.add_row().cells
        _cell_text(cells[0], etiqueta, bold=True)
        _cell_text(cells[1], valor)
    return tabla


def _tabla_dataframe(doc, df):
    if df is None or df.empty:
        _p(doc, "Sin información registrada.")
        return None

    cols = list(df.columns)
    tabla = doc.add_table(rows=1, cols=len(cols))
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(cols):
        _cell_text(tabla.rows[0].cells[i], col, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for _, fila in df.iterrows():
        cells = tabla.add_row().cells
        for i, col in enumerate(cols):
            _cell_text(cells[i], _valor_word(fila.get(col, ""), col), align=WD_ALIGN_PARAGRAPH.CENTER)
    return tabla


def _generar_word(
    generales,
    datos,
    df_modificaciones,
    df_prorrogas,
    df_suspensiones,
    df_adiciones,
    df_discriminacion,
    df_avance_fisico,
    df_avance_inversion,
    df_resumen_financiero,
):
    doc = Document()
    _set_doc_defaults(doc)

    tipos = ", ".join(datos.get("tipo_solicitud", []))
    _p(doc, "SOLICITUD DE ADICIÓN Y/O MODIFICACIÓN Y/O PRÓRROGA CONTRATO DE OBRA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    _p(doc, f"FECHA: {_fecha_texto(datos.get('fecha_solicitud'))}", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=8)
    _p(doc, f"TIPO DE SOLICITUD: {tipos}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

    _p(doc, "")
    _p(doc, "RESPONSABILIDAD", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "1. Mediante la suscripción del presente formato, el contratista y el interventor asumen plena responsabilidad por la veracidad de la información contenida.")
    _p(doc, "2. El Interventor debe justificar las causas que originan la solicitud de la adición y/o modificación y/o prórroga del contrato de obra.")

    _p(doc, "")
    _p(doc, "ANTECEDENTES", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_simple(
        doc,
        [
            ("CONTRATO DE OBRA No.", generales.get("numero_contrato", "")),
            ("CONTRATISTA", generales.get("contratista", "")),
            ("OBJETO DEL CONTRATO DE OBRA", generales.get("objeto", "")),
            ("CONTRATO DE INTERVENTORÍA No.", generales.get("contrato_interventoria", "")),
            ("INTERVENTOR", generales.get("interventor", "")),
            ("PLAZO INICIAL DEL CONTRATO DE OBRA", generales.get("plazo_inicial", "")),
            ("FECHA DE INICIO DEL CONTRATO", _fecha_texto(generales.get("fecha_inicio"))),
            ("FECHA DE VENCIMIENTO ACTUAL", _fecha_texto(datos.get("fecha_vencimiento_actual_manual")) or _fecha_texto(generales.get("fecha_vencimiento_actual"))),
            ("VALOR INICIAL DEL CONTRATO", _moneda(generales.get("valor_inicial", 0.0))),
            ("50% EN SMMLV", generales.get("cincuenta_smmlv", 0.0)),
            ("VALOR ACUMULADO DEL CONTRATO", _moneda(_safe_float(datos.get("valor_acumulado_manual"), generales.get("valor_acumulado", 0.0)))),
        ],
    )

    _p(doc, "")
    _p(doc, "MODIFICACIONES", bold=True)
    _tabla_dataframe(doc, df_modificaciones)

    _p(doc, "")
    _p(doc, "PRÓRROGAS", bold=True)
    _tabla_dataframe(doc, df_prorrogas)

    _p(doc, "")
    _p(doc, "SUSPENSIONES Y AMPLIACIONES DE SUSPENSIÓN", bold=True)
    _tabla_dataframe(doc, df_suspensiones)

    _p(doc, "")
    _p(doc, "VALOR ADICIONES", bold=True)
    _tabla_dataframe(doc, df_adiciones)

    _p(doc, "")
    _p(doc, "OBJETO DE LA SOLICITUD", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, f"Mediante comunicación escrita No. {datos.get('comunicacion_no', '')} de fecha {_fecha_texto(datos.get('comunicacion_fecha'))} el Contratista de Obra solicita lo siguiente:")
    _p(doc, datos.get("objeto_solicitud", ""))

    _p(doc, "")
    _p(doc, "ALCANCE DE LAS ACTIVIDADES A DESARROLLAR", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, datos.get("alcance_actividades", ""))

    _p(doc, "")
    _p(doc, "DISCRIMINACIÓN DEL VALOR ADICIONAL Y/O REDISTRIBUCIÓN DE RECURSOS SOLICITADOS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_dataframe(doc, df_discriminacion)

    _p(doc, "")
    _p(doc, "AVANCE FÍSICO DE OBRA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_dataframe(doc, df_avance_fisico)

    _p(doc, "")
    _p(doc, "AVANCE INVERSIÓN DE OBRA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_dataframe(doc, df_avance_inversion)

    _p(doc, "")
    _p(doc, "RESUMEN FINANCIERO CONTRATO DE OBRA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_dataframe(doc, df_resumen_financiero)

    _p(doc, "")
    _p(doc, "PROCESOS DE MULTAS Y SANCIONES", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, datos.get("procesos_multas_sanciones", ""))

    _p(doc, "")
    _p(doc, "JUSTIFICACIÓN DE LA INTERVENTORÍA PARA APROBACIÓN DE LA PRESENTE SOLICITUD", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "TÉCNICA:", bold=True)
    _p(doc, datos.get("justificacion_tecnica", ""))
    _p(doc, "JURÍDICA:", bold=True)
    _p(doc, datos.get("justificacion_juridica", ""))
    _p(doc, "FINANCIERA:", bold=True)
    _p(doc, datos.get("justificacion_financiera", ""))
    _p(doc, "PRESUPUESTAL:", bold=True)
    _p(doc, datos.get("justificacion_presupuestal", ""))
    _p(doc, "AMBIENTAL Y/O SOCIAL Y/O PREDIAL Y/O SOSTENIBILIDAD:", bold=True)
    _p(doc, datos.get("justificacion_ambiental_social_predial", ""))
    _p(doc, "OTRAS JUSTIFICACIONES:", bold=True)
    _p(doc, datos.get("otras_justificaciones", ""))

    _p(doc, "")
    _p(doc, "OTROS DOCUMENTOS QUE SOPORTAN LA SOLICITUD", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, datos.get("otros_documentos", ""))

    _p(doc, "")
    _p(doc, "NOTAS", bold=True)
    _p(doc, "3. Es responsabilidad de la Interventoría presentar las justificaciones técnica, jurídica, financiera, presupuestal, ambiental, social, predial, sostenibilidad y otros según el caso que respalden la solicitud de adición y/o prórroga y/o modificación, las cuales serán presentadas ante el Comité de Contratación del INVIAS.")
    _p(doc, "4. Cuando se trate de contrato adicional en valor, la justificación debe contener la distribución de los recursos a adicionar entre las actividades pactadas en el contrato, indicando el valor correspondiente al IVA. Es responsabilidad del Contratista y del Interventor garantizar que la distribución de recursos propuesta con la adición solicitada cubra todas las actividades requeridas.")
    _p(doc, "5. La justificación presentada por la Interventoría debe tener en cuenta: las cláusulas contractuales, el pliego de condiciones y la matriz de riesgos del contrato de obra.")

    _p(doc, "")
    _p(doc, "FIRMAS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    firmas = doc.add_table(rows=3, cols=3)
    firmas.style = "Table Grid"
    firmas.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cell_text(firmas.cell(0, 0), "CONTRATISTA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(0, 1), "INTERVENTORÍA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(0, 2), "SUPERVISOR", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(1, 0), "FIRMA", align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(1, 1), "FIRMA", align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(1, 2), "FIRMA", align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(2, 0), generales.get("contratista", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(2, 1), generales.get("interventor", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(2, 2), generales.get("supervisor", ""), align=WD_ALIGN_PARAGRAPH.CENTER)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ==========================================================
# Carga de datos
# ==========================================================
st.set_page_config(page_title="Modificación y prórroga", layout="wide")
st.title("28. Solicitud de adición, modificación y/o prórroga")

acta_inicio = _leer_estado("acta_inicio_obra")
contrato_obra = _leer_estado("contrato_obra")
contrato_interventoria = _leer_estado("contrato_interventoria")
control_obra = _leer_estado("control_obra")
seguimiento_fisico = _leer_estado("seguimiento_fisico")
plan_anticipo = _leer_estado("plan_inversion_anticipo")

_inicializar_estado()
datos = st.session_state["modificacion_prorroga_datos"]

generales = _datos_generales(acta_inicio, contrato_obra, contrato_interventoria, control_obra)
avance_fisico = _avance_fisico(seguimiento_fisico)
avance_inversion = _avance_inversion_obra(seguimiento_fisico)
resumen_financiero = _resumen_financiero(control_obra, plan_anticipo)

# ==========================================================
# Interfaz
# ==========================================================
st.markdown("### TIPO DE SOLICITUD")
col_fecha_solicitud, col_tipo_solicitud = st.columns([1, 2])
with col_fecha_solicitud:
    datos["fecha_solicitud"] = st.date_input(
        "FECHA",
        value=_fecha_input(datos.get("fecha_solicitud")),
        format="DD/MM/YYYY",
        key="mod_prorroga_fecha_solicitud",
    ).isoformat()
with col_tipo_solicitud:
    datos["tipo_solicitud"] = st.multiselect(
        "Seleccione el tipo de solicitud",
        options=["ADICIÓN", "MODIFICACIÓN", "PRÓRROGA"],
        default=datos.get("tipo_solicitud", ["ADICIÓN"]),
        key="mod_prorroga_tipo_solicitud",
    )

st.markdown("### RESPONSABILIDAD")
st.write("1. Mediante la suscripción del presente formato, el contratista y el interventor asumen plena responsabilidad por la veracidad de la información contenida.")
st.write("2. El Interventor debe justificar las causas que originan la solicitud de la adición y/o modificación y/o prórroga del contrato de obra.")

st.markdown("### ANTECEDENTES")
col1, col2 = st.columns(2)
with col1:
    st.text_input("CONTRATO DE OBRA No.", value=generales["numero_contrato"], disabled=True)
    st.text_input("CONTRATISTA", value=generales["contratista"], disabled=True)
    st.text_input("CONTRATO DE INTERVENTORÍA No.", value=generales["contrato_interventoria"], disabled=True)
    st.text_input("INTERVENTOR", value=generales["interventor"], disabled=True)
with col2:
    st.text_input("PLAZO INICIAL DEL CONTRATO DE OBRA", value=generales["plazo_inicial"], disabled=True)
    st.date_input("FECHA DE INICIO DEL CONTRATO", value=_fecha_input(generales["fecha_inicio"]), disabled=True, format="DD/MM/YYYY")
    datos["fecha_vencimiento_actual_manual"] = st.date_input(
        "FECHA DE VENCIMIENTO ACTUAL",
        value=_fecha_input(datos.get("fecha_vencimiento_actual_manual") or generales["fecha_vencimiento_actual"]),
        format="DD/MM/YYYY",
        key="mod_prorroga_fecha_vencimiento_actual",
    ).isoformat()

st.text_area("OBJETO DEL CONTRATO DE OBRA", value=generales["objeto"], disabled=True, height=100)

st.markdown("#### MODIFICACIONES")
df_modificaciones = pd.DataFrame(_filas_modificaciones(datos.get("modificaciones_rows", [])))
modificaciones_editadas = st.data_editor(
    df_modificaciones,
    hide_index=True,
    width="stretch",
    num_rows="dynamic",
    key="mod_prorroga_modificaciones_editor",
    column_config={
        "MODIFICACIÓN No.": st.column_config.TextColumn("MODIFICACIÓN No."),
        "FECHA": st.column_config.DateColumn("FECHA", format="DD/MM/YYYY"),
        "BREVE DESCRIPCIÓN": st.column_config.TextColumn("BREVE DESCRIPCIÓN"),
    },
)
datos["modificaciones_rows"] = _filas_modificaciones(modificaciones_editadas.to_dict("records"))

st.markdown("#### PRÓRROGAS")
df_prorrogas = pd.DataFrame(_tabla_prorrogas(control_obra))
st.dataframe(df_prorrogas, hide_index=True, width="stretch")

st.markdown("#### SUSPENSIONES Y AMPLIACIONES DE SUSPENSIÓN")
df_suspensiones = pd.DataFrame(_tabla_suspensiones(control_obra))
st.dataframe(df_suspensiones, hide_index=True, width="stretch")

col_valor1, col_valor2, col_valor3 = st.columns(3)
with col_valor1:
    st.number_input("VALOR INICIAL DEL CONTRATO", value=float(generales["valor_inicial"]), disabled=True, format="%.2f")
with col_valor2:
    st.number_input("50% EN SMMLV", value=float(generales["cincuenta_smmlv"]), disabled=True, format="%.4f")
with col_valor3:
    datos["valor_acumulado_manual"] = st.number_input(
        "VALOR ACUMULADO DEL CONTRATO",
        value=float(_safe_float(datos.get("valor_acumulado_manual"), generales["valor_acumulado"])),
        format="%.2f",
        key="mod_prorroga_valor_acumulado_manual",
    )

st.markdown("#### VALOR ADICIONES")
df_adiciones = pd.DataFrame(_tabla_adiciones(control_obra))
st.dataframe(df_adiciones, hide_index=True, width="stretch")

st.markdown("### OBJETO DE LA SOLICITUD")
col_com1, col_com2 = st.columns([1, 1])
with col_com1:
    datos["comunicacion_no"] = st.text_input(
        "Comunicación escrita No.",
        value=_texto(datos.get("comunicacion_no")),
        key="mod_prorroga_comunicacion_no",
    )
with col_com2:
    datos["comunicacion_fecha"] = st.date_input(
        "Fecha de la comunicación",
        value=_fecha_input(datos.get("comunicacion_fecha")),
        format="DD/MM/YYYY",
        key="mod_prorroga_comunicacion_fecha",
    ).isoformat()

datos["objeto_solicitud"] = st.text_area(
    "Mediante comunicación escrita el Contratista de Obra solicita lo siguiente",
    value=_texto(datos.get("objeto_solicitud")),
    height=150,
    key="mod_prorroga_objeto_solicitud",
)

st.markdown("### ALCANCE DE LAS ACTIVIDADES A DESARROLLAR")
datos["alcance_actividades"] = st.text_area(
    "Alcance de las actividades",
    value=_texto(datos.get("alcance_actividades")),
    height=150,
    key="mod_prorroga_alcance_actividades",
)

st.markdown("### DISCRIMINACIÓN DEL VALOR ADICIONAL Y/O REDISTRIBUCIÓN DE RECURSOS SOLICITADOS")
df_discriminacion = pd.DataFrame(_filas_discriminacion(datos.get("discriminacion_rows", [])))
discriminacion_editada = st.data_editor(
    df_discriminacion,
    hide_index=True,
    width="stretch",
    num_rows="dynamic",
    key="mod_prorroga_discriminacion_editor",
    column_config={
        "DESCRIPCIÓN": st.column_config.TextColumn("DESCRIPCIÓN"),
        "VALOR INICIAL": st.column_config.NumberColumn("VALOR INICIAL", format="$ %.2f"),
        "VALOR ACTUALIZADO": st.column_config.NumberColumn("VALOR ACTUALIZADO", format="$ %.2f"),
        "VALOR SOLICITADO": st.column_config.NumberColumn("VALOR SOLICITADO", format="$ %.2f"),
        "VALOR ACTUALIZADO PRESENTE SOLICITUD": st.column_config.NumberColumn("VALOR ACTUALIZADO PRESENTE SOLICITUD", format="$ %.2f", disabled=True),
    },
)
datos["discriminacion_rows"] = _filas_discriminacion(discriminacion_editada.to_dict("records"))
df_discriminacion = pd.DataFrame(datos["discriminacion_rows"])

st.markdown("### AVANCE FÍSICO DE OBRA")
st.write(
    f"Hasta la fecha {avance_fisico['FECHA DE CORTE']}, la obra presenta un avance físico ejecutado del "
    f"{avance_fisico['% EJECUTADO']:.4f}% contra el avance programado del {avance_fisico['% PROGRAMADO']:.4f}%."
)
df_avance_fisico = pd.DataFrame([avance_fisico])
st.dataframe(df_avance_fisico, hide_index=True, width="stretch")

st.markdown("### AVANCE INVERSIÓN DE OBRA")
df_avance_inversion = pd.DataFrame([avance_inversion])
st.dataframe(df_avance_inversion, hide_index=True, width="stretch")

st.markdown("### RESUMEN FINANCIERO CONTRATO DE OBRA")
resumen_financiero["VALOR AMORTIZADO"] = _safe_float(datos.get("valor_amortizado_manual"), resumen_financiero.get("VALOR AMORTIZADO", 0.0))
resumen_financiero["SALDO POR AMORTIZAR"] = round(resumen_financiero.get("VALOR ANTICIPO OTORGADO", 0.0) - resumen_financiero["VALOR AMORTIZADO"], 2)
resumen_financiero["VALOR FACTURADO"] = _safe_float(datos.get("valor_facturado_manual"), resumen_financiero.get("VALOR FACTURADO", 0.0))
resumen_financiero["ÚLTIMO MES FACTURADO"] = _fecha_texto(datos.get("ultimo_mes_facturado_manual"))
resumen_financiero["SALDO POR EJECUTAR"] = round(generales["valor_inicial"] - resumen_financiero["VALOR FACTURADO"], 2)

c_fin1, c_fin2, c_fin3 = st.columns(3)
with c_fin1:
    st.number_input("Valor anticipo otorgado", value=float(resumen_financiero["VALOR ANTICIPO OTORGADO"]), disabled=True, format="%.2f")
with c_fin2:
    datos["valor_amortizado_manual"] = st.number_input("Valor amortizado", value=float(resumen_financiero["VALOR AMORTIZADO"]), format="%.2f", key="mod_prorroga_valor_amortizado")
with c_fin3:
    st.number_input("Saldo por amortizar", value=float(resumen_financiero["SALDO POR AMORTIZAR"]), disabled=True, format="%.2f")

c_fin4, c_fin5, c_fin6 = st.columns(3)
with c_fin4:
    datos["valor_facturado_fecha"] = st.date_input("Fecha valor facturado", value=_fecha_input(datos.get("valor_facturado_fecha")), format="DD/MM/YYYY", key="mod_prorroga_valor_facturado_fecha").isoformat()
with c_fin5:
    datos["valor_facturado_manual"] = st.number_input("Valor facturado", value=float(resumen_financiero["VALOR FACTURADO"]), format="%.2f", key="mod_prorroga_valor_facturado")
with c_fin6:
    datos["ultimo_mes_facturado_manual"] = st.date_input("Último mes facturado", value=_fecha_input(datos.get("ultimo_mes_facturado_manual")), format="DD/MM/YYYY", key="mod_prorroga_ultimo_mes_facturado").isoformat()

st.number_input("Saldo por ejecutar (Incluido IVA)", value=float(resumen_financiero["SALDO POR EJECUTAR"]), disabled=True, format="%.2f")
df_resumen_financiero = pd.DataFrame([resumen_financiero])

st.markdown("### PROCESOS DE MULTAS Y SANCIONES")
datos["procesos_multas_sanciones"] = st.text_area(
    "Procesos de multas y sanciones",
    value=_texto(datos.get("procesos_multas_sanciones")),
    height=120,
    key="mod_prorroga_multas_sanciones",
)

st.markdown("### JUSTIFICACIÓN DE LA INTERVENTORÍA PARA APROBACIÓN DE LA PRESENTE SOLICITUD")
st.caption("La justificación debe describirse de manera precisa, detallada y completa.")
datos["justificacion_tecnica"] = st.text_area("TÉCNICA", value=_texto(datos.get("justificacion_tecnica")), height=100, key="mod_prorroga_just_tecnica")
datos["justificacion_juridica"] = st.text_area("JURÍDICA", value=_texto(datos.get("justificacion_juridica")), height=100, key="mod_prorroga_just_juridica")
datos["justificacion_financiera"] = st.text_area("FINANCIERA", value=_texto(datos.get("justificacion_financiera")), height=100, key="mod_prorroga_just_financiera")
datos["justificacion_presupuestal"] = st.text_area("PRESUPUESTAL", value=_texto(datos.get("justificacion_presupuestal")), height=100, key="mod_prorroga_just_presupuestal")
datos["justificacion_ambiental_social_predial"] = st.text_area("AMBIENTAL Y/O SOCIAL Y/O PREDIAL Y/O SOSTENIBILIDAD", value=_texto(datos.get("justificacion_ambiental_social_predial")), height=100, key="mod_prorroga_just_ambiental")
datos["otras_justificaciones"] = st.text_area("OTRAS JUSTIFICACIONES QUE AMERITEN LA SOLICITUD", value=_texto(datos.get("otras_justificaciones")), height=100, key="mod_prorroga_just_otras")

st.markdown("### OTROS DOCUMENTOS QUE SOPORTAN LA SOLICITUD")
datos["otros_documentos"] = st.text_area(
    "Otros documentos que soportan la solicitud",
    value=_texto(datos.get("otros_documentos")),
    height=120,
    key="mod_prorroga_otros_documentos",
)

st.markdown("### NOTAS")
st.write("3. Es responsabilidad de la Interventoría presentar las justificaciones técnica, jurídica, financiera, presupuestal, ambiental, social, predial, sostenibilidad y otros según el caso que respalden la solicitud de adición y/o prórroga y/o modificación, las cuales serán presentadas ante el Comité de Contratación del INVIAS.")
st.write("4. Cuando se trate de contrato adicional en valor, la justificación debe contener la distribución de los recursos a adicionar entre las actividades pactadas en el contrato, indicando el valor correspondiente al IVA. Es responsabilidad del Contratista y del Interventor garantizar que la distribución de recursos propuesta con la adición solicitada cubra todas las actividades requeridas.")
st.write("5. La justificación presentada por la Interventoría debe tener en cuenta: las cláusulas contractuales, el pliego de condiciones y la matriz de riesgos del contrato de obra.")

st.markdown("### FIRMAS")
df_firmas = pd.DataFrame([
    {
        "CONTRATISTA": generales.get("contratista", ""),
        "INTERVENTORÍA": generales.get("interventor", ""),
        "SUPERVISOR": generales.get("supervisor", ""),
    }
])
st.dataframe(df_firmas, hide_index=True, width="stretch")

col_guardar, col_word = st.columns([1, 1])
with col_guardar:
    if st.button("💾 Guardar solicitud", type="primary", key="mod_prorroga_guardar"):
        _guardar()
        st.success("Solicitud guardada correctamente.")

with col_word:
    word = _generar_word(
        generales,
        datos,
        pd.DataFrame(datos.get("modificaciones_rows", [])),
        df_prorrogas,
        df_suspensiones,
        df_adiciones,
        pd.DataFrame(datos.get("discriminacion_rows", [])),
        df_avance_fisico,
        df_avance_inversion,
        df_resumen_financiero,
    )
    st.download_button(
        "📄 Descargar Word",
        data=word,
        file_name="solicitud_modificacion_prorroga.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="mod_prorroga_descargar_word",
        width="stretch",
    )
