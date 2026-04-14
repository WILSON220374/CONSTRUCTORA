import io
import pandas as pd
import streamlit as st
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from supabase_state import cargar_estado, guardar_estado, cargar_apus_generados_obra


STORAGE_KEY = "informe_anexos_config"


def _safe_str(value) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _safe_bool(value, default=False) -> bool:
    if isinstance(value, bool):
        return value
    if value is None:
        return default
    if isinstance(value, str):
        return value.strip().lower() in ("true", "1", "si", "sí", "yes")
    return bool(value)


def _normalizar_cfg(data: dict) -> dict:
    if not isinstance(data, dict):
        data = {}

    return {
        "ultima_actualizacion": data.get("ultima_actualizacion"),
        "incluye_apus_obra": _safe_bool(data.get("incluye_apus_obra", True), True),
        "incluye_apus_consultoria": _safe_bool(data.get("incluye_apus_consultoria", True), True),
        "incluye_estudio_mercado": _safe_bool(data.get("incluye_estudio_mercado", True), True),
        "incluye_aiu": _safe_bool(data.get("incluye_aiu", True), True),
        "incluye_factor_multiplicador": _safe_bool(data.get("incluye_factor_multiplicador", True), True),
        "notas_apus_obra": _safe_str(data.get("notas_apus_obra")),
        "notas_apus_consultoria": _safe_str(data.get("notas_apus_consultoria")),
        "notas_estudio_mercado": _safe_str(data.get("notas_estudio_mercado")),
        "notas_aiu": _safe_str(data.get("notas_aiu")),
        "notas_factor_multiplicador": _safe_str(data.get("notas_factor_multiplicador")),
    }


def _cargar_cfg():
    if "informe_anexos_config" in st.session_state and isinstance(st.session_state["informe_anexos_config"], dict):
        st.session_state["informe_anexos_config"] = _normalizar_cfg(st.session_state["informe_anexos_config"])
        return

    try:
        data = cargar_estado(STORAGE_KEY) or {}
    except Exception:
        data = {}

    st.session_state["informe_anexos_config"] = _normalizar_cfg(data)


def _guardar_cfg():
    cfg = _normalizar_cfg(st.session_state.get("informe_anexos_config", {}))
    cfg["ultima_actualizacion"] = datetime.now().isoformat()
    st.session_state["informe_anexos_config"] = cfg
    guardar_estado(STORAGE_KEY, cfg)


def _tipo_proyecto() -> str:
    valor_directo = _safe_str(st.session_state.get("tipo_presupuesto_proyecto_crono"))
    if valor_directo in ("Obra", "Consultoría"):
        return valor_directo

    cronograma_datos = st.session_state.get("cronograma_datos", {}) or {}
    valor_cronograma = _safe_str(cronograma_datos.get("tipo_presupuesto_proyecto"))
    if valor_cronograma in ("Obra", "Consultoría"):
        return valor_cronograma

    return "Obra"


@st.cache_data
def _cargar_base_apu_obra():
    return pd.read_excel("data/Copia de APU.xlsx")


def _codigos_item_gober_usados_obra() -> list[str]:
    datos_obra = st.session_state.get("presupuesto_obra_datos")
    if not datos_obra:
        try:
            datos_obra = cargar_estado("presupuesto_obra") or {}
        except Exception:
            datos_obra = {}

    items_state = datos_obra.get("items", {}) or {}
    codigos = []

    if isinstance(items_state, dict):
        for _, item in items_state.items():
            if not isinstance(item, dict):
                continue

            codigo = _safe_str(item.get("item_catalogo"))
            if codigo:
                codigos.append(codigo)

    return sorted(set(codigos))


def _apus_filtrados_obra() -> tuple[list[str], pd.DataFrame]:
    codigos = _codigos_item_gober_usados_obra()

    try:
        df_base = _cargar_base_apu_obra().copy()
    except Exception:
        return codigos, pd.DataFrame()

    if df_base.empty or "cod_actividad" not in df_base.columns:
        return codigos, pd.DataFrame()

    df_base["cod_actividad"] = df_base["cod_actividad"].astype(str).str.strip()

    if not codigos:
        return codigos, pd.DataFrame(columns=df_base.columns)

    df_filtrado = df_base[df_base["cod_actividad"].isin(codigos)].copy()
    return codigos, df_filtrado


def _tabla_apu_tipo(df_fuente: pd.DataFrame, tipo_filtrar: str) -> tuple[pd.DataFrame, float]:
    if df_fuente.empty:
        return pd.DataFrame(), 0.0

    df_tabla = df_fuente[
        df_fuente["Tipo"].astype(str).str.strip().str.upper() == tipo_filtrar
    ][["Descripción", "Unidad", "Valor Unitario", "Cantidad"]].copy()

    if df_tabla.empty:
        return pd.DataFrame(), 0.0

    df_tabla["VALOR TOTAL"] = (
        pd.to_numeric(df_tabla["Valor Unitario"], errors="coerce").fillna(0.0)
        * pd.to_numeric(df_tabla["Cantidad"], errors="coerce").fillna(0.0)
    )

    total = float(df_tabla["VALOR TOTAL"].sum())
    return df_tabla, total


def _items_apu_generado_obra():
    datos_obra = st.session_state.get("presupuesto_obra_datos")
    if not datos_obra:
        try:
            datos_obra = cargar_estado("presupuesto_obra") or {}
        except Exception:
            datos_obra = {}

    items_state = datos_obra.get("items", {}) or {}
    salida = {}

    if isinstance(items_state, dict):
        for item_id, item_data in items_state.items():
            if not isinstance(item_data, dict):
                continue

            fuente = _safe_str(item_data.get("fuente"))
            if fuente == "APU generado":
                salida[str(item_id)] = item_data

    return salida


def _apus_generados_obra_filtrados():
    apus_generados = st.session_state.get("apus_generados_obra")
    if not apus_generados:
        try:
            apus_generados = cargar_apus_generados_obra() or {}
        except Exception:
            apus_generados = {}

    ids_validos = set(_items_apu_generado_obra().keys())
    salida = {}

    if isinstance(apus_generados, dict):
        for item_id, apu in apus_generados.items():
            if str(item_id) in ids_validos and isinstance(apu, dict):
                salida[str(item_id)] = apu

    return salida


def _generar_docx_apus_obra(codigos_usados, df_apus_filtrados, apus_generados_dict) -> io.BytesIO:
    doc = Document()

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tit.add_run("ANEXO APUS OBRA")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    p_cod = doc.add_paragraph()
    p_cod.add_run("Códigos ITEM GOBER usados: ").bold = True
    p_cod.add_run(", ".join(codigos_usados) if codigos_usados else "Ninguno")

    if not df_apus_filtrados.empty:
        codigos_disponibles = sorted(
            {
                str(x).strip()
                for x in df_apus_filtrados["cod_actividad"].dropna().tolist()
                if str(x).strip()
            }
        )

        for codigo_apu in codigos_disponibles:
            df_apu = df_apus_filtrados[
                df_apus_filtrados["cod_actividad"].astype(str).str.strip() == codigo_apu
            ].copy()

            if df_apu.empty:
                continue

            fila_ref = df_apu.iloc[0]
            doc.add_heading(
                f"{_safe_str(fila_ref.get('cod_actividad'))} - {_safe_str(fila_ref.get('actividad'))}",
                level=2,
            )

            p_info = doc.add_paragraph()
            p_info.add_run("Capítulo: ").bold = True
            p_info.add_run(_safe_str(fila_ref.get("capitulo")))
            p_info.add_run(" | Subcapítulo: ").bold = True
            p_info.add_run(_safe_str(fila_ref.get("subcapitulo")))
            p_info.add_run(" | Und. Act: ").bold = True
            p_info.add_run(_safe_str(fila_ref.get("Und. Act")))

            for titulo_tipo, clave_tipo in [("Materiales", "MATERIAL"), ("Equipos", "EQUIPO"), ("Mano de obra", "MANO DE OBRA")]:
                df_tabla, total_tipo = _tabla_apu_tipo(df_apu, clave_tipo)
                doc.add_paragraph()
                p_sub = doc.add_paragraph()
                p_sub.add_run(titulo_tipo).bold = True

                if not df_tabla.empty:
                    table = doc.add_table(rows=1, cols=len(df_tabla.columns))
                    table.style = "Table Grid"
                    hdr = table.rows[0].cells
                    for idx, col in enumerate(df_tabla.columns):
                        hdr[idx].text = str(col)
                        for paragraph in hdr[idx].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

                    for _, row in df_tabla.iterrows():
                        row_cells = table.add_row().cells
                        for idx, col in enumerate(df_tabla.columns):
                            row_cells[idx].text = str(row[col])

                p_total = doc.add_paragraph()
                p_total.add_run(f"Total {titulo_tipo.lower()}: ").bold = True
                p_total.add_run(f"{total_tipo:,.2f}")

            total_apu = (
                _tabla_apu_tipo(df_apu, "MATERIAL")[1]
                + _tabla_apu_tipo(df_apu, "EQUIPO")[1]
                + _tabla_apu_tipo(df_apu, "MANO DE OBRA")[1]
            )
            p_total_apu = doc.add_paragraph()
            p_total_apu.add_run("TOTAL APU: ").bold = True
            p_total_apu.add_run(f"{total_apu:,.2f}")

    if apus_generados_dict:
        doc.add_page_break()
        doc.add_heading("APUS GENERADOS", level=1)

        for _, apu in apus_generados_dict.items():
            doc.add_heading(_safe_str(apu.get("item_label")), level=2)

            p_info = doc.add_paragraph()
            p_info.add_run("APU base: ").bold = True
            p_info.add_run(f"{_safe_str(apu.get('apu_base_codigo'))} - {_safe_str(apu.get('apu_base_actividad'))}")
            p_info.add_run(" | Unidad: ").bold = True
            p_info.add_run(_safe_str(apu.get("unidad_apu")))

            for titulo, clave_lista, clave_total in [
                ("Materiales", "materiales", "total_materiales"),
                ("Equipos", "equipos", "total_equipos"),
                ("Mano de obra", "mano_obra", "total_mano_obra"),
            ]:
                lista = apu.get(clave_lista, []) or []
                doc.add_paragraph()
                p_sub = doc.add_paragraph()
                p_sub.add_run(titulo).bold = True

                if lista:
                    df_tabla = pd.DataFrame(lista)
                    table = doc.add_table(rows=1, cols=len(df_tabla.columns))
                    table.style = "Table Grid"
                    hdr = table.rows[0].cells
                    for idx, col in enumerate(df_tabla.columns):
                        hdr[idx].text = str(col)
                        for paragraph in hdr[idx].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

                    for _, row in df_tabla.iterrows():
                        row_cells = table.add_row().cells
                        for idx, col in enumerate(df_tabla.columns):
                            row_cells[idx].text = str(row[col])

                p_total = doc.add_paragraph()
                p_total.add_run(f"Total {titulo.lower()}: ").bold = True
                p_total.add_run(f"{float(apu.get(clave_total, 0.0) or 0.0):,.2f}")

            p_total_apu = doc.add_paragraph()
            p_total_apu.add_run("TOTAL APU: ").bold = True
            p_total_apu.add_run(f"{float(apu.get('total_apu', 0.0) or 0.0):,.2f}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _generar_excel_apus_obra(codigos_usados, df_apus_filtrados, apus_generados_dict) -> io.BytesIO:
    output = io.BytesIO()

    filas_detalle = []

    if codigos_usados:
        filas_detalle.append(
            {
                "BLOQUE": "CÓDIGOS ITEM GOBER USADOS",
                "ITEM / APU": ", ".join(codigos_usados),
                "TIPO": "",
                "DESCRIPCIÓN": "",
                "UNIDAD": "",
                "VALOR UNITARIO": "",
                "CANTIDAD": "",
                "VALOR TOTAL": "",
            }
        )
        filas_detalle.append(
            {
                "BLOQUE": "",
                "ITEM / APU": "",
                "TIPO": "",
                "DESCRIPCIÓN": "",
                "UNIDAD": "",
                "VALOR UNITARIO": "",
                "CANTIDAD": "",
                "VALOR TOTAL": "",
            }
        )

    if not df_apus_filtrados.empty:
        codigos_disponibles = sorted(
            {
                str(x).strip()
                for x in df_apus_filtrados["cod_actividad"].dropna().tolist()
                if str(x).strip()
            }
        )

        for codigo_apu in codigos_disponibles:
            df_apu = df_apus_filtrados[
                df_apus_filtrados["cod_actividad"].astype(str).str.strip() == codigo_apu
            ].copy()

            if df_apu.empty:
                continue

            fila_ref = df_apu.iloc[0]
            titulo_apu = f"{_safe_str(fila_ref.get('cod_actividad'))} - {_safe_str(fila_ref.get('actividad'))}"

            filas_detalle.append(
                {
                    "BLOQUE": "APU GOBERNACIÓN",
                    "ITEM / APU": titulo_apu,
                    "TIPO": "",
                    "DESCRIPCIÓN": "",
                    "UNIDAD": _safe_str(fila_ref.get("Und. Act")),
                    "VALOR UNITARIO": "",
                    "CANTIDAD": "",
                    "VALOR TOTAL": "",
                }
            )

            total_apu = 0.0

            for tipo_txt, tipo_filtrar in [("MATERIALES", "MATERIAL"), ("EQUIPOS", "EQUIPO"), ("MANO DE OBRA", "MANO DE OBRA")]:
                df_tabla, total_tipo = _tabla_apu_tipo(df_apu, tipo_filtrar)

                filas_detalle.append(
                    {
                        "BLOQUE": "",
                        "ITEM / APU": "",
                        "TIPO": tipo_txt,
                        "DESCRIPCIÓN": "",
                        "UNIDAD": "",
                        "VALOR UNITARIO": "",
                        "CANTIDAD": "",
                        "VALOR TOTAL": "",
                    }
                )

                if not df_tabla.empty:
                    for _, row in df_tabla.iterrows():
                        filas_detalle.append(
                            {
                                "BLOQUE": "",
                                "ITEM / APU": "",
                                "TIPO": tipo_txt,
                                "DESCRIPCIÓN": _safe_str(row.get("Descripción")),
                                "UNIDAD": _safe_str(row.get("Unidad")),
                                "VALOR UNITARIO": float(row.get("Valor Unitario", 0.0) or 0.0),
                                "CANTIDAD": float(row.get("Cantidad", 0.0) or 0.0),
                                "VALOR TOTAL": float(row.get("VALOR TOTAL", 0.0) or 0.0),
                            }
                        )

                filas_detalle.append(
                    {
                        "BLOQUE": "",
                        "ITEM / APU": "",
                        "TIPO": f"TOTAL {tipo_txt}",
                        "DESCRIPCIÓN": "",
                        "UNIDAD": "",
                        "VALOR UNITARIO": "",
                        "CANTIDAD": "",
                        "VALOR TOTAL": total_tipo,
                    }
                )

                total_apu += total_tipo

            filas_detalle.append(
                {
                    "BLOQUE": "",
                    "ITEM / APU": "",
                    "TIPO": "TOTAL APU",
                    "DESCRIPCIÓN": "",
                    "UNIDAD": "",
                    "VALOR UNITARIO": "",
                    "CANTIDAD": "",
                    "VALOR TOTAL": total_apu,
                }
            )
            filas_detalle.append(
                {
                    "BLOQUE": "",
                    "ITEM / APU": "",
                    "TIPO": "",
                    "DESCRIPCIÓN": "",
                    "UNIDAD": "",
                    "VALOR UNITARIO": "",
                    "CANTIDAD": "",
                    "VALOR TOTAL": "",
                }
            )

    if apus_generados_dict:
        for _, apu in apus_generados_dict.items():
            titulo_apu = _safe_str(apu.get("item_label"))

            filas_detalle.append(
                {
                    "BLOQUE": "APU GENERADO",
                    "ITEM / APU": titulo_apu,
                    "TIPO": "",
                    "DESCRIPCIÓN": _safe_str(apu.get("apu_base_codigo")) + " - " + _safe_str(apu.get("apu_base_actividad")),
                    "UNIDAD": _safe_str(apu.get("unidad_apu")),
                    "VALOR UNITARIO": "",
                    "CANTIDAD": "",
                    "VALOR TOTAL": "",
                }
            )

            for tipo_txt, clave_lista, clave_total in [
                ("MATERIALES", "materiales", "total_materiales"),
                ("EQUIPOS", "equipos", "total_equipos"),
                ("MANO DE OBRA", "mano_obra", "total_mano_obra"),
            ]:
                filas_detalle.append(
                    {
                        "BLOQUE": "",
                        "ITEM / APU": "",
                        "TIPO": tipo_txt,
                        "DESCRIPCIÓN": "",
                        "UNIDAD": "",
                        "VALOR UNITARIO": "",
                        "CANTIDAD": "",
                        "VALOR TOTAL": "",
                    }
                )

                lista = apu.get(clave_lista, []) or []
                if lista:
                    for row in lista:
                        filas_detalle.append(
                            {
                                "BLOQUE": "",
                                "ITEM / APU": "",
                                "TIPO": tipo_txt,
                                "DESCRIPCIÓN": _safe_str(row.get("Descripción")),
                                "UNIDAD": _safe_str(row.get("Unidad")),
                                "VALOR UNITARIO": float(row.get("Valor Unitario", 0.0) or 0.0),
                                "CANTIDAD": float(row.get("Cantidad", 0.0) or 0.0),
                                "VALOR TOTAL": float(row.get("VALOR TOTAL", 0.0) or 0.0),
                            }
                        )

                filas_detalle.append(
                    {
                        "BLOQUE": "",
                        "ITEM / APU": "",
                        "TIPO": f"TOTAL {tipo_txt}",
                        "DESCRIPCIÓN": "",
                        "UNIDAD": "",
                        "VALOR UNITARIO": "",
                        "CANTIDAD": "",
                        "VALOR TOTAL": float(apu.get(clave_total, 0.0) or 0.0),
                    }
                )

            filas_detalle.append(
                {
                    "BLOQUE": "",
                    "ITEM / APU": "",
                    "TIPO": "TOTAL APU",
                    "DESCRIPCIÓN": "",
                    "UNIDAD": "",
                    "VALOR UNITARIO": "",
                    "CANTIDAD": "",
                    "VALOR TOTAL": float(apu.get("total_apu", 0.0) or 0.0),
                }
            )
            filas_detalle.append(
                {
                    "BLOQUE": "",
                    "ITEM / APU": "",
                    "TIPO": "",
                    "DESCRIPCIÓN": "",
                    "UNIDAD": "",
                    "VALOR UNITARIO": "",
                    "CANTIDAD": "",
                    "VALOR TOTAL": "",
                }
            )

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df_detalle = pd.DataFrame(
            filas_detalle,
            columns=[
                "BLOQUE",
                "ITEM / APU",
                "TIPO",
                "DESCRIPCIÓN",
                "UNIDAD",
                "VALOR UNITARIO",
                "CANTIDAD",
                "VALOR TOTAL",
            ],
        )
        df_detalle.to_excel(writer, sheet_name="APUS OBRA", index=False)

    output.seek(0)
    return output


def _apus_consultoria_registros():
    try:
        apus_consultoria_datos = cargar_estado("apus_consultoria") or {}
    except Exception:
        apus_consultoria_datos = {}

    registros = apus_consultoria_datos.get("registros_por_oci", {}) or {}
    salida = []

    if isinstance(registros, dict):
        for oci_id, registro in registros.items():
            if not isinstance(registro, dict):
                continue

            salida.append(
                {
                    "oci_id": str(oci_id).strip(),
                    "nombre": _safe_str(registro.get("costo_indirecto_origen_nombre")),
                    "personal_profesional": registro.get("personal_profesional", []) or [],
                    "personal_tecnico": registro.get("personal_tecnico", []) or [],
                    "otro_personal": registro.get("otro_personal", []) or [],
                    "bienes": registro.get("bienes", []) or [],
                    "servicios": registro.get("servicios", []) or [],
                    "factor_multiplicador_personal": registro.get("factor_multiplicador_personal", 1.0),
                    "iva_bienes_pct": registro.get("iva_bienes_pct", 19.0),
                    "iva_servicios_pct": registro.get("iva_servicios_pct", 19.0),
                    "resumen_final": registro.get("resumen_final", []) or [],
                    "costo_directo_total": registro.get("costo_directo_total", 0.0),
                    "valor_total_final": registro.get("valor_total_final", 0.0),
                }
            )

    salida = sorted(salida, key=lambda x: x["nombre"])
    return salida


def _df_personal_consultoria(lista):
    df = pd.DataFrame(lista or [])
    if df.empty:
        return pd.DataFrame(columns=["ITEM", "PERFIL / ACTIVIDAD", "FUENTE", "UNIDAD", "CANT", "REND", "VR UNITARIO", "VR TOTAL"])

    for col in ["ITEM", "PERFIL / ACTIVIDAD", "FUENTE", "UNIDAD", "CANT", "REND", "VR UNITARIO", "VR TOTAL"]:
        if col not in df.columns:
            df[col] = ""

    return df[["ITEM", "PERFIL / ACTIVIDAD", "FUENTE", "UNIDAD", "CANT", "REND", "VR UNITARIO", "VR TOTAL"]].copy()


def _df_bs_consultoria(lista):
    df = pd.DataFrame(lista or [])
    if df.empty:
        return pd.DataFrame(columns=["ITEM", "BIEN / ACTIVIDAD", "FUENTE", "UNIDAD", "CANT", "VR UNITARIO SIN IVA", "VR TOTAL"])

    for col in ["ITEM", "BIEN / ACTIVIDAD", "FUENTE", "UNIDAD", "CANT", "VR UNITARIO SIN IVA", "VR TOTAL"]:
        if col not in df.columns:
            df[col] = ""

    return df[["ITEM", "BIEN / ACTIVIDAD", "FUENTE", "UNIDAD", "CANT", "VR UNITARIO SIN IVA", "VR TOTAL"]].copy()


def _generar_docx_apus_consultoria(registros_consultoria) -> io.BytesIO:
    doc = Document()

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tit.add_run("ANEXO APUS CONSULTORÍA")
    run.bold = True
    run.font.size = Pt(16)

    for registro in registros_consultoria:
        doc.add_paragraph()
        doc.add_heading(registro["nombre"] or "GRUPO SIN NOMBRE", level=2)

        bloques_personal = [
            ("1 Personal profesional y especializado", registro["personal_profesional"]),
            ("2 Personal técnico", registro["personal_tecnico"]),
            ("3 Otro personal", registro["otro_personal"]),
        ]

        for titulo, lista in bloques_personal:
            doc.add_paragraph()
            p_sub = doc.add_paragraph()
            p_sub.add_run(titulo).bold = True

            df = _df_personal_consultoria(lista)
            if not df.empty:
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                for idx, col in enumerate(df.columns):
                    hdr[idx].text = str(col)
                    for paragraph in hdr[idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for idx, col in enumerate(df.columns):
                        row_cells[idx].text = str(row[col])

        doc.add_paragraph()
        p_fac = doc.add_paragraph()
        p_fac.add_run("Factor multiplicador personal: ").bold = True
        p_fac.add_run(str(registro["factor_multiplicador_personal"]))

        for titulo, lista in [("4. BIENES", registro["bienes"]), ("5. SERVICIOS", registro["servicios"])]:
            doc.add_paragraph()
            p_sub = doc.add_paragraph()
            p_sub.add_run(titulo).bold = True

            df = _df_bs_consultoria(lista)
            if not df.empty:
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                for idx, col in enumerate(df.columns):
                    hdr[idx].text = str(col)
                    for paragraph in hdr[idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for idx, col in enumerate(df.columns):
                        row_cells[idx].text = str(row[col])

        resumen_df = pd.DataFrame(registro["resumen_final"] or [])
        if not resumen_df.empty:
            doc.add_paragraph()
            p_res = doc.add_paragraph()
            p_res.add_run("RESUMEN FINAL").bold = True

            table = doc.add_table(rows=1, cols=len(resumen_df.columns))
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for idx, col in enumerate(resumen_df.columns):
                hdr[idx].text = str(col)
                for paragraph in hdr[idx].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for _, row in resumen_df.iterrows():
                row_cells = table.add_row().cells
                for idx, col in enumerate(resumen_df.columns):
                    row_cells[idx].text = str(row[col])

        p_total = doc.add_paragraph()
        p_total.add_run("TOTAL APU CONSULTORÍA: ").bold = True
        p_total.add_run(f"{float(registro['valor_total_final'] or 0.0):,.2f}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _generar_excel_apus_consultoria(registros_consultoria) -> io.BytesIO:
    output = io.BytesIO()

    filas = []
    for registro in registros_consultoria:
        filas.append(
            {
                "GRUPO": registro["nombre"],
                "SECCIÓN": "",
                "ITEM": "",
                "DESCRIPCIÓN": "",
                "FUENTE": "",
                "UNIDAD": "",
                "CANTIDAD": "",
                "REND": "",
                "VR UNITARIO": "",
                "VR TOTAL": "",
            }
        )

        for titulo, lista in [
            ("1 Personal profesional y especializado", registro["personal_profesional"]),
            ("2 Personal técnico", registro["personal_tecnico"]),
            ("3 Otro personal", registro["otro_personal"]),
        ]:
            filas.append(
                {
                    "GRUPO": "",
                    "SECCIÓN": titulo,
                    "ITEM": "",
                    "DESCRIPCIÓN": "",
                    "FUENTE": "",
                    "UNIDAD": "",
                    "CANTIDAD": "",
                    "REND": "",
                    "VR UNITARIO": "",
                    "VR TOTAL": "",
                }
            )

            df = _df_personal_consultoria(lista)
            if not df.empty:
                for _, row in df.iterrows():
                    filas.append(
                        {
                            "GRUPO": "",
                            "SECCIÓN": titulo,
                            "ITEM": row.get("ITEM", ""),
                            "DESCRIPCIÓN": row.get("PERFIL / ACTIVIDAD", ""),
                            "FUENTE": row.get("FUENTE", ""),
                            "UNIDAD": row.get("UNIDAD", ""),
                            "CANTIDAD": row.get("CANT", ""),
                            "REND": row.get("REND", ""),
                            "VR UNITARIO": row.get("VR UNITARIO", ""),
                            "VR TOTAL": row.get("VR TOTAL", ""),
                        }
                    )

        filas.append(
            {
                "GRUPO": "",
                "SECCIÓN": "FACTOR MULTIPLICADOR",
                "ITEM": "",
                "DESCRIPCIÓN": "",
                "FUENTE": "",
                "UNIDAD": "",
                "CANTIDAD": "",
                "REND": "",
                "VR UNITARIO": registro["factor_multiplicador_personal"],
                "VR TOTAL": "",
            }
        )

        for titulo, lista in [("4. BIENES", registro["bienes"]), ("5. SERVICIOS", registro["servicios"])]:
            filas.append(
                {
                    "GRUPO": "",
                    "SECCIÓN": titulo,
                    "ITEM": "",
                    "DESCRIPCIÓN": "",
                    "FUENTE": "",
                    "UNIDAD": "",
                    "CANTIDAD": "",
                    "REND": "",
                    "VR UNITARIO": "",
                    "VR TOTAL": "",
                }
            )

            df = _df_bs_consultoria(lista)
            if not df.empty:
                for _, row in df.iterrows():
                    filas.append(
                        {
                            "GRUPO": "",
                            "SECCIÓN": titulo,
                            "ITEM": row.get("ITEM", ""),
                            "DESCRIPCIÓN": row.get("BIEN / ACTIVIDAD", ""),
                            "FUENTE": row.get("FUENTE", ""),
                            "UNIDAD": row.get("UNIDAD", ""),
                            "CANTIDAD": row.get("CANT", ""),
                            "REND": "",
                            "VR UNITARIO": row.get("VR UNITARIO SIN IVA", ""),
                            "VR TOTAL": row.get("VR TOTAL", ""),
                        }
                    )

        resumen_df = pd.DataFrame(registro["resumen_final"] or [])
        if not resumen_df.empty:
            filas.append(
                {
                    "GRUPO": "",
                    "SECCIÓN": "RESUMEN FINAL",
                    "ITEM": "",
                    "DESCRIPCIÓN": "",
                    "FUENTE": "",
                    "UNIDAD": "",
                    "CANTIDAD": "",
                    "REND": "",
                    "VR UNITARIO": "",
                    "VR TOTAL": "",
                }
            )

            for _, row in resumen_df.iterrows():
                filas.append(
                    {
                        "GRUPO": "",
                        "SECCIÓN": "RESUMEN FINAL",
                        "ITEM": "",
                        "DESCRIPCIÓN": row.get("Concepto", ""),
                        "FUENTE": row.get("IVA", ""),
                        "UNIDAD": "",
                        "CANTIDAD": "",
                        "REND": "",
                        "VR UNITARIO": row.get("Valor", ""),
                        "VR TOTAL": row.get("VALOR TOTAL", ""),
                    }
                )

        filas.append(
            {
                "GRUPO": "",
                "SECCIÓN": "TOTAL APU CONSULTORÍA",
                "ITEM": "",
                "DESCRIPCIÓN": "",
                "FUENTE": "",
                "UNIDAD": "",
                "CANTIDAD": "",
                "REND": "",
                "VR UNITARIO": "",
                "VR TOTAL": registro["valor_total_final"],
            }
        )
        filas.append(
            {
                "GRUPO": "",
                "SECCIÓN": "",
                "ITEM": "",
                "DESCRIPCIÓN": "",
                "FUENTE": "",
                "UNIDAD": "",
                "CANTIDAD": "",
                "REND": "",
                "VR UNITARIO": "",
                "VR TOTAL": "",
            }
        )

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        pd.DataFrame(
            filas,
            columns=["GRUPO", "SECCIÓN", "ITEM", "DESCRIPCIÓN", "FUENTE", "UNIDAD", "CANTIDAD", "REND", "VR UNITARIO", "VR TOTAL"],
        ).to_excel(writer, sheet_name="APUS CONSULTORIA", index=False)

    output.seek(0)
    return output


def _estudio_mercado_obra_datos():
    try:
        estudio_datos = cargar_estado("estudio_mercado") or {}
    except Exception:
        estudio_datos = {}

    items_df = pd.DataFrame(estudio_datos.get("items", []) or [])
    cot_df = pd.DataFrame(estudio_datos.get("cotizaciones", []) or [])

    for col in ["ID", "TIPO", "SUBTIPO", "NOMBRE", "CARACTERISTICAS", "UNIDAD", "IVA_PCT", "ACTIVO", "COSTO_INDIRECTO_ORIGEN_NOMBRE"]:
        if col not in items_df.columns:
            items_df[col] = "" if col not in ["IVA_PCT", "ACTIVO"] else (0.0 if col == "IVA_PCT" else False)

    for col in ["ID", "ITEM_ID", "PROVEEDOR", "VALOR_SIN_IVA"]:
        if col not in cot_df.columns:
            cot_df[col] = "" if col != "VALOR_SIN_IVA" else 0.0

    if not items_df.empty:
        items_df["IVA_PCT"] = pd.to_numeric(items_df["IVA_PCT"], errors="coerce").fillna(0.0)
        items_df["ACTIVO"] = items_df["ACTIVO"].fillna(False).astype(bool)
        items_df = items_df[items_df["ACTIVO"] == True].copy()

    cot_validas = cot_df.copy()
    if not cot_validas.empty:
        cot_validas["VALOR_SIN_IVA"] = pd.to_numeric(cot_validas["VALOR_SIN_IVA"], errors="coerce")
        cot_validas = cot_validas[cot_validas["ITEM_ID"].astype(str).str.strip() != ""].copy()
        cot_validas = cot_validas[cot_validas["VALOR_SIN_IVA"].notna()].copy()
        cot_validas = cot_validas[cot_validas["VALOR_SIN_IVA"] > 0].copy()

    if items_df.empty:
        resumen_df = pd.DataFrame(
            columns=[
                "TIPO",
                "SUBTIPO",
                "NOMBRE",
                "CARACTERISTICAS",
                "UNIDAD",
                "IVA_PCT",
                "COSTO_INDIRECTO_ORIGEN_NOMBRE",
                "PROMEDIO",
                "CANT_COT",
                "TOTAL",
            ]
        )
        tablas_visuales = {}
        return items_df, cot_df, resumen_df, tablas_visuales

    if cot_validas.empty:
        proms = pd.DataFrame(columns=["ITEM_ID", "PROMEDIO", "CANT_COT"])
    else:
        proms = (
            cot_validas.groupby("ITEM_ID", dropna=False)["VALOR_SIN_IVA"]
            .agg(["mean", "count"])
            .reset_index()
            .rename(columns={"mean": "PROMEDIO", "count": "CANT_COT"})
        )

    resumen_df = items_df.merge(proms, how="left", left_on="ID", right_on="ITEM_ID")
    resumen_df["PROMEDIO"] = pd.to_numeric(resumen_df["PROMEDIO"], errors="coerce").fillna(0.0)
    resumen_df["CANT_COT"] = pd.to_numeric(resumen_df["CANT_COT"], errors="coerce").fillna(0).astype(int)
    resumen_df["TOTAL"] = resumen_df["PROMEDIO"]

    mask_no_personal = resumen_df["TIPO"].astype(str).str.upper().ne("PERSONAL")
    resumen_df.loc[mask_no_personal, "TOTAL"] = resumen_df.loc[mask_no_personal, "PROMEDIO"] * (
        1 + (pd.to_numeric(resumen_df.loc[mask_no_personal, "IVA_PCT"], errors="coerce").fillna(0.0) / 100.0)
    )

    resumen_df = resumen_df[
        [
            "ID",
            "TIPO",
            "SUBTIPO",
            "NOMBRE",
            "CARACTERISTICAS",
            "UNIDAD",
            "IVA_PCT",
            "COSTO_INDIRECTO_ORIGEN_NOMBRE",
            "PROMEDIO",
            "CANT_COT",
            "TOTAL",
        ]
    ].copy()

    tablas_visuales = {}

    for tipo in ["BIENES", "SERVICIOS", "PERSONAL"]:
        base_tipo = resumen_df[resumen_df["TIPO"].astype(str).str.upper() == tipo].copy()
        if base_tipo.empty:
            continue

        filas_visuales = []

        for _, item in base_tipo.iterrows():
            item_id = _safe_str(item.get("ID"))
            cot_item = cot_validas[cot_validas["ITEM_ID"].astype(str).str.strip() == item_id].copy()

            cot_item = cot_item.sort_values(by=["PROVEEDOR", "VALOR_SIN_IVA"], ascending=[True, True]).reset_index(drop=True)

            fila = {
                "BIEN / ACTIVIDAD": _safe_str(item.get("NOMBRE")),
                "CARACTERISTICAS": _safe_str(item.get("CARACTERISTICAS")),
                "UNIDAD": _safe_str(item.get("UNIDAD")),
                "PROVEEDOR 1": "",
                "VALOR 1": 0.0,
                "PROVEEDOR 2": "",
                "VALOR 2": 0.0,
                "PROVEEDOR 3": "",
                "VALOR 3": 0.0,
                "PROMEDIO": float(item.get("PROMEDIO", 0.0) or 0.0),
                "IVA %": float(item.get("IVA_PCT", 0.0) or 0.0),
                "TOTAL": float(item.get("TOTAL", 0.0) or 0.0),
            }

            for i in range(min(len(cot_item), 3)):
                fila[f"PROVEEDOR {i+1}"] = _safe_str(cot_item.iloc[i].get("PROVEEDOR"))
                fila[f"VALOR {i+1}"] = float(cot_item.iloc[i].get("VALOR_SIN_IVA", 0.0) or 0.0)

            filas_visuales.append(fila)

        tablas_visuales[tipo] = pd.DataFrame(
            filas_visuales,
            columns=[
                "BIEN / ACTIVIDAD",
                "CARACTERISTICAS",
                "UNIDAD",
                "PROVEEDOR 1",
                "VALOR 1",
                "PROVEEDOR 2",
                "VALOR 2",
                "PROVEEDOR 3",
                "VALOR 3",
                "PROMEDIO",
                "IVA %",
                "TOTAL",
            ],
        )

    return items_df, cot_df, resumen_df, tablas_visuales


def _generar_docx_estudio_mercado_obra(tablas_visuales) -> io.BytesIO:
    doc = Document()

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tit.add_run("ANEXO ESTUDIO DE MERCADO OBRA")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    if not tablas_visuales:
        doc.add_paragraph("No hay información transformada del estudio de mercado obra.")
    else:
        for titulo in ["BIENES", "SERVICIOS", "PERSONAL"]:
            df = tablas_visuales.get(titulo)
            if df is None or df.empty:
                continue

            doc.add_heading(titulo, level=2)

            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = "Table Grid"

            hdr = table.rows[0].cells
            for idx, col in enumerate(df.columns):
                hdr[idx].text = str(col)
                for paragraph in hdr[idx].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for idx, col in enumerate(df.columns):
                    row_cells[idx].text = str(row[col])

            doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _generar_excel_estudio_mercado_obra(tablas_visuales) -> io.BytesIO:
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        if not tablas_visuales:
            pd.DataFrame({"MENSAJE": ["No hay información transformada del estudio de mercado obra."]}).to_excel(
                writer,
                sheet_name="Estudio mercado",
                index=False,
            )
        else:
            for titulo in ["BIENES", "SERVICIOS", "PERSONAL"]:
                df = tablas_visuales.get(titulo)
                if df is not None and not df.empty:
                    df.to_excel(writer, sheet_name=titulo[:31], index=False)

    output.seek(0)
    return output


def _estudio_mercado_consultoria_datos():
    try:
        datos = cargar_estado("estudio_mercado_consultoria") or {}
    except Exception:
        datos = {}

    items_df = pd.DataFrame(datos.get("items", []) or [])
    cot_df = pd.DataFrame(datos.get("cotizaciones", []) or [])
    num_cotizaciones = int(datos.get("num_cotizaciones", 3) or 3)

    if items_df.empty:
        return {}, pd.DataFrame()

    for col in ["ID", "TIPO", "SUBTIPO", "NOMBRE", "CARACTERISTICAS", "UNIDAD", "IVA_PCT", "ACTIVO", "APU_CONSULTORIA_ORIGEN_ID", "APU_CONSULTORIA_ORIGEN_NOMBRE"]:
        if col not in items_df.columns:
            items_df[col] = "" if col not in ["IVA_PCT", "ACTIVO"] else (0.0 if col == "IVA_PCT" else False)

    for col in ["ID", "ITEM_ID", "PROVEEDOR", "VALOR_SIN_IVA"]:
        if col not in cot_df.columns:
            cot_df[col] = "" if col != "VALOR_SIN_IVA" else 0.0

    items_df["TIPO"] = items_df["TIPO"].astype(str).str.upper()
    cot_df["ITEM_ID"] = cot_df["ITEM_ID"].astype(str)

    tablas = {}
    filas_consolidadas = []

    for tipo, titulo in [("PERSONAL", "SALARIOS"), ("BIENES", "BIENES"), ("SERVICIOS", "SERVICIOS")]:
        items_sec = items_df[items_df["TIPO"] == tipo].copy()
        if items_sec.empty:
            continue

        cot_sec = cot_df[cot_df["ITEM_ID"].isin(items_sec["ID"].astype(str))].copy()

        proms = pd.DataFrame(columns=["ITEM_ID", "PROMEDIO", "CANT_COT"])
        if not cot_sec.empty:
            cot_validas = cot_sec.copy()
            cot_validas["VALOR_SIN_IVA"] = pd.to_numeric(cot_validas["VALOR_SIN_IVA"], errors="coerce")
            cot_validas = cot_validas[cot_validas["ITEM_ID"].astype(str).str.strip() != ""].copy()
            cot_validas = cot_validas[cot_validas["VALOR_SIN_IVA"].notna()].copy()
            cot_validas = cot_validas[cot_validas["VALOR_SIN_IVA"] > 0].copy()

            if not cot_validas.empty:
                proms = (
                    cot_validas.groupby("ITEM_ID", dropna=False)["VALOR_SIN_IVA"]
                    .agg(["mean", "count"])
                    .reset_index()
                    .rename(columns={"mean": "PROMEDIO", "count": "CANT_COT"})
                )

        resumen = items_sec.merge(proms, how="left", left_on="ID", right_on="ITEM_ID")
        resumen["PROMEDIO"] = pd.to_numeric(resumen.get("PROMEDIO", 0.0), errors="coerce").fillna(0.0)
        resumen["CANT_COT"] = pd.to_numeric(resumen.get("CANT_COT", 0), errors="coerce").fillna(0).astype(int)
        resumen["TOTAL"] = resumen["PROMEDIO"]

        mask_no_personal = resumen["TIPO"].astype(str).str.upper().ne("PERSONAL")
        resumen.loc[mask_no_personal, "TOTAL"] = resumen.loc[mask_no_personal, "PROMEDIO"] * (
            1 + (pd.to_numeric(resumen.loc[mask_no_personal, "IVA_PCT"], errors="coerce").fillna(0.0) / 100.0)
        )

        filas_render = []

        for _, row in resumen.iterrows():
            item_id = str(row["ID"]).strip()
            cot_item = cot_sec[cot_sec["ITEM_ID"].astype(str).eq(item_id)].reset_index(drop=True).copy()

            fila = {
                "BLOQUE": titulo,
                "TIPO": str(row.get("TIPO", "")),
                "SUBTIPO": str(row.get("SUBTIPO", "")),
                "NOMBRE": str(row.get("NOMBRE", "")),
                "CARACTERISTICAS": str(row.get("CARACTERISTICAS", "")),
                "UNIDAD": str(row.get("UNIDAD", "")),
                "IVA_PCT": float(pd.to_numeric(row.get("IVA_PCT", 0.0), errors="coerce") or 0.0),
                "PROMEDIO": float(pd.to_numeric(row.get("PROMEDIO", 0.0), errors="coerce") or 0.0),
                "TOTAL": float(pd.to_numeric(row.get("TOTAL", 0.0), errors="coerce") or 0.0),
            }

            for idx in range(num_cotizaciones):
                if idx < len(cot_item):
                    fila[f"PROVEEDOR_{idx + 1}"] = str(cot_item.iloc[idx].get("PROVEEDOR", "") or "")
                    fila[f"VALOR_{idx + 1}"] = float(pd.to_numeric(cot_item.iloc[idx].get("VALOR_SIN_IVA", 0.0), errors="coerce") or 0.0)
                else:
                    fila[f"PROVEEDOR_{idx + 1}"] = ""
                    fila[f"VALOR_{idx + 1}"] = 0.0

            filas_render.append(fila)
            filas_consolidadas.append(fila)

        if filas_render:
            df_render = pd.DataFrame(filas_render)

            col_nombre = "PERFIL / ACTIVIDAD" if tipo == "PERSONAL" else "BIEN / ACTIVIDAD"
            df_render = df_render.rename(columns={"NOMBRE": col_nombre})

            columnas = [col_nombre, "CARACTERISTICAS", "UNIDAD"]
            for idx in range(num_cotizaciones):
                columnas.extend([f"PROVEEDOR_{idx + 1}", f"VALOR_{idx + 1}"])
            columnas.extend(["PROMEDIO", "IVA_PCT", "TOTAL"])

            df_render = df_render[columnas].copy()
            df_render = df_render.rename(columns={"IVA_PCT": "IVA %"})

            tablas[titulo] = df_render

    df_total = pd.DataFrame(filas_consolidadas) if filas_consolidadas else pd.DataFrame()
    return tablas, df_total


def _generar_docx_estudio_mercado_consultoria(tablas_emc, df_total) -> io.BytesIO:
    doc = Document()

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tit.add_run("ANEXO ESTUDIO DE MERCADO CONSULTORÍA")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    if not tablas_emc:
        doc.add_paragraph("No hay información de estudio de mercado consultoría.")
    else:
        for titulo in ["SALARIOS", "BIENES", "SERVICIOS"]:
            df = tablas_emc.get(titulo)
            if df is None or df.empty:
                continue

            doc.add_heading(titulo, level=2)

            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = "Table Grid"

            hdr = table.rows[0].cells
            for idx, col in enumerate(df.columns):
                hdr[idx].text = str(col)
                for paragraph in hdr[idx].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for idx, col in enumerate(df.columns):
                    row_cells[idx].text = str(row[col])

            doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def _generar_excel_estudio_mercado_consultoria(bloques, df_total) -> io.BytesIO:
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        if df_total.empty:
            pd.DataFrame({"MENSAJE": ["No hay información de estudio de mercado consultoría."]}).to_excel(
                writer,
                sheet_name="Estudio mercado",
                index=False,
            )
        else:
            df_total.to_excel(writer, sheet_name="Consolidado", index=False)

            for i, bloque in enumerate(bloques, start=1):
                if not isinstance(bloque, dict):
                    continue

                nombre_bloque = _safe_str(bloque.get("nombre")) or _safe_str(bloque.get("titulo")) or f"Bloque {i}"
                filas = bloque.get("filas", []) or []
                df = pd.DataFrame(filas) if filas else pd.DataFrame()

                if not df.empty:
                    hoja = f"{i}_{nombre_bloque}"[:31]
                    df.to_excel(writer, sheet_name=hoja, index=False)

    output.seek(0)
    return output


def _aiu_datos():
    try:
        aiu = cargar_estado("aiu") or {}
    except Exception:
        aiu = {}

    personal_df = pd.DataFrame(aiu.get("personal_administrativo", []) or [])
    equipos_df = pd.DataFrame(aiu.get("equipos_generales", []) or [])
    gastos_generales_df = pd.DataFrame(aiu.get("gastos_generales", []) or [])
    gastos_legales_df = pd.DataFrame(aiu.get("gastos_legales", []) or [])

    imprevistos_pct = float(aiu.get("imprevistos_porcentaje", 0.0) or 0.0)
    utilidad_pct = float(aiu.get("utilidad_porcentaje", 0.0) or 0.0)
    administracion_valor = float(aiu.get("administracion_valor", 0.0) or 0.0)
    imprevistos_valor = float(aiu.get("imprevistos_valor", 0.0) or 0.0)
    utilidad_valor = float(aiu.get("utilidad_valor", 0.0) or 0.0)
    aiu_total_valor = float(aiu.get("aiu_total_valor", 0.0) or 0.0)

    costo_directo = float(st.session_state.get("aiu_costo_directo", 0.0) or 0.0)
    if costo_directo <= 0:
        try:
            presupuesto_obra = cargar_estado("presupuesto_obra") or {}
        except Exception:
            presupuesto_obra = {}
        costo_directo = float(((presupuesto_obra.get("resumen", {}) or {}).get("costo_directo_total", 0.0)) or 0.0)

    administracion_pct = (administracion_valor / costo_directo * 100.0) if costo_directo > 0 else 0.0
    aiu_total_pct = administracion_pct + imprevistos_pct + utilidad_pct

    resumen_df = pd.DataFrame(
        [
            {"Componente": "Administración", "Valor": administracion_valor, "%": administracion_pct},
            {"Componente": "Imprevistos", "Valor": imprevistos_valor, "%": imprevistos_pct},
            {"Componente": "Utilidad", "Valor": utilidad_valor, "%": utilidad_pct},
        ]
    )

    return {
        "personal_df": personal_df,
        "equipos_df": equipos_df,
        "gastos_generales_df": gastos_generales_df,
        "gastos_legales_df": gastos_legales_df,
        "resumen_df": resumen_df,
        "administracion_valor": administracion_valor,
        "administracion_pct": administracion_pct,
        "imprevistos_valor": imprevistos_valor,
        "imprevistos_pct": imprevistos_pct,
        "utilidad_valor": utilidad_valor,
        "utilidad_pct": utilidad_pct,
        "aiu_total_valor": aiu_total_valor,
        "aiu_total_pct": aiu_total_pct,
    }


def _generar_docx_aiu(aiu_data) -> io.BytesIO:
    doc = Document()

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tit.add_run("ANEXO AIU")
    run.bold = True
    run.font.size = Pt(16)

    secciones = [
        ("1.1 Personal administrativo", aiu_data["personal_df"]),
        ("1.2 Equipos generales, movilización e instalación", aiu_data["equipos_df"]),
        ("1.3 Gastos generales", aiu_data["gastos_generales_df"]),
        ("1.4 Gastos legales, jurídicos, tributarios", aiu_data["gastos_legales_df"]),
        ("RESUMEN FINAL DEL AIU", aiu_data["resumen_df"]),
    ]

    for titulo, df in secciones:
        doc.add_paragraph()
        doc.add_heading(titulo, level=2)

        if df.empty:
            doc.add_paragraph("Sin información disponible.")
            continue

        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        for idx, col in enumerate(df.columns):
            hdr[idx].text = str(col)
            for paragraph in hdr[idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for idx, col in enumerate(df.columns):
                row_cells[idx].text = str(row[col])

    doc.add_paragraph()
    p_total = doc.add_paragraph()
    p_total.add_run("TOTAL AIU: ").bold = True
    p_total.add_run(f"{aiu_data['aiu_total_valor']:,.2f}")

    p_pct = doc.add_paragraph()
    p_pct.add_run("% TOTAL AIU: ").bold = True
    p_pct.add_run(f"{aiu_data['aiu_total_pct']:,.2f}%")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _generar_excel_aiu(aiu_data) -> io.BytesIO:
    output = io.BytesIO()

    filas = []

    def agregar_bloque(titulo, df):
        filas.append(
            {
                "SECCIÓN": titulo,
                "CONCEPTO": "",
                "VALOR": "",
                "PORCENTAJE": "",
            }
        )

        if df is None or df.empty:
            filas.append(
                {
                    "SECCIÓN": "",
                    "CONCEPTO": "Sin información disponible",
                    "VALOR": "",
                    "PORCENTAJE": "",
                }
            )
        else:
            df_local = df.copy()
            for _, row in df_local.iterrows():
                fila = {"SECCIÓN": ""}
                for col in df_local.columns:
                    fila[str(col)] = row[col]
                filas.append(fila)

        filas.append(
            {
                "SECCIÓN": "",
                "CONCEPTO": "",
                "VALOR": "",
                "PORCENTAJE": "",
            }
        )

    agregar_bloque("1.1 Personal administrativo", aiu_data["personal_df"])
    agregar_bloque("1.2 Equipos generales, movilización e instalación", aiu_data["equipos_df"])
    agregar_bloque("1.3 Gastos generales", aiu_data["gastos_generales_df"])
    agregar_bloque("1.4 Gastos legales, jurídicos, tributarios", aiu_data["gastos_legales_df"])
    agregar_bloque("RESUMEN FINAL DEL AIU", aiu_data["resumen_df"])

    filas.append(
        {
            "SECCIÓN": "TOTAL AIU",
            "CONCEPTO": "",
            "VALOR": aiu_data["aiu_total_valor"],
            "PORCENTAJE": aiu_data["aiu_total_pct"],
        }
    )

    df_export = pd.DataFrame(filas)

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df_export.to_excel(writer, sheet_name="AIU", index=False)

    output.seek(0)
    return output


def _factor_multiplicador_datos():
    try:
        datos = cargar_estado("factor_multiplicador") or {}
    except Exception:
        datos = {}

    def normalizar_df(lista):
        df = pd.DataFrame(lista or [])
        if df.empty:
            return pd.DataFrame(columns=["No.", "Factor", "Base", "Valor"])

        mapa = {}
        for col in df.columns:
            col_txt = str(col).strip().lower()
            if col_txt in ["no.", "no", "nro", "numero"]:
                mapa[col] = "No."
            elif col_txt == "factor":
                mapa[col] = "Factor"
            elif col_txt == "base":
                mapa[col] = "Base"
            elif col_txt in ["valor", "value"]:
                mapa[col] = "Valor"

        df = df.rename(columns=mapa).copy()

        for col in ["No.", "Factor", "Base", "Valor"]:
            if col not in df.columns:
                df[col] = ""

        return df[["No.", "Factor", "Base", "Valor"]].copy()

    df_prest = normalizar_df(datos.get("prestacional", []) or [])
    df_ind = normalizar_df(datos.get("costos_indirectos", []) or [])
    df_util = normalizar_df(datos.get("utilidad", []) or [])

    valor_base = float(datos.get("valor_base", 1.0) or 1.0)
    total_prest = float(datos.get("total_prestacional", 0.0) or 0.0)
    total_ind = float(datos.get("total_costos_indirectos", 0.0) or 0.0)
    total_util = float(datos.get("total_utilidad", 0.0) or 0.0)
    factor_final = float(datos.get("factor_multiplicador_final", 0.0) or 0.0)

    return {
        "valor_base": valor_base,
        "prestacional_df": df_prest,
        "costos_indirectos_df": df_ind,
        "utilidad_df": df_util,
        "total_prestacional": total_prest,
        "total_costos_indirectos": total_ind,
        "total_utilidad": total_util,
        "factor_multiplicador_final": factor_final,
    }


def _generar_docx_factor_multiplicador(fm_data) -> io.BytesIO:
    doc = Document()

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tit.add_run("ANEXO FACTOR MULTIPLICADOR")
    run.bold = True
    run.font.size = Pt(16)

    secciones = [
        ("1. Factor prestacional", fm_data["prestacional_df"], fm_data["total_prestacional"]),
        ("2. Costos indirectos", fm_data["costos_indirectos_df"], fm_data["total_costos_indirectos"]),
        ("3. Utilidad", fm_data["utilidad_df"], fm_data["total_utilidad"]),
    ]

    for titulo, df, total in secciones:
        doc.add_paragraph()
        doc.add_heading(titulo, level=2)

        if df.empty:
            doc.add_paragraph("Sin información disponible.")
        else:
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = "Table Grid"

            hdr = table.rows[0].cells
            for idx, col in enumerate(df.columns):
                hdr[idx].text = str(col)
                for paragraph in hdr[idx].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for idx, col in enumerate(df.columns):
                    row_cells[idx].text = str(row[col])

        p_total = doc.add_paragraph()
        p_total.add_run(f"Total {titulo.lower()}: ").bold = True
        p_total.add_run(f"{total:,.4f}")

    doc.add_paragraph()
    p_base = doc.add_paragraph()
    p_base.add_run("Valor base: ").bold = True
    p_base.add_run(f"{fm_data['valor_base']:,.4f}")

    p_final = doc.add_paragraph()
    p_final.add_run("FACTOR MULTIPLICADOR FINAL: ").bold = True
    p_final.add_run(f"{fm_data['factor_multiplicador_final']:,.4f}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tit.add_run("ANEXO FACTOR MULTIPLICADOR")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    df = fm_data["df"]
    if df.empty:
        doc.add_paragraph("No hay información de factor multiplicador.")
    else:
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        for idx, col in enumerate(df.columns):
            hdr[idx].text = str(col)
            for paragraph in hdr[idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for idx, col in enumerate(df.columns):
                row_cells[idx].text = str(row[col])

    doc.add_paragraph()
    p_total = doc.add_paragraph()
    p_total.add_run("VALOR TOTAL: ").bold = True
    p_total.add_run(f"{fm_data['valor_total']:,.2f}")

    p_pct = doc.add_paragraph()
    p_pct.add_run("% TOTAL: ").bold = True
    p_pct.add_run(f"{fm_data['porcentaje_total']:,.2f}%")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _generar_excel_factor_multiplicador(fm_data) -> io.BytesIO:
    output = io.BytesIO()

    filas = []

    def agregar_bloque(titulo, df, total):
        filas.append(
            {
                "SECCIÓN": titulo,
                "No.": "",
                "Factor": "",
                "Base": "",
                "Valor": "",
            }
        )

        if df.empty:
            filas.append(
                {
                    "SECCIÓN": "",
                    "No.": "",
                    "Factor": "Sin información disponible",
                    "Base": "",
                    "Valor": "",
                }
            )
        else:
            for _, row in df.iterrows():
                filas.append(
                    {
                        "SECCIÓN": "",
                        "No.": row.get("No.", ""),
                        "Factor": row.get("Factor", ""),
                        "Base": row.get("Base", ""),
                        "Valor": row.get("Valor", ""),
                    }
                )

        filas.append(
            {
                "SECCIÓN": "",
                "No.": "",
                "Factor": f"Total {titulo.lower()}",
                "Base": "",
                "Valor": total,
            }
        )
        filas.append(
            {
                "SECCIÓN": "",
                "No.": "",
                "Factor": "",
                "Base": "",
                "Valor": "",
            }
        )

    agregar_bloque("1. Factor prestacional", fm_data["prestacional_df"], fm_data["total_prestacional"])
    agregar_bloque("2. Costos indirectos", fm_data["costos_indirectos_df"], fm_data["total_costos_indirectos"])
    agregar_bloque("3. Utilidad", fm_data["utilidad_df"], fm_data["total_utilidad"])

    filas.append(
        {
            "SECCIÓN": "FACTOR MULTIPLICADOR FINAL",
            "No.": "",
            "Factor": "",
            "Base": "",
            "Valor": fm_data["factor_multiplicador_final"],
        }
    )

    df_export = pd.DataFrame(filas)

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df_export.to_excel(writer, sheet_name="Factor multiplicador", index=False)

    output.seek(0)
    return output


def _append_doc_content(doc_destino: Document, doc_fuente_buffer: io.BytesIO):
    from copy import deepcopy

    doc_fuente_buffer.seek(0)
    doc_fuente = Document(doc_fuente_buffer)

    if len(doc_destino.paragraphs) > 0 and any(p.text.strip() for p in doc_destino.paragraphs):
        doc_destino.add_page_break()

    for element in doc_fuente.element.body:
        doc_destino.element.body.append(deepcopy(element))


def _generar_docx_anexos_combinado(modulos: list[str]) -> io.BytesIO:
    doc = Document()

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tit.add_run("DOCUMENTO COMBINADO - ANEXOS")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    for modulo in modulos:
        if modulo == "apus_obra":
            codigos_usados, df_apus_filtrados = _apus_filtrados_obra()
            apus_generados_dict = _apus_generados_obra_filtrados()
            buffer = _generar_docx_apus_obra(codigos_usados, df_apus_filtrados, apus_generados_dict)
            _append_doc_content(doc, buffer)

        elif modulo == "apus_consultoria":
            registros_consultoria = _apus_consultoria_registros()
            buffer = _generar_docx_apus_consultoria(registros_consultoria)
            _append_doc_content(doc, buffer)

        elif modulo == "em_obra":
            _, _, _, tablas_visuales = _estudio_mercado_obra_datos()
            buffer = _generar_docx_estudio_mercado_obra(tablas_visuales)
            _append_doc_content(doc, buffer)

        elif modulo == "em_consultoria":
            tablas_emc, df_emc = _estudio_mercado_consultoria_datos()
            buffer = _generar_docx_estudio_mercado_consultoria(tablas_emc, df_emc)
            _append_doc_content(doc, buffer)

        elif modulo == "aiu":
            aiu_data = _aiu_datos()
            buffer = _generar_docx_aiu(aiu_data)
            _append_doc_content(doc, buffer)

        elif modulo == "factor_multiplicador":
            fm_data = _factor_multiplicador_datos()
            buffer = _generar_docx_factor_multiplicador(fm_data)
            _append_doc_content(doc, buffer)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


st.title("🗂️ Informe Anexos")

_cargar_cfg()
cfg = st.session_state["informe_anexos_config"]

tipo_proyecto = _tipo_proyecto()

st.info(
    "Esta hoja organiza los anexos técnicos y financieros del proyecto. "
    "Aquí se separan del informe principal para evitar sobrecargar la hoja de informes."
)

col_ctrl_1, col_ctrl_2 = st.columns([1, 1])

with col_ctrl_1:
    if st.button("💾 Guardar configuración", type="primary", width="stretch"):
        _guardar_cfg()
        st.success("Configuración de anexos guardada correctamente.")

with col_ctrl_2:
    if st.button("🔄 Recargar desde nube", width="stretch"):
        if "informe_anexos_config" in st.session_state:
            del st.session_state["informe_anexos_config"]
        st.rerun()

st.divider()

st.markdown(f"**Tipo de proyecto detectado:** {tipo_proyecto}")

tab1, tab2, tab3, tab4, tab5 = st.tabs(
    [
        "1. APUS obra",
        "2. Estudio de mercado obra",
        "3. AIU",
        "4. Factor multiplicador",
        "5. Documento combinado",
    ]
)

with tab1:
    st.subheader("APUS obra")

    with st.container(border=True):
        cfg["incluye_apus_obra"] = st.checkbox(
            "Incluir anexo APUS obra",
            value=bool(cfg.get("incluye_apus_obra", True)),
            key="incluye_apus_obra_cfg",
        )

        cfg["notas_apus_obra"] = st.text_area(
            "Notas del anexo APUS obra",
            value=cfg.get("notas_apus_obra", ""),
            height=120,
            key="notas_apus_obra_cfg",
        )

    with st.container(border=True):
        if tipo_proyecto != "Obra":
            st.info("Este proyecto no está marcado como Obra. La carga de APUS obra no aplica en este caso.")
        else:
            codigos_usados, df_apus_filtrados = _apus_filtrados_obra()
            apus_generados_dict = _apus_generados_obra_filtrados()

            st.markdown("**Códigos ITEM GOBER usados en el presupuesto de obra**")
            if codigos_usados:
                st.code(", ".join(codigos_usados))
            else:
                st.warning("No se encontraron códigos ITEM GOBER asociados en el presupuesto de obra.")

            if df_apus_filtrados.empty:
                st.info("No se encontraron APUS de gobernación coincidentes con los códigos usados en el presupuesto.")
            else:
                st.markdown("## APUS de gobernación usados en el proyecto")

                codigos_disponibles = sorted(
                    {
                        str(x).strip()
                        for x in df_apus_filtrados["cod_actividad"].dropna().tolist()
                        if str(x).strip()
                    }
                )

                for codigo_apu in codigos_disponibles:
                    df_apu = df_apus_filtrados[
                        df_apus_filtrados["cod_actividad"].astype(str).str.strip() == codigo_apu
                    ].copy()

                    if df_apu.empty:
                        continue

                    fila_ref = df_apu.iloc[0]

                    st.markdown(
                        f"### {str(fila_ref.get('cod_actividad', '')).strip()} - {str(fila_ref.get('actividad', '')).strip()}"
                    )

                    c1, c2 = st.columns([1, 3])
                    with c1:
                        st.text_input(
                            "cod_capitulo",
                            value=str(fila_ref.get("cod_capitulo", "") or ""),
                            disabled=True,
                            key=f"apu_obra_cod_capitulo_{codigo_apu}",
                        )
                    with c2:
                        st.text_input(
                            "capitulo",
                            value=str(fila_ref.get("capitulo", "") or ""),
                            disabled=True,
                            key=f"apu_obra_capitulo_{codigo_apu}",
                        )

                    c3, c4, c5 = st.columns([1, 3, 1])
                    with c3:
                        st.text_input(
                            "cod_subcapitulo",
                            value=str(fila_ref.get("cod_subcapitulo", "") or ""),
                            disabled=True,
                            key=f"apu_obra_cod_subcapitulo_{codigo_apu}",
                        )
                    with c4:
                        st.text_input(
                            "subcapitulo",
                            value=str(fila_ref.get("subcapitulo", "") or ""),
                            disabled=True,
                            key=f"apu_obra_subcapitulo_{codigo_apu}",
                        )
                    with c5:
                        st.text_input(
                            "Und. Act",
                            value=str(fila_ref.get("Und. Act", "") or ""),
                            disabled=True,
                            key=f"apu_obra_und_act_{codigo_apu}",
                        )

                    st.markdown("**Materiales**")
                    df_materiales, total_materiales = _tabla_apu_tipo(df_apu, "MATERIAL")
                    if not df_materiales.empty:
                        st.dataframe(df_materiales, width="stretch", hide_index=True)
                    else:
                        st.info("Sin materiales para este APU.")
                    st.markdown(f"**Total materiales:** {total_materiales:,.2f}")

                    st.markdown("**Equipos**")
                    df_equipos, total_equipos = _tabla_apu_tipo(df_apu, "EQUIPO")
                    if not df_equipos.empty:
                        st.dataframe(df_equipos, width="stretch", hide_index=True)
                    else:
                        st.info("Sin equipos para este APU.")
                    st.markdown(f"**Total equipos:** {total_equipos:,.2f}")

                    st.markdown("**Mano de obra**")
                    df_mo, total_mo = _tabla_apu_tipo(df_apu, "MANO DE OBRA")
                    if not df_mo.empty:
                        st.dataframe(df_mo, width="stretch", hide_index=True)
                    else:
                        st.info("Sin mano de obra para este APU.")
                    st.markdown(f"**Total mano de obra:** {total_mo:,.2f}")

                    total_apu = total_materiales + total_equipos + total_mo
                    st.markdown(f"## TOTAL APU: {total_apu:,.2f}")
                    st.divider()

            st.markdown("## APUS generados")
            if not apus_generados_dict:
                st.info("No se encontraron APUS generados asociados al presupuesto de obra.")
            else:
                for _, apu in apus_generados_dict.items():
                    st.markdown(f"### {_safe_str(apu.get('item_label'))}")

                    st.markdown(
                        f"**APU base:** {_safe_str(apu.get('apu_base_codigo'))} - {_safe_str(apu.get('apu_base_actividad'))}"
                    )
                    st.markdown(f"**Unidad:** {_safe_str(apu.get('unidad_apu'))}")

                    for titulo, clave_lista, clave_total, clave_widget in [
                        ("Materiales", "materiales", "total_materiales", "gen_mat"),
                        ("Equipos", "equipos", "total_equipos", "gen_eq"),
                        ("Mano de obra", "mano_obra", "total_mano_obra", "gen_mo"),
                    ]:
                        st.markdown(f"**{titulo}**")
                        lista = apu.get(clave_lista, []) or []
                        if lista:
                            st.dataframe(pd.DataFrame(lista), width="stretch", hide_index=True)
                        else:
                            st.info(f"Sin {titulo.lower()} para este APU generado.")
                        st.markdown(f"**Total {titulo.lower()}:** {float(apu.get(clave_total, 0.0) or 0.0):,.2f}")

                    st.markdown(f"## TOTAL APU GENERADO: {float(apu.get('total_apu', 0.0) or 0.0):,.2f}")
                    st.divider()

            col_apu_1, col_apu_2 = st.columns(2)

            with col_apu_1:
                if st.button("📥 Generar Word APUS obra", key="btn_docx_apus_obra", width="stretch"):
                    st.session_state["archivo_apus_obra_docx"] = _generar_docx_apus_obra(
                        codigos_usados,
                        df_apus_filtrados,
                        apus_generados_dict,
                    )
                    st.success("Documento Word de APUS obra generado.")

            with col_apu_2:
                if st.button("📊 Generar Excel APUS obra", key="btn_excel_apus_obra", width="stretch"):
                    st.session_state["archivo_apus_obra_excel"] = _generar_excel_apus_obra(
                        codigos_usados,
                        df_apus_filtrados,
                        apus_generados_dict,
                    )
                    st.success("Archivo Excel de APUS obra generado.")

            if "archivo_apus_obra_docx" in st.session_state:
                st.download_button(
                    label="⬇️ Descargar Word APUS obra",
                    data=st.session_state["archivo_apus_obra_docx"],
                    file_name="anexo_apus_obra.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_docx_apus_obra",
                    width="stretch",
                )

            if "archivo_apus_obra_excel" in st.session_state:
                st.download_button(
                    label="⬇️ Descargar Excel APUS obra",
                    data=st.session_state["archivo_apus_obra_excel"],
                    file_name="anexo_apus_obra.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="download_excel_apus_obra",
                    width="stretch",
                )
with tab2:
    st.subheader("Estudio de mercado obra")

    with st.container(border=True):
        cfg["incluye_estudio_mercado"] = st.checkbox(
            "Incluir anexo Estudio de mercado obra",
            value=bool(cfg.get("incluye_estudio_mercado", True)),
            key="incluye_estudio_mercado_cfg",
        )

        cfg["notas_estudio_mercado"] = st.text_area(
            "Notas del anexo Estudio de mercado obra",
            value=cfg.get("notas_estudio_mercado", ""),
            height=120,
            key="notas_estudio_mercado_cfg",
        )

        with st.container(border=True):
            tablas_visuales = {}

            if tipo_proyecto != "Obra":
                st.info("Este proyecto no está marcado como Obra. La hoja Estudio de mercado obra no aplica en este caso.")
            else:
                items_df, cot_df, resumen_df, tablas_visuales = _estudio_mercado_obra_datos()

            if not tablas_visuales:
                st.info("No hay información transformada del estudio de mercado obra.")
            else:
                for titulo in ["BIENES", "SERVICIOS", "PERSONAL"]:
                    df_visual = tablas_visuales.get(titulo)
                    if df_visual is not None and not df_visual.empty:
                        st.markdown(f"## {titulo}")
                        st.dataframe(df_visual, width="stretch", hide_index=True)
                        st.divider()

            col_em_1, col_em_2 = st.columns(2)

            with col_em_1:
                if st.button("📥 Generar Word Estudio de mercado obra", key="btn_docx_em_obra", width="stretch"):
                    st.session_state["archivo_em_obra_docx"] = _generar_docx_estudio_mercado_obra(
                        tablas_visuales,
                    )
                    st.success("Documento Word de Estudio de mercado obra generado.")

            with col_em_2:
                if st.button("📊 Generar Excel Estudio de mercado obra", key="btn_excel_em_obra", width="stretch"):
                    st.session_state["archivo_em_obra_excel"] = _generar_excel_estudio_mercado_obra(
                        tablas_visuales,
                    )
                    st.success("Archivo Excel de Estudio de mercado obra generado.")

            if "archivo_em_obra_docx" in st.session_state:
                st.download_button(
                    label="⬇️ Descargar Word Estudio de mercado obra",
                    data=st.session_state["archivo_em_obra_docx"],
                    file_name="anexo_estudio_mercado_obra.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_docx_em_obra",
                    width="stretch",
                )

            if "archivo_em_obra_excel" in st.session_state:
                st.download_button(
                    label="⬇️ Descargar Excel Estudio de mercado obra",
                    data=st.session_state["archivo_em_obra_excel"],
                    file_name="anexo_estudio_mercado_obra.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="download_excel_em_obra",
                    width="stretch",
                )

with tab3:
    st.subheader("AIU")

    with st.container(border=True):
        cfg["incluye_aiu"] = st.checkbox(
            "Incluir anexo AIU",
            value=bool(cfg.get("incluye_aiu", True)),
            key="incluye_aiu_cfg",
        )

        cfg["notas_aiu"] = st.text_area(
            "Notas del anexo AIU",
            value=cfg.get("notas_aiu", ""),
            height=120,
            key="notas_aiu_cfg",
        )

    with st.container(border=True):
        if tipo_proyecto != "Obra":
            st.info("El AIU normalmente aplica a Obra. En este proyecto la inclusión puede no ser necesaria.")
        else:
            aiu_data = _aiu_datos()

            st.markdown("## 1.1 Personal administrativo")
            if not aiu_data["personal_df"].empty:
                st.dataframe(aiu_data["personal_df"], width="stretch", hide_index=True)
            else:
                st.info("Sin información disponible.")

            st.markdown("## 1.2 Equipos generales, movilización e instalación")
            if not aiu_data["equipos_df"].empty:
                st.dataframe(aiu_data["equipos_df"], width="stretch", hide_index=True)
            else:
                st.info("Sin información disponible.")

            st.markdown("## 1.3 Gastos generales")
            if not aiu_data["gastos_generales_df"].empty:
                st.dataframe(aiu_data["gastos_generales_df"], width="stretch", hide_index=True)
            else:
                st.info("Sin información disponible.")

            st.markdown("## 1.4 Gastos legales, jurídicos, tributarios")
            if not aiu_data["gastos_legales_df"].empty:
                st.dataframe(aiu_data["gastos_legales_df"], width="stretch", hide_index=True)
            else:
                st.info("Sin información disponible.")

            st.markdown("## RESUMEN FINAL DEL AIU")
            if not aiu_data["resumen_df"].empty:
                st.dataframe(aiu_data["resumen_df"], width="stretch", hide_index=True)
            else:
                st.info("Sin resumen disponible.")

            c_aiu_1, c_aiu_2 = st.columns([5, 1])
            with c_aiu_1:
                st.markdown("### TOTAL AIU")
            with c_aiu_2:
                st.markdown(f"## $ {aiu_data['aiu_total_valor']:,.2f}")

            c_aiu_3, c_aiu_4 = st.columns([5, 1])
            with c_aiu_3:
                st.markdown("### % TOTAL AIU")
            with c_aiu_4:
                st.markdown(f"## {aiu_data['aiu_total_pct']:,.2f}%")

            col_aiu_1, col_aiu_2 = st.columns(2)

            with col_aiu_1:
                if st.button("📥 Generar Word AIU", key="btn_docx_aiu", width="stretch"):
                    st.session_state["archivo_aiu_docx"] = _generar_docx_aiu(aiu_data)
                    st.success("Documento Word de AIU generado.")

            with col_aiu_2:
                if st.button("📊 Generar Excel AIU", key="btn_excel_aiu", width="stretch"):
                    st.session_state["archivo_aiu_excel"] = _generar_excel_aiu(aiu_data)
                    st.success("Archivo Excel de AIU generado.")

            if "archivo_aiu_docx" in st.session_state:
                st.download_button(
                    label="⬇️ Descargar Word AIU",
                    data=st.session_state["archivo_aiu_docx"],
                    file_name="anexo_aiu.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_docx_aiu",
                    width="stretch",
                )

            if "archivo_aiu_excel" in st.session_state:
                st.download_button(
                    label="⬇️ Descargar Excel AIU",
                    data=st.session_state["archivo_aiu_excel"],
                    file_name="anexo_aiu.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="download_excel_aiu",
                    width="stretch",
                )

with tab4:
    st.subheader("Factor multiplicador")

    with st.container(border=True):
        cfg["incluye_factor_multiplicador"] = st.checkbox(
            "Incluir anexo Factor multiplicador",
            value=bool(cfg.get("incluye_factor_multiplicador", True)),
            key="incluye_factor_multiplicador_cfg",
        )

        cfg["notas_factor_multiplicador"] = st.text_area(
            "Notas del anexo Factor multiplicador",
            value=cfg.get("notas_factor_multiplicador", ""),
            height=120,
            key="notas_factor_multiplicador_cfg",
        )

    with st.container(border=True):
        fm_data = _factor_multiplicador_datos()

        st.markdown("## 1. Factor prestacional")
        if not fm_data["prestacional_df"].empty:
            st.dataframe(fm_data["prestacional_df"], width="stretch", hide_index=True)
        else:
            st.info("Sin información disponible.")
        c_fm_1, c_fm_2 = st.columns([5, 1])
        with c_fm_1:
            st.markdown("### Total factor prestacional")
        with c_fm_2:
            st.markdown(f"## {fm_data['total_prestacional']:,.4f}")

        st.divider()

        st.markdown("## 2. Costos indirectos")
        if not fm_data["costos_indirectos_df"].empty:
            st.dataframe(fm_data["costos_indirectos_df"], width="stretch", hide_index=True)
        else:
            st.info("Sin información disponible.")
        c_fm_3, c_fm_4 = st.columns([5, 1])
        with c_fm_3:
            st.markdown("### Total costos indirectos")
        with c_fm_4:
            st.markdown(f"## {fm_data['total_costos_indirectos']:,.4f}")

        st.divider()

        st.markdown("## 3. Utilidad")
        if not fm_data["utilidad_df"].empty:
            st.dataframe(fm_data["utilidad_df"], width="stretch", hide_index=True)
        else:
            st.info("Sin información disponible.")
        c_fm_5, c_fm_6 = st.columns([5, 1])
        with c_fm_5:
            st.markdown("### Total utilidad")
        with c_fm_6:
            st.markdown(f"## {fm_data['total_utilidad']:,.4f}")

        st.divider()

        c_fm_7, c_fm_8 = st.columns([5, 1])
        with c_fm_7:
            st.markdown("## FACTOR MULTIPLICADOR FINAL")
        with c_fm_8:
            st.markdown(f"## {fm_data['factor_multiplicador_final']:,.4f}")
        col_fm_1, col_fm_2 = st.columns(2)

        with col_fm_1:
            if st.button("📥 Generar Word Factor multiplicador", key="btn_docx_factor_multiplicador", width="stretch"):
                st.session_state["archivo_factor_multiplicador_docx"] = _generar_docx_factor_multiplicador(fm_data)
                st.success("Documento Word de Factor multiplicador generado.")

        with col_fm_2:
            if st.button("📊 Generar Excel Factor multiplicador", key="btn_excel_factor_multiplicador", width="stretch"):
                st.session_state["archivo_factor_multiplicador_excel"] = _generar_excel_factor_multiplicador(fm_data)
                st.success("Archivo Excel de Factor multiplicador generado.")

        if "archivo_factor_multiplicador_docx" in st.session_state:
            st.download_button(
                label="⬇️ Descargar Word Factor multiplicador",
                data=st.session_state["archivo_factor_multiplicador_docx"],
                file_name="anexo_factor_multiplicador.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_docx_factor_multiplicador",
                width="stretch",
            )

        if "archivo_factor_multiplicador_excel" in st.session_state:
            st.download_button(
                label="⬇️ Descargar Excel Factor multiplicador",
                data=st.session_state["archivo_factor_multiplicador_excel"],
                file_name="anexo_factor_multiplicador.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="download_excel_factor_multiplicador",
                width="stretch",
            )
with tab5:
    st.subheader("Documento combinado")

    with st.container(border=True):
        st.markdown("**Seleccione qué anexos unir**")

        incluir_apus_obra = st.checkbox(
            "Incluir APUS obra",
            value=True,
            key="comb_anexos_apus_obra",
        )
        incluir_em_obra = st.checkbox(
            "Incluir Estudio de mercado obra",
            value=True,
            key="comb_anexos_em_obra",
        )
        incluir_aiu = st.checkbox(
            "Incluir AIU",
            value=True,
            key="comb_anexos_aiu",
        )
        incluir_factor = st.checkbox(
            "Incluir Factor multiplicador",
            value=True,
            key="comb_anexos_factor",
        )

    modulos = []
    nombres = []

    if incluir_apus_obra:
        modulos.append("apus_obra")
        nombres.append("APUS obra")
    if incluir_em_obra:
        modulos.append("em_obra")
        nombres.append("Estudio de mercado obra")
    if incluir_aiu:
        modulos.append("aiu")
        nombres.append("AIU")
    if incluir_factor:
        modulos.append("factor_multiplicador")
        nombres.append("Factor multiplicador")

    with st.container(border=True):
        st.markdown("**Resumen del documento combinado**")
        if nombres:
            st.write("Módulos seleccionados:", ", ".join(nombres))
        else:
            st.info("No ha seleccionado módulos para unir.")

    if st.button("📥 Generar documento combinado de anexos", key="btn_docx_anexos_combinado", width="stretch"):
        if not modulos:
            st.warning("Seleccione al menos un módulo.")
        else:
            st.session_state["archivo_anexos_combinado_docx"] = _generar_docx_anexos_combinado(modulos)
            st.success("Documento combinado de anexos generado.")

    if "archivo_anexos_combinado_docx" in st.session_state:
        st.download_button(
            label="⬇️ Descargar documento combinado de anexos",
            data=st.session_state["archivo_anexos_combinado_docx"],
            file_name="documento_combinado_anexos.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_docx_anexos_combinado",
            width="stretch",
        )
st.session_state["informe_anexos_config"] = _normalizar_cfg(cfg)
