
from io import BytesIO
from datetime import date, datetime, timedelta
import re

import pandas as pd
import streamlit as st

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

from supabase_state import cargar_estado
from supabase_state import guardar_estado as guardar_estado_bd


CLAVE_GUARDADO = "acta_entrega_recibo_definitivo_obra"


# ==========================================================
# Persistencia
# ==========================================================
def _serializar(obj):
    if isinstance(obj, dict):
        return {k: _serializar(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_serializar(x) for x in obj]
    if isinstance(obj, (date, datetime)):
        return obj.isoformat()
    if pd.isna(obj) if not isinstance(obj, (list, dict, str, bytes)) else False:
        return ""
    return obj


def guardar_estado(clave, datos):
    guardar_estado_bd(clave, _serializar(datos))


def _leer_estado(clave):
    datos = cargar_estado(clave) or {}
    return datos if isinstance(datos, dict) else {}


# ==========================================================
# Helpers
# ==========================================================
def _texto(valor):
    if valor is None:
        return ""
    if isinstance(valor, float) and pd.isna(valor):
        return ""
    return str(valor).strip()


def _primero_no_vacio(*valores):
    for valor in valores:
        txt = _texto(valor)
        if txt:
            return txt
    return ""


def _safe_float(valor, default=0.0):
    try:
        if valor is None or valor == "":
            return float(default)
        if isinstance(valor, (int, float)):
            if pd.isna(valor):
                return float(default)
            return float(valor)
        txt = str(valor).strip().replace("$", "").replace(" ", "")
        txt = txt.replace("COP", "").replace("cop", "")
        if "," in txt and "." in txt:
            txt = txt.replace(".", "").replace(",", ".")
        elif "," in txt:
            txt = txt.replace(",", ".")
        return float(txt)
    except Exception:
        return float(default)


def _parse_fecha(valor, default=None):
    if isinstance(valor, datetime):
        return valor.date()
    if isinstance(valor, date):
        return valor
    if isinstance(valor, str) and valor.strip():
        txt = valor.strip()
        try:
            return datetime.fromisoformat(txt).date()
        except Exception:
            pass
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
            try:
                return datetime.strptime(txt, fmt).date()
            except Exception:
                continue
    return default if default is not None else date.today()


def _parse_fecha_opcional(valor):
    if valor in (None, ""):
        return None
    return _parse_fecha(valor, None)


def _fecha_texto(valor):
    fecha = _parse_fecha_opcional(valor)
    return "" if fecha is None else fecha.strftime("%d/%m/%Y")


def _moneda(valor):
    numero = _safe_float(valor, 0.0)
    return "$ " + f"{numero:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _numero(valor, decimales=2):
    numero = _safe_float(valor, 0.0)
    return f"{numero:,.{decimales}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _key_codigo_natural(valor):
    partes = []
    for chunk in _texto(valor).split("."):
        try:
            partes.append(int(chunk))
        except Exception:
            partes.append(chunk)
    return tuple(partes)


def _extraer_dias_plazo(valor):
    txt = _texto(valor).lower()
    nums = re.findall(r"\d+", txt)
    if not nums:
        return 0
    n = int(nums[0])
    if "mes" in txt:
        return n * 30
    if "semana" in txt:
        return n * 7
    return n


# ==========================================================
# Lectura de fuentes
# ==========================================================
def _leer_acta_inicio():
    return _leer_estado("acta_inicio_obra")


def _leer_contrato_obra():
    return _leer_estado("contrato_obra")


def _leer_contrato_interventoria():
    return _leer_estado("contrato_interventoria")


def _leer_control_obra():
    return _leer_estado("control_obra")


def _leer_presupuesto_obra():
    return _leer_estado("presupuesto_obra")


def _leer_plan_anticipo():
    return _leer_estado("plan_inversion_anticipo")


# ==========================================================
# Datos generales
# ==========================================================
def _valor_contrato(contrato_obra, acta_inicio):
    for valor in [
        acta_inicio.get("valor_contrato"),
        acta_inicio.get("valor_inicial"),
        contrato_obra.get("valor_total_numeros"),
        contrato_obra.get("valor_contrato"),
        contrato_obra.get("valor"),
    ]:
        numero = _safe_float(valor, 0.0)
        if numero > 0:
            return round(numero, 2)
    return 0.0


def _fecha_inicio(acta_inicio):
    return _parse_fecha(
        _primero_no_vacio(
            acta_inicio.get("fecha_inicio"),
            acta_inicio.get("fecha_presente_acta"),
            acta_inicio.get("fecha_acta_inicio"),
        ),
        date.today(),
    )


def _fecha_vencimiento_inicial(acta_inicio, contrato_obra):
    directa = _primero_no_vacio(
        acta_inicio.get("fecha_terminacion"),
        acta_inicio.get("fecha_terminacion_contrato"),
        acta_inicio.get("fecha_finalizacion"),
    )
    if directa:
        return _parse_fecha(directa, date.today())

    inicio = _fecha_inicio(acta_inicio)
    plazo_dias = int(_safe_float(acta_inicio.get("plazo_ejecucion_dias"), 0.0))
    if plazo_dias <= 0:
        plazo_dias = int(_safe_float(contrato_obra.get("plazo_ejecucion_dias"), 0.0))
    if plazo_dias <= 0:
        plazo_dias = _extraer_dias_plazo(
            _primero_no_vacio(acta_inicio.get("plazo_ejecucion"), contrato_obra.get("plazo_ejecucion"))
        )
    return inicio + timedelta(days=plazo_dias)


def _datos_generales(acta_inicio, contrato_obra, contrato_interventoria):
    return {
        "numero_contrato": _primero_no_vacio(acta_inicio.get("numero_contrato"), contrato_obra.get("numero_contrato")),
        "objeto": _primero_no_vacio(
            acta_inicio.get("objeto_contrato"),
            contrato_obra.get("objeto_general"),
            contrato_obra.get("objeto_contrato"),
            contrato_obra.get("objeto"),
        ),
        "contratista": _primero_no_vacio(
            acta_inicio.get("nombre_firma_contratista"),
            contrato_obra.get("nombre_contratista"),
            contrato_obra.get("contratista"),
        ),
        "fecha_inicio": _fecha_inicio(acta_inicio),
        "plazo_inicial": _primero_no_vacio(acta_inicio.get("plazo_ejecucion"), contrato_obra.get("plazo_ejecucion")),
        "valor_inicial": _valor_contrato(contrato_obra, acta_inicio),
        "fecha_vencimiento_inicial": _fecha_vencimiento_inicial(acta_inicio, contrato_obra),
        "contrato_interventoria": _primero_no_vacio(
            contrato_interventoria.get("numero_contrato"),
            contrato_interventoria.get("numero_proceso_contratacion"),
            contrato_interventoria.get("contrato_interventoria"),
        ),
        "interventor": _primero_no_vacio(
            acta_inicio.get("nombre_firma_interventor"),
            contrato_obra.get("nombre_interventor"),
            contrato_obra.get("interventor"),
        ),
    }


# ==========================================================
# AIU, pagos, anticipo y modificaciones desde fuentes
# ==========================================================
def _aiu_desde_presupuesto(presupuesto_obra):
    config = presupuesto_obra.get("configuracion") or {}
    administracion = _safe_float(_primero_no_vacio(config.get("administracion"), config.get("administracion_pct"), config.get("a_pct"), presupuesto_obra.get("administracion")), 0.0)
    imprevistos = _safe_float(_primero_no_vacio(config.get("imprevistos"), config.get("imprevistos_pct"), config.get("i_pct"), presupuesto_obra.get("imprevistos")), 0.0)
    utilidad = _safe_float(_primero_no_vacio(config.get("utilidad"), config.get("utilidad_pct"), config.get("u_pct"), presupuesto_obra.get("utilidad")), 0.0)
    aiu_total = _safe_float(_primero_no_vacio(config.get("aiu_pct_global"), config.get("aiu_total"), presupuesto_obra.get("aiu_pct_global"), presupuesto_obra.get("aiu_total")), 0.0)
    if aiu_total <= 0:
        aiu_total = administracion + imprevistos + utilidad
    return {
        "aiu_total": round(aiu_total, 2),
        "administracion": round(administracion, 2),
        "imprevistos": round(imprevistos, 2),
        "utilidad": round(utilidad, 2),
    }


def _valor_anticipo_desde_fuentes(plan_anticipo, control_obra, valor_contrato):
    for valor in [
        control_obra.get("valor_anticipo"),
        control_obra.get("anticipo_valor"),
        plan_anticipo.get("valor_anticipo"),
    ]:
        numero = _safe_float(valor, 0.0)
        if numero > 0:
            return round(numero, 2)

    porcentaje = _safe_float(plan_anticipo.get("porcentaje_anticipo"), 0.0)
    if porcentaje > 0 and valor_contrato > 0:
        return round(valor_contrato * porcentaje / 100.0, 2)
    return 0.0


def _anticipo_desde_control(plan_anticipo, control_obra, valor_contrato):
    concedido = _valor_anticipo_desde_fuentes(plan_anticipo, control_obra, valor_contrato)
    rows = control_obra.get("anticipo_rows", []) or []
    amortizado = 0.0
    for fila in rows:
        if isinstance(fila, dict):
            amortizado += _safe_float(fila.get("VALOR AMORTIZADO"), 0.0)
    return {
        "valor_total_concedido": round(concedido, 2),
        "valor_total_amortizado": round(amortizado, 2),
        "saldo_por_amortizar": round(concedido - amortizado, 2),
    }


def _pagos_desde_control(control_obra):
    rows = control_obra.get("pagos_rows", []) or []
    salida = []
    for fila in rows:
        if not isinstance(fila, dict):
            continue
        salida.append(
            {
                "FECHA": _parse_fecha_opcional(fila.get("FECHA")),
                "VALOR INICIAL": round(_safe_float(fila.get("VALOR INICIAL"), 0.0), 2),
                "VALOR FACTURADO": round(_safe_float(fila.get("VALOR FACTURADO"), 0.0), 2),
                "PENDIENTE POR FACTURAR": round(_safe_float(fila.get("PENDIENTE POR FACTURAR"), 0.0), 2),
            }
        )
    return salida


def _adiciones_desde_control(control_obra):
    rows = control_obra.get("adiciones_rows", []) or []
    salida = []
    for fila in rows:
        if not isinstance(fila, dict):
            continue
        if not _texto(fila.get("ADICIONAL No.")) and _safe_float(fila.get("VALOR"), 0.0) <= 0:
            continue
        salida.append(
            {
                "ADICIONAL No.": _texto(fila.get("ADICIONAL No.")),
                "FECHA": _parse_fecha_opcional(fila.get("FECHA")),
                "VALOR": round(_safe_float(fila.get("VALOR"), 0.0), 2),
                "VALOR ACUMULADO DEL CONTRATO": round(_safe_float(fila.get("VALOR ACUMULADO DEL CONTRATO"), 0.0), 2),
            }
        )
    return salida


def _suspensiones_desde_control(control_obra):
    rows = control_obra.get("suspensiones_rows", []) or []
    salida = []
    for fila in rows:
        if not isinstance(fila, dict):
            continue
        numero = _primero_no_vacio(fila.get("ACTA DE SUSPENSIÓN No."), fila.get("ACTA DE AMPLIACIÓN SUSPENSIÓN No."))
        if not numero and not fila.get("DESDE") and not fila.get("HASTA"):
            continue
        salida.append(
            {
                "ACTA": numero,
                "TIPO": "AMPLIACIÓN" if _texto(fila.get("ACTA DE AMPLIACIÓN SUSPENSIÓN No.")) else "SUSPENSIÓN",
                "FECHA DEL ACTA": _parse_fecha_opcional(fila.get("FECHA DEL ACTA")),
                "DESDE": _parse_fecha_opcional(fila.get("DESDE")),
                "HASTA": _parse_fecha_opcional(fila.get("HASTA")),
                "PERIODO DE SUSPENSIÓN": int(_safe_float(fila.get("PERIODO DE SUSPENSIÓN"), 0.0)),
                "NUEVA FECHA DE FINALIZACIÓN": _parse_fecha_opcional(fila.get("NUEVA FECHA DE FINALIZACIÓN")),
            }
        )
    return salida


def _valor_acumulado(valor_inicial, adiciones):
    total = _safe_float(valor_inicial, 0.0)
    for fila in adiciones:
        total += _safe_float(fila.get("VALOR"), 0.0)
    return round(total, 2)


# ==========================================================
# Presupuesto: items para la tabla de cantidades
# ==========================================================
def _items_desde_presupuesto(presupuesto_obra):
    filas = []
    vistos = set()

    tablas = presupuesto_obra.get("__tablas__", {}) or {}
    grupos = tablas.get("grupos_presupuesto_obra", []) or presupuesto_obra.get("grupos_presupuesto_obra", []) or []
    flujo_directos = presupuesto_obra.get("flujo_fondos_directos", []) or []
    items_dict = presupuesto_obra.get("items", {}) or {}

    def tomar_valor_unitario(fila):
        return _safe_float(
            _primero_no_vacio(
                fila.get("VR AFECTADO POR FACTOR"),
                fila.get("VALOR UNITARIO AJUSTADO POR DISTANCIA"),
                fila.get("VALOR AFECTADO POR FACTOR"),
                fila.get("VR_AFECTADO_POR_FACTOR"),
                fila.get("VALOR_UNITARIO_AJUSTADO_POR_DISTANCIA"),
                fila.get("VALOR UNITARIO"),
                fila.get("VR UNITARIO"),
                fila.get("VALOR_UNITARIO"),
            ),
            0.0,
        )

    def agregar(item, descripcion, unidad, valor_unitario):
        item = _texto(item)
        if not item or item in vistos:
            return
        vistos.add(item)
        filas.append(
            {
                "No. ORDEN": item,
                "DESCRIPCIÓN ITEM": _texto(descripcion),
                "UNIDAD": _texto(unidad),
                "CANTIDAD": 0.0,
                "VALOR UNITARIO AJUSTADO POR DISTANCIA": round(_safe_float(valor_unitario, 0.0), 2),
                "VALOR TOTAL EJECUTADO": 0.0,
            }
        )

    for grupo in grupos:
        if not isinstance(grupo, dict):
            continue
        for fila in grupo.get("rows", []) or []:
            if not isinstance(fila, dict):
                continue
            item = _texto(fila.get("ITEM") or fila.get("ITEM GOBER") or fila.get("No. ORDEN"))
            descripcion = _primero_no_vacio(fila.get("DESCRIPCIÓN"), fila.get("DESCRIPCION"), fila.get("DESCRIPCIÓN ITEM"))
            unidad = _primero_no_vacio(fila.get("UNIDAD"), fila.get("unidad"), fila.get("UND"))
            valor_unitario = tomar_valor_unitario(fila)
            agregar(item, descripcion, unidad, valor_unitario)

    for fila in flujo_directos:
        if not isinstance(fila, dict):
            continue
        item = _texto(fila.get("ITEM") or fila.get("No. ORDEN"))
        descripcion = _primero_no_vacio(fila.get("DESCRIPCIÓN"), fila.get("DESCRIPCION"), fila.get("DESCRIPCIÓN ITEM"))
        unidad = _primero_no_vacio(fila.get("UNIDAD"), fila.get("unidad"), fila.get("UND"))
        valor_unitario = tomar_valor_unitario(fila)
        if valor_unitario <= 0:
            valor_base = _safe_float(fila.get("VALOR BASE"), 0.0)
            cantidad = _safe_float(_primero_no_vacio(fila.get("CANT"), fila.get("CANTIDAD"), fila.get("CANTIDAD TOTAL")), 0.0)
            if cantidad > 0:
                valor_unitario = valor_base / cantidad
        agregar(item, descripcion, unidad, valor_unitario)

    return sorted(filas, key=lambda x: _key_codigo_natural(x.get("No. ORDEN")))


def _mapa_items(items):
    return {_texto(f.get("No. ORDEN")): f for f in items if _texto(f.get("No. ORDEN"))}


def _normalizar_cantidades(rows, mapa_items):
    salida = []
    for fila in rows or []:
        if not isinstance(fila, dict):
            continue
        item = _texto(fila.get("No. ORDEN"))
        if not item:
            continue
        base = mapa_items.get(item, {})
        descripcion = _primero_no_vacio(fila.get("DESCRIPCIÓN ITEM"), base.get("DESCRIPCIÓN ITEM"))
        unidad = _primero_no_vacio(fila.get("UNIDAD"), base.get("UNIDAD"))
        valor_unitario = _safe_float(
            _primero_no_vacio(fila.get("VALOR UNITARIO AJUSTADO POR DISTANCIA"), base.get("VALOR UNITARIO AJUSTADO POR DISTANCIA")),
            0.0,
        )
        valor_total = _safe_float(fila.get("VALOR TOTAL EJECUTADO"), 0.0)
        cantidad = round(valor_total / valor_unitario, 4) if valor_unitario > 0 else 0.0
        salida.append(
            {
                "No. ORDEN": item,
                "DESCRIPCIÓN ITEM": descripcion,
                "UNIDAD": unidad,
                "CANTIDAD": cantidad,
                "VALOR UNITARIO AJUSTADO POR DISTANCIA": round(valor_unitario, 2),
                "VALOR TOTAL EJECUTADO": round(valor_total, 2),
            }
        )
    return salida


# ==========================================================
# Word
# ==========================================================
def _doc_parrafo(doc, texto="", bold=False, size=9, align=None):
    p = doc.add_paragraph()
    r = p.add_run(str(texto))
    r.bold = bold
    r.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    return p


def _doc_titulo(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _doc_tabla_kv(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for etiqueta, valor in rows:
        cells = table.add_row().cells
        cells[0].text = str(etiqueta)
        cells[1].text = str(valor)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def _formatear_valor_word(valor):
    if isinstance(valor, (date, datetime)):
        return _fecha_texto(valor)
    if isinstance(valor, (int, float)):
        return _numero(valor, 2)
    return _texto(valor)


def _doc_tabla_df(doc, df, columnas=None):
    if df is None or df.empty:
        df = pd.DataFrame([{}])
    if columnas is None:
        columnas = list(df.columns)
    table = doc.add_table(rows=1, cols=len(columnas))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(columnas):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columnas):
            cells[i].text = _formatear_valor_word(row.get(col, ""))
    return table


def _generar_word(payload):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    generales = payload.get("generales", {})
    aiu = payload.get("aiu", {})
    anticipo = payload.get("anticipo", {})

    _doc_titulo(doc, "ACTA DE ENTREGA Y RECIBO DEFINITIVO DE OBRA")
    _doc_tabla_kv(
        doc,
        [
            ("FECHA", _fecha_texto(payload.get("fecha"))),
            ("CONTRATO DE OBRA No.", generales.get("numero_contrato", "")),
            ("OBJETO DEL CONTRATO", generales.get("objeto", "")),
            ("CONTRATISTA", generales.get("contratista", "")),
            ("FECHA DE INICIO", _fecha_texto(generales.get("fecha_inicio"))),
            ("PLAZO INICIAL", generales.get("plazo_inicial", "")),
            ("VALOR INICIAL", _moneda(generales.get("valor_inicial", 0.0))),
            ("FECHA DE VENCIMIENTO", _fecha_texto(payload.get("fecha_vencimiento"))),
            ("VALOR ACUMULADO", _moneda(payload.get("valor_acumulado", 0.0))),
            ("CONTRATO DE INTERVENTORÍA No.", generales.get("contrato_interventoria", "")),
            ("INTERVENTOR", generales.get("interventor", "")),
        ],
    )

    doc.add_paragraph()
    _doc_parrafo(doc, "ADICIONALES", bold=True)
    _doc_tabla_df(doc, pd.DataFrame(payload.get("adiciones", [])))

    doc.add_paragraph()
    _doc_parrafo(doc, "RELACIÓN SUSPENSIONES Y AMPLIACIONES DE SUSPENSIÓN", bold=True)
    _doc_tabla_df(doc, pd.DataFrame(payload.get("suspensiones", [])))

    doc.add_paragraph()
    _doc_parrafo(doc, "DESCRIPCIÓN GENERAL DE LA OBRA EJECUTADA", bold=True)
    _doc_parrafo(doc, "LOCALIZACIÓN Y LONGITUD DE LA META FÍSICA EJECUTADA (Indicar PR):", bold=True)
    _doc_parrafo(doc, payload.get("localizacion_meta", ""))
    _doc_parrafo(doc, "CARACTERÍSTICAS TÉCNICAS GENERALES Y ALCANCE DE LA OBRA QUE SE ENTREGA Y RECIBE", bold=True)
    _doc_parrafo(doc, payload.get("caracteristicas_tecnicas", ""))
    _doc_parrafo(doc, "RELACIÓN DE SITIOS CRÍTICOS ATENDIDOS Y/O SEÑALIZADOS (Indicar PR):", bold=True)
    _doc_parrafo(doc, payload.get("sitios_criticos", ""))

    doc.add_paragraph()
    _doc_parrafo(doc, "Con la suscripción del Acta de Entrega y Recibo Definitivo de Obra el Contratista y el Interventor certifican que las obras ejecutadas cumplen con la calidad, normas y especificaciones contractuales.")

    doc.add_paragraph()
    _doc_parrafo(doc, "VALOR TOTAL EJECUTADO DEL CONTRATO DE OBRA", bold=True)
    _doc_tabla_kv(
        doc,
        [
            ("Valor total ejecutado del contrato", _moneda(payload.get("valor_total_ejecutado_contrato", 0.0))),
            ("AIU", f"{_numero(aiu.get('aiu_total', 0.0), 2)}%"),
            ("A", f"{_numero(aiu.get('administracion', 0.0), 2)}%"),
            ("I", f"{_numero(aiu.get('imprevistos', 0.0), 2)}%"),
            ("U", f"{_numero(aiu.get('utilidad', 0.0), 2)}%"),
        ],
    )

    doc.add_paragraph()
    _doc_parrafo(doc, "DESCRIPCIÓN CANTIDADES DE OBRA Y PROVISIONES EJECUTADAS", bold=True)
    _doc_parrafo(doc, "El Interventor debe ajustar y discriminar de forma independiente los valores ejecutados de los conceptos determinados en la propuesta económica aprobada.")
    _doc_tabla_df(
        doc,
        pd.DataFrame(payload.get("cantidades_rows", [])),
        ["No. ORDEN", "DESCRIPCIÓN ITEM", "UNIDAD", "VALOR UNITARIO AJUSTADO POR DISTANCIA", "CANTIDAD", "VALOR TOTAL EJECUTADO"],
    )
    _doc_tabla_kv(doc, [("VALOR TOTAL EJECUTADO", _moneda(payload.get("valor_total_cantidades", 0.0)))])

    doc.add_paragraph()
    _doc_parrafo(doc, "ANTICIPO", bold=True)
    _doc_tabla_kv(
        doc,
        [
            ("Valor total concedido", _moneda(anticipo.get("valor_total_concedido", 0.0))),
            ("Valor total amortizado", _moneda(anticipo.get("valor_total_amortizado", 0.0))),
            ("Saldo por amortizar", _moneda(anticipo.get("saldo_por_amortizar", 0.0))),
        ],
    )
    _doc_parrafo(doc, "OBSERVACIONES DEL INTERVENTOR RESPECTO DEL ESTADO DE EJECUCIÓN Y AMORTIZACIÓN DEL ANTICIPO Y/O EJECUCIÓN Y LEGALIZACIÓN DEL PAGO ANTICIPADO", bold=True)
    _doc_parrafo(doc, payload.get("observaciones_anticipo", ""))

    doc.add_paragraph()
    _doc_parrafo(doc, "RESUMEN FINANCIERO DEL CONTRATO", bold=True)
    _doc_tabla_df(doc, pd.DataFrame(payload.get("pagos_rows", [])))
    _doc_tabla_kv(doc, [("VALOR TOTAL EJECUTADO", _moneda(payload.get("valor_total_ejecutado_resumen", 0.0)))])

    doc.add_paragraph()
    _doc_parrafo(doc, "CONCEPTO DE LA INTERVENTORÍA RECIBO DE LAS OBRAS", bold=True)
    _doc_parrafo(doc, payload.get("concepto_interventoria", ""))
    _doc_parrafo(doc, "CONCEPTO DE LA INTERVENTORÍA SOBRE LA APROBACIÓN DE LOS PLANOS RECORD", bold=True)
    _doc_parrafo(doc, payload.get("concepto_planos_record", ""))
    _doc_parrafo(doc, "GESTIÓN AMBIENTAL, SOCIAL, PREDIAL Y DE SOSTENIBILIDAD", bold=True)
    _doc_parrafo(doc, payload.get("gestion_ambiental_social_predial", ""))
    _doc_parrafo(doc, "OBSERVACIONES DEL INTERVENTOR RESPECTO DEL ESTADO DE LAS GARANTÍAS Y SEGUROS CONTRACTUALES EXIGIDOS AL CONTRATISTA DE OBRA", bold=True)
    _doc_parrafo(doc, payload.get("observaciones_garantias", ""))
    _doc_parrafo(doc, "OBSERVACIONES", bold=True)
    _doc_parrafo(doc, payload.get("observaciones_generales", ""))

    doc.add_paragraph()
    _doc_parrafo(doc, "Mediante la suscripción del Acta de entrega y recibo Definitivo de obra se asume plena responsabilidad por la veracidad de los valores registrados en los formatos descritos, así como, por las operaciones aritméticas contenidas en los mismos, pero no exonera al Contratista ni al Interventor de las obligaciones y responsabilidades estipuladas en el contrato; en consecuencia, si dentro del periodo de vigencia de la garantía de estabilidad y/o calidad, se presentan fallas imputables a la calidad de la obra, el CONTRATANTE debe exigir al constructor, las reparaciones del caso o en su defecto hará efectivas las garantías de estabilidad y/o calidad correspondientes.")
    _doc_parrafo(doc, "Para constancia de lo anterior, suscriben la presente Acta de Entrega y Recibo Definitivo de Obra quienes en ella intervienen.")

    doc.add_paragraph()
    _doc_parrafo(doc, "FIRMAS", bold=True)
    _doc_tabla_kv(doc, [("Contratista", payload.get("firma_contratista", "")), ("Interventoría", payload.get("firma_interventoria", ""))])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==========================================================
# Estado inicial
# ==========================================================
def _estado_inicial(generales):
    return {
        "fecha": date.today(),
        "fecha_vencimiento": generales.get("fecha_vencimiento_inicial", date.today()),
        "valor_acumulado": generales.get("valor_inicial", 0.0),
        "localizacion_meta": "",
        "caracteristicas_tecnicas": "",
        "sitios_criticos": "",
        "valor_total_ejecutado_contrato": 0.0,
        "cantidades_rows": [],
        "observaciones_anticipo": "",
        "valor_total_ejecutado_resumen": 0.0,
        "concepto_interventoria": "",
        "concepto_planos_record": "",
        "gestion_ambiental_social_predial": "",
        "observaciones_garantias": "",
        "observaciones_generales": "",
        "firma_contratista": generales.get("contratista", ""),
        "firma_interventoria": generales.get("interventor", ""),
    }


# ==========================================================
# Interfaz
# ==========================================================
st.set_page_config(page_title="Acta entrega y recibo definitivo obra", layout="wide")
st.title("ACTA DE ENTREGA Y RECIBO DEFINITIVO DE OBRA")

acta_inicio = _leer_acta_inicio()
contrato_obra = _leer_contrato_obra()
contrato_interventoria = _leer_contrato_interventoria()
control_obra = _leer_control_obra()
presupuesto_obra = _leer_presupuesto_obra()
plan_anticipo = _leer_plan_anticipo()

generales = _datos_generales(acta_inicio, contrato_obra, contrato_interventoria)
aiu = _aiu_desde_presupuesto(presupuesto_obra)
adiciones = _adiciones_desde_control(control_obra)
suspensiones = _suspensiones_desde_control(control_obra)
pagos_rows = _pagos_desde_control(control_obra)
anticipo = _anticipo_desde_control(plan_anticipo, control_obra, generales.get("valor_inicial", 0.0))
items_presupuesto = _items_desde_presupuesto(presupuesto_obra)
mapa_items = _mapa_items(items_presupuesto)

guardado = _leer_estado(CLAVE_GUARDADO)
if not guardado:
    guardado = _estado_inicial(generales)

valor_acumulado_fuente = _valor_acumulado(generales.get("valor_inicial", 0.0), adiciones)
if _safe_float(guardado.get("valor_acumulado"), 0.0) <= 0:
    guardado["valor_acumulado"] = valor_acumulado_fuente

st.markdown("### INFORMACIÓN GENERAL")

c1, c2, c3 = st.columns(3)
with c1:
    fecha_acta = st.date_input("FECHA", value=_parse_fecha(guardado.get("fecha"), date.today()), format="DD/MM/YYYY")
with c2:
    st.text_input("CONTRATO DE OBRA No.", value=generales.get("numero_contrato", ""), disabled=True)
with c3:
    st.text_input("CONTRATO DE INTERVENTORÍA No.", value=generales.get("contrato_interventoria", ""), disabled=True)

st.text_area("OBJETO DEL CONTRATO", value=generales.get("objeto", ""), disabled=True, height=90)

c4, c5, c6 = st.columns(3)
with c4:
    st.text_input("CONTRATISTA", value=generales.get("contratista", ""), disabled=True)
with c5:
    st.date_input("FECHA DE INICIO", value=generales.get("fecha_inicio", date.today()), format="DD/MM/YYYY", disabled=True)
with c6:
    st.text_input("PLAZO INICIAL", value=generales.get("plazo_inicial", ""), disabled=True)

c7, c8, c9 = st.columns(3)
with c7:
    st.number_input("VALOR INICIAL", value=float(generales.get("valor_inicial", 0.0)), format="%.2f", disabled=True)
with c8:
    fecha_vencimiento = st.date_input("FECHA DE VENCIMIENTO", value=_parse_fecha(guardado.get("fecha_vencimiento"), generales.get("fecha_vencimiento_inicial", date.today())), format="DD/MM/YYYY")
with c9:
    valor_acumulado = st.number_input("VALOR ACUMULADO", value=float(_safe_float(guardado.get("valor_acumulado"), valor_acumulado_fuente)), step=1000.0, format="%.2f")

st.text_input("INTERVENTOR", value=generales.get("interventor", ""), disabled=True)

st.markdown("### ADICIONALES")
df_adiciones = pd.DataFrame(adiciones)
if df_adiciones.empty:
    df_adiciones = pd.DataFrame(columns=["ADICIONAL No.", "FECHA", "VALOR", "VALOR ACUMULADO DEL CONTRATO"])
st.dataframe(
    df_adiciones,
    hide_index=True,
    width="stretch",
    column_config={
        "FECHA": st.column_config.DateColumn("FECHA", format="DD/MM/YYYY"),
        "VALOR": st.column_config.NumberColumn("VALOR", format="$ %.2f"),
        "VALOR ACUMULADO DEL CONTRATO": st.column_config.NumberColumn("VALOR ACUMULADO DEL CONTRATO", format="$ %.2f"),
    },
)

st.markdown("### RELACIÓN SUSPENSIONES Y AMPLIACIONES DE SUSPENSIÓN")
df_suspensiones = pd.DataFrame(suspensiones)
if df_suspensiones.empty:
    df_suspensiones = pd.DataFrame(columns=["ACTA", "TIPO", "FECHA DEL ACTA", "DESDE", "HASTA", "PERIODO DE SUSPENSIÓN", "NUEVA FECHA DE FINALIZACIÓN"])
st.dataframe(
    df_suspensiones,
    hide_index=True,
    width="stretch",
    column_config={
        "FECHA DEL ACTA": st.column_config.DateColumn("FECHA DEL ACTA", format="DD/MM/YYYY"),
        "DESDE": st.column_config.DateColumn("DESDE", format="DD/MM/YYYY"),
        "HASTA": st.column_config.DateColumn("HASTA", format="DD/MM/YYYY"),
        "NUEVA FECHA DE FINALIZACIÓN": st.column_config.DateColumn("NUEVA FECHA DE FINALIZACIÓN", format="DD/MM/YYYY"),
    },
)

st.markdown("### DESCRIPCIÓN GENERAL DE LA OBRA EJECUTADA")
localizacion_meta = st.text_area("LOCALIZACIÓN Y LONGITUD DE LA META FÍSICA EJECUTADA (Indicar PR)", value=guardado.get("localizacion_meta", ""), height=90)
caracteristicas_tecnicas = st.text_area("CARACTERÍSTICAS TÉCNICAS GENERALES Y ALCANCE DE LA OBRA QUE SE ENTREGA Y RECIBE", value=guardado.get("caracteristicas_tecnicas", ""), height=120)
sitios_criticos = st.text_area("RELACIÓN DE SITIOS CRÍTICOS ATENDIDOS Y/O SEÑALIZADOS (Indicar PR)", value=guardado.get("sitios_criticos", ""), height=90)

st.markdown("### VALOR TOTAL EJECUTADO DEL CONTRATO DE OBRA")
valor_total_ejecutado_contrato = st.number_input("Valor total ejecutado del contrato", value=float(_safe_float(guardado.get("valor_total_ejecutado_contrato"), 0.0)), min_value=0.0, step=1000.0, format="%.2f")

st.number_input(
    "AIU %",
    value=float(aiu.get("aiu_total", 0.0)),
    format="%.2f",
    disabled=True,
)

st.markdown("### DESCRIPCIÓN CANTIDADES DE OBRA Y PROVISIONES EJECUTADAS")
st.caption("El valor unitario se alimenta de la columna del presupuesto: VR AFECTADO POR FACTOR. La cantidad se calcula como VALOR TOTAL EJECUTADO / VALOR UNITARIO AJUSTADO POR DISTANCIA.")

opciones_items = [""] + [fila["No. ORDEN"] for fila in items_presupuesto]
item_sel = st.selectbox("Seleccionar ítem desde presupuesto", options=opciones_items, key="selector_item_acta_definitiva")

if st.button("Agregar ítem", key="agregar_item_acta_definitiva"):
    if item_sel:
        filas_actuales = guardado.get("cantidades_rows", []) or []
        existentes = [_texto(f.get("No. ORDEN")) for f in filas_actuales if isinstance(f, dict)]
        if item_sel not in existentes:
            base = mapa_items.get(item_sel, {}).copy()
            filas_actuales.append(base)
            guardado["cantidades_rows"] = _normalizar_cantidades(filas_actuales, mapa_items)
            guardar_estado(CLAVE_GUARDADO, guardado)
            st.rerun()

columnas_cantidades = [
    "No. ORDEN",
    "DESCRIPCIÓN ITEM",
    "UNIDAD",
    "VALOR UNITARIO AJUSTADO POR DISTANCIA",
    "CANTIDAD",
    "VALOR TOTAL EJECUTADO",
]

df_cantidades = pd.DataFrame(
    _normalizar_cantidades(guardado.get("cantidades_rows", []), mapa_items),
    columns=columnas_cantidades,
)

cantidades_editadas = st.data_editor(
    df_cantidades,
    hide_index=True,
    width="stretch",
    num_rows="dynamic",
    key="cantidades_definitivas_editor",
    disabled=["DESCRIPCIÓN ITEM", "UNIDAD", "VALOR UNITARIO AJUSTADO POR DISTANCIA", "CANTIDAD"],
    column_config={
        "No. ORDEN": st.column_config.TextColumn("No. ORDEN"),
        "DESCRIPCIÓN ITEM": st.column_config.TextColumn("DESCRIPCIÓN ITEM"),
        "UNIDAD": st.column_config.TextColumn("UNIDAD"),
        "VALOR UNITARIO AJUSTADO POR DISTANCIA": st.column_config.NumberColumn("VALOR UNITARIO AJUSTADO POR DISTANCIA", format="$ %.2f"),
        "CANTIDAD": st.column_config.NumberColumn("CANTIDAD", format="%.4f"),
        "VALOR TOTAL EJECUTADO": st.column_config.NumberColumn("VALOR TOTAL EJECUTADO", format="$ %.2f"),
    },
)

cantidades_recalculadas = _normalizar_cantidades(cantidades_editadas.to_dict("records"), mapa_items)
if cantidades_recalculadas != guardado.get("cantidades_rows", []):
    guardado["cantidades_rows"] = cantidades_recalculadas
    guardar_estado(CLAVE_GUARDADO, guardado)
    st.rerun()

total_cantidades = round(sum(_safe_float(f.get("VALOR TOTAL EJECUTADO"), 0.0) for f in cantidades_recalculadas), 2)
st.metric("VALOR TOTAL EJECUTADO", _moneda(total_cantidades))

st.markdown("### ANTICIPO")
c_ant1, c_ant2, c_ant3 = st.columns(3)
with c_ant1:
    st.number_input("Valor total concedido", value=float(anticipo.get("valor_total_concedido", 0.0)), format="%.2f", disabled=True)
with c_ant2:
    st.number_input("Valor total amortizado", value=float(anticipo.get("valor_total_amortizado", 0.0)), format="%.2f", disabled=True)
with c_ant3:
    st.number_input("Saldo por amortizar", value=float(anticipo.get("saldo_por_amortizar", 0.0)), format="%.2f", disabled=True)

observaciones_anticipo = st.text_area("OBSERVACIONES DEL INTERVENTOR RESPECTO DEL ESTADO DE EJECUCIÓN Y AMORTIZACIÓN DEL ANTICIPO Y/O EJECUCIÓN Y LEGALIZACIÓN DEL PAGO ANTICIPADO", value=guardado.get("observaciones_anticipo", ""), height=120)

st.markdown("### RESUMEN FINANCIERO DEL CONTRATO")
df_pagos = pd.DataFrame(pagos_rows)
if df_pagos.empty:
    df_pagos = pd.DataFrame(columns=["FECHA", "VALOR INICIAL", "VALOR FACTURADO", "PENDIENTE POR FACTURAR"])
st.dataframe(
    df_pagos,
    hide_index=True,
    width="stretch",
    column_config={
        "FECHA": st.column_config.DateColumn("FECHA", format="DD/MM/YYYY"),
        "VALOR INICIAL": st.column_config.NumberColumn("VALOR INICIAL", format="$ %.2f"),
        "VALOR FACTURADO": st.column_config.NumberColumn("VALOR FACTURADO", format="$ %.2f"),
        "PENDIENTE POR FACTURAR": st.column_config.NumberColumn("PENDIENTE POR FACTURAR", format="$ %.2f"),
    },
)
valor_total_ejecutado_resumen = st.number_input("VALOR TOTAL EJECUTADO", value=float(_safe_float(guardado.get("valor_total_ejecutado_resumen"), 0.0)), min_value=0.0, step=1000.0, format="%.2f", key="valor_total_ejecutado_resumen")

st.markdown("### CONCEPTOS Y OBSERVACIONES")
st.markdown("#### CONCEPTO DE LA INTERVENTORÍA RECIBO DE LAS OBRAS")
st.markdown(
    "La interventoría deja constancia que las obras recibidas cumplen con los requerimientos de calidad, "
    "con las normas, especificaciones generales y particulares de construcción y demás condiciones contractuales, "
    "de acuerdo con los diseños, planos y especificaciones estipuladas para el proyecto."
)
concepto_interventoria = st.text_area(
    "Observaciones adicionales sobre el recibo de las obras",
    value=guardado.get("concepto_interventoria", ""),
    height=90,
)

st.markdown("#### CONCEPTO DE LA INTERVENTORÍA SOBRE LA APROBACIÓN DE LOS PLANOS RECORD")
st.markdown(
    "Con la suscripción de la presente Acta la interventoría deja constancia de la revisión y aprobación "
    "de los planos record elaborados y entregados por el Contratista."
)
concepto_planos_record = st.text_area(
    "Observaciones adicionales sobre planos record",
    value=guardado.get("concepto_planos_record", ""),
    height=90,
)

st.markdown("#### GESTIÓN AMBIENTAL, SOCIAL, PREDIAL Y DE SOSTENIBILIDAD")
st.markdown(
    "Las actividades del componente ambiental, social, predial y de sostenibilidad son recibidas y aprobadas "
    "por la Interventoría, teniendo en cuenta que el Contratista cumplió con las actividades de conformidad "
    "con los instructivos y formatos."
)
gestion_ambiental_social_predial = st.text_area(
    "Observaciones adicionales sobre gestión ambiental, social, predial y de sostenibilidad",
    value=guardado.get("gestion_ambiental_social_predial", ""),
    height=90,
)
observaciones_generales = st.text_area("OBSERVACIONES", value=guardado.get("observaciones_generales", ""), height=90)

observaciones_garantias = st.text_area(
    "OBSERVACIONES DEL INTERVENTOR RESPECTO DEL ESTADO DE LAS GARANTÍAS Y SEGUROS CONTRACTUALES EXIGIDOS AL CONTRATISTA DE OBRA",
    value=guardado.get("observaciones_garantias", ""),
    height=90,
    key="acta_def_observaciones_garantias",
)

observaciones_generales = st.text_area(
    "OBSERVACIONES",
    value=guardado.get("observaciones_generales", ""),
    height=90,
    key="acta_def_observaciones_generales",
)

st.markdown("### FIRMAS")
c_f1, c_f2 = st.columns(2)
with c_f1:
    firma_contratista = st.text_input("Contratista", value=guardado.get("firma_contratista", generales.get("contratista", "")))
with c_f2:
    firma_interventoria = st.text_input("Interventoría", value=guardado.get("firma_interventoria", generales.get("interventor", "")))

payload = {
    "fecha": fecha_acta,
    "generales": generales,
    "fecha_vencimiento": fecha_vencimiento,
    "valor_acumulado": round(_safe_float(valor_acumulado, 0.0), 2),
    "adiciones": adiciones,
    "suspensiones": suspensiones,
    "localizacion_meta": localizacion_meta,
    "caracteristicas_tecnicas": caracteristicas_tecnicas,
    "sitios_criticos": sitios_criticos,
    "valor_total_ejecutado_contrato": round(_safe_float(valor_total_ejecutado_contrato, 0.0), 2),
    "aiu": aiu,
    "cantidades_rows": cantidades_recalculadas,
    "valor_total_cantidades": total_cantidades,
    "anticipo": anticipo,
    "observaciones_anticipo": observaciones_anticipo,
    "pagos_rows": pagos_rows,
    "valor_total_ejecutado_resumen": round(_safe_float(valor_total_ejecutado_resumen, 0.0), 2),
    "concepto_interventoria": concepto_interventoria,
    "concepto_planos_record": concepto_planos_record,
    "gestion_ambiental_social_predial": gestion_ambiental_social_predial,
    "observaciones_garantias": observaciones_garantias,
    "observaciones_generales": observaciones_generales,
    "firma_contratista": firma_contratista,
    "firma_interventoria": firma_interventoria,
}

c_btn1, c_btn2 = st.columns(2)
with c_btn1:
    if st.button("Guardar acta definitiva", key="guardar_acta_entrega_recibo_definitivo"):
        guardar_estado(CLAVE_GUARDADO, payload)
        st.success("Acta de entrega y recibo definitivo de obra guardada correctamente.")

with c_btn2:
    st.download_button(
        "Descargar Word",
        data=_generar_word(payload),
        file_name="acta_entrega_recibo_definitivo_obra.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="descargar_acta_entrega_recibo_definitivo_word",
    )
