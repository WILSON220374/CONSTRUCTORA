from io import BytesIO
from datetime import date, datetime, timedelta

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, Inches

from supabase_state import cargar_estado
from supabase_state import guardar_estado as guardar_estado_bd

# ==========================================================
# Helpers
# ==========================================================
def _texto(valor):
    if valor is None:
        return ""
    return str(valor).strip()


def _safe_float(valor, default=0.0):
    try:
        if valor is None or valor == "":
            return float(default)
        if isinstance(valor, (int, float)):
            return float(valor)
        txt = str(valor).strip().replace("$", "").replace(" ", "")
        txt = txt.replace(".", "").replace(",", ".")
        return float(txt)
    except Exception:
        return float(default)


def _parse_fecha(valor):
    if isinstance(valor, date):
        return valor
    if isinstance(valor, datetime):
        return valor.date()
    if isinstance(valor, str) and valor.strip():
        txt = valor.strip()
        try:
            return datetime.fromisoformat(txt).date()
        except Exception:
            pass
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
            try:
                return datetime.strptime(txt, fmt).date()
            except Exception:
                continue
    return None


def _fecha_texto(valor):
    fecha = _parse_fecha(valor)
    if not fecha:
        return ""
    return fecha.strftime("%d/%m/%Y")


def _leer_estado(clave):
    datos = cargar_estado(clave) or {}
    return datos if isinstance(datos, dict) else {}

def guardar_estado(clave, datos):
    def serializar(obj):
        if isinstance(obj, dict):
            return {k: serializar(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [serializar(x) for x in obj]
        if isinstance(obj, (date, datetime)):
            return obj.isoformat()
        return obj

    return guardar_estado_bd(clave, serializar(datos))


def _clave_acta_reanudacion(fila):
    tipo = "ampliacion" if _texto(fila.get("ACTA DE AMPLIACIÓN SUSPENSIÓN No.")) else "suspension"
    numero = _texto(fila.get("ACTA DE AMPLIACIÓN SUSPENSIÓN No.")) or _texto(fila.get("ACTA DE SUSPENSIÓN No.")) or "0"
    fecha = _fecha_texto(fila.get("FECHA DEL ACTA")).replace("/", "_")
    return f"{tipo}_{numero}_{fecha}"


def _cargar_actas_reanudacion_guardadas():
    datos = cargar_estado("acta_reanudacion") or {}
    return datos if isinstance(datos, dict) else {}


def _guardar_acta_reanudacion(clave_acta, payload):
    datos = _cargar_actas_reanudacion_guardadas()
    actas = datos.get("actas", {})

    if not isinstance(actas, dict):
        actas = {}

    actas[clave_acta] = payload
    datos["actas"] = actas
    datos["acta_activa"] = clave_acta

    respuesta = guardar_estado("acta_reanudacion", datos)

    if respuesta is None:
        st.error("No se pudo guardar el acta de reanudación en la nube.")
        return False

    st.success("Acta de reanudación guardada correctamente.")
    return True


def _primero_no_vacio(*valores):
    for valor in valores:
        txt = _texto(valor)
        if txt:
            return txt
    return ""


def _extraer_dias_plazo(valor):
    try:
        if isinstance(valor, (int, float)):
            return int(valor)
        txt = str(valor or "").strip()
        numeros = "".join(ch if ch.isdigit() else " " for ch in txt).split()
        return int(numeros[0]) if numeros else 0
    except Exception:
        return 0


def _valor_contrato(acta_inicio, contrato_obra):
    for valor in [
        acta_inicio.get("valor_total_contrato_obra"),
        acta_inicio.get("valor_contrato"),
        acta_inicio.get("valor"),
        contrato_obra.get("valor_total_numeros"),
        contrato_obra.get("valor_contrato"),
        contrato_obra.get("valor"),
    ]:
        numero = _safe_float(valor, 0.0)
        if numero > 0:
            return numero
    return 0.0


def _fecha_inicio(acta_inicio):
    return _parse_fecha(
        _primero_no_vacio(
            acta_inicio.get("fecha_inicio"),
            acta_inicio.get("fecha_presente_acta"),
        )
    )


def _fecha_vencimiento_inicial(acta_inicio, contrato_obra):
    fecha_inicio = _fecha_inicio(acta_inicio)
    if not fecha_inicio:
        return None

    plazo_dias = int(acta_inicio.get("plazo_ejecucion_dias", 0) or 0)
    if plazo_dias <= 0:
        plazo_dias = int(contrato_obra.get("plazo_ejecucion_dias", 0) or 0)
    if plazo_dias <= 0:
        plazo_dias = _extraer_dias_plazo(
            _primero_no_vacio(
                acta_inicio.get("plazo_ejecucion"),
                contrato_obra.get("plazo_ejecucion"),
            )
        )
    if plazo_dias <= 0:
        return None

    return fecha_inicio + timedelta(days=plazo_dias)


def _datos_generales(acta_inicio, contrato_obra, contrato_interventoria):
    return {
        "numero_contrato": _primero_no_vacio(
            acta_inicio.get("numero_contrato"),
            contrato_obra.get("numero_contrato"),
        ),
        "contratante": _primero_no_vacio(
            contrato_obra.get("nombre_entidad"),
            acta_inicio.get("contratante"),
        ),
        "contratista": _primero_no_vacio(
            acta_inicio.get("nombre_firma_contratista"),
            contrato_obra.get("nombre_contratista"),
        ),
        "interventor": _primero_no_vacio(
            acta_inicio.get("nombre_firma_interventor"),
            contrato_obra.get("nombre_interventor"),
            contrato_obra.get("nombre_supervisor"),
        ),
        "supervisor": _primero_no_vacio(
            acta_inicio.get("nombre_firma_supervisor"),
            contrato_obra.get("nombre_supervisor"),
        ),
        "objeto": _primero_no_vacio(
            acta_inicio.get("objeto_contrato"),
            contrato_obra.get("objeto_general"),
            contrato_obra.get("objeto_contrato"),
            contrato_obra.get("objeto"),
        ),
        "valor_contrato": _valor_contrato(acta_inicio, contrato_obra),
        "plazo": _primero_no_vacio(
            acta_inicio.get("plazo_ejecucion"),
            contrato_obra.get("plazo_ejecucion"),
        ),
        "fecha_inicio": _fecha_inicio(acta_inicio),
        "fecha_vencimiento_inicial": _fecha_vencimiento_inicial(acta_inicio, contrato_obra),
        "contrato_interventoria": _primero_no_vacio(
            contrato_interventoria.get("numero_proceso_contratacion"),
        ),
    }
def _suspensiones(control_obra):
    rows = control_obra.get("suspensiones_rows", []) or []
    return [fila for fila in rows if isinstance(fila, dict)]


def _tipo_acta(fila):
    if _texto(fila.get("ACTA DE AMPLIACIÓN SUSPENSIÓN No.")):
        return "AMPLIACIÓN DE SUSPENSIÓN"
    return "SUSPENSIÓN"


def _numero_acta(fila):
    if _tipo_acta(fila) == "AMPLIACIÓN DE SUSPENSIÓN":
        return _texto(fila.get("ACTA DE AMPLIACIÓN SUSPENSIÓN No."))
    return _texto(fila.get("ACTA DE SUSPENSIÓN No."))


def _etiqueta_suspension(fila, idx):
    tipo = _tipo_acta(fila)
    numero = _numero_acta(fila)
    desde = _fecha_texto(fila.get("DESDE"))
    hasta = _fecha_texto(fila.get("HASTA"))
    return f"{idx + 1}. ACTA DE {tipo} No. {numero} | Desde {desde} hasta {hasta}"


# ==========================================================
# Word helpers
# ==========================================================
def _set_doc_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(8)


def _p(doc, texto="", bold=False, align=None, size=8):
    par = doc.add_paragraph()
    run = par.add_run(str(texto))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if align:
        par.alignment = align
    return par


def _cell_text(cell, texto, bold=False, align=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    par = cell.paragraphs[0]
    run = par.add_run(str(texto))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(8)
    if align:
        par.alignment = align


def _tabla_simple(doc, filas):
    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    for etiqueta, valor in filas:
        cells = tabla.add_row().cells
        _cell_text(cells[0], etiqueta, bold=True)
        _cell_text(cells[1], valor)
    return tabla


def _campo_linea_word(doc, etiqueta, valor):
    tabla = doc.add_table(rows=1, cols=2)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = tabla.rows[0].cells
    _cell_text(cells[0], etiqueta, bold=True)
    _cell_text(cells[1], _fecha_texto(valor), align=WD_ALIGN_PARAGRAPH.CENTER)
    return tabla


def _generar_word(generales, fila, fecha_reanudacion, nueva_fecha_vencimiento):
    doc = Document()
    _set_doc_defaults(doc)

    numero = _numero_acta(fila)
    tipo = _tipo_acta(fila)

    _p(doc, "ACTA DE REANUDACIÓN CONTRATO DE OBRA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    _p(doc, f"Acta de reanudación asociada a acta de {tipo.lower()} No. {numero}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=8)

    _p(doc, "")
    _tabla_simple(
        doc,
        [
            ("CONTRATO DE OBRA No.", generales.get("numero_contrato", "")),
            ("CONTRATISTA:", generales.get("contratista", "")),
            ("OBJETO DEL CONTRATO DE OBRA:", generales.get("objeto", "")),
            ("CONTRATO DE INTERVENTORÍA No.", generales.get("contrato_interventoria", "")),
            ("INTERVENTOR:", generales.get("interventor", "")),
        ],
    )

    _p(doc, "")
    _p(
        doc,
        "Se suscribe la presente acta de reanudación del contrato de obra considerando que están dadas las condiciones para continuar con la ejecución del contrato.",
    )

    _p(doc, "")
    _campo_linea_word(doc, "FECHA DE SUSPENSIÓN INICIAL:", fila.get("DESDE"))
    _p(doc, "")
    _campo_linea_word(doc, "FECHA DE REANUDACIÓN:", fecha_reanudacion)
    _p(doc, "")
    _campo_linea_word(doc, "NUEVA FECHA DE VENCIMIENTO DEL CONTRATO:", nueva_fecha_vencimiento)

    _p(doc, "")
    _p(doc, "NOTAS:", bold=True)
    _p(
        doc,
        "1. Con la suscripción de la presente acta se reanuda la ejecución del contrato de obra asociado a la suspensión seleccionada.",
    )
    _p(
        doc,
        "2. Dentro de los tres (3) días hábiles siguientes a la fecha de reanudación, el CONTRATISTA se obliga a presentar al Instituto para su aprobación los certificados de modificación de la garantía única y el seguro constituidos, en los cuales se tenga en cuenta la nueva fecha de vencimiento del plazo contractual. El incumplimiento de esta obligación acarreará el inicio de las acciones administrativas y/o judiciales a que haya lugar.",
    )

    _p(doc, "")
    _p(doc, "Para constancia de lo anterior, firman quienes en ella intervinieron:")

    firmas = doc.add_table(rows=3, cols=3)
    firmas.style = "Table Grid"
    firmas.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cell_text(firmas.cell(0, 0), "Contratista", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(0, 1), "Supervisor", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(0, 2), "Interventoría", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(1, 0), "FIRMA", align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(1, 1), "FIRMA", align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(1, 2), "FIRMA", align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(2, 0), generales.get("contratista", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(2, 1), generales.get("supervisor", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(2, 2), generales.get("interventor", ""), align=WD_ALIGN_PARAGRAPH.CENTER)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ==========================================================
# Carga de estados
# ==========================================================
acta_inicio = _leer_estado("acta_inicio_obra")
contrato_obra = _leer_estado("contrato_obra")
contrato_interventoria = _leer_estado("contrato_interventoria")
control_obra = _leer_estado("control_obra")

generales = _datos_generales(acta_inicio, contrato_obra, contrato_interventoria)
suspensiones = _suspensiones(control_obra)

st.markdown("# ACTA DE REANUDACIÓN")

if not suspensiones:
    st.warning("No hay registros en la tabla de suspensiones de la hoja 24.")
    st.stop()

opciones = {_etiqueta_suspension(fila, idx): idx for idx, fila in enumerate(suspensiones)}
seleccion = st.selectbox(
    "Seleccione el acta de suspensión o ampliación a la que corresponde la reanudación",
    options=list(opciones.keys()),
)
fila = suspensiones[opciones[seleccion]]

clave_acta = _clave_acta_reanudacion(fila)
actas_guardadas = _cargar_actas_reanudacion_guardadas().get("actas", {})
acta_guardada = actas_guardadas.get(clave_acta, {}) if isinstance(actas_guardadas, dict) else {}

fecha_suspension_inicial = _parse_fecha(fila.get("DESDE")) or date.today()
fecha_reanudacion_base = (
    _parse_fecha(acta_guardada.get("fecha_reanudacion"))
    or _parse_fecha(fila.get("HASTA"))
    or date.today()
)
nueva_fecha_base = _parse_fecha(fila.get("NUEVA FECHA DE FINALIZACIÓN")) or generales.get("fecha_vencimiento_inicial") or date.today()

st.markdown("### DATOS GENERALES")
col1, col2 = st.columns(2)
with col1:
    st.text_input("CONTRATO DE OBRA No.", value=generales["numero_contrato"], disabled=True)
    st.text_input("CONTRATISTA", value=generales["contratista"], disabled=True)
    st.text_input("CONTRATO DE INTERVENTORÍA No.", value=generales["contrato_interventoria"], disabled=True)
with col2:
    st.text_input("INTERVENTOR", value=generales["interventor"], disabled=True)
    st.date_input("FECHA DE INICIO DEL CONTRATO", value=generales["fecha_inicio"] or date.today(), disabled=True, format="DD/MM/YYYY")
    st.date_input("FECHA DE VENCIMIENTO ACTUAL", value=nueva_fecha_base, disabled=True, format="DD/MM/YYYY")

st.text_area("OBJETO DEL CONTRATO DE OBRA", value=generales["objeto"], disabled=True, height=100)

st.markdown("### ACTA DE SUSPENSIÓN SELECCIONADA")
df_suspension = pd.DataFrame(
    [
        {
            "TIPO": _tipo_acta(fila),
            "No.": _numero_acta(fila),
            "FECHA DEL ACTA": _parse_fecha(fila.get("FECHA DEL ACTA")),
            "DESDE": _parse_fecha(fila.get("DESDE")),
            "HASTA": _parse_fecha(fila.get("HASTA")),
            "PERIODO DE SUSPENSIÓN": _texto(fila.get("PERIODO DE SUSPENSIÓN")),
            "NUEVA FECHA DE FINALIZACIÓN": _parse_fecha(fila.get("NUEVA FECHA DE FINALIZACIÓN")),
        }
    ]
)
st.dataframe(df_suspension, hide_index=True, width="stretch")

st.markdown("### CONDICIONES DE REANUDACIÓN")
col_r1, col_r2, col_r3 = st.columns(3)
with col_r1:
    st.text_input(
        "FECHA DE SUSPENSIÓN INICIAL:",
        value=_fecha_texto(fecha_suspension_inicial),
        disabled=True,
    )
with col_r2:
    fecha_reanudacion = st.date_input(
        "FECHA DE REANUDACIÓN:",
        value=fecha_reanudacion_base,
        format="DD/MM/YYYY",
        key="acta_reanudacion_fecha_reanudacion",
    )
with col_r3:
    st.text_input(
        "NUEVA FECHA DE VENCIMIENTO DEL CONTRATO:",
        value=_fecha_texto(nueva_fecha_base),
        disabled=True,
    )

st.markdown("### ")
st.write(
    "Se suscribe la presente acta de reanudación del contrato de obra considerando que están dadas las condiciones para continuar con la ejecución del contrato."
)

st.markdown("### NOTAS:")
st.write(
    "1. Con la suscripción de la presente acta se reanuda la ejecución del contrato de obra asociado a la suspensión seleccionada."
)
st.write(
    "2. Dentro de los tres (3) días hábiles siguientes a la fecha de reanudación, el CONTRATISTA se obliga a presentar al Instituto para su aprobación los certificados de modificación de la garantía única y el seguro constituidos, en los cuales se tenga en cuenta la nueva fecha de vencimiento del plazo contractual. El incumplimiento de esta obligación acarreará el inicio de las acciones administrativas y/o judiciales a que haya lugar."
)

st.markdown("### FIRMAS")
df_firmas = pd.DataFrame(
    [
        {
            "Contratista": generales.get("contratista", ""),
            "Supervisor": generales.get("supervisor", ""),
            "Interventoría": generales.get("interventor", ""),
        }
    ]
)
st.dataframe(df_firmas, hide_index=True, width="stretch")

if st.button("💾 Guardar acta de reanudación", key=f"guardar_acta_reanudacion_{clave_acta}"):
    _guardar_acta_reanudacion(
        clave_acta,
        {
            "clave_acta": clave_acta,
            "fecha_reanudacion": fecha_reanudacion,
            "fila_origen": fila,
        },
    )

word = _generar_word(generales, fila, fecha_reanudacion, nueva_fecha_base)
st.download_button(
    "📄 Descargar Word",
    data=word,
    file_name="acta_reanudacion.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)
