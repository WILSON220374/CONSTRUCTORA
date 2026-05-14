from io import BytesIO
from datetime import date, datetime, timedelta
import re

import pandas as pd
import streamlit as st

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

from supabase_state import cargar_estado
from supabase_state import guardar_estado as guardar_estado_bd


CLAVE_GUARDADO = "acta_liquidacion_obra"


# ==========================================================
# Persistencia
# ==========================================================
def guardar_estado(clave, datos):
    def serializar(obj):
        if isinstance(obj, dict):
            return {k: serializar(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [serializar(x) for x in obj]
        if isinstance(obj, (date, datetime)):
            return obj.isoformat()
        return obj

    return guardar_estado_bd(clave, serializar(datos))


def _leer_estado(clave):
    datos = cargar_estado(clave) or {}
    return datos if isinstance(datos, dict) else {}


# ==========================================================
# Helpers
# ==========================================================
def _texto(valor):
    if valor is None:
        return ""
    return str(valor).strip()


def _primero_no_vacio(*valores):
    for valor in valores:
        txt = _texto(valor)
        if txt:
            return txt
    return ""


def _safe_float(valor, default=0.0):
    try:
        if valor is None or valor == "":
            return None if default is None else float(default)

        if isinstance(valor, (int, float)):
            return float(valor)

        txt = str(valor).strip()
        txt = txt.replace("$", "").replace(" ", "")

        if "," in txt and "." in txt:
            txt = txt.replace(".", "").replace(",", ".")
        elif "," in txt:
            txt = txt.replace(",", ".")

        return float(txt)
    except Exception:
        return None if default is None else float(default)


def _parse_fecha(valor, default=None):
    if isinstance(valor, datetime):
        return valor.date()
    if isinstance(valor, date):
        return valor
    if isinstance(valor, str) and valor.strip():
        txt = valor.strip()
        try:
            return datetime.fromisoformat(txt).date()
        except Exception:
            pass
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
            try:
                return datetime.strptime(txt, fmt).date()
            except Exception:
                continue

    return default if default is not None else date.today()


def _fecha_texto(valor):
    fecha = _parse_fecha(valor, None)
    if not fecha:
        return ""
    return fecha.strftime("%d/%m/%Y")


def _moneda(valor):
    numero = _safe_float(valor, 0.0)
    return "$ " + f"{numero:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _numero(valor, decimales=2):
    numero = _safe_float(valor, 0.0)
    return f"{numero:,.{decimales}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _extraer_dias_plazo(valor):
    txt = _texto(valor).lower()
    if not txt:
        return 0

    nums = re.findall(r"\d+", txt)
    if not nums:
        return 0

    n = int(nums[0])

    if "mes" in txt:
        return n * 30
    if "semana" in txt:
        return n * 7

    return n


# ==========================================================
# Lectura de fuentes
# ==========================================================
def _leer_acta_inicio():
    return _leer_estado("acta_inicio_obra")


def _leer_contrato_obra():
    return _leer_estado("contrato_obra")


def _leer_contrato_interventoria():
    return _leer_estado("contrato_interventoria")


def _leer_control_obra():
    return _leer_estado("control_obra")


def _leer_presupuesto_obra():
    return _leer_estado("presupuesto_obra")


def _leer_plan_anticipo():
    return _leer_estado("plan_inversion_anticipo")


def _leer_acta_definitiva():
    return _leer_estado("acta_entrega_recibo_definitivo_obra")


# ==========================================================
# Datos generales
# ==========================================================
def _valor_contrato(contrato_obra, acta_inicio):
    for valor in [
        acta_inicio.get("valor_contrato"),
        acta_inicio.get("valor_inicial"),
        contrato_obra.get("valor_total_numeros"),
        contrato_obra.get("valor_contrato"),
        contrato_obra.get("valor"),
    ]:
        numero = _safe_float(valor, None)
        if numero is not None and numero > 0:
            return round(numero, 2)
    return 0.0


def _fecha_inicio(acta_inicio):
    return _parse_fecha(
        _primero_no_vacio(
            acta_inicio.get("fecha_inicio"),
            acta_inicio.get("fecha_presente_acta"),
        ),
        date.today(),
    )


def _fecha_vencimiento_inicial(acta_inicio, contrato_obra):
    fecha_directa = _primero_no_vacio(
        acta_inicio.get("fecha_terminacion"),
        acta_inicio.get("fecha_terminacion_contrato"),
        acta_inicio.get("fecha_finalizacion"),
    )

    if fecha_directa:
        return _parse_fecha(fecha_directa, date.today())

    fecha_inicio = _fecha_inicio(acta_inicio)
    plazo_dias = int(acta_inicio.get("plazo_ejecucion_dias", 0) or 0)

    if plazo_dias <= 0:
        plazo_dias = int(contrato_obra.get("plazo_ejecucion_dias", 0) or 0)

    if plazo_dias <= 0:
        plazo_dias = _extraer_dias_plazo(
            _primero_no_vacio(
                acta_inicio.get("plazo_ejecucion"),
                contrato_obra.get("plazo_ejecucion"),
            )
        )

    return fecha_inicio + timedelta(days=plazo_dias)


def _datos_generales(acta_inicio, contrato_obra, contrato_interventoria):
    return {
        "numero_contrato": _primero_no_vacio(
            acta_inicio.get("numero_contrato"),
            contrato_obra.get("numero_contrato"),
        ),
        "contratista": _primero_no_vacio(
            acta_inicio.get("nombre_firma_contratista"),
            contrato_obra.get("nombre_contratista"),
            contrato_obra.get("contratista"),
        ),
        "objeto": _primero_no_vacio(
            acta_inicio.get("objeto_contrato"),
            contrato_obra.get("objeto_general"),
            contrato_obra.get("objeto_contrato"),
            contrato_obra.get("objeto"),
        ),
        "contrato_interventoria": _primero_no_vacio(
            contrato_interventoria.get("numero_proceso_contratacion"),
            contrato_interventoria.get("numero_contrato"),
            contrato_interventoria.get("contrato_interventoria"),
        ),
        "fecha_inicio": _fecha_inicio(acta_inicio),
        "fecha_vencimiento_inicial": _fecha_vencimiento_inicial(acta_inicio, contrato_obra),
        "valor_contrato": _valor_contrato(contrato_obra, acta_inicio),
        "supervisor": _primero_no_vacio(
            acta_inicio.get("nombre_firma_supervisor"),
            contrato_obra.get("nombre_supervisor"),
            contrato_obra.get("supervisor"),
        ),
        "interventor": _primero_no_vacio(
            acta_inicio.get("nombre_firma_interventor"),
            contrato_obra.get("nombre_interventor"),
            contrato_obra.get("interventor"),
        ),
    }


# ==========================================================
# Control: adiciones / suspensiones / anticipo
# ==========================================================
def _adiciones(control_obra):
    rows = control_obra.get("adiciones_rows", []) or []
    salida = []

    for fila in rows:
        if not isinstance(fila, dict):
            continue

        valor = _safe_float(fila.get("VALOR"), 0.0)
        numero = _texto(fila.get("ADICIONAL No."))

        if not numero and valor <= 0:
            continue

        salida.append(
            {
                "ADICIONAL No.": numero,
                "FECHA": _parse_fecha(fila.get("FECHA"), None),
                "VALOR": round(valor, 2),
                "VALOR ACUMULADO DEL CONTRATO": round(
                    _safe_float(fila.get("VALOR ACUMULADO DEL CONTRATO"), 0.0),
                    2,
                ),
            }
        )

    return salida


def _valor_adiciones(adiciones):
    return round(sum(_safe_float(fila.get("VALOR"), 0.0) for fila in adiciones), 2)


def _suspensiones(control_obra):
    rows = control_obra.get("suspensiones_rows", []) or []
    salida = []

    for fila in rows:
        if not isinstance(fila, dict):
            continue

        numero = _primero_no_vacio(
            fila.get("ACTA DE SUSPENSIÓN No."),
            fila.get("ACTA DE AMPLIACIÓN SUSPENSIÓN No."),
            fila.get("ACTA"),
        )

        desde = _parse_fecha(fila.get("DESDE"), None)
        hasta = _parse_fecha(fila.get("HASTA"), None)

        if not numero and not desde and not hasta:
            continue

        salida.append(
            {
                "ACTA": numero,
                "TIPO": "AMPLIACIÓN" if _texto(fila.get("ACTA DE AMPLIACIÓN SUSPENSIÓN No.")) else "SUSPENSIÓN",
                "FECHA DEL ACTA": _parse_fecha(fila.get("FECHA DEL ACTA"), None),
                "DESDE": desde,
                "HASTA": hasta,
                "PERIODO DE SUSPENSIÓN": int(_safe_float(fila.get("PERIODO DE SUSPENSIÓN"), 0.0)),
                "NUEVA FECHA DE FINALIZACIÓN": _parse_fecha(fila.get("NUEVA FECHA DE FINALIZACIÓN"), None),
            }
        )

    return salida


def _valor_anticipo(plan_anticipo, control_obra, valor_contrato):
    for valor in [
        control_obra.get("valor_anticipo"),
        control_obra.get("anticipo_valor_total_concedido"),
        plan_anticipo.get("valor_anticipo"),
    ]:
        numero = _safe_float(valor, None)
        if numero is not None and numero > 0:
            return round(numero, 2)

    porcentaje = _safe_float(
        _primero_no_vacio(
            plan_anticipo.get("porcentaje_anticipo"),
            control_obra.get("porcentaje_anticipo"),
        ),
        0.0,
    )

    if porcentaje > 0 and valor_contrato > 0:
        return round(valor_contrato * porcentaje / 100.0, 2)

    return 0.0


def _anticipo(control_obra, plan_anticipo, valor_contrato):
    valor_total = _valor_anticipo(plan_anticipo, control_obra, valor_contrato)

    rows = (
        control_obra.get("anticipo_rows", [])
        or control_obra.get("control_anticipo_rows", [])
        or []
    )

    amortizado = 0.0
    for fila in rows:
        if not isinstance(fila, dict):
            continue

        amortizado += _safe_float(
            _primero_no_vacio(
                fila.get("VALOR AMORTIZADO"),
                fila.get("AMORTIZADO"),
                fila.get("VALOR A AMORTIZAR"),
            ),
            0.0,
        )

    saldo = max(0.0, valor_total - amortizado)

    return {
        "valor_total": round(valor_total, 2),
        "valor_amortizado": round(amortizado, 2),
        "saldo_por_amortizar": round(saldo, 2),
    }


# ==========================================================
# AIU
# ==========================================================
def _aiu_desde_presupuesto(presupuesto_obra):
    config = presupuesto_obra.get("configuracion") or {}

    administracion = _safe_float(
        _primero_no_vacio(
            config.get("administracion"),
            config.get("administracion_pct"),
            config.get("a_pct"),
            presupuesto_obra.get("administracion"),
        ),
        0.0,
    )

    imprevistos = _safe_float(
        _primero_no_vacio(
            config.get("imprevistos"),
            config.get("imprevistos_pct"),
            config.get("i_pct"),
            presupuesto_obra.get("imprevistos"),
        ),
        0.0,
    )

    utilidad = _safe_float(
        _primero_no_vacio(
            config.get("utilidad"),
            config.get("utilidad_pct"),
            config.get("u_pct"),
            presupuesto_obra.get("utilidad"),
        ),
        0.0,
    )

    aiu_total = _safe_float(
        _primero_no_vacio(
            config.get("aiu_pct_global"),
            config.get("aiu_total"),
            presupuesto_obra.get("aiu_pct_global"),
            presupuesto_obra.get("aiu_total"),
        ),
        0.0,
    )

    if aiu_total <= 0:
        aiu_total = administracion + imprevistos + utilidad

    return {
        "aiu_total": round(aiu_total, 2),
        "administracion": round(administracion, 2),
        "imprevistos": round(imprevistos, 2),
        "utilidad": round(utilidad, 2),
    }


# ==========================================================
# Balance general
# ==========================================================
def _filas_balance_inicial(guardado, valor_total_ejecutado=0.0):
    rows = guardado.get("balance_rows", []) if isinstance(guardado, dict) else []

    valor_pagado = 0.0

    if isinstance(rows, list):
        for fila in rows:
            if not isinstance(fila, dict):
                continue

            descripcion = _texto(fila.get("DESCRIPCIÓN")).upper()

            if "PAGADO" in descripcion:
                valor_pagado = _safe_float(fila.get("PAGADO"), 0.0)
                break

            posible_pagado = _safe_float(fila.get("PAGADO"), 0.0)
            if posible_pagado > 0:
                valor_pagado = posible_pagado

    return [
        {
            "DESCRIPCIÓN": "Valor total ejecutado",
            "EJECUTADO": round(_safe_float(valor_total_ejecutado, 0.0), 2),
            "PAGADO": 0.0,
        },
        {
            "DESCRIPCIÓN": "Valor pagado en actas parciales",
            "EJECUTADO": 0.0,
            "PAGADO": round(valor_pagado, 2),
        },
    ]


def _normalizar_balance(rows):
    salida = []

    for fila in rows or []:
        if not isinstance(fila, dict):
            continue

        descripcion = _texto(fila.get("DESCRIPCIÓN"))
        if not descripcion:
            continue

        salida.append(
            {
                "DESCRIPCIÓN": descripcion,
                "EJECUTADO": round(_safe_float(fila.get("EJECUTADO"), 0.0), 2),
                "PAGADO": round(_safe_float(fila.get("PAGADO"), 0.0), 2),
            }
        )

    return salida


def _totales_balance(rows):
    total_ejecutado = 0.0
    total_pagado = 0.0

    for fila in rows or []:
        total_ejecutado += _safe_float(fila.get("EJECUTADO"), 0.0)
        total_pagado += _safe_float(fila.get("PAGADO"), 0.0)

    saldo = total_ejecutado - total_pagado

    return round(total_ejecutado, 2), round(total_pagado, 2), round(saldo, 2)


# ==========================================================
# Documentos de liquidación
# ==========================================================
def _documentos_base(guardado):
    rows = guardado.get("documentos_rows", [])
    if isinstance(rows, list) and rows:
        return rows

    descripciones = [
        "Acta de entrega y recibo definitivo de obra.",
        "Informe final de interventoría.",
        "Balance financiero del contrato.",
        "Relación de actas de recibo parcial de obra.",
        "Certificación de pagos y amortización del anticipo.",
        "Certificación de cumplimiento de obligaciones laborales y de seguridad social.",
        "Certificación del estado de garantías y seguros contractuales.",
        "Planos record, memorias, ensayos, registros fotográficos y demás soportes técnicos aplicables.",
    ]

    return [
        {
            "No.": i,
            "DESCRIPCIÓN": descripcion,
            "FOLIOS": 0,
        }
        for i, descripcion in enumerate(descripciones, start=1)
    ]


def _normalizar_documentos(rows):
    salida = []

    for i, fila in enumerate(rows or [], start=1):
        if not isinstance(fila, dict):
            continue

        descripcion = _texto(fila.get("DESCRIPCIÓN"))
        if not descripcion:
            continue

        salida.append(
            {
                "No.": int(_safe_float(fila.get("No."), i)),
                "DESCRIPCIÓN": descripcion,
                "FOLIOS": int(_safe_float(fila.get("FOLIOS"), 0)),
            }
        )

    return salida


# ==========================================================
# Word
# ==========================================================
def _doc_titulo(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _doc_parrafo(doc, texto="", bold=False, size=10, align=None):
    p = doc.add_paragraph()
    r = p.add_run(str(texto))
    r.bold = bold
    r.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    return p


def _doc_tabla_kv(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for etiqueta, valor in rows:
        cells = table.add_row().cells
        cells[0].text = str(etiqueta)
        cells[1].text = str(valor)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    return table


def _doc_tabla_df(doc, df, columnas=None):
    if columnas is None:
        columnas = list(df.columns)

    table = doc.add_table(rows=1, cols=len(columnas))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    encabezado = table.rows[0].cells
    for i, col in enumerate(columnas):
        encabezado[i].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columnas):
            valor = row.get(col, "")

            if isinstance(valor, (date, datetime)):
                valor = _fecha_texto(valor)
            elif isinstance(valor, float):
                valor = _numero(valor, 2)

            cells[i].text = str(valor)

    return table


def _generar_word(payload):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    generales = payload.get("generales", {})
    anticipo = payload.get("anticipo", {})
    aiu = payload.get("aiu", {})

    _doc_titulo(doc, "ACTA DE LIQUIDACIÓN CONTRATO DE OBRA")

    _doc_tabla_kv(
        doc,
        [
            ("FECHA", _fecha_texto(payload.get("fecha"))),
            ("CONTRATO DE OBRA No.", generales.get("numero_contrato", "")),
            ("CONTRATISTA", generales.get("contratista", "")),
            ("OBJETO DEL CONTRATO", generales.get("objeto", "")),
            ("CONTRATO DE INTERVENTORÍA No.", generales.get("contrato_interventoria", "")),
            ("FECHA DE INICIO CONTRATO DE OBRA", _fecha_texto(generales.get("fecha_inicio"))),
            ("VALOR CONTRATO", _moneda(payload.get("valor_contrato", 0.0))),
            ("VALOR ADICIONES", _moneda(payload.get("valor_adiciones", 0.0))),
            ("VALOR ACUMULADO", _moneda(payload.get("valor_acumulado", 0.0))),
            ("FECHA DE VENCIMIENTO", _fecha_texto(payload.get("fecha_vencimiento"))),
        ],
    )

    doc.add_paragraph()
    _doc_parrafo(doc, "RELACIÓN SUSPENSIONES Y AMPLIACIONES DE SUSPENSIÓN", bold=True)
    _doc_tabla_df(doc, pd.DataFrame(payload.get("suspensiones", [])))

    doc.add_paragraph()
    _doc_parrafo(doc, "ANTICIPO", bold=True)
    _doc_tabla_kv(
        doc,
        [
            ("Valor total anticipo", _moneda(anticipo.get("valor_total", 0.0))),
            ("Valor total amortizado", _moneda(anticipo.get("valor_amortizado", 0.0))),
            ("Saldo por amortizar", _moneda(anticipo.get("saldo_por_amortizar", 0.0))),
        ],
    )

    doc.add_paragraph()
    _doc_parrafo(doc, "VALOR TOTAL EJECUTADO DEL CONTRATO DE OBRA", bold=True)
    _doc_tabla_kv(
        doc,
        [
            ("Valor total ejecutado del contrato", _moneda(payload.get("valor_total_ejecutado", 0.0))),
            ("AIU", f"{_numero(aiu.get('aiu_total', 0.0), 2)}%"),
            ("Administración", f"{_numero(aiu.get('administracion', 0.0), 2)}%"),
            ("Imprevistos", f"{_numero(aiu.get('imprevistos', 0.0), 2)}%"),
            ("Utilidad", f"{_numero(aiu.get('utilidad', 0.0), 2)}%"),
        ],
    )

    doc.add_paragraph()
    _doc_parrafo(doc, "BALANCE GENERAL DEL CONTRATO", bold=True)
    _doc_tabla_df(doc, pd.DataFrame(payload.get("balance_rows", [])))
    _doc_tabla_kv(
        doc,
        [
            ("TOTAL EJECUTADO", _moneda(payload.get("total_balance_ejecutado", 0.0))),
            ("TOTAL PAGADO", _moneda(payload.get("total_balance_pagado", 0.0))),
            ("SALDO", _moneda(payload.get("saldo_balance", 0.0))),
        ],
    )

    doc.add_paragraph()
    _doc_parrafo(doc, "DOCUMENTOS APORTADOS PARA LA LIQUIDACIÓN", bold=True)
    _doc_tabla_df(doc, pd.DataFrame(payload.get("documentos_rows", [])))
    _doc_parrafo(doc, f"TOTAL FOLIOS: {payload.get('total_folios', 0)}", bold=True)

    doc.add_paragraph()
    _doc_parrafo(doc, "OBSERVACIONES", bold=True)
    _doc_parrafo(doc, payload.get("observaciones", ""))

    doc.add_paragraph()
    _doc_parrafo(
        doc,
        "Las partes declaran liquidado el contrato de obra en los términos aquí consignados, de conformidad con la información contractual, técnica, administrativa y financiera disponible.",
    )

    doc.add_paragraph()
    _doc_parrafo(doc, "FIRMAS", bold=True)
    _doc_tabla_kv(
        doc,
        [
            ("Contratista", payload.get("firma_contratista", "")),
            ("Supervisor", payload.get("firma_supervisor", "")),
        ],
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==========================================================
# Estado inicial
# ==========================================================
def _estado_inicial(generales, valor_adiciones, valor_acumulado, anticipo):
    return {
        "fecha": date.today(),
        "fecha_vencimiento": generales.get("fecha_vencimiento_inicial", date.today()),
        "valor_total_ejecutado": 0.0,
        "balance_rows": _filas_balance_inicial({}),
        "documentos_rows": _documentos_base({}),
        "observaciones": "",
        "firma_contratista": generales.get("contratista", ""),
        "firma_supervisor": generales.get("supervisor", ""),
    }


# ==========================================================
# Interfaz
# ==========================================================
st.set_page_config(page_title="Acta liquidación obra", layout="wide")
st.title("📄 ACTA DE LIQUIDACIÓN CONTRATO DE OBRA")

acta_inicio = _leer_acta_inicio()
contrato_obra = _leer_contrato_obra()
contrato_interventoria = _leer_contrato_interventoria()
control_obra = _leer_control_obra()
presupuesto_obra = _leer_presupuesto_obra()
plan_anticipo = _leer_plan_anticipo()

generales = _datos_generales(acta_inicio, contrato_obra, contrato_interventoria)
adiciones = _adiciones(control_obra)
suspensiones = _suspensiones(control_obra)
aiu = _aiu_desde_presupuesto(presupuesto_obra)

valor_contrato = generales.get("valor_contrato", 0.0)
valor_adiciones = _valor_adiciones(adiciones)
valor_acumulado_fuente = round(valor_contrato + valor_adiciones, 2)
anticipo = _anticipo(control_obra, plan_anticipo, valor_contrato)

guardado = _leer_estado(CLAVE_GUARDADO)
if not guardado:
    guardado = _estado_inicial(generales, valor_adiciones, valor_acumulado_fuente, anticipo)

st.markdown("### INFORMACIÓN GENERAL")

c1, c2, c3 = st.columns(3)
with c1:
    fecha_acta = st.date_input(
        "FECHA",
        value=_parse_fecha(guardado.get("fecha"), date.today()),
        format="DD/MM/YYYY",
    )
with c2:
    st.text_input("CONTRATO DE OBRA No.", value=generales.get("numero_contrato", ""), disabled=True)
with c3:
    st.text_input("CONTRATO DE INTERVENTORÍA No.", value=generales.get("contrato_interventoria", ""), disabled=True)

c4, c5 = st.columns(2)
with c4:
    st.text_input("CONTRATISTA", value=generales.get("contratista", ""), disabled=True)
with c5:
    st.date_input(
        "FECHA DE INICIO CONTRATO DE OBRA",
        value=generales.get("fecha_inicio", date.today()),
        format="DD/MM/YYYY",
        disabled=True,
    )

st.text_area("OBJETO DEL CONTRATO", value=generales.get("objeto", ""), disabled=True, height=90)

st.markdown("### CONTRATO Y ADICIONES")

c6, c7, c8 = st.columns(3)
with c6:
    st.number_input("VALOR CONTRATO", value=float(valor_contrato), format="%.2f", disabled=True)
with c7:
    st.number_input("VALOR ADICIONES", value=float(valor_adiciones), format="%.2f", disabled=True)
with c8:
    valor_acumulado = st.number_input(
        "VALOR ACUMULADO",
        value=float(valor_acumulado_fuente),
        format="%.2f",
        disabled=True,
    )

df_adiciones = pd.DataFrame(adiciones)
if df_adiciones.empty:
    df_adiciones = pd.DataFrame(columns=["ADICIONAL No.", "FECHA", "VALOR", "VALOR ACUMULADO DEL CONTRATO"])

st.dataframe(
    df_adiciones,
    hide_index=True,
    width="stretch",
    column_config={
        "FECHA": st.column_config.DateColumn("FECHA", format="DD/MM/YYYY"),
        "VALOR": st.column_config.NumberColumn("VALOR", format="$ %.2f"),
        "VALOR ACUMULADO DEL CONTRATO": st.column_config.NumberColumn("VALOR ACUMULADO DEL CONTRATO", format="$ %.2f"),
    },
)

st.markdown("### RELACIÓN SUSPENSIONES Y AMPLIACIONES DE SUSPENSIÓN")

df_suspensiones = pd.DataFrame(suspensiones)
if df_suspensiones.empty:
    df_suspensiones = pd.DataFrame(
        columns=[
            "ACTA",
            "TIPO",
            "FECHA DEL ACTA",
            "DESDE",
            "HASTA",
            "PERIODO DE SUSPENSIÓN",
            "NUEVA FECHA DE FINALIZACIÓN",
        ]
    )

st.dataframe(
    df_suspensiones,
    hide_index=True,
    width="stretch",
    column_config={
        "FECHA DEL ACTA": st.column_config.DateColumn("FECHA DEL ACTA", format="DD/MM/YYYY"),
        "DESDE": st.column_config.DateColumn("DESDE", format="DD/MM/YYYY"),
        "HASTA": st.column_config.DateColumn("HASTA", format="DD/MM/YYYY"),
        "NUEVA FECHA DE FINALIZACIÓN": st.column_config.DateColumn("NUEVA FECHA DE FINALIZACIÓN", format="DD/MM/YYYY"),
    },
)

fecha_vencimiento = st.date_input(
    "FECHA DE VENCIMIENTO",
    value=_parse_fecha(guardado.get("fecha_vencimiento"), generales.get("fecha_vencimiento_inicial", date.today())),
    format="DD/MM/YYYY",
)

st.markdown("### ANTICIPO")

c_ant1, c_ant2, c_ant3 = st.columns(3)
with c_ant1:
    st.number_input("VALOR TOTAL ANTICIPO", value=float(anticipo.get("valor_total", 0.0)), format="%.2f", disabled=True)
with c_ant2:
    st.number_input("VALOR TOTAL AMORTIZADO", value=float(anticipo.get("valor_amortizado", 0.0)), format="%.2f", disabled=True)
with c_ant3:
    st.number_input("SALDO POR AMORTIZAR", value=float(anticipo.get("saldo_por_amortizar", 0.0)), format="%.2f", disabled=True)

st.markdown("### VALOR TOTAL EJECUTADO DEL CONTRATO DE OBRA")

valor_total_ejecutado = st.number_input(
    "VALOR TOTAL EJECUTADO DEL CONTRATO",
    value=float(_safe_float(guardado.get("valor_total_ejecutado"), 0.0)),
    min_value=0.0,
    step=1000.0,
    format="%.2f",
)

st.number_input(
    "AIU %",
    value=float(aiu.get("aiu_total", 0.0)),
    format="%.2f",
    disabled=True,
)

st.markdown("### BALANCE GENERAL DEL CONTRATO")

df_balance = pd.DataFrame(
    _filas_balance_inicial(guardado, valor_total_ejecutado),
    columns=["DESCRIPCIÓN", "EJECUTADO", "PAGADO"],
)

df_balance = df_balance[["DESCRIPCIÓN", "EJECUTADO"]]

balance_editado = st.data_editor(
    df_balance,
    hide_index=True,
    width="stretch",
    num_rows="fixed",
    key="balance_liquidacion_obra_editor",
    disabled=["DESCRIPCIÓN", "EJECUTADO"],
    column_config={
        "DESCRIPCIÓN": st.column_config.TextColumn("DESCRIPCIÓN"),
        "EJECUTADO": st.column_config.NumberColumn("EJECUTADO", format="$ %.2f"),
    },
)

balance_rows = _normalizar_balance(balance_editado.to_dict("records"))
balance_rows = _filas_balance_inicial({"balance_rows": balance_rows}, valor_total_ejecutado)

for fila in balance_rows:
    fila["PAGADO"] = 0.0

total_balance_ejecutado, total_balance_pagado, saldo_balance = _totales_balance(balance_rows)

c_bal1, c_bal2, c_bal3 = st.columns(3)
with c_bal1:
    st.metric("TOTAL EJECUTADO", _moneda(total_balance_ejecutado))
with c_bal2:
    st.metric("TOTAL PAGADO", _moneda(total_balance_pagado))
with c_bal3:
    st.metric("SALDO", _moneda(saldo_balance))

st.markdown("### DOCUMENTOS APORTADOS PARA LA LIQUIDACIÓN")

df_documentos = pd.DataFrame(
    _documentos_base(guardado),
    columns=["No.", "DESCRIPCIÓN", "FOLIOS"],
)

documentos_editados = st.data_editor(
    df_documentos,
    hide_index=True,
    width="stretch",
    num_rows="dynamic",
    key="documentos_liquidacion_obra_editor",
    column_config={
        "No.": st.column_config.NumberColumn("No.", disabled=True),
        "DESCRIPCIÓN": st.column_config.TextColumn("DESCRIPCIÓN"),
        "FOLIOS": st.column_config.NumberColumn("FOLIOS"),
    },
)

documentos_rows = _normalizar_documentos(documentos_editados.to_dict("records"))
total_folios = sum(int(_safe_float(fila.get("FOLIOS"), 0)) for fila in documentos_rows)
st.metric("TOTAL FOLIOS", f"{total_folios}")

st.markdown("### OBSERVACIONES")

observaciones = st.text_area(
    "OBSERVACIONES",
    value=guardado.get("observaciones", ""),
    height=120,
)

st.markdown("### FIRMAS")

c_f1, c_f2 = st.columns(2)
with c_f1:
    firma_contratista = st.text_input(
        "Contratista",
        value=guardado.get("firma_contratista", generales.get("contratista", "")),
    )
with c_f2:
    firma_supervisor = st.text_input(
        "Supervisor",
        value=guardado.get("firma_supervisor", generales.get("supervisor", "")),
    )

payload = {
    "fecha": fecha_acta,
    "generales": generales,
    "valor_contrato": round(valor_contrato, 2),
    "valor_adiciones": round(valor_adiciones, 2),
    "valor_acumulado": round(valor_acumulado, 2),
    "suspensiones": suspensiones,
    "fecha_vencimiento": fecha_vencimiento,
    "anticipo": anticipo,
    "valor_total_ejecutado": round(valor_total_ejecutado, 2),
    "aiu": aiu,
    "balance_rows": balance_rows,
    "total_balance_ejecutado": total_balance_ejecutado,
    "total_balance_pagado": total_balance_pagado,
    "saldo_balance": saldo_balance,
    "documentos_rows": documentos_rows,
    "total_folios": total_folios,
    "observaciones": observaciones,
    "firma_contratista": firma_contratista,
    "firma_supervisor": firma_supervisor,
}

c_btn1, c_btn2 = st.columns([1, 1])

with c_btn1:
    if st.button("💾 Guardar acta de liquidación", key="guardar_acta_liquidacion_obra"):
        guardar_estado(CLAVE_GUARDADO, payload)
        st.success("Acta de liquidación de obra guardada correctamente.")

with c_btn2:
    buffer = _generar_word(payload)
    st.download_button(
        "📥 Descargar Word",
        data=buffer,
        file_name="acta_liquidacion_obra.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="descargar_word_acta_liquidacion_obra",
    )
