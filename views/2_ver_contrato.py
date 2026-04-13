import os
from io import BytesIO

import streamlit as st
from docx import Document

from supabase_state import cargar_estado


def obtener_datos_contrato():
    datos = st.session_state.get("contrato_obra_datos")
    if not datos:
        datos = cargar_estado("contrato_obra") or {}
        st.session_state["contrato_obra_datos"] = datos
    return datos


def texto_si_vacio(valor, pendiente="PENDIENTE"):
    if valor is None:
        return pendiente
    if isinstance(valor, str):
        return valor.strip() if valor.strip() else pendiente
    return str(valor)


def escapar_tabla(valor):
    return texto_si_vacio(valor).replace("|", "\\|")


def construir_tabla_garantias_markdown(garantias):
    if not garantias or not isinstance(garantias, list):
        return (
            "| Amparo | Suficiencia | Vigencia |\n"
            "|---|---|---|\n"
            "| PENDIENTE | PENDIENTE | PENDIENTE |"
        )

    filas = []
    for fila in garantias:
        if not isinstance(fila, dict):
            continue
        amparo = escapar_tabla(fila.get("amparo", ""))
        suficiencia = escapar_tabla(fila.get("suficiencia", ""))
        vigencia = escapar_tabla(fila.get("vigencia", ""))
        if amparo == "PENDIENTE" and suficiencia == "PENDIENTE" and vigencia == "PENDIENTE":
            continue
        filas.append(f"| {amparo} | {suficiencia} | {vigencia} |")

    if not filas:
        filas.append("| PENDIENTE | PENDIENTE | PENDIENTE |")

    encabezado = "| Amparo | Suficiencia | Vigencia |\n|---|---|---|"
    return encabezado + "\n" + "\n".join(filas)


def construir_tabla_garantias_doc(doc, garantias):
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = "Table Grid"
    encabezado = tabla.rows[0].cells
    encabezado[0].text = "Amparo"
    encabezado[1].text = "Suficiencia"
    encabezado[2].text = "Vigencia"

    filas_validas = []
    if garantias and isinstance(garantias, list):
        for fila in garantias:
            if not isinstance(fila, dict):
                continue
            amparo = texto_si_vacio(fila.get("amparo", ""))
            suficiencia = texto_si_vacio(fila.get("suficiencia", ""))
            vigencia = texto_si_vacio(fila.get("vigencia", ""))
            if amparo == "PENDIENTE" and suficiencia == "PENDIENTE" and vigencia == "PENDIENTE":
                continue
            filas_validas.append((amparo, suficiencia, vigencia))

    if not filas_validas:
        filas_validas.append(("PENDIENTE", "PENDIENTE", "PENDIENTE"))

    for amparo, suficiencia, vigencia in filas_validas:
        celdas = tabla.add_row().cells
        celdas[0].text = amparo
        celdas[1].text = suficiencia
        celdas[2].text = vigencia


def construir_bloque_anexos(datos):
    anexos = []

    if datos.get("anexos_estudios_previos"):
        anexos.append("25.1. Los estudios previos.")
    if datos.get("anexos_pliego"):
        anexos.append("25.2. El Pliego de Condiciones del proceso de selección, sus anexos, adendas o cualquier otro Documento del Proceso.")
    if datos.get("anexos_oferta"):
        anexos.append("25.3. La oferta presentada por el Contratista.")
    if datos.get("anexos_actas_informes"):
        anexos.append("25.4. Las actas, acuerdos, informes y documentos precontractuales.")
    if datos.get("anexos_cdp"):
        anexos.append("25.5. Certificado de Disponibilidad Presupuestal.")

    if not anexos:
        anexos.append("25.1. PENDIENTE.")

    return "\n".join(anexos)


def construir_contrato(datos):
    numero_contrato = texto_si_vacio(datos.get("numero_contrato"))
    nombre_proyecto = texto_si_vacio(datos.get("nombre_proyecto"))
    fecha_contrato = texto_si_vacio(datos.get("fecha_contrato"))
    lugar_celebracion = texto_si_vacio(datos.get("lugar_celebracion"))

    nombre_entidad = texto_si_vacio(datos.get("nombre_entidad"))
    nit_entidad = texto_si_vacio(datos.get("nit_entidad"))
    mision_entidad = texto_si_vacio(datos.get("mision_entidad"))
    justificacion_general = texto_si_vacio(datos.get("justificacion_general"))
    necesidad_contratar = texto_si_vacio(datos.get("necesidad_contratar"))

    rep_entidad_nombre = texto_si_vacio(datos.get("rep_entidad_nombre"))
    rep_entidad_tipo_doc = texto_si_vacio(datos.get("rep_entidad_tipo_doc"))
    rep_entidad_num_doc = texto_si_vacio(datos.get("rep_entidad_num_doc"))
    rep_entidad_municipio_expedicion = texto_si_vacio(datos.get("rep_entidad_municipio_expedicion"))
    rep_entidad_cargo = texto_si_vacio(datos.get("rep_entidad_cargo"))

    tipo_contratista = texto_si_vacio(datos.get("tipo_contratista"))
    nombre_contratista = texto_si_vacio(datos.get("nombre_contratista"))
    nit_contratista = texto_si_vacio(datos.get("nit_contratista"))

    rep_contratista_nombre = texto_si_vacio(datos.get("rep_contratista_nombre"))
    rep_contratista_tipo_doc = texto_si_vacio(datos.get("rep_contratista_tipo_doc"))
    rep_contratista_num_doc = texto_si_vacio(datos.get("rep_contratista_num_doc"))
    rep_contratista_ciudad_expedicion = texto_si_vacio(datos.get("rep_contratista_ciudad_expedicion"))

    modalidad_seleccion = texto_si_vacio(datos.get("modalidad_seleccion"))
    objeto_general = texto_si_vacio(datos.get("objeto_general"))
    objeto_especifico = texto_si_vacio(datos.get("objeto_especifico"))

    valor_total_numeros = texto_si_vacio(datos.get("valor_total_numeros"))
    valor_total_letras = texto_si_vacio(datos.get("valor_total_letras"))
    periodicidad_pago = texto_si_vacio(datos.get("periodicidad_pago"))
    dias_pago = texto_si_vacio(datos.get("dias_pago"))

    plazo_ejecucion = texto_si_vacio(datos.get("plazo_ejecucion"))

    clausula_penal_numeros = texto_si_vacio(datos.get("clausula_penal_numeros"))
    clausula_penal_letras = texto_si_vacio(datos.get("clausula_penal_letras"))

    tabla_garantias = construir_tabla_garantias_markdown(datos.get("garantias", []))
    plazo_garantias_dias = texto_si_vacio(datos.get("plazo_garantias_dias"))

    not_entidad_direccion = texto_si_vacio(datos.get("not_entidad_direccion"))
    not_entidad_telefono = texto_si_vacio(datos.get("not_entidad_telefono"))
    not_entidad_correo = texto_si_vacio(datos.get("not_entidad_correo"))

    not_contratista_direccion = texto_si_vacio(datos.get("not_contratista_direccion"))
    not_contratista_telefono = texto_si_vacio(datos.get("not_contratista_telefono"))
    not_contratista_correo = texto_si_vacio(datos.get("not_contratista_correo"))

    tipo_seguimiento = texto_si_vacio(datos.get("tipo_seguimiento"))
    nombre_supervisor = texto_si_vacio(datos.get("nombre_supervisor"))
    nombre_interventor = texto_si_vacio(datos.get("nombre_interventor"))

    lugar_ejecucion = texto_si_vacio(datos.get("lugar_ejecucion"))

    anexos_txt = construir_bloque_anexos(datos)

    bloque_supervision = f"La supervisión de la ejecución y cumplimiento de las obligaciones contraídas por el Contratista a favor de la Entidad Estatal Contratante, estará a cargo de {nombre_supervisor}."
    bloque_interventoria = f"La interventoría de la ejecución y cumplimiento de las obligaciones contraídas por el Contratista a favor de la Entidad Estatal Contratante, estará a cargo de {nombre_interventor}."

    if tipo_seguimiento == "Solo supervisión":
        bloque_interventoria = "La interventoría no aplica para el presente contrato."
    elif tipo_seguimiento == "Solo interventoría":
        bloque_supervision = "La supervisión no aplica para el presente contrato."

    contrato = f"""# Contrato de obra pública {numero_contrato} para la ejecución del proyecto {nombre_proyecto} del {fecha_contrato}, celebrado entre {nombre_entidad} y {nombre_contratista}.

Entre los suscritos: {rep_entidad_nombre}, identificado con {rep_entidad_tipo_doc} No. {rep_entidad_num_doc}, expedida en {rep_entidad_municipio_expedicion}, en su calidad de {rep_entidad_cargo}, actuando en nombre y representación de {nombre_entidad}, con NIT {nit_entidad}, quien para los efectos del presente contrato se denomina la Entidad Estatal contratante; y por la otra, {rep_contratista_nombre}, identificado con {rep_contratista_tipo_doc} No. {rep_contratista_num_doc}, expedida en {rep_contratista_ciudad_expedicion}, en representación de {nombre_contratista}, identificado con NIT {nit_contratista}, quien para los efectos del presente contrato se denominará el Contratista, hemos convenido en celebrar el presente Contrato de obra pública, teniendo en cuenta las siguientes consideraciones:

I. Que la misión de {nombre_entidad} es {mision_entidad} y el contrato a celebrarse se relaciona con esta misión porque {justificacion_general}.

II. Que la necesidad a satisfacer por parte de la Entidad Estatal contratante es {necesidad_contratar}.

III. Que la modalidad de selección corresponde a {modalidad_seleccion}.

Por lo anterior, las partes celebran el presente contrato, el cual se regirá por las siguientes cláusulas:

## Cláusula 1 – Objeto del contrato

El objeto del contrato es {objeto_general} que incluye {objeto_especifico}.

Los Documentos del Proceso forman parte del presente Contrato y definen igualmente las actividades, alcance y obligaciones del Contrato.

## Cláusula 2 – Valor del Contrato y Forma de pago

El valor del presente Contrato corresponde a la suma de {valor_total_numeros} ({valor_total_letras}).

La Entidad Estatal Contratante pagará al Contratista el valor del contrato en los siguientes periodos: {periodicidad_pago}.

Los pagos se realizarán dentro de los {dias_pago} siguientes a la fecha de presentación del certificado de cumplimiento firmado por el supervisor del Contrato.

## Cláusula 3 – Declaraciones del contratista

El Contratista hace las siguientes declaraciones:

3.1. Conoce y acepta los Documentos del Proceso.  
3.2. Tuvo la oportunidad de solicitar aclaraciones y modificaciones a los Documentos del Proceso y recibió de {nombre_entidad} respuesta oportuna a cada una de las solicitudes.  
3.3. Se encuentra debidamente facultado para suscribir el presente Contrato.  
3.4. El Contratista al momento de la celebración del presente Contrato no se encuentra en ninguna causal de inhabilidad, incompatibilidad.  
3.5. El Contratista está a paz y salvo con sus obligaciones laborales frente al sistema de seguridad social integral.  
3.6. El valor del Contrato incluye todos los gastos, costos, derechos, impuestos, tasas y demás contribuciones relacionados con el cumplimiento del objeto del presente contrato.  
3.7. El Contratista durante la ejecución del presente Contrato realizará todas las actividades necesarias para la ejecución final de la obra, cumpliendo con el plazo establecido en la cláusula 4 del presente Contrato.  
3.8. El Contratista manifiesta que los recursos que componen su patrimonio no provienen de lavado de activos, financiación del terrorismo, narcotráfico, captación ilegal de dineros y en general de cualquier actividad ilícita; de igual manera manifiesta que los recursos recibidos en desarrollo de este contrato, no serán destinados a ninguna de las actividades antes descritas.  
3.9. El Contratista se compromete a no contratar menores de edad para el ejercicio del objeto contractual, así como a no permitir que se subcontrate a menores de edad para tales efectos, dando aplicación a la normativa vigente.

## Cláusula 4 – Plazo y Cronograma de Obra

El plazo de ejecución del Contrato es {plazo_ejecucion}.

El Cronograma Estimado de Obra del presente Contrato resulta del análisis conjunto del Contratista y de la Entidad Estatal contratante y forma parte del presente Contrato como anexo cuando aplique.

La fecha de inicio del plazo de ejecución de la obra es la fecha en la cual se suscriba entre las partes el Acta de Inicio de obra.

La fecha de terminación del plazo de ejecución de la obra es la fecha en la cual se suscriba el Acta de Recibo Final. Para que se pueda suscribir el Acta de Recibo Final, el Contratista debe cumplir a cabalidad con los compromisos y obligaciones contenidos en el presente Contrato y sus anexos.

## Cláusula 5 – Derechos del Contratista

5.1. Recibir una remuneración por la ejecución de la obra en los términos pactados en la cláusula 2 del presente Contrato.  
5.2. Los demás derechos que resulten aplicables conforme al documento tipo y a la normatividad vigente.

## Cláusula 6 – Obligaciones particulares del Contratista

6.1. Desarrollar y cumplir el objeto del Contrato, en las condiciones de calidad, oportunidad y obligaciones definidas en el presente Contrato, incluyendo su Anexo Técnico y sus Pliegos de Condiciones.  
6.2. Entregar el Cronograma estimado de obra que constituirá el anexo correspondiente del presente Contrato.  
6.3. Colaborar con {nombre_entidad} en cualquier requerimiento que ella haga.  
6.4. Garantizar la calidad de los bienes y servicios prestados, de acuerdo con el Anexo Técnico, el pliego de condiciones y la oferta presentada a {nombre_entidad}.  
6.5. Dar a conocer a {nombre_entidad} cualquier reclamación que indirecta o directamente pueda tener algún efecto sobre el objeto del Contrato o sobre sus obligaciones.  
6.6. Comunicarle a {nombre_entidad} cualquier circunstancia política, jurídica, social, económica, técnica, ambiental o de cualquier tipo, que pueda afectar la ejecución del contrato.  
6.7. Elaborar, suscribir y presentar a {nombre_entidad} las respectivas Actas parciales de Obra. Estas Actas parciales de Obra deben estar aprobadas por el Interventor y/o Supervisor del Contrato, según corresponda.  
6.8. Cumplir las obligaciones en materia ambiental, predial y de responsabilidad social que le competen conforme a normas aplicables y a las especificaciones técnicas de la obra.  
6.9. Las demás obligaciones que resulten aplicables conforme al documento tipo y al proceso de contratación.

## Cláusula 7 – Derechos particulares de la Entidad Estatal contratante

7.1. Revisar, rechazar, corregir o modificar las Actas de Obra y solicitar las correcciones o modificaciones que la obra necesite.  
7.2. Hacer uso de las cláusulas excepcionales del contrato.  
7.3. Hacer uso de la cláusula de imposición de multas, la cláusula penal o cualquier otro derecho consagrado a la Entidad Estatal contratante de manera legal o contractual.  
7.4. Los demás derechos que resulten aplicables conforme al documento tipo y a la normatividad vigente.

## Cláusula 8 – Obligaciones Generales de la Entidad Estatal contratante

8.1. Ejercer control sobre el presente Contrato, de manera directa o indirecta.  
8.2. Pagar el valor de la obra pública, de acuerdo con los términos establecidos en el presente Contrato.  
8.3. Prestar su colaboración para el cumplimiento de las obligaciones del Contratista.  
8.4. Acoger y ejecutar respecto del Contratista las directrices y lineamientos sobre la ejecución, seguimiento y monitoreo del Contrato que resulten aplicables.  
8.5. Las demás obligaciones que resulten aplicables conforme al documento tipo y a la normatividad vigente.

## Cláusula 9 – Responsabilidad

{nombre_contratista} es responsable por el cumplimiento del objeto establecido en la cláusula 1 del presente Contrato. {nombre_contratista} será responsable por los daños que ocasionen sus empleados y/o consultores, los empleados y/o consultores de sus subcontratistas, a {nombre_entidad} en la ejecución del objeto del presente Contrato.

## Cláusula 10 – Confidencialidad

En caso de que exista información sujeta a reserva legal, las partes deben mantener la confidencialidad de esta información. Para ello, la parte interesada debe comunicar a la otra parte que la información suministrada tiene el carácter de confidencial.

La Entidad Estatal contratante puede definir qué documentos o asuntos están sometidos a confidencialidad.

## Cláusula 11 – Terminación, modificación e interpretación unilaterales del Contrato

{nombre_entidad} puede terminar, modificar y/o interpretar unilateralmente el Contrato, de acuerdo con la normatividad aplicable, cuando lo considere necesario para que el Contratista cumpla con el objeto del presente contrato.

## Cláusula 12 – Caducidad

{nombre_entidad} estará facultada para declarar la caducidad cuando exista un incumplimiento del contrato por parte del Contratista en la forma y de acuerdo con el procedimiento previsto por la ley.

## Cláusula 13 – Multas

En caso de incumplimiento a las obligaciones del Contratista derivadas del presente Contrato, {nombre_entidad} puede adelantar el procedimiento establecido en la ley e imponer las multas previstas en el documento tipo de contrato de obra pública.

## Cláusula 14 – Cláusula Penal

En caso de declaratoria de caducidad o de incumplimiento total o parcial de las obligaciones del presente Contrato, {nombre_contratista} debe pagar a {nombre_entidad}, a título de indemnización, una suma equivalente a {clausula_penal_numeros} ({clausula_penal_letras}). El valor pactado de la presente cláusula penal es el de la estimación anticipada de perjuicios; no obstante, la presente cláusula no impide el cobro de todos los perjuicios adicionales que se causen sobre el citado valor.

## Cláusula 15 – Garantías y Mecanismos de cobertura del riesgo

El Contratista se obliga a garantizar el cumplimiento de las obligaciones surgidas a favor de la Entidad Estatal contratante, con ocasión de la ejecución del contrato, de acuerdo con la información de la siguiente tabla:

{tabla_garantias}

El Contratista se compromete a mantener vigente la garantía durante todo el tiempo de ejecución del contrato.

El Contratista debe presentar dentro de los {plazo_garantias_dias} días hábiles siguientes a la firma del presente contrato las garantías a favor de {nombre_entidad}.

## Cláusula 16 – Independencia del Contratista

El Contratista es una entidad independiente de {nombre_entidad}, y en consecuencia, el Contratista no es su representante, agente o mandatario.

## Cláusula 17 – Cesión

El Contratista no puede ceder parcial ni totalmente sus obligaciones o derechos derivados del presente Contrato sin la autorización previa y por escrito de {nombre_entidad}.

## Cláusula 18 – Subcontratación

{nombre_contratista} puede subcontratar con cualquier tercero la ejecución de las actividades relacionadas con el objeto del presente contrato. Sin embargo, el Contratista debe comunicar estas contrataciones a la Entidad Estatal contratante y debe tener el debido registro de este tipo de negocios jurídicos. El Contratista debe mantener indemne a la Entidad Estatal contratante de acuerdo con la cláusula 19.

## Cláusula 19 – Indemnidad

El Contratista se obliga a indemnizar a {nombre_entidad} con ocasión de la violación o el incumplimiento de las obligaciones previstas en el presente Contrato.

## Cláusula 20 – Caso Fortuito y Fuerza Mayor

Las partes quedan exoneradas de responsabilidad por el incumplimiento de cualquiera de sus obligaciones o por la demora en la satisfacción de cualquiera de las prestaciones a su cargo derivadas del presente Contrato cuando el incumplimiento sea resultado o consecuencia de la ocurrencia de un evento de fuerza mayor y caso fortuito debidamente invocado y constatado de acuerdo con la ley y la jurisprudencia colombiana.

## Cláusula 21 – Solución de Controversias

Las controversias o diferencias que surjan entre el Contratista y la Entidad Estatal Contratante con ocasión de la firma, ejecución, interpretación, prórroga o terminación del Contrato, así como de cualquier otro asunto relacionado con el presente Contrato, serán sometidas a la revisión de las partes para buscar un arreglo directo, en un término no mayor a cinco (5) días hábiles a partir de la fecha en que cualquiera de las partes comunique por escrito a la otra parte la existencia de una diferencia y la explique someramente.

Las controversias que no puedan ser resueltas de forma directa entre las partes, se resolverán mediante los mecanismos previstos en el documento tipo de contrato de obra pública, incluidos, según aplique, amigable composición, conciliación, tribunal de arbitramento o jurisdicción contenciosa administrativa.

## Cláusula 22 – Notificaciones

Los avisos, solicitudes, comunicaciones y notificaciones que las Partes deban hacer en desarrollo del presente Contrato deben constar por escrito y se entenderán debidamente efectuadas solo si son entregadas personalmente o por correo electrónico a las siguientes direcciones:

**{nombre_entidad}**
- Dirección: {not_entidad_direccion}
- Teléfono: {not_entidad_telefono}
- Correo electrónico: {not_entidad_correo}

**{nombre_contratista}**
- Dirección: {not_contratista_direccion}
- Teléfono: {not_contratista_telefono}
- Correo electrónico: {not_contratista_correo}

## Cláusula 23 – Supervisión

{bloque_supervision}

## Cláusula 24 – Interventoría

{bloque_interventoria}

## Cláusula 25 – Anexos del Contrato

{anexos_txt}

## Cláusula 26 – Perfeccionamiento y ejecución

El presente contrato requiere para su perfeccionamiento de la firma de las partes. Para su ejecución requiere el registro presupuestal y la acreditación de encontrarse el Contratista a paz y salvo por concepto de aportes al sistema de seguridad social integral.

## Cláusula 27 – Lugar de ejecución y domicilio contractual

Las actividades previstas en el presente Contrato se desarrollarán en {lugar_ejecucion}.

Para constancia, se firma en {lugar_celebracion} el {fecha_contrato}.
"""
    return contrato


def generar_word(datos):
    numero_contrato = texto_si_vacio(datos.get("numero_contrato"))
    nombre_proyecto = texto_si_vacio(datos.get("nombre_proyecto"))
    fecha_contrato = texto_si_vacio(datos.get("fecha_contrato"))
    lugar_celebracion = texto_si_vacio(datos.get("lugar_celebracion"))

    rep_entidad_nombre = texto_si_vacio(datos.get("rep_entidad_nombre"))
    rep_entidad_tipo_doc = texto_si_vacio(datos.get("rep_entidad_tipo_doc"))
    rep_entidad_num_doc = texto_si_vacio(datos.get("rep_entidad_num_doc"))
    rep_entidad_cargo = texto_si_vacio(datos.get("rep_entidad_cargo"))

    tipo_contratista = texto_si_vacio(datos.get("tipo_contratista"))
    rep_contratista_nombre = texto_si_vacio(datos.get("rep_contratista_nombre"))
    rep_contratista_tipo_doc = texto_si_vacio(datos.get("rep_contratista_tipo_doc"))
    rep_contratista_num_doc = texto_si_vacio(datos.get("rep_contratista_num_doc"))

    contrato_txt = construir_contrato(datos)

    doc = Document()
    doc.add_heading("Contrato de obra pública", level=0)

    for bloque in contrato_txt.split("\n\n"):
        if bloque.strip().startswith("# "):
            doc.add_heading(bloque.replace("# ", "").strip(), level=1)
        elif bloque.strip().startswith("## "):
            doc.add_heading(bloque.replace("## ", "").strip(), level=2)
        elif "| Amparo | Suficiencia | Vigencia |" in bloque:
            construir_tabla_garantias_doc(doc, datos.get("garantias", []))
        else:
            doc.add_paragraph(bloque.strip())

    doc.add_paragraph("")
    tabla_firmas = doc.add_table(rows=1, cols=2)
    tabla_firmas.style = "Table Grid"

    izquierda = tabla_firmas.rows[0].cells[0]
    derecha = tabla_firmas.rows[0].cells[1]

    izquierda.text = (
        f"{rep_entidad_nombre}\n"
        f"{rep_entidad_cargo}\n"
        f"{rep_entidad_tipo_doc} No. {rep_entidad_num_doc}"
    )

    derecha.text = (
        f"{rep_contratista_nombre}\n"
        f"Representante del contratista {tipo_contratista}\n"
        f"{rep_contratista_tipo_doc} No. {rep_contratista_num_doc}"
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    nombre_archivo = f"contrato_{numero_contrato}_{nombre_proyecto}".replace(" ", "_")
    return buffer, f"{nombre_archivo}.docx"


datos = obtener_datos_contrato()

st.markdown("""
    <style>
    .titulo-seccion { font-size: 32px !important; font-weight: 800 !important; color: #7A0019; }
    .subtitulo-gris { font-size: 16px !important; color: #666; margin-bottom: 15px; }
    div[data-testid="stProgress"] > div > div > div > div { background-color: #C62828 !important; }
    section[data-testid="stSidebar"] { background-color: #f4f4f4; }
    .stButton > button { width: 100%; border-radius: 6px; height: 3em; font-weight: bold; }
    button[kind="primary"] {
        background-color: #7A0019 !important;
        border-color: #7A0019 !important;
        color: white !important;
    }
    button[kind="primary"]:hover {
        background-color: #5C0013 !important;
        border-color: #5C0013 !important;
        color: white !important;
    }
    </style>
""", unsafe_allow_html=True)

col_t, col_l = st.columns([4, 1], vertical_alignment="center")
with col_t:
    st.markdown('<div class="titulo-seccion">📄 Contrato armado</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="subtitulo-gris">Vista de lectura del contrato construido automáticamente con la información diligenciada.</div>',
        unsafe_allow_html=True
    )
with col_l:
    if os.path.exists("unnamed.jpg"):
        st.image("unnamed.jpg", use_container_width=True)

st.divider()

with st.sidebar:
    st.header("🧭 Resumen")
    st.markdown(f"**Contrato:** {texto_si_vacio(datos.get('numero_contrato'))}")
    st.markdown(f"**Proyecto:** {texto_si_vacio(datos.get('nombre_proyecto'))}")
    st.markdown(f"**Entidad:** {texto_si_vacio(datos.get('nombre_entidad'))}")
    st.markdown(f"**Contratista:** {texto_si_vacio(datos.get('nombre_contratista'))}")
    st.markdown(f"**Valor:** {texto_si_vacio(datos.get('valor_total_numeros'))}")
    st.markdown(f"**Plazo:** {texto_si_vacio(datos.get('plazo_ejecucion'))}")

contrato_armado = construir_contrato(datos)
st.markdown(contrato_armado)

st.markdown("---")
st.markdown("### Firmas")

col1, col2 = st.columns(2)
with col1:
    st.markdown(
        f"""**{texto_si_vacio(datos.get("rep_entidad_nombre"))}**  
{texto_si_vacio(datos.get("rep_entidad_cargo"))}  
{texto_si_vacio(datos.get("rep_entidad_tipo_doc"))} No. {texto_si_vacio(datos.get("rep_entidad_num_doc"))}"""
    )

with col2:
    st.markdown(
        f"""**{texto_si_vacio(datos.get("rep_contratista_nombre"))}**  
Representante del contratista {texto_si_vacio(datos.get("tipo_contratista"))}  
{texto_si_vacio(datos.get("rep_contratista_tipo_doc"))} No. {texto_si_vacio(datos.get("rep_contratista_num_doc"))}"""
    )

word_buffer, nombre_word = generar_word(datos)
st.download_button(
    "Descargar contrato en Word",
    data=word_buffer,
    file_name=nombre_word,
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    type="primary",
    key="descargar_contrato_word",
)
