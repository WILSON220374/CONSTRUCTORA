from io import BytesIO
from datetime import date, datetime

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt

from supabase_state import cargar_estado
from supabase_state import guardar_estado as guardar_estado_bd


CLAVE_GUARDADO = "items_no_previstos"


# ==========================================================
# Persistencia
# ==========================================================
def guardar_estado(clave, datos):
    def serializar(obj):
        if isinstance(obj, dict):
            return {k: serializar(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [serializar(x) for x in obj]
        if isinstance(obj, (date, datetime)):
            return obj.isoformat()
        return obj

    guardar_estado_bd(clave, serializar(datos))


def _leer_estado(clave):
    datos = cargar_estado(clave) or {}
    return datos if isinstance(datos, dict) else {}


# ==========================================================
# Helpers
# ==========================================================
def _texto(valor):
    if valor is None:
        return ""
    return str(valor).strip()


def _safe_float(valor, default=0.0):
    try:
        if valor is None or valor == "":
            return float(default)

        if isinstance(valor, (int, float)):
            return float(valor)

        txt = str(valor).strip().replace("$", "").replace(" ", "")

        if "," in txt and "." in txt:
            txt = txt.replace(".", "").replace(",", ".")
        elif "," in txt:
            txt = txt.replace(",", ".")

        return float(txt)
    except Exception:
        return float(default)


def _parse_fecha(valor):
    if isinstance(valor, date):
        return valor
    if isinstance(valor, datetime):
        return valor.date()
    if isinstance(valor, str) and valor.strip():
        txt = valor.strip()
        try:
            return datetime.fromisoformat(txt).date()
        except Exception:
            pass
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
            try:
                return datetime.strptime(txt, fmt).date()
            except Exception:
                continue
    return None


def _fecha_input(valor):
    return _parse_fecha(valor) or date.today()


def _fecha_texto(valor):
    fecha = _parse_fecha(valor)
    return fecha.strftime("%d/%m/%Y") if fecha else ""


def _moneda(valor):
    return f"$ {_safe_float(valor, 0.0):,.2f}"


def _primero_no_vacio(*valores):
    for valor in valores:
        txt = _texto(valor)
        if txt:
            return txt
    return ""


def _key_codigo_natural(value):
    partes = []
    for chunk in _texto(value).split("."):
        try:
            partes.append(int(chunk))
        except Exception:
            partes.append(chunk)
    return tuple(partes)


# ==========================================================
# Datos base
# ==========================================================
def _valor_contrato_obra(acta_inicio, contrato_obra):
    for valor in [
        contrato_obra.get("valor_total_numeros"),
        contrato_obra.get("valor_contrato"),
        contrato_obra.get("valor"),
    ]:
        numero = _safe_float(valor, 0.0)
        if numero > 0:
            return numero
    return 0.0


def _fecha_inicio(acta_inicio):
    return _parse_fecha(
        _primero_no_vacio(
            acta_inicio.get("fecha_inicio"),
            acta_inicio.get("fecha_presente_acta"),
        )
    )


def _datos_generales(acta_inicio, contrato_obra):
    return {
        "contrato_no": _primero_no_vacio(acta_inicio.get("numero_contrato"), contrato_obra.get("numero_contrato")),
        "objeto_contrato": _primero_no_vacio(
            acta_inicio.get("objeto_contrato"),
            contrato_obra.get("objeto_general"),
            contrato_obra.get("objeto_contrato"),
            contrato_obra.get("objeto"),
        ),
        "contratista": _primero_no_vacio(
            acta_inicio.get("nombre_firma_contratista"),
            contrato_obra.get("nombre_contratista"),
        ),
        "interventor": _primero_no_vacio(
            acta_inicio.get("nombre_firma_interventor"),
            contrato_obra.get("nombre_interventor"),
            contrato_obra.get("nombre_supervisor"),
        ),
        "fecha_inicio": _fecha_inicio(acta_inicio),
        "valor_inicial": _valor_contrato_obra(acta_inicio, contrato_obra),
    }


# ==========================================================
# Catálogo presupuesto
# ==========================================================
def _catalogo_presupuesto(presupuesto_obra):
    catalogo = []
    mapa = {}

    if not isinstance(presupuesto_obra, dict):
        return catalogo, mapa

    tablas = presupuesto_obra.get("__tablas__", {}) or {}
    grupos = tablas.get("grupos_presupuesto_obra", []) or []

    for grupo in grupos:
        if not isinstance(grupo, dict):
            continue

        for fila in grupo.get("rows", []) or []:
            if not isinstance(fila, dict):
                continue

            item = _texto(fila.get("ITEM"))
            if not item:
                continue

            registro = {
                "ÍTEM": item,
                "DESCRIPCIÓN DEL ÍTEM": _texto(fila.get("DESCRIPCIÓN", fila.get("DESCRIPCION", ""))),
                "UNIDAD": _texto(fila.get("UNIDAD")),
                "CANTIDAD": _safe_float(fila.get("CANT"), 0.0),
            }

            if item not in mapa:
                catalogo.append(registro)
                mapa[item] = registro

    catalogo = sorted(catalogo, key=lambda x: _key_codigo_natural(x.get("ÍTEM")))
    return catalogo, mapa


# ==========================================================
# Estado actas
# ==========================================================
def _fila_item_vacia():
    return {
        "ÍTEM": "",
        "DESCRIPCIÓN DEL ÍTEM": "",
        "UNIDAD": "",
        "CANTIDAD": 0.0,
        "PRECIO UNITARIO CONTRATISTA": 0.0,
        "PRECIO UNITARIO INTERVENTORÍA": 0.0,
        "PRECIO ACORDADO": 0.0,
    }


def _normalizar_items(rows, mapa_catalogo):
    filas = []

    for fila in rows or []:
        base = _fila_item_vacia()
        if isinstance(fila, dict):
            item = _texto(fila.get("ÍTEM"))
            base["ÍTEM"] = item

            if item and item in mapa_catalogo:
                base["DESCRIPCIÓN DEL ÍTEM"] = _texto(mapa_catalogo[item].get("DESCRIPCIÓN DEL ÍTEM"))
                base["UNIDAD"] = _texto(mapa_catalogo[item].get("UNIDAD"))
                base["CANTIDAD"] = _safe_float(mapa_catalogo[item].get("CANTIDAD"), 0.0)
            else:
                base["DESCRIPCIÓN DEL ÍTEM"] = _texto(fila.get("DESCRIPCIÓN DEL ÍTEM"))
                base["UNIDAD"] = _texto(fila.get("UNIDAD"))
                base["CANTIDAD"] = _safe_float(fila.get("CANTIDAD"), 0.0)

            base["PRECIO UNITARIO CONTRATISTA"] = _safe_float(fila.get("PRECIO UNITARIO CONTRATISTA"), 0.0)
            base["PRECIO UNITARIO INTERVENTORÍA"] = _safe_float(fila.get("PRECIO UNITARIO INTERVENTORÍA"), 0.0)
            base["PRECIO ACORDADO"] = _safe_float(fila.get("PRECIO ACORDADO"), 0.0)

        filas.append(base)

    if not filas:
        filas.append(_fila_item_vacia())

    return filas


def _acta_vacia(consecutivo, generales):
    return {
        "consecutivo": int(consecutivo),
        "fecha": date.today().isoformat(),
        "fecha_vencimiento": date.today().isoformat(),
        "valor_acumulado": generales.get("valor_inicial", 0.0),
        "items": [_fila_item_vacia()],
        "observaciones": "",
        "nombre_contratista_firma": generales.get("contratista", ""),
        "nombre_interventor_firma": generales.get("interventor", ""),
    }


def _normalizar_acta(acta, consecutivo, generales, mapa_catalogo):
    base = _acta_vacia(consecutivo, generales)
    if isinstance(acta, dict):
        base.update(acta)
        base["consecutivo"] = int(acta.get("consecutivo") or consecutivo or 1)
        base["fecha"] = (_parse_fecha(acta.get("fecha")) or date.today()).isoformat()
        base["fecha_vencimiento"] = (_parse_fecha(acta.get("fecha_vencimiento")) or date.today()).isoformat()
        valor_acumulado_guardado = _safe_float(acta.get("valor_acumulado"), generales.get("valor_inicial", 0.0))
        valor_inicial_base = _safe_float(generales.get("valor_inicial"), 0.0)

        if valor_acumulado_guardado > valor_inicial_base * 10 and valor_inicial_base > 0:
            valor_acumulado_guardado = valor_inicial_base

        base["valor_acumulado"] = valor_acumulado_guardado
        base["items"] = _normalizar_items(acta.get("items", []), mapa_catalogo)
        base["observaciones"] = _texto(acta.get("observaciones"))
        base["nombre_contratista_firma"] = _primero_no_vacio(acta.get("nombre_contratista_firma"), generales.get("contratista"))
        base["nombre_interventor_firma"] = _primero_no_vacio(acta.get("nombre_interventor_firma"), generales.get("interventor"))
    return base


def _normalizar_estado(cargado, generales, mapa_catalogo):
    if not isinstance(cargado, dict):
        cargado = {}

    if isinstance(cargado.get("actas"), list):
        actas = [
            _normalizar_acta(acta, i, generales, mapa_catalogo)
            for i, acta in enumerate(cargado.get("actas", []), start=1)
            if isinstance(acta, dict)
        ]
    else:
        migrada = _normalizar_acta(cargado, int(cargado.get("consecutivo") or 1), generales, mapa_catalogo)
        actas = [migrada]

    if not actas:
        actas = [_acta_vacia(1, generales)]

    actas = sorted(actas, key=lambda x: int(x.get("consecutivo") or 0))
    activa = int(cargado.get("acta_activa") or actas[-1].get("consecutivo") or 1)
    consecutivos = [int(x.get("consecutivo") or 0) for x in actas]
    if activa not in consecutivos:
        activa = int(actas[-1].get("consecutivo") or 1)

    return {"actas": actas, "acta_activa": activa}


def _inicializar_estado(generales, mapa_catalogo):
    group_id_actual = _texto(st.session_state.get("group_id"))
    cache_group = _texto(st.session_state.get("_items_no_previstos_group"))

    if cache_group != group_id_actual or "items_no_previstos_datos" not in st.session_state:
        cargado = cargar_estado(CLAVE_GUARDADO) or {}
        st.session_state["items_no_previstos_datos"] = _normalizar_estado(cargado, generales, mapa_catalogo)
        st.session_state["_items_no_previstos_group"] = group_id_actual


def _obtener_acta_activa(generales, mapa_catalogo):
    estado = st.session_state["items_no_previstos_datos"]
    actas = estado.get("actas", [])
    activa = int(estado.get("acta_activa") or 1)

    for acta in actas:
        if int(acta.get("consecutivo") or 0) == activa:
            return acta

    nueva = _acta_vacia(1, generales)
    nueva["items"] = _normalizar_items(nueva.get("items", []), mapa_catalogo)
    estado["actas"] = [nueva]
    estado["acta_activa"] = 1
    return nueva


def _crear_nueva_acta(generales):
    estado = st.session_state["items_no_previstos_datos"]
    actas = estado.get("actas", [])
    ultimo = max([int(x.get("consecutivo") or 0) for x in actas], default=0)
    nueva = _acta_vacia(ultimo + 1, generales)
    actas.append(nueva)
    estado["actas"] = actas
    estado["acta_activa"] = int(nueva["consecutivo"])
    return int(nueva["consecutivo"])


def _guardar():
    estado = st.session_state["items_no_previstos_datos"]
    actas_actuales = estado.get("actas", []) if isinstance(estado.get("actas"), list) else []

    guardado = cargar_estado(CLAVE_GUARDADO) or {}
    actas_guardadas = guardado.get("actas", []) if isinstance(guardado, dict) else []
    if not isinstance(actas_guardadas, list):
        actas_guardadas = []

    if actas_guardadas and len(actas_actuales) < len(actas_guardadas):
        st.error(
            "No se guardó porque el estado actual tiene menos actas que el estado ya guardado. "
            "Esto evita sobrescribir actas anteriores."
        )
        return False

    guardar_estado(CLAVE_GUARDADO, estado)
    st.success("Ítems no previstos guardados correctamente.")
    return True


# ==========================================================
# Word
# ==========================================================
def _set_doc_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(8)


def _p(doc, texto="", bold=False, align=None, size=8):
    p = doc.add_paragraph()
    run = p.add_run(str(texto))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def _cell_text(cell, texto, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(texto))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(7)
    if align:
        p.alignment = align


def _tabla_simple(doc, filas):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for etiqueta, valor in filas:
        cells = table.add_row().cells
        _cell_text(cells[0], etiqueta, bold=True)
        _cell_text(cells[1], valor)
    return table


def _tabla_items_word(doc, rows):
    cols = [
        "ÍTEM",
        "DESCRIPCIÓN DEL ÍTEM",
        "UNIDAD",
        "CANTIDAD",
        "PRECIO UNITARIO CONTRATISTA",
        "PRECIO UNITARIO INTERVENTORÍA",
        "PRECIO ACORDADO",
    ]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(cols):
        _cell_text(table.rows[0].cells[i], col, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for fila in rows or []:
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            valor = fila.get(col, "")
            if col in ["CANTIDAD"]:
                valor = f"{_safe_float(valor, 0.0):,.4f}"
            elif col.startswith("PRECIO"):
                valor = _moneda(valor)
            _cell_text(cells[i], valor, align=WD_ALIGN_PARAGRAPH.CENTER)

    return table


def _generar_word(generales, acta):
    doc = Document()
    _set_doc_defaults(doc)

    _p(doc, "ÍTEMS NO PREVISTOS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    _p(doc, f"ACTA No. {int(acta.get('consecutivo') or 0)}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    _p(doc, "")

    _tabla_simple(
        doc,
        [
            ("FECHA", _fecha_texto(acta.get("fecha"))),
            ("CONTRATO No.", generales.get("contrato_no", "")),
            ("OBJETO DEL CONTRATO", generales.get("objeto_contrato", "")),
            ("CONTRATISTA", generales.get("contratista", "")),
            ("INTERVENTOR", generales.get("interventor", "")),
            ("FECHA DE INICIO", _fecha_texto(generales.get("fecha_inicio"))),
            ("VALOR INICIAL", _moneda(generales.get("valor_inicial", 0.0))),
            ("FECHA DE VENCIMIENTO", _fecha_texto(acta.get("fecha_vencimiento"))),
            ("VALOR ACUMULADO", _moneda(acta.get("valor_acumulado", 0.0))),
        ],
    )

    _p(doc, "")
    _p(doc, "DATOS ESPECÍFICOS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_items_word(doc, acta.get("items", []))

    _p(doc, "")
    _p(doc, "OBSERVACIONES", bold=True)
    _p(doc, acta.get("observaciones", ""))

    _p(doc, "")
    _p(doc, "NOTAS:", bold=True)
    _p(doc, "1. El Interventor certifica que realizó un análisis comparativo de los análisis de precios no previstos presentados por el Contratista.")
    _p(doc, "2. El Interventor es el responsable de la revisión, análisis y aprobación del análisis o los análisis de precios no previstos mediante los cuales se acuerda el precio o los precios entre el Contratista y el Interventor.")
    _p(doc, "3. El Contratista y el Interventor certifican que los precios unitarios del ítem o de los ítems no previstos corresponden a los precios del mercado.")
    _p(doc, "4. En la casilla de Observaciones el Interventor debe consignar si es del caso las aclaraciones a que haya lugar del precio o de los precios que se acordaron.")

    _p(doc, "")
    tabla_firmas = doc.add_table(rows=2, cols=2)
    tabla_firmas.style = "Table Grid"
    tabla_firmas.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cell_text(tabla_firmas.cell(0, 0), "CONTRATISTA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(tabla_firmas.cell(0, 1), "INTERVENTOR", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(tabla_firmas.cell(1, 0), f"Nombre: {acta.get('nombre_contratista_firma', '')}", align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(tabla_firmas.cell(1, 1), f"Nombre: {acta.get('nombre_interventor_firma', '')}", align=WD_ALIGN_PARAGRAPH.CENTER)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==========================================================
# Interfaz
# ==========================================================
st.set_page_config(page_title="Ítems no previstos", layout="wide")
st.title("31. Ítems no previstos")

acta_inicio = _leer_estado("acta_inicio_obra")
contrato_obra = _leer_estado("contrato_obra")
presupuesto_obra = _leer_estado("presupuesto_obra")

generales = _datos_generales(acta_inicio, contrato_obra)
catalogo, mapa_catalogo = _catalogo_presupuesto(presupuesto_obra)
items_catalogo = [x["ÍTEM"] for x in catalogo]

_inicializar_estado(generales, mapa_catalogo)
estado = st.session_state["items_no_previstos_datos"]

st.markdown("### ACTAS GENERADAS")
col_nueva, col_selector = st.columns([1, 3])

with col_nueva:
    if st.button("Nueva acta", key="items_no_previstos_nueva_acta"):
        nueva = _crear_nueva_acta(generales)
        _guardar()
        st.session_state["items_no_previstos_selector"] = nueva
        st.rerun()

actas = estado.get("actas", [])
opciones = [int(x.get("consecutivo") or 0) for x in actas]
activa_actual = int(estado.get("acta_activa") or (opciones[-1] if opciones else 1))

with col_selector:
    acta_activa = st.selectbox(
        "Seleccione acta",
        options=opciones,
        index=opciones.index(activa_actual) if activa_actual in opciones else 0,
        format_func=lambda x: f"Acta No. {x}",
        key="items_no_previstos_selector",
    )

estado["acta_activa"] = int(acta_activa)
acta = _obtener_acta_activa(generales, mapa_catalogo)
acta_no = int(acta.get("consecutivo") or acta_activa or 1)

st.markdown("### DATOS GENERALES")

acta["fecha"] = st.date_input(
    "FECHA",
    value=_fecha_input(acta.get("fecha")),
    format="DD/MM/YYYY",
    key=f"items_no_previstos_fecha_{acta_no}",
)

c1, c2 = st.columns(2)
with c1:
    st.text_input("CONTRATO No.", value=generales.get("contrato_no", ""), disabled=True)
    st.text_area("OBJETO DEL CONTRATO", value=generales.get("objeto_contrato", ""), disabled=True, height=100)
    st.text_input("CONTRATISTA", value=generales.get("contratista", ""), disabled=True)
    st.text_input("INTERVENTOR", value=generales.get("interventor", ""), disabled=True)
with c2:
    st.date_input("FECHA DE INICIO", value=_fecha_input(generales.get("fecha_inicio")), disabled=True, format="DD/MM/YYYY")
    st.number_input("VALOR INICIAL", value=float(generales.get("valor_inicial", 0.0)), disabled=True, format="%.2f")
    acta["fecha_vencimiento"] = st.date_input(
        "FECHA DE VENCIMIENTO",
        value=_fecha_input(acta.get("fecha_vencimiento")),
        format="DD/MM/YYYY",
        key=f"items_no_previstos_fecha_vencimiento_{acta_no}",
    )
    key_valor_acumulado = f"items_no_previstos_valor_acumulado_{acta_no}"
    valor_inicial_base = _safe_float(generales.get("valor_inicial"), 0.0)
    valor_acumulado_mostrar = _safe_float(acta.get("valor_acumulado"), valor_inicial_base)

    if valor_acumulado_mostrar > valor_inicial_base * 10 and valor_inicial_base > 0:
        valor_acumulado_mostrar = valor_inicial_base
        acta["valor_acumulado"] = valor_acumulado_mostrar

    if key_valor_acumulado in st.session_state:
        valor_widget = _safe_float(st.session_state.get(key_valor_acumulado), 0.0)
        if valor_widget > valor_inicial_base * 10 and valor_inicial_base > 0:
            st.session_state[key_valor_acumulado] = valor_inicial_base

    acta["valor_acumulado"] = st.number_input(
        "VALOR ACUMULADO",
        value=float(valor_acumulado_mostrar),
        format="%.2f",
        key=key_valor_acumulado,
    )

st.markdown("### DATOS ESPECÍFICOS")

items_normalizados = _normalizar_items(acta.get("items", []), mapa_catalogo)
df_items = pd.DataFrame(
    items_normalizados,
    columns=[
        "ÍTEM",
        "DESCRIPCIÓN DEL ÍTEM",
        "UNIDAD",
        "CANTIDAD",
        "PRECIO UNITARIO CONTRATISTA",
        "PRECIO UNITARIO INTERVENTORÍA",
        "PRECIO ACORDADO",
    ],
)

items_editados = st.data_editor(
    df_items,
    hide_index=True,
    width="stretch",
    num_rows="dynamic",
    key=f"items_no_previstos_editor_{acta_no}",
    column_config={
        "ÍTEM": st.column_config.SelectboxColumn("ÍTEM", options=[""] + items_catalogo, required=False),
        "DESCRIPCIÓN DEL ÍTEM": st.column_config.TextColumn("DESCRIPCIÓN DEL ÍTEM", disabled=True),
        "UNIDAD": st.column_config.TextColumn("UNIDAD", disabled=True),
        "CANTIDAD": st.column_config.NumberColumn("CANTIDAD", format="%.4f", disabled=True),
        "PRECIO UNITARIO CONTRATISTA": st.column_config.NumberColumn("PRECIO UNITARIO CONTRATISTA", format="$ %.2f"),
        "PRECIO UNITARIO INTERVENTORÍA": st.column_config.NumberColumn("PRECIO UNITARIO INTERVENTORÍA", format="$ %.2f"),
        "PRECIO ACORDADO": st.column_config.NumberColumn("PRECIO ACORDADO", format="$ %.2f"),
    },
)

items_recalculados = _normalizar_items(items_editados.to_dict("records"), mapa_catalogo)
if items_recalculados != acta.get("items", []):
    acta["items"] = items_recalculados
    st.rerun()
acta["items"] = items_recalculados

st.markdown("### OBSERVACIONES")
acta["observaciones"] = st.text_area(
    "OBSERVACIONES",
    value=acta.get("observaciones", ""),
    height=160,
    label_visibility="collapsed",
    key=f"items_no_previstos_observaciones_{acta_no}",
)

st.markdown("### NOTAS:")
st.write("1. El Interventor certifica que realizó un análisis comparativo de los análisis de precios no previstos presentados por el Contratista.")
st.write("2. El Interventor es el responsable de la revisión, análisis y aprobación del análisis o los análisis de precios no previstos mediante los cuales se acuerda el precio o los precios entre el Contratista y el Interventor.")
st.write("3. El Contratista y el Interventor certifican que los precios unitarios del ítem o de los ítems no previstos corresponden a los precios del mercado.")
st.write("4. En la casilla de Observaciones el Interventor debe consignar si es del caso las aclaraciones a que haya lugar del precio o de los precios que se acordaron.")

st.markdown("### FIRMAS")
col_f1, col_f2 = st.columns(2)
with col_f1:
    acta["nombre_contratista_firma"] = st.text_input(
        "CONTRATISTA - Nombre",
        value=acta.get("nombre_contratista_firma", generales.get("contratista", "")),
        key=f"items_no_previstos_firma_contratista_{acta_no}",
    )
with col_f2:
    acta["nombre_interventor_firma"] = st.text_input(
        "INTERVENTOR - Nombre",
        value=acta.get("nombre_interventor_firma", generales.get("interventor", "")),
        key=f"items_no_previstos_firma_interventor_{acta_no}",
    )

col_b1, col_b2 = st.columns(2)
with col_b1:
    if st.button("Guardar", key=f"items_no_previstos_guardar_{acta_no}"):
        _guardar()

with col_b2:
    word = _generar_word(generales, acta)
    st.download_button(
        "Descargar Word",
        data=word,
        file_name=f"items_no_previstos_{acta_no}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        width="stretch",
        key=f"items_no_previstos_word_{acta_no}",
    )
