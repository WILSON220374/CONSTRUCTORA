from io import BytesIO
from datetime import date, datetime, timedelta

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, Inches

from supabase_state import cargar_estado


# ==========================================================
# Helpers
# ==========================================================
def _texto(valor):
    if valor is None:
        return ""
    return str(valor).strip()


def _safe_float(valor, default=0.0):
    try:
        if valor is None or valor == "":
            return float(default)
        if isinstance(valor, (int, float)):
            return float(valor)
        txt = str(valor).strip().replace("$", "").replace(" ", "")
        txt = txt.replace(".", "").replace(",", ".")
        return float(txt)
    except Exception:
        return float(default)


def _parse_fecha(valor):
    if isinstance(valor, date):
        return valor
    if isinstance(valor, datetime):
        return valor.date()
    if isinstance(valor, str) and valor.strip():
        txt = valor.strip()
        try:
            return datetime.fromisoformat(txt).date()
        except Exception:
            pass
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
            try:
                return datetime.strptime(txt, fmt).date()
            except Exception:
                continue
    return None


def _fecha_texto(valor):
    fecha = _parse_fecha(valor)
    if not fecha:
        return ""
    return fecha.strftime("%d/%m/%Y")


def _moneda(valor):
    return f"$ {_safe_float(valor, 0.0):,.2f}"


def _leer_estado(clave):
    datos = cargar_estado(clave) or {}
    return datos if isinstance(datos, dict) else {}


def _primero_no_vacio(*valores):
    for valor in valores:
        txt = _texto(valor)
        if txt:
            return txt
    return ""


def _extraer_dias_plazo(valor):
    try:
        if isinstance(valor, (int, float)):
            return int(valor)
        txt = str(valor or "").strip()
        numeros = "".join(ch if ch.isdigit() else " " for ch in txt).split()
        return int(numeros[0]) if numeros else 0
    except Exception:
        return 0


def _valor_contrato(acta_inicio, contrato_obra):
    for valor in [
        acta_inicio.get("valor_total_contrato_obra"),
        acta_inicio.get("valor_contrato"),
        acta_inicio.get("valor"),
        contrato_obra.get("valor_total_numeros"),
        contrato_obra.get("valor_contrato"),
        contrato_obra.get("valor"),
    ]:
        numero = _safe_float(valor, 0.0)
        if numero > 0:
            return numero
    return 0.0


def _fecha_inicio(acta_inicio):
    return _parse_fecha(
        _primero_no_vacio(
            acta_inicio.get("fecha_inicio"),
            acta_inicio.get("fecha_presente_acta"),
        )
    )


def _fecha_vencimiento_inicial(acta_inicio, contrato_obra):
    fecha_inicio = _fecha_inicio(acta_inicio)
    if not fecha_inicio:
        return None

    plazo_dias = int(acta_inicio.get("plazo_ejecucion_dias", 0) or 0)
    if plazo_dias <= 0:
        plazo_dias = int(contrato_obra.get("plazo_ejecucion_dias", 0) or 0)
    if plazo_dias <= 0:
        plazo_dias = _extraer_dias_plazo(
            _primero_no_vacio(
                acta_inicio.get("plazo_ejecucion"),
                contrato_obra.get("plazo_ejecucion"),
            )
        )
    if plazo_dias <= 0:
        return None

    return fecha_inicio + timedelta(days=plazo_dias)


def _datos_generales(acta_inicio, contrato_obra):
    return {
        "numero_contrato": _primero_no_vacio(
            acta_inicio.get("numero_contrato"),
            contrato_obra.get("numero_contrato"),
        ),
        "contratante": _primero_no_vacio(
            contrato_obra.get("nombre_entidad"),
            acta_inicio.get("contratante"),
        ),
        "nit_contratante": _primero_no_vacio(
            contrato_obra.get("nit_entidad"),
            acta_inicio.get("nit_contratante"),
        ),
        "contratista": _primero_no_vacio(
            acta_inicio.get("nombre_firma_contratista"),
            contrato_obra.get("nombre_contratista"),
        ),
        "nit_contratista": _primero_no_vacio(
            contrato_obra.get("nit_contratista"),
            acta_inicio.get("nit_contratista"),
        ),
        "interventor": _primero_no_vacio(
            acta_inicio.get("nombre_firma_interventor"),
            contrato_obra.get("nombre_interventor"),
            contrato_obra.get("nombre_supervisor"),
        ),
        "supervisor": _primero_no_vacio(
            acta_inicio.get("nombre_firma_supervisor"),
            contrato_obra.get("nombre_supervisor"),
        ),
        "objeto": _primero_no_vacio(
            acta_inicio.get("objeto_contrato"),
            contrato_obra.get("objeto_general"),
            contrato_obra.get("objeto_contrato"),
            contrato_obra.get("objeto"),
        ),
        "valor_contrato": _valor_contrato(acta_inicio, contrato_obra),
        "plazo": _primero_no_vacio(
            acta_inicio.get("plazo_ejecucion"),
            contrato_obra.get("plazo_ejecucion"),
        ),
        "fecha_inicio": _fecha_inicio(acta_inicio),
        "fecha_vencimiento_inicial": _fecha_vencimiento_inicial(acta_inicio, contrato_obra),
        "contrato_interventoria": _primero_no_vacio(
            contrato_obra.get("numero_contrato_interventoria"),
            acta_inicio.get("numero_contrato_interventoria"),
        ),
        "concepto_interventoria": _primero_no_vacio(
            contrato_obra.get("concepto_interventoria"),
            acta_inicio.get("concepto_interventoria"),
        ),
    }


def _suspensiones(control_obra):
    rows = control_obra.get("suspensiones_rows", []) or []
    return [fila for fila in rows if isinstance(fila, dict)]


def _tipo_acta(fila):
    if _texto(fila.get("ACTA DE AMPLIACIÓN SUSPENSIÓN No.")):
        return "AMPLIACIÓN DE LA SUSPENSIÓN"
    return "SUSPENSIÓN"


def _numero_acta(fila):
    if _tipo_acta(fila) == "AMPLIACIÓN DE LA SUSPENSIÓN":
        return _texto(fila.get("ACTA DE AMPLIACIÓN SUSPENSIÓN No."))
    return _texto(fila.get("ACTA DE SUSPENSIÓN No."))


def _etiqueta_suspension(fila, idx):
    tipo = _tipo_acta(fila)
    numero = _numero_acta(fila)
    desde = _fecha_texto(fila.get("DESDE"))
    hasta = _fecha_texto(fila.get("HASTA"))
    return f"{idx + 1}. ACTA DE {tipo} No. {numero} | Desde {desde} hasta {hasta}"


def _filas_relacion_suspensiones(suspensiones, fecha_corte=None):
    filas = []
    fecha_corte = _parse_fecha(fecha_corte) if fecha_corte else None

    for fila in suspensiones:
        fecha_acta = _parse_fecha(fila.get("FECHA DEL ACTA"))

        if fecha_corte and fecha_acta > fecha_corte:
            continue

        filas.append(
            {
                "ACTA DE SUSPENSIÓN No.": _texto(fila.get("ACTA DE SUSPENSIÓN No.")),
                "ACTA DE AMPLIACIÓN SUSPENSIÓN No.": _texto(fila.get("ACTA DE AMPLIACIÓN SUSPENSIÓN No.")),
                "FECHA DEL ACTA": _fecha_texto(fila.get("FECHA DEL ACTA")),
                "DESDE": _fecha_texto(fila.get("DESDE")),
                "HASTA": _fecha_texto(fila.get("HASTA")),
                "PERIODO DE SUSPENSIÓN": _texto(fila.get("PERIODO DE SUSPENSIÓN")),
                "FECHA REANUDACIÓN": _fecha_texto(fila.get("HASTA")),
                "NUEVA FECHA DE VENCIMIENTO": _fecha_texto(fila.get("NUEVA FECHA DE FINALIZACIÓN")),
            }
        )
    return filas

# ==========================================================
# Word helpers
# ==========================================================
def _set_doc_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(8)


def _p(doc, texto="", bold=False, align=None, size=8):
    par = doc.add_paragraph()
    run = par.add_run(str(texto))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if align:
        par.alignment = align
    return par


def _cell_text(cell, texto, bold=False, align=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    par = cell.paragraphs[0]
    run = par.add_run(str(texto))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(8)
    if align:
        par.alignment = align


def _tabla_simple(doc, filas, widths=None):
    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    for etiqueta, valor in filas:
        cells = tabla.add_row().cells
        _cell_text(cells[0], etiqueta, bold=True)
        _cell_text(cells[1], valor)
    return tabla


def _tabla_relacion(doc, filas):
    cols = [
        "ACTA DE SUSPENSIÓN No.",
        "ACTA DE AMPLIACIÓN SUSPENSIÓN No.",
        "FECHA DEL ACTA",
        "DESDE",
        "HASTA",
        "PERIODO DE SUSPENSIÓN",
        "FECHA REANUDACIÓN",
        "NUEVA FECHA DE VENCIMIENTO",
    ]
    tabla = doc.add_table(rows=1, cols=len(cols))
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(cols):
        _cell_text(tabla.rows[0].cells[i], col, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for fila in filas:
        cells = tabla.add_row().cells
        for i, col in enumerate(cols):
            _cell_text(cells[i], fila.get(col, ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    return tabla

def _campo_linea_word(doc, etiqueta, valor):
    tabla = doc.add_table(rows=1, cols=2)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = tabla.rows[0].cells

    _cell_text(cells[0], etiqueta, bold=True)
    _cell_text(cells[1], _fecha_texto(valor), align=WD_ALIGN_PARAGRAPH.CENTER)

    for cell in cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(12)

    return tabla


def _generar_word(generales, fila, suspensiones, causas_suspension="", concepto_interventoria=""):
    doc = Document()
    _set_doc_defaults(doc)

    tipo = _tipo_acta(fila)
    numero = _numero_acta(fila)
    titulo = "ACTA DE SUSPENSIÓN O AMPLIACIÓN DE LA SUSPENSIÓN CONTRATO DE OBRA"

    _p(doc, titulo, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    _p(doc, f"ACTA DE {tipo} No. {numero}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

    _p(doc, "")
    _p(doc, "RESPONSABILIDAD", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(
        doc,
        "1. Mediante la suscripción de la presente acta de suspensión o ampliación de la suspensión, "
        "el Contratista y el Interventor asumen plena responsabilidad por la veracidad de la información "
        "en ella contenida, especialmente respecto a:"
    )
    _p(doc, "- Condiciones que justifiquen plenamente la necesidad de suspender o ampliar la suspensión del plazo contractual.")
    _p(doc, "- Período solicitado de la suspensión o ampliación de la suspensión.")
    _p(
        doc,
        "2. El Contratista de obra suspende las labores a partir de la fecha indicada, comprometiéndose "
        "a reanudar la ejecución del contrato en la fecha acordada, bajo la responsabilidad exclusiva del mismo."
    )
    _p(
        doc,
        "3. El Interventor del contrato como representante del CONTRATANTE debe hacer el análisis de las causas "
        "que conllevan a la suspensión del contrato de obra y será solidariamente responsable con el Contratista "
        "de obra, lo cual refrenda con su firma."
    )

    _p(doc, "")
    _p(doc, "CONDICIONES DEL CONTRATO DE OBRA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_simple(
        doc,
        [
            ("CONTRATO DE OBRA No.", generales.get("numero_contrato", "")),
            ("CONTRATANTE", generales.get("contratante", "")),
            ("NIT/C.C. CONTRATANTE", generales.get("nit_contratante", "")),
            ("CONTRATISTA", generales.get("contratista", "")),
            ("NIT/C.C. CONTRATISTA", generales.get("nit_contratista", "")),
            ("OBJETO DEL CONTRATO DE OBRA", generales.get("objeto", "")),
            ("VALOR ACUMULADO DEL CONTRATO", _moneda(generales.get("valor_contrato"))),
            ("FECHA DE INICIO DEL CONTRATO", _fecha_texto(generales.get("fecha_inicio"))),
            ("FECHA DE VENCIMIENTO ACTUAL", _fecha_texto(generales.get("fecha_vencimiento_inicial"))),
            ("PLAZO DE EJECUCIÓN", generales.get("plazo", "")),
            ("CONTRATO DE INTERVENTORÍA No.", generales.get("contrato_interventoria", "")),
            ("CONCEPTO DE LA INTERVENTORÍA", generales.get("concepto_interventoria", "")),
        ],
    )

    _p(doc, "")
    _p(doc, "RELACIÓN SUSPENSIONES Y AMPLIACIONES DE LA SUSPENSIÓN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_relacion(doc, _filas_relacion_suspensiones(suspensiones, fila.get("FECHA DEL ACTA")))

    _p(doc, "")
    _p(doc, "NUEVAS CONDICIONES DEL CONTRATO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "")
    _campo_linea_word(doc, "FECHA DE SUSPENSIÓN INICIAL:", fila.get("DESDE"))
    _p(doc, "")
    _campo_linea_word(doc, "FECHA DE REANUDACIÓN:", fila.get("HASTA"))
    _p(doc, "")
    _campo_linea_word(doc, "NUEVA FECHA DE VENCIMIENTO:", fila.get("NUEVA FECHA DE FINALIZACIÓN"))

    _p(doc, "")
    _p(doc, "CAUSAS QUE DAN ORIGEN A LA SUSPENSIÓN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, causas_suspension)

    _p(doc, "")
    _p(doc, "CONCEPTO DE LA INTERVENTORÍA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, concepto_interventoria)

    _p(doc, "")
    _p(doc, "RESPONSABILIDAD", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(
        doc,
        "1. Mediante la suscripción de la presente acta de suspensión o ampliación de la suspensión, el Contratista y el Interventor asumen plena responsabilidad por la veracidad de la información en ella contenida, especialmente respecto a:",
    )
    _p(doc, " - Condiciones que justifiquen plenamente la necesidad de suspender o ampliar la suspensión del plazo contractual.")
    _p(doc, " - Periodo solicitado de la suspensión o ampliación de la suspensión.")
    _p(
        doc,
        "2. El Contratista de obra suspende las labores a partir de la fecha indicada, comprometiéndose a reanudar la ejecución del contrato en la fecha acordada, bajo la responsabilidad exclusiva del mismo.",
    )
    _p(
        doc,
        "3. El contrato debe reanudarse en la fecha indicada en la presente acta sin que sea necesaria la elaboración y suscripción de un acta de reanudación.",
    )
    _p(
        doc,
        "4. En el evento en que las causales de suspensión cesen con anterioridad a la fecha de reanudación indicada en la presente Acta, se debe diligenciar el Acta de Reanudación del Contrato de Obra.",
    )

    _p(doc, "")
    _p(doc, "NOTAS:", bold=True)
    _p(
        doc,
        "1. Los motivos que dan origen a la presente suspensión y que imposibilitan de manera temporal "
        "el desarrollo del contrato no son imputables a las partes contratantes."
    )
    _p(
        doc,
        "2. Dentro de los tres (3) días hábiles siguientes a la fecha de reanudación, el CONTRATISTA "
        "se obliga a presentar para su aprobación los certificados de modificación de la garantía única "
        "y el seguro constituidos, en los cuales se tenga en cuenta la nueva fecha de vencimiento del plazo "
        "contractual. El incumplimiento de esta obligación acarrea el inicio de las acciones administrativas "
        "y judiciales a que haya lugar."
    )
    _p(
        doc,
        "3. El contrato debe reanudarse en la fecha indicada en la presente acta sin que sea necesaria "
        "la elaboración y suscripción de un acta de reanudación."
    )
    _p(
        doc,
        "4. En el evento en que las causales de suspensión cesen con anterioridad a la fecha de reanudación "
        "indicada en la presente Acta, se debe diligenciar el Acta de Reanudación del Contrato de Obra."
    )

    _p(doc, "")
    _p(doc, "Para constancia de lo anterior, firman quienes en ella intervinieron:")
    firmas = doc.add_table(rows=3, cols=3)
    firmas.style = "Table Grid"
    firmas.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cell_text(firmas.cell(0, 0), "CONTRATISTA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(0, 1), "INTERVENTOR", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(0, 2), "CONTRATANTE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(1, 0), "FIRMA DIGITAL A TRAVÉS DE LA PLATAFORMA DEL SECOP", align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(1, 1), "FIRMA DIGITAL A TRAVÉS DE LA PLATAFORMA DEL SECOP", align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(1, 2), "FIRMA DIGITAL A TRAVÉS DE LA PLATAFORMA DEL SECOP", align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(2, 0), generales.get("contratista", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(2, 1), generales.get("interventor", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firmas.cell(2, 2), generales.get("contratante", ""), align=WD_ALIGN_PARAGRAPH.CENTER)

    _p(doc, "")
    _p(doc, "Original: Archivo de Gestión.")
    _p(doc, "Copias: Unidad Ejecutora, Contratista, Interventor y entidad contratante.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ==========================================================
# Carga de estados
# ==========================================================
acta_inicio = _leer_estado("acta_inicio_obra")
contrato_obra = _leer_estado("contrato_obra")
control_obra = _leer_estado("control_obra")

generales = _datos_generales(acta_inicio, contrato_obra)
suspensiones = _suspensiones(control_obra)

st.markdown("# ACTA DE SUSPENSION")

if not suspensiones:
    st.warning("No hay registros en la tabla de suspensiones de la hoja 24.")
    st.stop()

opciones = {_etiqueta_suspension(fila, idx): idx for idx, fila in enumerate(suspensiones)}
seleccion = st.selectbox("Seleccione el acta de suspensión o ampliación", options=list(opciones.keys()))
fila = suspensiones[opciones[seleccion]]

st.markdown("### RESPONSABILIDAD")
st.write(
    "1. Mediante la suscripción de la presente acta de suspensión o ampliación de la suspensión, "
    "el Contratista y el Interventor asumen plena responsabilidad por la veracidad de la información "
    "en ella contenida, especialmente respecto a:"
)
st.write("- Condiciones que justifiquen plenamente la necesidad de suspender o ampliar la suspensión del plazo contractual.")
st.write("- Período solicitado de la suspensión o ampliación de la suspensión.")
st.write(
    "2. El Contratista de obra suspende las labores a partir de la fecha indicada, comprometiéndose "
    "a reanudar la ejecución del contrato en la fecha acordada, bajo la responsabilidad exclusiva del mismo."
)
st.write(
    "3. El Interventor del contrato como representante del CONTRATANTE debe hacer el análisis de las causas "
    "que conllevan a la suspensión del contrato de obra y será solidariamente responsable con el Contratista "
    "de obra, lo cual refrenda con su firma."
)

st.markdown("### CONDICIONES DEL CONTRATO DE OBRA")
col1, col2 = st.columns(2)
with col1:
    st.text_input("CONTRATO DE OBRA No.", value=generales["numero_contrato"], disabled=True)
    st.text_input("CONTRATANTE", value=generales["contratante"], disabled=True)
    st.text_input("CONTRATISTA", value=generales["contratista"], disabled=True)
    st.number_input("VALOR ACUMULADO DEL CONTRATO", value=generales["valor_contrato"], disabled=True, format="%.2f")
with col2:
    st.text_input("NIT/C.C. CONTRATANTE", value=generales["nit_contratante"], disabled=True)
    st.text_input("NIT/C.C. CONTRATISTA", value=generales["nit_contratista"], disabled=True)
    st.date_input("FECHA DE INICIO DEL CONTRATO", value=generales["fecha_inicio"] or date.today(), disabled=True, format="DD/MM/YYYY")
    st.date_input("FECHA DE VENCIMIENTO ACTUAL", value=generales["fecha_vencimiento_inicial"] or date.today(), disabled=True, format="DD/MM/YYYY")

st.text_area("OBJETO DEL CONTRATO DE OBRA", value=generales["objeto"], disabled=True, height=100)

st.markdown("### RELACIÓN SUSPENSIONES Y AMPLIACIONES DE LA SUSPENSIÓN")
st.dataframe(
    pd.DataFrame(_filas_relacion_suspensiones(suspensiones, fila.get("FECHA DEL ACTA"))),
    hide_index=True,
    width="stretch",
)

st.markdown("### NUEVAS CONDICIONES DEL CONTRATO")

c_nc1, c_nc2 = st.columns([1, 1])
with c_nc1:
    st.text_input(
        "FECHA DE SUSPENSIÓN INICIAL:",
        value=_fecha_texto(fila.get("DESDE")),
        disabled=True,
    )
with c_nc2:
    st.text_input(
        "FECHA DE REANUDACIÓN:",
        value=_fecha_texto(fila.get("HASTA")),
        disabled=True,
    )

st.text_input(
    "NUEVA FECHA DE VENCIMIENTO:",
    value=_fecha_texto(fila.get("NUEVA FECHA DE FINALIZACIÓN")),
    disabled=True,
)

st.markdown("### CAUSAS QUE DAN ORIGEN A LA SUSPENSIÓN")
causas_suspension = st.text_area(
    "Describa de manera clara y completa las causales que dan origen a la suspensión, justificadas por el Contratista y aprobadas por la interventoría, teniendo en cuenta que las mismas no sean imputables a las partes contratantes.",
    value="",
    height=180,
    key="acta_suspension_causas",
)

st.markdown("### CONCEPTO DE LA INTERVENTORÍA")
concepto_interventoria = st.text_area(
    "Concepto de la interventoría",
    value="",
    height=180,
    key="acta_suspension_concepto_interventoria",
)

st.markdown("### NOTAS:")

st.write(
    "1. Los motivos que dan origen a la presente suspensión y que imposibilitan de manera temporal "
    "el desarrollo del contrato no son imputables a las partes contratantes."
)
st.write(
    "2. Dentro de los tres (3) días hábiles siguientes a la fecha de reanudación, el CONTRATISTA "
    "se obliga a presentar para su aprobación los certificados de modificación de la garantía única "
    "y el seguro constituidos, en los cuales se tenga en cuenta la nueva fecha de vencimiento del plazo "
    "contractual. El incumplimiento de esta obligación acarrea el inicio de las acciones administrativas "
    "y judiciales a que haya lugar."
)
st.write(
    "3. El contrato debe reanudarse en la fecha indicada en la presente acta sin que sea necesaria "
    "la elaboración y suscripción de un acta de reanudación."
)
st.write(
    "4. En el evento en que las causales de suspensión cesen con anterioridad a la fecha de reanudación "
    "indicada en la presente Acta, se debe diligenciar el Acta de Reanudación del Contrato de Obra."
)

st.markdown("### FIRMAS")
df_firmas = pd.DataFrame(
    [
        {
            "CONTRATISTA": generales.get("contratista", ""),
            "INTERVENTOR": generales.get("interventor", ""),
            "CONTRATANTE": generales.get("contratante", ""),
        }
    ]
)
st.dataframe(df_firmas, hide_index=True, width="stretch")
word = _generar_word(generales, fila, suspensiones, causas_suspension, concepto_interventoria)
st.download_button(
    "📄 Descargar Word",
    data=word,
    file_name="acta_suspension.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)
