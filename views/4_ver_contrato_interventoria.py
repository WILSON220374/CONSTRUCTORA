import os
from datetime import datetime
from io import BytesIO

import streamlit as st
from docx import Document

from supabase_state import cargar_estado


def obtener_datos_contrato():
    datos = cargar_estado("contrato_interventoria") or {}
    return datos if isinstance(datos, dict) else {}


def texto_si_vacio(valor, pendiente="PENDIENTE"):
    if valor is None:
        return pendienteF
    if isinstance(valor, str):
        return valor.strip() if valor.strip() else pendiente
    return str(valor)


def escapar_tabla(valor):
    return texto_si_vacio(valor).replace("|", "\\|")

def _numero_a_letras_basico(n):
    unidades = {
        0: "",
        1: "uno",
        2: "dos",
        3: "tres",
        4: "cuatro",
        5: "cinco",
        6: "seis",
        7: "siete",
        8: "ocho",
        9: "nueve",
        10: "diez",
        11: "once",
        12: "doce",
        13: "trece",
        14: "catorce",
        15: "quince",
        16: "dieciseis",
        17: "diecisiete",
        18: "dieciocho",
        19: "diecinueve",
        20: "veinte",
        21: "veintiuno",
        22: "veintidos",
        23: "veintitres",
        24: "veinticuatro",
        25: "veinticinco",
        26: "veintiseis",
        27: "veintisiete",
        28: "veintiocho",
        29: "veintinueve",
    }

    decenas = {
        30: "treinta",
        40: "cuarenta",
        50: "cincuenta",
        60: "sesenta",
        70: "setenta",
        80: "ochenta",
        90: "noventa",
    }

    centenas = {
        100: "cien",
        200: "doscientos",
        300: "trescientos",
        400: "cuatrocientos",
        500: "quinientos",
        600: "seiscientos",
        700: "setecientos",
        800: "ochocientos",
        900: "novecientos",
    }

    if n in unidades:
        return unidades[n]

    if n < 100:
        d = (n // 10) * 10
        u = n % 10
        return decenas[d] if u == 0 else f"{decenas[d]} y {unidades[u]}"

    if n == 100:
        return "cien"

    if n < 200:
        return f"ciento {_numero_a_letras_basico(n - 100)}"

    if n < 1000:
        c = (n // 100) * 100
        r = n % 100
        return centenas[c] if r == 0 else f"{centenas[c]} {_numero_a_letras_basico(r)}"

    if n == 1000:
        return "mil"

    if n < 2000:
        return f"mil {_numero_a_letras_basico(n - 1000)}"

    if n < 1000000:
        m = n // 1000
        r = n % 1000
        base = f"{_numero_a_letras_basico(m)} mil"
        return base if r == 0 else f"{base} {_numero_a_letras_basico(r)}"

    return str(n)


def fecha_en_letras(valor):
    txt = texto_si_vacio(valor, "")
    if not txt:
        return ""

    meses = {
        1: "enero",
        2: "febrero",
        3: "marzo",
        4: "abril",
        5: "mayo",
        6: "junio",
        7: "julio",
        8: "agosto",
        9: "septiembre",
        10: "octubre",
        11: "noviembre",
        12: "diciembre",
    }

    formatos = ("%d-%m-%Y", "%d/%m/%Y", "%Y-%m-%d")
    for fmt in formatos:
        try:
            f = datetime.strptime(txt, fmt)
            dia_letras = _numero_a_letras_basico(f.day)
            anio_letras = _numero_a_letras_basico(f.year)
            return f"a los {dia_letras} dias del mes de {meses[f.month]} de {anio_letras}"
        except Exception:
            continue

    return txt


def construir_tabla_garantias_markdown(garantias):
    if not garantias or not isinstance(garantias, list):
        return (
            "| Amparo | Vigencia | Valor asegurado |\n"
            "|---|---|---|\n"
            "| PENDIENTE | PENDIENTE | PENDIENTE |"
        )

    filas = []
    for fila in garantias:
        if not isinstance(fila, dict):
            continue
        amparo = escapar_tabla(fila.get("amparo", ""))
        vigencia = escapar_tabla(fila.get("vigencia", ""))
        valor_asegurado = escapar_tabla(fila.get("valor_asegurado", ""))
        if amparo == "PENDIENTE" and vigencia == "PENDIENTE" and valor_asegurado == "PENDIENTE":
            continue
        filas.append(f"| {amparo} | {vigencia} | {valor_asegurado} |")

    if not filas:
        filas.append("| PENDIENTE | PENDIENTE | PENDIENTE |")

    encabezado = "| Amparo | Vigencia | Valor asegurado |\n|---|---|---|"
    return encabezado + "\n" + "\n".join(filas)


def construir_tabla_garantias_doc(doc, garantias):
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = "Table Grid"
    encabezado = tabla.rows[0].cells
    encabezado[0].text = "Amparo"
    encabezado[1].text = "Vigencia"
    encabezado[2].text = "Valor asegurado"

    filas_validas = []
    if garantias and isinstance(garantias, list):
        for fila in garantias:
            if not isinstance(fila, dict):
                continue
            amparo = texto_si_vacio(fila.get("amparo", ""))
            vigencia = texto_si_vacio(fila.get("vigencia", ""))
            valor_asegurado = texto_si_vacio(fila.get("valor_asegurado", ""))
            if amparo == "PENDIENTE" and vigencia == "PENDIENTE" and valor_asegurado == "PENDIENTE":
                continue
            filas_validas.append((amparo, vigencia, valor_asegurado))

    if not filas_validas:
        filas_validas.append(("PENDIENTE", "PENDIENTE", "PENDIENTE"))

    for amparo, vigencia, valor_asegurado in filas_validas:
        celdas = tabla.add_row().cells
        celdas[0].text = amparo
        celdas[1].text = vigencia
        celdas[2].text = valor_asegurado


def construir_bloque_documentos():
    return """1. Estudios y documentos previos.
2. Pliego de Condiciones, Adendas, Anexos, Formatos, Matrices, Formularios.
3. Propuesta presentada por el interventor.
4. Las Garantías debidamente aprobadas.
5. Toda la correspondencia que se surta entre las partes durante el término de ejecución del contrato."""


def construir_bloque_multas(multas):
    valores = {
        "Atraso o incumplimiento del Cronograma": "XXX",
        "No mantener en vigor las Garantías": "XXX",
        "No entrega la información completa que le solicite el supervisor": "XXX",
        "Atraso imputable al Interventor": "XX",
        "Por incumplir, sin justa causa, las órdenes que el supervisor dé": "XX",
        "Por cambiar el equipo de trabajo presentado en la oferta, sin la aprobación previa del supervisor": "XXX",
    }

    if multas and isinstance(multas, list):
        for fila in multas:
            if isinstance(fila, dict):
                causal = str(fila.get("causal", "")).strip()
                porcentaje = str(fila.get("porcentaje", "")).strip()
                if causal and porcentaje:
                    valores[causal] = porcentaje

    return f"""Causales: 
1. Por atraso o incumplimiento del Cronograma de Interventoría se causará una multa equivalente al [{valores["Atraso o incumplimiento del Cronograma"]}%] del valor del contrato, por cada día calendario de atraso.
2. Por no mantener en vigor, renovar, prorrogar, corregir o adicionar las Garantías, en los plazos y por los montos establecidos de acuerdo con el contrato o sus modificaciones, se causará una multa equivalente al [{valores["No mantener en vigor las Garantías"]}%] del contrato, por cada día calendario de atraso en el cumplimiento. 
3. Si el Interventor no entrega la información completa que le solicite el supervisor, que se relacione con el objeto del contrato o con el cumplimiento de las actividades del proyecto a ejecutar, dentro de los plazos y en los términos de cada requerimiento siempre y cuando sean razonables en función de la información exigida, se causará una multa equivalente al [{valores["No entrega la información completa que le solicite el supervisor"]}%] del contrato. Estas multas se causarán sucesivamente por cada día de atraso, hasta cuando el Interventor demuestre que corrija el incumplimiento respectivo a satisfacción del supervisor. 
4. Por atraso imputable al Interventor, se causará una multa diaria equivalente al [{valores["Atraso imputable al Interventor"]}%] del contrato, por cada día calendario de atraso. Igual sanción se aplicará en caso de que el Interventor no inicie efectivamente con la ejecución del contrato en la fecha acordada.
5. Por incumplir, sin justa causa, las órdenes que el supervisor dé en ejercicio de sus funciones y en el marco del ordenamiento jurídico, el Interventor se hará acreedor a una multa equivalente al [{valores["Por incumplir, sin justa causa, las órdenes que el supervisor dé"]}%] del contrato, por cada orden incumplida.
6. Por cambiar el equipo de trabajo presentado en la oferta, sin la aprobación previa del supervisor, al Interventor se le impondrá una multa equivalente al [{valores["Por cambiar el equipo de trabajo presentado en la oferta, sin la aprobación previa del supervisor"]}%] del contrato.

Parágrafo 1. Las multas son apremios al Interventor para el cumplimiento de sus obligaciones y, por lo tanto, no tienen el carácter de estimación anticipada de perjuicios, de manera que pueden acumularse con cualquier forma de indemnización, en los términos previstos en el artículo 1600 del Código Civil. 

Parágrafo 2. En caso de que el Interventor incurra en una de las causales de multa, este autoriza a la Entidad para descontar el valor de la misma, la cual se tomará directamente de cualquier suma que se le adeude, sin perjuicio de hacer efectiva la Garantía de cumplimiento del contrato.

Parágrafo 3. El pago en cualquier forma, incluyendo la deducción de los valores adeudados al Interventor, realizado con fundamento en las multas impuestas, no lo exonerará de continuar con la ejecución del contrato ni de las demás responsabilidades y obligaciones que emanen del mismo, amén de la obligación incumplida.

Parágrafo 4. En caso de que el Interventor reincida en el incumplimiento de una o de varias obligaciones se podrán imponer nuevas multas.

Parágrafo 5. Para efectos de la imposición de las multas el salario mínimo diario o mensual vigente, será aquel que rija para el momento de la expedición del acto administrativo que lo declara.

Parágrafo 6. El monto de ninguna de las sanciones asociadas a cada causal de multa, aplicada de forma independiente, podrá ser superior al cinco por ciento (5 %) del valor del contrato, particularmente frente a aquellas que se imponen de forma sucesiva. Lo anterior, sin perjuicio de que se inicie un nuevo procedimiento sancionatorio para efectos de imponer nuevas multas."""


def construir_contrato(datos):
    nombre_entidad = texto_si_vacio(datos.get("nombre_entidad"))
    nombre_representante_entidad = texto_si_vacio(datos.get("nombre_representante_entidad"))
    nombre_empresa_interventora = texto_si_vacio(datos.get("nombre_empresa_interventora"))
    nombre_interventor = texto_si_vacio(datos.get("nombre_interventor"))
    numero_proceso_contratacion = texto_si_vacio(datos.get("numero_proceso_contratacion"))

    objeto_general = texto_si_vacio(datos.get("objeto_general"))
    alcance_objeto = texto_si_vacio(datos.get("alcance_objeto"))

    plazo_contrato = texto_si_vacio(datos.get("plazo_contrato"))

    valor_contrato_numeros = texto_si_vacio(datos.get("valor_contrato_numeros"))
    valor_contrato_letras = texto_si_vacio(datos.get("valor_contrato_letras"))
    numero_smmlv = texto_si_vacio(datos.get("numero_smmlv"))
    anio_suscripcion = texto_si_vacio(datos.get("anio_suscripcion"))
    dias_habiles_pago = texto_si_vacio(datos.get("dias_habiles_pago"))

    obligaciones_especificas_interventor = texto_si_vacio(datos.get("obligaciones_especificas_interventor"))
    clausula_penal_porcentaje_valor = texto_si_vacio(datos.get("clausula_penal_porcentaje_valor"))

    dias_presentacion_garantia = texto_si_vacio(datos.get("dias_presentacion_garantia"))
    lugar_ejecucion = texto_si_vacio(datos.get("lugar_ejecucion"))
    lugar_perfeccionamiento = texto_si_vacio(datos.get("lugar_perfeccionamiento"))
    fecha_suscripcion = fecha_en_letras(datos.get("fecha_suscripcion"))
    termino_liquidacion = texto_si_vacio(datos.get("termino_liquidacion"))

    firmante_entidad = texto_si_vacio(datos.get("firmante_entidad"))
    firmante_interventor = texto_si_vacio(datos.get("firmante_interventor"))

    tabla_garantias = construir_tabla_garantias_markdown(datos.get("garantias_interventoria", []))
    documentos_txt = construir_bloque_documentos()
    bloque_multas = construir_bloque_multas(datos.get("multas_interventoria", []))

    contrato = f"""# CONTRATO DE INTERVENTORÍA

Entre {nombre_entidad} (en adelante la “Entidad”) por medio de su representante legal {nombre_representante_entidad}, por una parte; y por la otra {nombre_interventor} en representación de {nombre_empresa_interventora} (en adelante el “Interventor”), hemos convenido celebrar el presente Contrato de Interventoría, previas las siguientes consideraciones:

Con base en las anteriores consideraciones, la Entidad y el Interventor (individualmente la “Parte”, conjuntamente las “Partes”), convienen las siguientes cláusulas.

## OBJETO

El objeto del contrato es {objeto_general}.

## ALCANCE DEL OBJETO

El Interventor debe ejecutar el contrato de conformidad con las especificaciones y características técnicas señaladas en los Documentos del Proceso de Contratación No. {numero_proceso_contratacion}, los cuales hacen parte integral de este.

El Interventor se obliga con la Entidad a ejecutar, a los precios cotizados en la propuesta y con sus propios medios –materiales, maquinaria, laboratorios, equipos y personal– en forma independiente y con plena autonomía técnica y administrativa, hasta su total terminación y aceptación final, las actividades propias de interventoría según lo establece la legislación vigente, el Pliego de Condiciones, el Anexo Técnico y el contrato al cual se ejercerá la interventoría.

{alcance_objeto}

El Interventor y la Entidad asumen de forma obligatoria los Riesgos previsibles identificados y plasmados en el Pliego de Condiciones.

## PLAZO DEL CONTRATO

El plazo estimado para la ejecución del contrato será de {plazo_contrato}, contados a partir del cumplimiento de los requisitos de perfeccionamiento y ejecución del mismo y aprobación de los documentos previstos en el Pliego de Condiciones.

Excepcionalmente, por causas que constituyan fuerza mayor o caso fortuito, las partes de común acuerdo podrán suspender el plazo de ejecución del contrato, siempre que la naturaleza de las obligaciones lo admita, no se contraríen normas de orden público y se adopten medidas para superar las causas que motivaron la suspensión en el menor tiempo posible.

## VALOR DEL CONTRATO

El valor del contrato es hasta por la suma de {valor_contrato_letras} ({valor_contrato_numeros}), equivalentes a {numero_smmlv} SMMLV para el año de suscripción del contrato {anio_suscripcion}.

El Interventor con la suscripción del contrato acepta que en el evento en que el valor total a pagar tenga centavos, estos se ajusten o aproximen al Peso, ya sea por exceso o por defecto, si la suma es mayor o menor a cincuenta (50) centavos. Lo anterior, sin que sobrepase el valor total establecido en el contrato.

## FORMA DE PAGO

La Entidad pagará al Interventor el valor del contrato en pagos parciales mensuales en Pesos Colombianos, de acuerdo con la ejecución del contrato hasta el noventa y cinco por ciento (95 %) de su monto. El cinco por ciento (5 %) restante se pagará una vez finalizada la ejecución del contrato.

El pago al Interventor se efectuará dentro de los {dias_habiles_pago} días hábiles siguientes a la presentación de la factura, el acta de recibo suscrita por el supervisor designado para el recibo a satisfacción de las actividades y de la certificación de encontrarse al día con los aportes al Sistema de la Seguridad Social y Parafiscales.

La Entidad no se hace responsable por las demoras presentadas en el trámite para el pago al Interventor cuando ellas fueren ocasionadas por encontrarse incompleta la documentación de soporte o no ajustarse a cualquiera de las condiciones establecidas en el contrato.

La Entidad hará las retenciones a que haya lugar sobre cada pago, de acuerdo con las disposiciones legales vigentes sobre la materia.

El Interventor deberá acreditar para cada pago derivado del contrato, que se encuentra al día en el pago de aportes parafiscales relativos al Sistema de Seguridad Social Integral, así como los propios del Sena, ICBF y Cajas de Compensación Familiar, cuando corresponda.

## OBLIGACIONES GENERALES DEL INTERVENTOR

Además de las derivadas de la esencia y naturaleza del contrato, la ley, las obligaciones y condiciones señaladas en el Pliego de Condiciones y demás Documentos del Proceso, vigente durante la ejecución del contrato, el Interventor se obliga a:

- Dar cumplimiento al objeto y alcance del contrato de acuerdo con lo establecido en el presente documento y en sus anexos.
- Estar en comunicación con el supervisor del contrato.
- Permitir la labor de seguimiento y control que realiza el supervisor, atendiendo y dando respuesta oportuna a las observaciones o requerimientos que se realicen.
- Disponer del personal idóneo, así como de los recursos logísticos, materiales y/o equipos, para desarrollar el contrato dentro de la oportunidad y con la calidad establecidos.
- Acreditar el cumplimiento de la formación académica y la experiencia del equipo de trabajo definidos en el documento base y en el anexo técnico en los plazos acordados con la Entidad.
- Identificar las oportunidades para promover el empleo local durante la ejecución del contrato.
- Aportar todo su conocimiento y experiencia para desarrollar adecuadamente el objeto del contrato de conformidad con lo requerido por el contratante.
- Cumplir con las normas de gestión ambiental, así como con las normas del Sistema de Seguridad y Salud en el Trabajo que rijan durante la vigencia del contrato.
- Realizar todos los pagos de honorarios y/o salarios, parafiscales e indemnizaciones a que haya lugar.
- Manejar con la debida confidencialidad la información a que tenga acceso, así como la producida a lo largo de la ejecución del contrato.
- Reportar la información relacionada con la ejecución del contrato o que tenga incidencia en ella cuando sea requerida por la Entidad.
- Cumplir los protocolos de bioseguridad y demás exigencias aplicables.

## OBLIGACIONES ESPECÍFICAS DEL INTERVENTOR

{obligaciones_especificas_interventor}

## DERECHOS DEL INTERVENTOR

El Interventor tiene derecho a recibir una remuneración por la ejecución del Contrato de Interventoría en los términos pactados en la cláusula de valor y forma de pago del presente contrato.

## OBLIGACIONES DE LA ENTIDAD

La Entidad está obligada a:

- Cumplir con las condiciones establecidas en los Documentos del Proceso de Contratación.
- Pagar la remuneración por la ejecución de la interventoría en los términos pactados en la cláusula de valor y forma de pago del presente contrato.

## RESPONSABILIDAD

El Interventor es responsable de cumplir las obligaciones pactadas en el contrato. Además, responderá por los daños generados a la Entidad en la ejecución del contrato, causados por sus contratistas o empleados y sus subcontratistas.

## MULTAS

{bloque_multas}

## CLÁUSULA PENAL

Las Partes acuerdan que la aplicación de la cláusula penal no exime el cumplimiento de las obligaciones contractuales y podrá exigirse al Interventor la pena y la indemnización de perjuicios.

En caso de presentarse por parte del Interventor incumplimiento parcial o total del Contrato, o por incurrir en mora o retardo en el cumplimiento de sus obligaciones, este pagará a título de cláusula penal pecuniaria a la Entidad una suma equivalente a {clausula_penal_porcentaje_valor}. La imposición de esta pena pecuniaria se considerará como una estimación anticipada de perjuicios que el Interventor cause a la Entidad.

El pago o deducción de la cláusula penal no exonerará al Interventor del cumplimiento de sus obligaciones contractuales, incluyendo las que dieron lugar a la imposición de la pena.

## GARANTÍAS

### GARANTÍA DE CUMPLIMIENTO

Para cubrir cualquier hecho constitutivo de incumplimiento, el Interventor deberá presentar la Garantía de cumplimiento en original a la Entidad dentro de los {dias_presentacion_garantia} contados a partir de la firma del contrato y requerirá de su aprobación.

La garantía tendrá a {nombre_entidad} como asegurado o beneficiario.

La información de amparos, vigencia y valor asegurado corresponde a la siguiente tabla:

{tabla_garantias}

El Interventor está obligado a restablecer el valor de la Garantía cuando esta se vea reducida por razón de las reclamaciones que efectúe la Entidad, así como a ampliar las Garantías en los eventos de adición, suspensión y/o prórroga del contrato.

### DEL AMPARO DE CALIDAD DEL SERVICIO EN LA GARANTÍA ÚNICA DE CUMPLIMIENTO

En relación con el amparo de calidad del servicio de la Garantía única de cumplimiento, se tendrá en cuenta que el Interventor será responsable por los perjuicios causados a la Entidad contratante que se produzcan con posterioridad a la terminación del contrato y que se compruebe tienen su causa en la mala calidad de los productos entregados o de los servicios prestados imputables al Interventor.

## INDEPENDENCIA DEL INTERVENTOR

El Interventor es independiente de la Entidad y, en consecuencia, no es su representante, agente o mandatario. El Interventor no tiene la facultad de hacer declaraciones, representaciones o compromisos en nombre de la Entidad, ni de tomar decisiones o iniciar acciones que generen obligaciones a su cargo.

## INEXISTENCIA DE RELACIÓN LABORAL ENTRE LA ENTIDAD Y EL INTERVENTOR

El Interventor ejecutará el contrato con sus propios medios y con plena autonomía técnica y administrativa y el personal que vincule durante la ejecución del contrato será de su libre escogencia. Entre el Interventor, el equipo de trabajo que éste contrate y la Entidad no existe, ni existirá vínculo laboral alguno.

El Interventor solo podrá subcontratar, con la autorización previa y expresa de la Entidad, cuando dicha subcontratación implique modificaciones al equipo de trabajo inicialmente ofertado y con el que inició la ejecución del contrato.

## CESIÓN

El Interventor no podrá ceder los derechos y obligaciones emanados del contrato sin el consentimiento previo y expreso de la Entidad.

## LIQUIDACIÓN

El contrato será objeto de liquidación de acuerdo con lo establecido en las normas que regulan la materia. El término para la liquidación del contrato será de {termino_liquidacion}.

Para la liquidación se exigirá al Interventor la ampliación de la Garantía, si es del caso, a fin de avalar las obligaciones que éste deba cumplir con posterioridad a la terminación del contrato.

Si el Interventor no se presenta para efectos de la liquidación del contrato durante el término indicado, previa notificación o convocatoria que haga la Entidad para la liquidación bilateral, o las partes no llegan a ningún acuerdo o se logra parcialmente, la Entidad procederá a su liquidación por medio de resolución motivada susceptible del recurso de reposición.

## SUSCRIPCIÓN, PERFECCIONAMIENTO Y EJECUCIÓN

El contrato quedará perfeccionado con la firma de la minuta por las Partes. Para la ejecución se requerirá que el Interventor se encuentre al día en los pagos al Sistema de Seguridad Social Integral y demás aportes parafiscales correspondientes, la aprobación de las Garantías y el Registro Presupuestal correspondiente.

## LUGAR DE EJECUCIÓN Y DOMICILIO CONTRACTUAL

Las actividades previstas en el contrato se deben desarrollar en {lugar_ejecucion} y el domicilio contractual será el previsto en la minuta tipo de interventoría.

## DOCUMENTOS

Los documentos que a continuación se relacionan hacen parte integral del contrato y, en consecuencia, producen sus mismos efectos y obligaciones jurídicas y contractuales:

{documentos_txt}

En constancia se firma el presente contrato en {lugar_perfeccionamiento}, el {fecha_suscripcion}.
"""
    return contrato


def generar_word(datos):
    nombre_entidad = texto_si_vacio(datos.get("nombre_entidad"))
    nombre_representante_entidad = texto_si_vacio(datos.get("nombre_representante_entidad"))
    nombre_empresa_interventora = texto_si_vacio(datos.get("nombre_empresa_interventora"))
    nombre_interventor = texto_si_vacio(datos.get("nombre_interventor"))
    contrato_txt = construir_contrato(datos)

    doc = Document()
    doc.add_heading("Contrato de interventoría", level=0)

    for bloque in contrato_txt.split("\n\n"):
        if bloque.strip().startswith("# "):
            doc.add_heading(bloque.replace("# ", "").strip(), level=1)
        elif bloque.strip().startswith("## "):
            doc.add_heading(bloque.replace("## ", "").strip(), level=2)
        elif "### GARANTÍA DE CUMPLIMIENTO" in bloque:
            doc.add_heading("GARANTÍA DE CUMPLIMIENTO", level=3)
        elif "### DEL AMPARO DE CALIDAD DEL SERVICIO EN LA GARANTÍA ÚNICA DE CUMPLIMIENTO" in bloque:
            doc.add_heading("DEL AMPARO DE CALIDAD DEL SERVICIO EN LA GARANTÍA ÚNICA DE CUMPLIMIENTO", level=3)
        elif "| Amparo | Vigencia | Valor asegurado |" in bloque:
            construir_tabla_garantias_doc(doc, datos.get("garantias_interventoria", []))
        else:
            doc.add_paragraph(bloque.strip())

    doc.add_paragraph("")
    tabla_firmas = doc.add_table(rows=1, cols=2)
    tabla_firmas.style = "Table Grid"

    izquierda = tabla_firmas.rows[0].cells[0]
    derecha = tabla_firmas.rows[0].cells[1]

    izquierda.text = (
        "Por la Entidad,\n\n"
        f"{texto_si_vacio(datos.get('firmante_entidad'))}"
    )
    derecha.text = (
        "Por el Interventor,\n\n"
        f"{texto_si_vacio(datos.get('firmante_interventor'))}"
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    nombre_archivo = f"contrato_interventoria_{nombre_entidad}_{nombre_interventor}".replace(" ", "_")
    return buffer, f"{nombre_archivo}.docx"


datos = obtener_datos_contrato()

st.markdown("""
    <style>
    .titulo-seccion { font-size: 32px !important; font-weight: 800 !important; color: #7A0019; }
    .subtitulo-gris { font-size: 16px !important; color: #666; margin-bottom: 15px; }
    div[data-testid="stProgress"] > div > div > div > div { background-color: #C62828 !important; }
    section[data-testid="stSidebar"] { background-color: #f4f4f4; }
    .stButton > button { width: 100%; border-radius: 6px; height: 3em; font-weight: bold; }
    button[kind="primary"] {
        background-color: #7A0019 !important;
        border-color: #7A0019 !important;
        color: white !important;
    }
    button[kind="primary"]:hover {
        background-color: #5C0013 !important;
        border-color: #5C0013 !important;
        color: white !important;
    }
    </style>
""", unsafe_allow_html=True)

col_t, col_l = st.columns([4, 1], vertical_alignment="center")
with col_t:
    st.markdown('<div class="titulo-seccion">📄 Contrato de interventoría</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="subtitulo-gris">Vista de lectura del contrato de interventoría construido automáticamente con la información diligenciada.</div>',
        unsafe_allow_html=True
    )
with col_l:
    if os.path.exists("unnamed.jpg"):
        st.image("unnamed.jpg", use_container_width=True)

st.divider()

with st.sidebar:
    st.header("🧭 Resumen")
    st.markdown(f"**Entidad:** {texto_si_vacio(datos.get('nombre_entidad'))}")
    st.markdown(f"**Representante entidad:** {texto_si_vacio(datos.get('nombre_representante_entidad'))}")
    st.markdown(f"**Interventor:** {texto_si_vacio(datos.get('nombre_interventor'))}")
    st.markdown(f"**Empresa interventora:** {texto_si_vacio(datos.get('nombre_empresa_interventora'))}")
    st.markdown(f"**Proceso:** {texto_si_vacio(datos.get('numero_proceso_contratacion'))}")
    st.markdown(f"**Valor:** {texto_si_vacio(datos.get('valor_contrato_numeros'))}")
    st.markdown(f"**Plazo:** {texto_si_vacio(datos.get('plazo_contrato'))}")
    st.markdown(f"**Fecha:** {texto_si_vacio(datos.get('fecha_suscripcion'))}")

contrato_armado = construir_contrato(datos)
st.markdown(contrato_armado)

st.markdown("---")
st.markdown("### Firmas")

col1, col2 = st.columns(2)
with col1:
    st.markdown(
        f"**Por la Entidad**  \n"
        f"{texto_si_vacio(datos.get('firmante_entidad'))}"
    )

with col2:
    st.markdown(
        f"**Por el Interventor**  \n"
        f"{texto_si_vacio(datos.get('firmante_interventor'))}"
    )

word_buffer, nombre_word = generar_word(datos)
st.download_button(
    "Descargar contrato de interventoría en Word",
    data=word_buffer,
    file_name=nombre_word,
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    type="primary",
    key="descargar_contrato_interventoria_word",
)
