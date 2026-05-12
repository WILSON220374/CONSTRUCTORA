from io import BytesIO
from datetime import date, datetime, timedelta

import pandas as pd
import plotly.express as px
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt

from supabase_state import cargar_estado
from supabase_state import guardar_estado as guardar_estado_bd


CLAVE_GUARDADO = "informe_interventoria"


# ==========================================================
# Persistencia
# ==========================================================
def guardar_estado(clave, datos):
    def serializar(obj):
        if isinstance(obj, dict):
            return {k: serializar(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [serializar(x) for x in obj]
        if isinstance(obj, (date, datetime)):
            return obj.isoformat()
        return obj

    guardar_estado_bd(clave, serializar(datos))


def _leer_estado(clave):
    datos = cargar_estado(clave) or {}
    return datos if isinstance(datos, dict) else {}


# ==========================================================
# Helpers generales
# ==========================================================
def _texto(valor):
    if valor is None:
        return ""
    return str(valor).strip()


def _safe_float(valor, default=0.0):
    try:
        if valor is None or valor == "":
            return float(default)

        if isinstance(valor, (int, float)):
            return float(valor)

        txt = str(valor).strip().replace("$", "").replace(" ", "")

        if "," in txt and "." in txt:
            txt = txt.replace(".", "").replace(",", ".")
        elif "," in txt:
            txt = txt.replace(",", ".")

        return float(txt)
    except Exception:
        return float(default)


def _parse_fecha(valor):
    if isinstance(valor, date):
        return valor
    if isinstance(valor, datetime):
        return valor.date()
    if isinstance(valor, str) and valor.strip():
        txt = valor.strip()
        try:
            return datetime.fromisoformat(txt).date()
        except Exception:
            pass
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
            try:
                return datetime.strptime(txt, fmt).date()
            except Exception:
                continue
    return None


def _fecha_input(valor):
    return _parse_fecha(valor) or date.today()


def _fecha_texto(valor):
    fecha = _parse_fecha(valor)
    return fecha.strftime("%d/%m/%Y") if fecha else ""


def _moneda(valor):
    return f"$ {_safe_float(valor, 0.0):,.2f}"


def _primero_no_vacio(*valores):
    for valor in valores:
        txt = _texto(valor)
        if txt:
            return txt
    return ""


def _sumar_valores(rows, campo):
    total = 0.0
    for fila in rows or []:
        if isinstance(fila, dict):
            total += _safe_float(fila.get(campo), 0.0)
    return round(total, 2)


def _filtrar_df_hasta_fecha(df, columna_fecha, fecha_corte):
    if df is None or df.empty or columna_fecha not in df.columns:
        return df
    fecha_corte = _parse_fecha(fecha_corte)
    if not fecha_corte:
        return df
    out = df.copy()
    out[columna_fecha] = pd.to_datetime(out[columna_fecha], errors="coerce").dt.date
    return out[out[columna_fecha].notna() & (out[columna_fecha] <= fecha_corte)]


def _fechas_disponibles_df(df, columna_fecha):
    if df is None or df.empty or columna_fecha not in df.columns:
        return []
    fechas = pd.to_datetime(df[columna_fecha], errors="coerce").dt.date.dropna().unique().tolist()
    return sorted(fechas)


# ==========================================================
# Datos base desde otras hojas
# ==========================================================
def _valor_contrato_obra(acta_inicio, contrato_obra):
    for valor in [
        contrato_obra.get("valor_total_numeros"),
        contrato_obra.get("valor_contrato"),
        contrato_obra.get("valor"),
    ]:
        numero = _safe_float(valor, 0.0)
        if numero > 0:
            return numero
    return 0.0


def _fecha_inicio(acta_inicio):
    return _parse_fecha(
        _primero_no_vacio(
            acta_inicio.get("fecha_inicio"),
            acta_inicio.get("fecha_presente_acta"),
        )
    )


def _fecha_terminacion_acta(acta_inicio):
    return _parse_fecha(
        _primero_no_vacio(
            acta_inicio.get("fecha_terminacion"),
            acta_inicio.get("fecha_terminacion_contrato"),
        )
    )


def _datos_generales(acta_inicio, contrato_obra, contrato_interventoria, control_obra):
    valor_obra = _valor_contrato_obra(acta_inicio, contrato_obra)
    valor_adiciones = _sumar_valores(control_obra.get("adiciones_rows", []), "VALOR")
    return {
        "contrato_obra_no": _primero_no_vacio(acta_inicio.get("numero_contrato"), contrato_obra.get("numero_contrato")),
        "contrato_interventoria_no": _primero_no_vacio(
            contrato_interventoria.get("numero_proceso_contratacion"),
            acta_inicio.get("numero_contrato_interventoria"),
            contrato_obra.get("numero_contrato_interventoria"),
        ),
        "objeto_obra": _primero_no_vacio(
            acta_inicio.get("objeto_contrato"),
            contrato_obra.get("objeto_general"),
            contrato_obra.get("objeto_contrato"),
            contrato_obra.get("objeto"),
        ),
        "objeto_interventoria": _primero_no_vacio(
            contrato_interventoria.get("objeto_general"),
            contrato_interventoria.get("alcance_objeto"),
        ),
        "valor_obra": valor_obra,
        "valor_interventoria": _safe_float(contrato_interventoria.get("valor_contrato_numeros"), 0.0),
        "plazo_ejecucion": _primero_no_vacio(acta_inicio.get("plazo_ejecucion"), contrato_obra.get("plazo_ejecucion")),
        "fecha_inicio": _fecha_inicio(acta_inicio),
        "fecha_terminacion_acta": _fecha_terminacion_acta(acta_inicio),
        "valor_inicial_obra": valor_obra,
        "valor_adiciones": valor_adiciones,
        "valor_actual_obra": round(valor_obra + valor_adiciones, 2),
        "interventor": _primero_no_vacio(
            acta_inicio.get("nombre_firma_interventor"),
            contrato_obra.get("nombre_interventor"),
            contrato_obra.get("nombre_supervisor"),
        ),
    }


# ==========================================================
# Tablas desde hojas fuente
# ==========================================================
def _df_suspensiones(control_obra):
    filas = []
    for fila in control_obra.get("suspensiones_rows", []) or []:
        if not isinstance(fila, dict):
            continue
        filas.append(
            {
                "ACTA DE SUSPENSIÓN No.": _texto(fila.get("ACTA DE SUSPENSIÓN No.")),
                "ACTA DE AMPLIACIÓN SUSPENSIÓN No.": _texto(fila.get("ACTA DE AMPLIACIÓN SUSPENSIÓN No.")),
                "FECHA DEL ACTA": _parse_fecha(fila.get("FECHA DEL ACTA")),
                "DESDE": _parse_fecha(fila.get("DESDE")),
                "HASTA": _parse_fecha(fila.get("HASTA")),
                "PERIODO DE SUSPENSIÓN": _safe_float(fila.get("PERIODO DE SUSPENSIÓN"), 0.0),
                "NUEVA FECHA DE FINALIZACIÓN": _parse_fecha(fila.get("NUEVA FECHA DE FINALIZACIÓN")),
            }
        )
    return pd.DataFrame(filas)


def _df_adiciones(control_obra):
    filas = []
    for fila in control_obra.get("adiciones_rows", []) or []:
        if not isinstance(fila, dict):
            continue
        filas.append(
            {
                "ADICIONAL No.": _texto(fila.get("ADICIONAL No.")),
                "FECHA": _parse_fecha(fila.get("FECHA")),
                "VALOR": _safe_float(fila.get("VALOR"), 0.0),
                "SMMLV DEL AÑO DE LA ADICIÓN": _safe_float(fila.get("SMMLV DEL AÑO DE LA ADICIÓN"), 0.0),
                "ADICIÓN EN SALARIOS MÍNIMOS": _safe_float(fila.get("ADICIÓN EN SALARIOS MÍNIMOS"), 0.0),
                "VALOR ACUMULADO DEL CONTRATO": _safe_float(fila.get("VALOR ACUMULADO DEL CONTRATO"), 0.0),
            }
        )
    return pd.DataFrame(filas)


def _df_modificaciones(modificacion_prorroga):
    filas = []
    estado = modificacion_prorroga if isinstance(modificacion_prorroga, dict) else {}
    solicitudes = estado.get("solicitudes", []) if isinstance(estado.get("solicitudes"), list) else []

    if solicitudes:
        for solicitud in solicitudes:
            consecutivo = int(solicitud.get("consecutivo") or 0) if isinstance(solicitud, dict) else 0
            for fila in solicitud.get("modificaciones_rows", []) or []:
                if not isinstance(fila, dict):
                    continue
                filas.append(
                    {
                        "SOLICITUD No.": consecutivo,
                        "MODIFICACIÓN No.": _texto(fila.get("Modificación No.")),
                        "FECHA": _parse_fecha(fila.get("Fecha")),
                        "BREVE DESCRIPCIÓN": _texto(fila.get("Breve descripción")),
                    }
                )
    else:
        for fila in estado.get("modificaciones_rows", []) or []:
            if not isinstance(fila, dict):
                continue
            filas.append(
                {
                    "SOLICITUD No.": "",
                    "MODIFICACIÓN No.": _texto(fila.get("Modificación No.")),
                    "FECHA": _parse_fecha(fila.get("Fecha")),
                    "BREVE DESCRIPCIÓN": _texto(fila.get("Breve descripción")),
                }
            )

    return pd.DataFrame(filas)


def _df_garantias_iniciales(contrato_obra):
    filas = []
    for fila in contrato_obra.get("garantias", []) or []:
        if not isinstance(fila, dict):
            continue
        filas.append(
            {
                "TIPO": "CONDICIONES INICIALES",
                "AMPARO": _texto(fila.get("amparo")),
                "SUFICIENCIA": _texto(fila.get("suficiencia")),
                "%": _safe_float(fila.get("%"), 0.0),
                "COBERTURA": _safe_float(fila.get("cobertura"), 0.0),
                "DESDE": _parse_fecha(fila.get("desde")),
                "HASTA": _parse_fecha(fila.get("hasta")),
            }
        )
    return pd.DataFrame(filas)


def _df_garantias_modificadas(control_obra):
    filas = []
    for bloque in control_obra.get("garantias_modificaciones_bloques", []) or []:
        if not isinstance(bloque, dict):
            continue

        numero = int(bloque.get("numero") or 0)

        for fila in bloque.get("rows", []) or []:
            if not isinstance(fila, dict):
                continue

            filas.append(
                {
                    "TIPO": f"MODIFICACIÓN DE GARANTÍAS No. {numero}",
                    "AMPARO": _texto(_primero_no_vacio(fila.get("AMPARO"), fila.get("amparo"))),
                    "SUFICIENCIA": _texto(_primero_no_vacio(fila.get("SUFICIENCIA"), fila.get("suficiencia"))),
                    "%": _safe_float(_primero_no_vacio(fila.get("%"), fila.get("porcentaje")), 0.0),
                    "COBERTURA": _safe_float(_primero_no_vacio(fila.get("COBERTURA"), fila.get("cobertura")), 0.0),
                    "DESDE": _parse_fecha(_primero_no_vacio(fila.get("DESDE"), fila.get("desde"))),
                    "HASTA": _parse_fecha(_primero_no_vacio(fila.get("HASTA"), fila.get("hasta"))),
                }
            )

    return pd.DataFrame(filas)


def _df_pagos(control_obra):
    filas = []
    for i, fila in enumerate(control_obra.get("pagos_rows", []) or [], start=1):
        if not isinstance(fila, dict):
            continue
        filas.append(
            {
                "ACTA PARCIAL No.": i,
                "FECHA": _parse_fecha(fila.get("FECHA")),
                "VALOR": _safe_float(fila.get("VALOR FACTURADO"), 0.0),
                "PENDIENTE POR FACTURAR": _safe_float(fila.get("PENDIENTE POR FACTURAR"), 0.0),
            }
        )
    return pd.DataFrame(filas)


def _valor_anticipo(control_obra):
    filas = control_obra.get("anticipo_rows", []) or []
    if isinstance(filas, list) and filas:
        return _safe_float(filas[0].get("VALOR INICIAL"), 0.0) if isinstance(filas[0], dict) else 0.0
    return 0.0


# ==========================================================
# Seguimiento físico e indicadores
# ==========================================================
def _fechas_corte_seguimiento(seguimiento_fisico):
    seguimientos = seguimiento_fisico.get("seguimientos_fisicos", {}) or {}
    if not isinstance(seguimientos, dict):
        return []
    fechas = []
    for clave, corte in seguimientos.items():
        if isinstance(corte, dict):
            fecha = _parse_fecha(corte.get("fecha_corte", clave))
            if fecha:
                fechas.append(fecha)
    return sorted(set(fechas))


def _corte_seguimiento(seguimiento_fisico, fecha_corte):
    seguimientos = seguimiento_fisico.get("seguimientos_fisicos", {}) or {}
    fecha_corte = _parse_fecha(fecha_corte)
    if not isinstance(seguimientos, dict) or not fecha_corte:
        return {}
    for clave, corte in seguimientos.items():
        if not isinstance(corte, dict):
            continue
        fecha = _parse_fecha(corte.get("fecha_corte", clave))
        if fecha == fecha_corte:
            return corte
    return {}


def _avance_general_desde_corte(corte, fecha_corte):
    avance_general = corte.get("avance_general", []) or []
    fila = avance_general[0] if isinstance(avance_general, list) and avance_general and isinstance(avance_general[0], dict) else {}
    return {
        "FECHA DE CORTE": _fecha_texto(fecha_corte),
        "% PROGRAMADO": _safe_float(fila.get("% PROGRAMADO"), 0.0),
        "% EJECUTADO": _safe_float(fila.get("% EJECUTADO"), 0.0),
        "$ PROGRAMADO": _safe_float(fila.get("$ PROGRAMADO"), 0.0),
        "$ EJECUTADO": _safe_float(fila.get("$ EJECUTADO"), 0.0),
    }


def _indicadores_desde_corte(corte):
    indices = corte.get("indices_valor_ganado", []) or corte.get("df_indices_valor_ganado", []) or []
    salida = {"CPI": "", "SPI": "", "RETRASO": ""}
    if isinstance(indices, list):
        for fila in indices:
            if not isinstance(fila, dict):
                continue
            indice = _texto(fila.get("ÍNDICE", fila.get("INDICE"))).upper()
            valor = fila.get("VALOR", "")
            if indice in salida:
                salida[indice] = valor
    return salida


def _df_avance_actividad(corte):
    filas = corte.get("avance_actividad", []) or corte.get("avance_fisico_actividad", []) or []
    if isinstance(filas, list) and filas:
        return pd.DataFrame(filas)
    return pd.DataFrame()


def _df_valor_ganado_seguimiento(seguimiento_fisico, flujo_fondos, fecha_corte, fecha_inicio, valor_contrato):
    seguimientos = seguimiento_fisico.get("seguimientos_fisicos", {}) or {}
    fecha_corte = _parse_fecha(fecha_corte)
    fecha_inicio = _parse_fecha(fecha_inicio)

    if not isinstance(seguimientos, dict) or not fecha_corte or not fecha_inicio:
        return pd.DataFrame(columns=["FECHA DE CORTE", "TIPO DE AVANCE", "VALOR"])

    filas_ev_ac = []

    for clave, corte in seguimientos.items():
        if not isinstance(corte, dict):
            continue

        fecha = _parse_fecha(corte.get("fecha_corte", clave))
        if not fecha or fecha > fecha_corte:
            continue

        avance_general = corte.get("avance_general", []) or []
        fila = avance_general[0] if isinstance(avance_general, list) and avance_general and isinstance(avance_general[0], dict) else {}

        ev = _safe_float(valor_contrato, 0.0) * (_safe_float(fila.get("% EJECUTADO"), 0.0) / 100.0)
        ac = _safe_float(fila.get("$ EJECUTADO"), 0.0)

        filas_ev_ac.extend(
            [
                {
                    "FECHA DE CORTE": fecha,
                    "TIPO DE AVANCE": "EV - VALOR GANADO",
                    "VALOR": ev,
                },
                {
                    "FECHA DE CORTE": fecha,
                    "TIPO DE AVANCE": "AC - COSTO REAL",
                    "VALOR": ac,
                },
            ]
        )

    tablas_flujo = flujo_fondos.get("__tablas__", {}) or {}
    resumen_flujo = tablas_flujo.get("df_resumen", []) or []

    fila_acumulado = {}
    if isinstance(resumen_flujo, list):
        for fila in resumen_flujo:
            if isinstance(fila, dict) and _texto(fila.get("CONCEPTO")).upper() == "ACUMULADO":
                fila_acumulado = fila
                break

    filas_pv = [
        {
            "FECHA DE CORTE": fecha_inicio,
            "TIPO DE AVANCE": "PV - VALOR PLANEADO",
            "VALOR": 0.0,
        }
    ]

    periodos_programacion = []
    for columna in fila_acumulado.keys():
        nombre = _texto(columna)
        if nombre.startswith("Periodo "):
            try:
                numero_periodo = int(nombre.replace("Periodo ", "").strip())
                periodos_programacion.append((numero_periodo, nombre))
            except Exception:
                pass

    for numero_periodo, columna in sorted(periodos_programacion, key=lambda x: x[0]):
        filas_pv.append(
            {
                "FECHA DE CORTE": fecha_inicio + timedelta(days=(numero_periodo * 30) - 1),
                "TIPO DE AVANCE": "PV - VALOR PLANEADO",
                "VALOR": _safe_float(fila_acumulado.get(columna), 0.0),
            }
        )

    return pd.DataFrame(filas_pv + filas_ev_ac, columns=["FECHA DE CORTE", "TIPO DE AVANCE", "VALOR"])


def _indicadores_calculados_seguimiento(seguimiento_fisico, flujo_fondos, fecha_corte, fecha_inicio, valor_contrato):
    corte = _corte_seguimiento(seguimiento_fisico, fecha_corte)
    avance_general = corte.get("avance_general", []) or []
    fila = avance_general[0] if isinstance(avance_general, list) and avance_general and isinstance(avance_general[0], dict) else {}

    pv = _safe_float(fila.get("$ PROGRAMADO"), 0.0)
    ev = _safe_float(valor_contrato, 0.0) * (_safe_float(fila.get("% EJECUTADO"), 0.0) / 100.0)
    ac = _safe_float(fila.get("$ EJECUTADO"), 0.0)

    spi = ev / pv if pv > 0 else 0.0
    cpi = ev / ac if ac > 0 else 0.0

    df_grafica = _df_valor_ganado_seguimiento(
        seguimiento_fisico,
        flujo_fondos,
        fecha_corte,
        fecha_inicio,
        valor_contrato,
    )
    fecha_corte_parseada = _parse_fecha(fecha_corte)
    retraso = 0

    if not df_grafica.empty and ev > 0 and fecha_corte_parseada:
        df_pv = df_grafica[
            df_grafica["TIPO DE AVANCE"] == "PV - VALOR PLANEADO"
        ].sort_values("FECHA DE CORTE")

        puntos = [
            {
                "fecha": _parse_fecha(fila_pv.get("FECHA DE CORTE")),
                "valor": _safe_float(fila_pv.get("VALOR"), 0.0),
            }
            for _, fila_pv in df_pv.iterrows()
        ]

        puntos = [p for p in puntos if p["fecha"]]
        fecha_programacion_ganada = None

        if puntos:
            for i in range(1, len(puntos)):
                anterior = puntos[i - 1]
                actual = puntos[i]

                if anterior["valor"] <= ev <= actual["valor"] and actual["valor"] > anterior["valor"]:
                    proporcion = (ev - anterior["valor"]) / (actual["valor"] - anterior["valor"])
                    dias = (actual["fecha"] - anterior["fecha"]).days
                    fecha_programacion_ganada = anterior["fecha"] + timedelta(days=round(dias * proporcion))
                    break

            if fecha_programacion_ganada is None:
                fecha_programacion_ganada = puntos[0]["fecha"] if ev <= puntos[0]["valor"] else puntos[-1]["fecha"]

        if fecha_programacion_ganada:
            retraso = (fecha_corte_parseada - fecha_programacion_ganada).days

    return {"CPI": round(cpi, 4), "SPI": round(spi, 4), "RETRASO": retraso}


# ==========================================================
# Estado propio del informe
# ==========================================================
def _estado_vacio(generales):
    return {
        "consecutivo": 1,
        "fecha_informe": date.today().isoformat(),
        "fecha_terminacion_real": _fecha_texto(generales.get("fecha_terminacion_acta")),
        "valor_adiciones_editable": generales.get("valor_adiciones", 0.0),
        "aseguradora": "",
        "poliza_no": "",
        "anexo_no": "",
        "anticipo_fecha": date.today().isoformat(),
        "anticipo_valor": 0.0,
        "actividades_obra": "",
        "actividades_interventoria": "",
        "seguimiento": "",
        "nombre_interventor_firma": generales.get("interventor", ""),
    }


def _normalizar_informe(informe, consecutivo, generales):
    base = _estado_vacio(generales)
    base["consecutivo"] = int(consecutivo or 1)

    if isinstance(informe, dict):
        base.update(informe)
        base["consecutivo"] = int(informe.get("consecutivo") or consecutivo or 1)

    return base


def _normalizar_estado_informes(cargado, generales):
    if not isinstance(cargado, dict):
        cargado = {}

    if isinstance(cargado.get("informes"), list):
        informes = [
            _normalizar_informe(informe, i, generales)
            for i, informe in enumerate(cargado.get("informes", []), start=1)
            if isinstance(informe, dict)
        ]
    else:
        informe_migrado = _estado_vacio(generales)
        informe_migrado.update(cargado)
        informe_migrado["consecutivo"] = int(cargado.get("consecutivo") or 1)
        informes = [informe_migrado]

    if not informes:
        informes = [_normalizar_informe({}, 1, generales)]

    informes = sorted(informes, key=lambda x: int(x.get("consecutivo") or 0))
    activo = int(cargado.get("informe_activo") or informes[-1].get("consecutivo") or 1)

    if activo not in [int(x.get("consecutivo") or 0) for x in informes]:
        activo = int(informes[-1].get("consecutivo") or 1)

    return {"informes": informes, "informe_activo": activo}


def _inicializar_estado(generales):
    group_id_actual = _texto(st.session_state.get("group_id"))
    cache_group = _texto(st.session_state.get("_informe_interventoria_group"))
    if cache_group != group_id_actual or "informe_interventoria_datos" not in st.session_state:
        cargado = cargar_estado(CLAVE_GUARDADO) or {}
        st.session_state["informe_interventoria_datos"] = _normalizar_estado_informes(cargado, generales)
        st.session_state["_informe_interventoria_group"] = group_id_actual


def _obtener_informe_activo(generales):
    estado = st.session_state["informe_interventoria_datos"]
    informes = estado.get("informes", [])
    activo = int(estado.get("informe_activo") or 1)

    for informe in informes:
        if int(informe.get("consecutivo") or 0) == activo:
            return informe

    nuevo = _normalizar_informe({}, 1, generales)
    estado["informes"] = [nuevo]
    estado["informe_activo"] = 1
    return nuevo


def _crear_nuevo_informe(generales):
    estado = st.session_state["informe_interventoria_datos"]
    informes = estado.get("informes", [])
    ultimo = max([int(x.get("consecutivo") or 0) for x in informes], default=0)
    nuevo = _normalizar_informe({}, ultimo + 1, generales)
    nuevo["fecha_informe"] = date.today().isoformat()
    informes.append(nuevo)
    estado["informes"] = informes
    estado["informe_activo"] = int(nuevo["consecutivo"])
    return int(nuevo["consecutivo"])


def _guardar(datos=None):
    estado = st.session_state["informe_interventoria_datos"]

    informes_actuales = estado.get("informes", [])
    if not isinstance(informes_actuales, list):
        informes_actuales = []

    guardado = cargar_estado(CLAVE_GUARDADO) or {}
    informes_guardados = guardado.get("informes", []) if isinstance(guardado, dict) else []

    if not isinstance(informes_guardados, list):
        informes_guardados = []

    if informes_guardados and len(informes_actuales) < len(informes_guardados):
        st.error(
            "No se guardó el informe porque el estado actual tiene menos informes que el estado ya guardado. "
            "Esto evita sobrescribir informes anteriores."
        )
        return False

    guardar_estado(CLAVE_GUARDADO, estado)
    st.success("Informe de interventoría guardado correctamente.")
    return True


# ==========================================================
# Word helpers
# ==========================================================
def _set_doc_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(8)


def _p(doc, texto="", bold=False, align=None, size=8):
    p = doc.add_paragraph()
    run = p.add_run(str(texto))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def _cell_text(cell, texto, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(texto))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(7)
    if align:
        p.alignment = align


def _tabla_simple(doc, filas):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for etiqueta, valor in filas:
        cells = table.add_row().cells
        _cell_text(cells[0], etiqueta, bold=True)
        _cell_text(cells[1], valor)
    return table


def _tabla_dataframe(doc, df):
    if df is None or df.empty:
        _p(doc, "Sin registros.")
        return None
    df_out = df.copy()
    for col in df_out.columns:
        df_out[col] = df_out[col].apply(lambda x: _fecha_texto(x) if isinstance(x, (date, datetime, pd.Timestamp)) else x)
    table = doc.add_table(rows=1, cols=len(df_out.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(df_out.columns):
        _cell_text(table.rows[0].cells[i], col, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _, row in df_out.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df_out.columns):
            _cell_text(cells[i], row[col], align=WD_ALIGN_PARAGRAPH.CENTER)
    return table


def _generar_word(generales, datos, df_suspensiones, df_modificaciones, df_adiciones, df_garantias, df_pagos_seleccionados, df_avance, df_avance_actividad, indicadores):
    doc = Document()
    _set_doc_defaults(doc)

    valor_adiciones = _safe_float(datos.get("valor_adiciones_editable"), generales.get("valor_adiciones", 0.0))
    valor_actual = _safe_float(generales.get("valor_inicial_obra"), 0.0) + valor_adiciones

    _p(doc, "INFORME DE INTERVENTORÍA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    _p(doc, f"INFORME No. {datos.get('consecutivo', '')}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    _p(doc, "")
    _tabla_simple(
        doc,
        [
            ("FECHA", _fecha_texto(datos.get("fecha_informe"))),
            ("CONTRATO DE OBRA No", generales.get("contrato_obra_no", "")),
            ("CONTRATO DE INTERVENTORÍA No", generales.get("contrato_interventoria_no", "")),
            ("OBJETO CONTRATO DE OBRA", generales.get("objeto_obra", "")),
            ("OBJETO CONTRATO DE INTERVENTORÍA", generales.get("objeto_interventoria", "")),
            ("VALOR CONTRATO DE OBRA", _moneda(generales.get("valor_obra", 0.0))),
            ("VALOR CONTRATO DE INTERVENTORÍA", _moneda(generales.get("valor_interventoria", 0.0))),
            ("PLAZO DE EJECUCIÓN", generales.get("plazo_ejecucion", "")),
            ("FECHA DE INICIO DEL CONTRATO", _fecha_texto(generales.get("fecha_inicio"))),
            ("FECHA DE TERMINACIÓN SEGÚN ACTA DE INICIO", _fecha_texto(generales.get("fecha_terminacion_acta"))),
            ("FECHA DE TERMINACIÓN REAL", _fecha_texto(datos.get("fecha_terminacion_real"))),
            ("VALOR INICIAL CONTRATO DE OBRA", _moneda(generales.get("valor_inicial_obra", 0.0))),
            ("VALOR ADICIONES", _moneda(valor_adiciones)),
            ("VALOR ACTUAL CONTRATO DE OBRA", _moneda(valor_actual)),
        ],
    )

    _p(doc, "")
    _p(doc, "SUSPENSIONES Y AMPLIACIONES DE SUSPENSIÓN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_dataframe(doc, df_suspensiones)

    _p(doc, "")
    _p(doc, "MODIFICACIONES", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_dataframe(doc, df_modificaciones)

    _p(doc, "")
    _p(doc, "ADICIONES", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_dataframe(doc, df_adiciones)

    _p(doc, "")
    _p(doc, "PÓLIZAS VIGENTES", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_simple(
        doc,
        [
            ("ASEGURADORA", datos.get("aseguradora", "")),
            ("PÓLIZA No", datos.get("poliza_no", "")),
            ("ANEXO No", datos.get("anexo_no", "")),
        ],
    )
    _tabla_dataframe(doc, df_garantias)

    _p(doc, "")
    _p(doc, "SEGUIMIENTO FINANCIERO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_dataframe(doc, df_pagos_seleccionados)

    _p(doc, "")
    _p(doc, "ACTIVIDADES DESARROLLADAS CONTRATISTA DE OBRA", bold=True)
    _p(doc, datos.get("actividades_obra", ""))

    _p(doc, "")
    _p(doc, "ACTIVIDADES DESARROLLADAS CONTRATISTA DE INTERVENTORÍA", bold=True)
    _p(doc, datos.get("actividades_interventoria", ""))

    _p(doc, "")
    _p(doc, "SEGUIMIENTO", bold=True)
    _p(doc, datos.get("seguimiento", ""))

    _p(doc, "")
    _p(doc, "AVANCE FÍSICO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_dataframe(doc, df_avance)
    if df_avance_actividad is not None and not df_avance_actividad.empty:
        _p(doc, "AVANCE POR ACTIVIDAD", bold=True)
        _tabla_dataframe(doc, df_avance_actividad)

    _p(doc, "")
    _p(doc, "INDICADORES", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_simple(doc, [("CPI", indicadores.get("CPI", "")), ("SPI", indicadores.get("SPI", "")), ("RETRASO", indicadores.get("RETRASO", ""))])

    _p(doc, "")
    _p(doc, "NOMBRE DEL INTERVENTOR", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, f"Nombre: {datos.get('nombre_interventor_firma', '')}", align=WD_ALIGN_PARAGRAPH.CENTER)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==========================================================
# Carga de datos
# ==========================================================
st.set_page_config(page_title="Informe de interventoría", layout="wide")
st.title("29. Informe de interventoría")

acta_inicio = _leer_estado("acta_inicio_obra")
contrato_obra = _leer_estado("contrato_obra")
contrato_interventoria = _leer_estado("contrato_interventoria")
control_obra = _leer_estado("control_obra")
seguimiento_fisico = _leer_estado("seguimiento_fisico")
flujo_fondos = _leer_estado("flujo_fondos")
modificacion_prorroga = _leer_estado("modificacion_prorroga")

generales = _datos_generales(acta_inicio, contrato_obra, contrato_interventoria, control_obra)
_inicializar_estado(generales)
estado_informes = st.session_state["informe_interventoria_datos"]

# ==========================================================
# Selector de informes
# ==========================================================
st.markdown("### INFORMES GENERADOS")

col_nuevo_informe, col_selector_informe = st.columns([1, 3])

with col_nuevo_informe:
    if st.button("➕ Nuevo informe", key="informe_nuevo"):
        nuevo_consecutivo = _crear_nuevo_informe(generales)
        _guardar()
        st.session_state["informe_selector_activo"] = nuevo_consecutivo
        st.rerun()

informes = estado_informes.get("informes", [])
opciones_informes = [int(x.get("consecutivo") or 0) for x in informes]
activo_actual = int(estado_informes.get("informe_activo") or (opciones_informes[-1] if opciones_informes else 1))

with col_selector_informe:
    informe_activo = st.selectbox(
        "Seleccione informe",
        options=opciones_informes,
        index=opciones_informes.index(activo_actual) if activo_actual in opciones_informes else 0,
        format_func=lambda x: f"Informe No. {x}",
        key="informe_selector_activo",
    )

estado_informes["informe_activo"] = int(informe_activo)
datos = _obtener_informe_activo(generales)
consecutivo_informe = int(datos.get("consecutivo") or informe_activo or 1)

# ==========================================================
# Encabezado
# ==========================================================
st.markdown("### DATOS GENERALES DEL INFORME")
fechas_corte = _fechas_corte_seguimiento(seguimiento_fisico)
fecha_actual_informe = _fecha_input(datos.get("fecha_informe", date.today()))

if fechas_corte:
    fecha_base_informe = fecha_actual_informe if fecha_actual_informe in fechas_corte else fechas_corte[-1]
else:
    fecha_base_informe = fecha_actual_informe

fecha_informe = st.date_input(
    "FECHA",
    value=fecha_base_informe,
    format="DD/MM/YYYY",
    key=f"informe_fecha_{consecutivo_informe}",
)

datos["fecha_informe"] = fecha_informe

c1, c2 = st.columns(2)
with c1:
    st.text_input("CONTRATO DE OBRA No", value=generales["contrato_obra_no"], disabled=True)
    st.text_input("CONTRATO DE INTERVENTORÍA No", value=generales["contrato_interventoria_no"], disabled=True)
    st.text_area("OBJETO CONTRATO DE OBRA", value=generales["objeto_obra"], disabled=True, height=90)
    st.text_area("OBJETO CONTRATO DE INTERVENTORÍA", value=generales["objeto_interventoria"], disabled=True, height=90)
with c2:
    st.number_input("VALOR CONTRATO DE OBRA", value=float(generales["valor_obra"]), disabled=True, format="%.2f")
    st.number_input("VALOR CONTRATO DE INTERVENTORÍA", value=float(generales["valor_interventoria"]), disabled=True, format="%.2f")
    st.text_input("PLAZO DE EJECUCIÓN", value=generales["plazo_ejecucion"], disabled=True)
    st.date_input("FECHA DE INICIO DEL CONTRATO", value=_fecha_input(generales["fecha_inicio"]), disabled=True, format="DD/MM/YYYY")
    st.date_input("FECHA DE TERMINACIÓN SEGÚN ACTA DE INICIO", value=_fecha_input(generales["fecha_terminacion_acta"]), disabled=True, format="DD/MM/YYYY")

c3, c4, c5 = st.columns(3)
with c3:
    datos["fecha_terminacion_real"] = st.date_input(
        "FECHA DE TERMINACIÓN REAL",
        value=_fecha_input(datos.get("fecha_terminacion_real", generales["fecha_terminacion_acta"])),
        format="DD/MM/YYYY",
        key=f"informe_fecha_terminacion_real_{consecutivo_informe}",
    )
with c4:
    st.number_input("VALOR INICIAL CONTRATO DE OBRA", value=float(generales["valor_inicial_obra"]), disabled=True, format="%.2f")
with c5:
    datos["valor_adiciones_editable"] = st.number_input(
        "VALOR ADICIONES",
        value=float(datos.get("valor_adiciones_editable", generales["valor_adiciones"])),
        format="%.2f",
        key=f"informe_valor_adiciones_{consecutivo_informe}",
    )

valor_actual_obra = float(generales["valor_inicial_obra"]) + float(datos.get("valor_adiciones_editable", 0.0))
st.number_input("VALOR ACTUAL CONTRATO DE OBRA", value=valor_actual_obra, disabled=True, format="%.2f")

# ==========================================================
# Tablas contractuales
# ==========================================================
st.markdown("### SUSPENSIONES Y AMPLIACIONES DE SUSPENSIÓN")
df_susp_base = _df_suspensiones(control_obra)
fechas_susp = _fechas_disponibles_df(df_susp_base, "FECHA DEL ACTA")
if fechas_susp:
    fecha_susp = st.date_input("Mostrar suspensiones hasta", value=fechas_susp[-1], format="DD/MM/YYYY", key=f"informe_fecha_suspensiones_{consecutivo_informe}")
else:
    fecha_susp = date.today()
df_suspensiones = _filtrar_df_hasta_fecha(df_susp_base, "FECHA DEL ACTA", fecha_susp)
st.dataframe(df_suspensiones, hide_index=True, width="stretch")

st.markdown("### MODIFICACIONES")
df_mod_base = _df_modificaciones(modificacion_prorroga)
fechas_mod = _fechas_disponibles_df(df_mod_base, "FECHA")
if fechas_mod:
    fecha_mod = st.date_input("Mostrar modificaciones hasta", value=fechas_mod[-1], format="DD/MM/YYYY", key=f"informe_fecha_modificaciones_{consecutivo_informe}")
else:
    fecha_mod = date.today()
df_modificaciones = _filtrar_df_hasta_fecha(df_mod_base, "FECHA", fecha_mod)
st.dataframe(df_modificaciones, hide_index=True, width="stretch")

st.markdown("### ADICIONES")
df_adic_base = _df_adiciones(control_obra)
fechas_adic = _fechas_disponibles_df(df_adic_base, "FECHA")
if fechas_adic:
    fecha_adic = st.date_input("Mostrar adiciones hasta", value=fechas_adic[-1], format="DD/MM/YYYY", key=f"informe_fecha_adiciones_{consecutivo_informe}")
else:
    fecha_adic = date.today()
df_adiciones = _filtrar_df_hasta_fecha(df_adic_base, "FECHA", fecha_adic)
st.dataframe(df_adiciones, hide_index=True, width="stretch")

# ==========================================================
# Pólizas y garantías
# ==========================================================
st.markdown("### PÓLIZAS VIGENTES")
col_pol1, col_pol2, col_pol3 = st.columns(3)
with col_pol1:
    datos["aseguradora"] = st.text_input("ASEGURADORA", value=datos.get("aseguradora", ""), key=f"informe_aseguradora_{consecutivo_informe}")
with col_pol2:
    datos["poliza_no"] = st.text_input("PÓLIZA No", value=datos.get("poliza_no", ""), key=f"informe_poliza_no_{consecutivo_informe}")
with col_pol3:
    datos["anexo_no"] = st.text_input("ANEXO No", value=datos.get("anexo_no", ""), key=f"informe_anexo_no_{consecutivo_informe}")

df_garantias = pd.concat(
    [_df_garantias_iniciales(contrato_obra), _df_garantias_modificadas(control_obra)],
    ignore_index=True,
)
st.dataframe(df_garantias, hide_index=True, width="stretch")

# ==========================================================
# Seguimiento financiero
# ==========================================================
st.markdown("### SEGUIMIENTO FINANCIERO")
datos["anticipo_fecha"] = st.date_input(
    "FECHA ANTICIPO",
    value=_fecha_input(datos.get("anticipo_fecha", date.today())),
    format="DD/MM/YYYY",
    key=f"informe_anticipo_fecha_{consecutivo_informe}",
)
datos["anticipo_valor"] = st.number_input(
    "VALOR ANTICIPO",
    value=float(datos.get("anticipo_valor", _valor_anticipo(control_obra))),
    format="%.2f",
    key=f"informe_anticipo_valor_{consecutivo_informe}",
)

df_pagos_base = _df_pagos(control_obra)
fechas_pagos = _fechas_disponibles_df(df_pagos_base, "FECHA")
if fechas_pagos:
    fecha_pago = st.date_input("Mostrar actas parciales hasta", value=fechas_pagos[-1], format="DD/MM/YYYY", key=f"informe_fecha_pagos_{consecutivo_informe}")
else:
    fecha_pago = date.today()
df_pagos_filtrados = _filtrar_df_hasta_fecha(df_pagos_base, "FECHA", fecha_pago)

descripcion_anticipo = pd.DataFrame(
    [
        {
            "DESCRIPCIÓN": "ANTICIPO",
            "FECHA": datos.get("anticipo_fecha"),
            "VALOR": float(datos.get("anticipo_valor", 0.0)),
        }
    ]
)
actas_financieras = df_pagos_filtrados.rename(columns={"ACTA PARCIAL No.": "DESCRIPCIÓN"})
if not actas_financieras.empty:
    actas_financieras["DESCRIPCIÓN"] = actas_financieras["DESCRIPCIÓN"].apply(lambda x: f"ACTA PARCIAL No. {x}")
    actas_financieras = actas_financieras[["DESCRIPCIÓN", "FECHA", "VALOR"]]

df_pagos_seleccionados = pd.concat([descripcion_anticipo, actas_financieras], ignore_index=True)
st.dataframe(df_pagos_seleccionados, hide_index=True, width="stretch")
st.number_input("TOTAL PAGOS AUTORIZADOS", value=float(df_pagos_seleccionados["VALOR"].apply(lambda x: _safe_float(x, 0.0)).sum()), disabled=True, format="%.2f")

# ==========================================================
# Actividades y seguimiento
# ==========================================================
st.markdown("### ACTIVIDADES DESARROLLADAS CONTRATISTA DE OBRA")
datos["actividades_obra"] = st.text_area(
    "Actividades desarrolladas contratista de obra",
    value=datos.get("actividades_obra", ""),
    height=180,
    label_visibility="collapsed",
    key=f"informe_actividades_obra_{consecutivo_informe}",
)

st.markdown("### ACTIVIDADES DESARROLLADAS CONTRATISTA DE INTERVENTORÍA")
datos["actividades_interventoria"] = st.text_area(
    "Actividades desarrolladas contratista de interventoría",
    value=datos.get("actividades_interventoria", ""),
    height=180,
    label_visibility="collapsed",
    key=f"informe_actividades_interventoria_{consecutivo_informe}",
)

st.markdown("### SEGUIMIENTO")

# ==========================================================
# Avance físico e indicadores
# ==========================================================
st.markdown("### AVANCE FÍSICO")
corte = _corte_seguimiento(seguimiento_fisico, datos.get("fecha_informe"))
avance_general = _avance_general_desde_corte(corte, datos.get("fecha_informe"))
df_avance = pd.DataFrame([avance_general])
st.dataframe(df_avance, hide_index=True, width="stretch")

df_avance_actividad = _df_avance_actividad(corte)

if not df_avance_actividad.empty:
    columnas_ordenadas = [
        "ITEM",
        "DESCRIPCIÓN",
        "% EJECUTADO",
        "$ EJECUTADO",
        "% PROGRAMADO",
        "$ PROGRAMADO",
    ]

    columnas_existentes = [
        col for col in columnas_ordenadas
        if col in df_avance_actividad.columns
    ]

    columnas_restantes = [
        col for col in df_avance_actividad.columns
        if col not in columnas_existentes
    ]

    df_avance_actividad = df_avance_actividad[columnas_existentes + columnas_restantes]

    st.markdown("#### AVANCE POR ACTIVIDAD")
    st.dataframe(df_avance_actividad, hide_index=True, width="stretch")

st.markdown("### GRÁFICO DE VALOR GANADO")
df_valor_ganado_grafico = _df_valor_ganado_seguimiento(
    seguimiento_fisico,
    flujo_fondos,
    datos.get("fecha_informe"),
    generales.get("fecha_inicio"),
    generales.get("valor_obra", 0.0),
)

if not df_valor_ganado_grafico.empty:
    fig_valor_ganado = px.line(
        df_valor_ganado_grafico,
        x="FECHA DE CORTE",
        y="VALOR",
        color="TIPO DE AVANCE",
        markers=True,
        title="Valor ganado",
        color_discrete_map={
            "PV - VALOR PLANEADO": "blue",
            "EV - VALOR GANADO": "green",
            "AC - COSTO REAL": "orange",
        },
    )

    fig_valor_ganado.update_layout(
        xaxis_title="Fecha de programación",
        yaxis_title="Valor",
        legend_title="Tipo de avance",
        yaxis_tickprefix="$ ",
        yaxis_tickformat=",",
    )

    st.plotly_chart(fig_valor_ganado, width="stretch")
else:
    st.info("No hay cortes de seguimiento físico suficientes para generar el gráfico de valor ganado.")

indicadores = _indicadores_calculados_seguimiento(
    seguimiento_fisico,
    flujo_fondos,
    datos.get("fecha_informe"),
    generales.get("fecha_inicio"),
    generales.get("valor_obra", 0.0),
)
st.markdown("### INDICADORES")
col_i1, col_i2, col_i3 = st.columns(3)
with col_i1:
    st.text_input("CPI - ÍNDICE DE DESEMPEÑO EN COSTO", value=_texto(indicadores.get("CPI")), disabled=True)
with col_i2:
    st.text_input("SPI - ÍNDICE DE DESEMPEÑO DE LA PROGRAMACIÓN", value=_texto(indicadores.get("SPI")), disabled=True)
with col_i3:
    st.text_input("RETRASO", value=_texto(indicadores.get("RETRASO")), disabled=True)

st.markdown("### FIRMA")
datos["nombre_interventor_firma"] = st.text_input(
    "Nombre del interventor",
    value=datos.get("nombre_interventor_firma", generales.get("interventor", "")),
    key=f"informe_nombre_interventor_firma_{consecutivo_informe}",
)

# ==========================================================
# Guardar y Word
# ==========================================================
col_b1, col_b2 = st.columns([1, 1])
with col_b1:
    if st.button("💾 Guardar informe", key=f"informe_guardar_{consecutivo_informe}"):
        _guardar(datos)

with col_b2:
    word = _generar_word(
        generales,
        datos,
        df_suspensiones,
        df_modificaciones,
        df_adiciones,
        df_garantias,
        df_pagos_seleccionados,
        df_avance,
        df_avance_actividad,
        indicadores,
    )
    st.download_button(
        "📄 Descargar informe en Word",
        data=word,
        file_name=f"informe_interventoria_{consecutivo_informe}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        width="stretch",
        key=f"informe_descargar_word_{consecutivo_informe}",
    )
