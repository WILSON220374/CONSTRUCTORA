from __future__ import annotations

from datetime import date, datetime
from io import BytesIO

import pandas as pd
import streamlit as st
from docx import Document

from supabase_state import cargar_estado
from supabase_state import guardar_estado as guardar_estado_bd


CLAVE_GUARDADO = "plan_inversion_anticipo"
PORCENTAJE_ANTICIPO = 30.0
MIN_MESES = 3
FILAS_INICIALES = 3


# ==========================================================
# Helpers base
# ==========================================================
def guardar_estado(clave, datos):
    def serializar(obj):
        if isinstance(obj, dict):
            return {k: serializar(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [serializar(x) for x in obj]
        if isinstance(obj, (date, datetime)):
            return obj.isoformat()
        return obj

    guardar_estado_bd(clave, serializar(datos))


def _texto(valor) -> str:
    if valor is None:
        return ""
    return str(valor).strip()


def _parse_fecha(valor):
    if isinstance(valor, date):
        return valor
    if isinstance(valor, datetime):
        return valor.date()
    if isinstance(valor, str) and valor.strip():
        txt = valor.strip()
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
            try:
                return datetime.strptime(txt, fmt).date()
            except Exception:
                continue
        try:
            return datetime.fromisoformat(txt).date()
        except Exception:
            return None
    return None


def _fecha_texto(valor) -> str:
    f = _parse_fecha(valor)
    if not f:
        return ""
    return f.strftime("%d/%m/%Y")


def _safe_float(valor, default=0.0) -> float:
    if valor is None:
        return 0.0 if default is None else float(default)

    if isinstance(valor, (int, float)):
        return float(valor)

    txt = str(valor).strip()
    if not txt:
        return 0.0 if default is None else float(default)

    txt = txt.replace("$", "").replace("%", "").replace(" ", "")
    txt = txt.replace(",", "")

    try:
        return float(txt)
    except Exception:
        return 0.0 if default is None else float(default)


def _formato_moneda(valor) -> str:
    return f"${_safe_float(valor):,.2f}"


def _primero_no_vacio(*valores):
    for valor in valores:
        if isinstance(valor, str):
            if valor.strip():
                return valor.strip()
        elif valor not in (None, "", []):
            return valor
    return ""


# ==========================================================
# Lecturas seguras por grupo actual
# ==========================================================
def _leer_estado_directo(clave: str) -> dict:
    datos = cargar_estado(clave) or {}
    return datos if isinstance(datos, dict) else {}


def _leer_acta_inicio() -> dict:
    return _leer_estado_directo("acta_inicio_obra")


def _leer_contrato_interventoria() -> dict:
    return _leer_estado_directo("contrato_interventoria")


def _leer_presupuesto_obra() -> dict:
    return _leer_estado_directo("presupuesto_obra")


def _leer_contrato_obra() -> dict:
    return _leer_estado_directo("contrato_obra")


# ==========================================================
# Presupuesto base para selector
# ==========================================================
def _construir_catalogo_presupuesto(presupuesto_obra: dict):
    catalogo = []
    mapa = {}

    if not isinstance(presupuesto_obra, dict):
        return catalogo, mapa

    tablas = presupuesto_obra.get("__tablas__", {}) or {}
    grupos = tablas.get("grupos_presupuesto_obra", []) or []

    for grupo in grupos:
        if not isinstance(grupo, dict):
            continue

        for fila in grupo.get("rows", []) or []:
            if not isinstance(fila, dict):
                continue

            item_no = _texto(fila.get("ITEM"))
            descripcion = _texto(fila.get("DESCRIPCIÓN"))
            valor_total = _safe_float(fila.get("VR TOTAL"), 0.0)

            if not descripcion:
                continue

            registro = {
                "descripcion": descripcion,
                "item_no": item_no,
                "valor_referencia": valor_total,
            }

            if descripcion not in mapa:
                catalogo.append(registro)
                mapa[descripcion] = registro

    catalogo = sorted(catalogo, key=lambda x: (x["descripcion"], x["item_no"]))
    return catalogo, mapa


# ==========================================================
# Estado propio de la hoja
# ==========================================================
def _columnas_meses(cantidad_meses: int):
    cantidad = max(MIN_MESES, int(cantidad_meses or MIN_MESES))
    return [f"VR. MES {i}" for i in range(1, cantidad + 1)]


def _fila_vacia(columnas_meses):
    fila = {
        "ÍTEM No.": "",
        "DESCRIPCIÓN DEL ÍTEM": "",
        "VALOR": 0.0,
        "VALOR PROGRAMA APROBADO": 0.0,
        "%": 0.0,
    }
    for col in columnas_meses:
        fila[col] = 0.0
    return fila


def _normalizar_filas(rows, columnas_meses):
    rows_out = []

    for fila in rows or []:
        base = _fila_vacia(columnas_meses)

        if isinstance(fila, dict):
            for k in base.keys():
                if k in fila:
                    base[k] = fila[k]

        elif isinstance(fila, str):
            txt = _texto(fila)
            if txt:
                base["DESCRIPCIÓN DEL ÍTEM"] = txt

        rows_out.append(base)

    if not rows_out:
        rows_out = [_fila_vacia(columnas_meses) for _ in range(FILAS_INICIALES)]

    return rows_out

def _inicializar_estado():
    group_id_actual = _texto(st.session_state.get("group_id"))
    cache_group = _texto(st.session_state.get("_plan_inversion_anticipo_group"))

    if cache_group != group_id_actual or "plan_inversion_anticipo_datos" not in st.session_state:
        cargado = cargar_estado(CLAVE_GUARDADO) or {}
        if not isinstance(cargado, dict):
            cargado = {}

        cantidad_meses = int(cargado.get("cantidad_meses", MIN_MESES) or MIN_MESES)
        columnas_meses = _columnas_meses(cantidad_meses)

        st.session_state["plan_inversion_anticipo_datos"] = {
            "fecha_documento": _parse_fecha(cargado.get("fecha_documento")) or date.today(),
            "cantidad_meses": cantidad_meses,
            "porcentaje_anticipo": _safe_float(cargado.get("porcentaje_anticipo"), PORCENTAJE_ANTICIPO),
            "rows": _normalizar_filas(cargado.get("rows", []), columnas_meses),
        }

        st.session_state["_plan_inversion_anticipo_group"] = group_id_actual


def _guardar():
    guardar_estado(CLAVE_GUARDADO, st.session_state["plan_inversion_anticipo_datos"])
    st.success("Plan de inversión del anticipo guardado correctamente.")

# ==========================================================
# Reglas de negocio
# ==========================================================
def _valor_contrato_desde_fuentes(acta_inicio: dict, contrato_obra: dict) -> float:
    candidatos = [
        acta_inicio.get("valor_total_contrato_obra"),
        acta_inicio.get("valor_contrato"),
        acta_inicio.get("valor"),
        contrato_obra.get("valor_total_numeros"),
        contrato_obra.get("valor_contrato"),
    ]
    for valor in candidatos:
        numero = _safe_float(valor, None)
        if numero > 0:
            return numero
    return 0.0


def _recalcular_filas(datos, mapa_catalogo, valor_anticipo):
    columnas_meses = _columnas_meses(datos.get("cantidad_meses", MIN_MESES))
    filas_base = _normalizar_filas(datos.get("rows", []), columnas_meses)
    filas_out = []

    for fila in filas_base:
        descripcion = _texto(fila.get("DESCRIPCIÓN DEL ÍTEM"))
        fila_nueva = _fila_vacia(columnas_meses)

        fila_nueva["DESCRIPCIÓN DEL ÍTEM"] = descripcion

        if descripcion and descripcion in mapa_catalogo:
            fila_nueva["ÍTEM No."] = _texto(mapa_catalogo[descripcion].get("item_no"))
            fila_nueva["VALOR"] = _safe_float(mapa_catalogo[descripcion].get("valor_referencia"), 0.0)

        suma_meses = 0.0
        for col in columnas_meses:
            valor_mes = _safe_float(fila.get(col), 0.0)
            fila_nueva[col] = valor_mes
            suma_meses += valor_mes

        fila_nueva["VALOR PROGRAMA APROBADO"] = round(suma_meses, 2)
        fila_nueva["%"] = round((suma_meses / valor_anticipo) * 100.0, 4) if valor_anticipo > 0 else 0.0

        filas_out.append(fila_nueva)

    datos["rows"] = filas_out
    return columnas_meses


def _dataframe_para_editor(datos, columnas_meses):
    columnas = ["ÍTEM No.", "DESCRIPCIÓN DEL ÍTEM", "VALOR"] + columnas_meses + ["VALOR PROGRAMA APROBADO", "%"]
    return pd.DataFrame(datos.get("rows", []), columns=columnas)


def _sumas_totales(df, columnas_meses):
    total_programado = _safe_float(df["VALOR PROGRAMA APROBADO"].sum(), 0.0) if "VALOR PROGRAMA APROBADO" in df.columns else 0.0
    total_porcentaje = _safe_float(df["%"].sum(), 0.0) if "%" in df.columns else 0.0
    return total_programado, total_porcentaje


def _duplicados_descripcion(rows):
    vistos = set()
    repetidos = set()

    for fila in rows:
        descripcion = _texto(fila.get("DESCRIPCIÓN DEL ÍTEM"))
        if not descripcion:
            continue
        if descripcion in vistos:
            repetidos.add(descripcion)
        vistos.add(descripcion)

    return repetidos


# ==========================================================
# Word
# ==========================================================
def _generar_word(
    datos,
    contrato_obra_no,
    contratista,
    objeto_contrato,
    contrato_interventoria_no,
    interventor,
    valor_contrato,
    valor_anticipo,
    porcentaje_anticipo,
    nombre_firma_contratista,
    nombre_firma_interventor,
):
    columnas_meses = _columnas_meses(datos.get("cantidad_meses", MIN_MESES))
    df = _dataframe_para_editor(datos, columnas_meses)

    doc = Document()
    doc.add_heading("PLAN DE INVERSIÓN DEL ANTICIPO", level=1)

    doc.add_paragraph(f"Fecha: {_fecha_texto(datos.get('fecha_documento'))}")
    doc.add_paragraph(f"Contrato de obra No.: {_texto(contrato_obra_no)}")
    doc.add_paragraph(f"Contratista: {_texto(contratista)}")
    doc.add_paragraph(f"Objeto del contrato: {_texto(objeto_contrato)}")
    doc.add_paragraph(f"Contrato de interventoría No.: {_texto(contrato_interventoria_no)}")
    doc.add_paragraph(f"Interventor: {_texto(interventor)}")

    doc.add_paragraph("")
    doc.add_paragraph(f"Valor del contrato: {_formato_moneda(valor_contrato)}")
    doc.add_paragraph(f"Porcentaje del anticipo: {porcentaje_anticipo:.2f}%")
    doc.add_paragraph(f"Valor del anticipo: {_formato_moneda(valor_anticipo)}")

    doc.add_paragraph("")
    encabezados = ["ÍTEM No.", "DESCRIPCIÓN DEL ÍTEM", "VALOR"] + columnas_meses + ["VALOR PROGRAMA APROBADO", "%"]
    tabla = doc.add_table(rows=1, cols=len(encabezados))
    tabla.style = "Table Grid"

    for i, enc in enumerate(encabezados):
        tabla.rows[0].cells[i].text = enc

    for _, fila in df.iterrows():
        celdas = tabla.add_row().cells
        idx = 0

        celdas[idx].text = _texto(fila.get("ÍTEM No.", "")); idx += 1
        celdas[idx].text = _texto(fila.get("DESCRIPCIÓN DEL ÍTEM", "")); idx += 1
        celdas[idx].text = _formato_moneda(fila.get("VALOR", 0.0)); idx += 1

        for col in columnas_meses:
            celdas[idx].text = _formato_moneda(fila.get(col, 0.0))
            idx += 1

        celdas[idx].text = _formato_moneda(fila.get("VALOR PROGRAMA APROBADO", 0.0)); idx += 1
        celdas[idx].text = f"{_safe_float(fila.get('%', 0.0)):.4f}%"

    total_programado, total_porcentaje = _sumas_totales(df, columnas_meses)

    doc.add_paragraph("")
    doc.add_paragraph(f"Total valor programa aprobado: {_formato_moneda(total_programado)}")
    doc.add_paragraph(f"Total porcentaje: {total_porcentaje:.4f}%")

    doc.add_paragraph("")
    tabla_firmas = doc.add_table(rows=2, cols=2)
    tabla_firmas.style = "Table Grid"
    tabla_firmas.cell(0, 0).text = "CONTRATISTA"
    tabla_firmas.cell(0, 1).text = "INTERVENTOR"
    tabla_firmas.cell(1, 0).text = _texto(nombre_firma_contratista)
    tabla_firmas.cell(1, 1).text = _texto(nombre_firma_interventor)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ==========================================================
# Inicialización
# ==========================================================
_inicializar_estado()
datos = st.session_state["plan_inversion_anticipo_datos"]

acta_inicio = _leer_acta_inicio()
contrato_interventoria = _leer_contrato_interventoria()
presupuesto_obra = _leer_presupuesto_obra()
contrato_obra = _leer_contrato_obra()

catalogo_presupuesto, mapa_catalogo = _construir_catalogo_presupuesto(presupuesto_obra)

valor_contrato = _valor_contrato_desde_fuentes(acta_inicio, contrato_obra)
valor_anticipo = round(valor_contrato * (PORCENTAJE_ANTICIPO / 100.0), 2)

contrato_obra_no = _primero_no_vacio(
    acta_inicio.get("numero_contrato"),
    contrato_obra.get("numero_contrato"),
)
contratista = _primero_no_vacio(
    acta_inicio.get("nombre_firma_contratista"),
    contrato_obra.get("nombre_contratista"),
)
objeto_contrato = _primero_no_vacio(
    acta_inicio.get("objeto_contrato_obra"),
    contrato_obra.get("objeto_general"),
)
contrato_interventoria_no = _primero_no_vacio(
    contrato_interventoria.get("numero_proceso_contratacion"),
    contrato_interventoria.get("numero_contrato"),
)
interventor = _primero_no_vacio(
    acta_inicio.get("nombre_firma_interventor"),
    contrato_interventoria.get("nombre_interventor"),
    contrato_interventoria.get("firmante_interventor"),
)

nombre_firma_contratista = _primero_no_vacio(acta_inicio.get("nombre_firma_contratista"), contratista)
nombre_firma_interventor = _primero_no_vacio(acta_inicio.get("nombre_firma_interventor"), interventor)


# ==========================================================
# UI
# ==========================================================
st.markdown(
    """
    <style>
    .titulo-plan {
        font-size: 30px;
        font-weight: 800;
        margin-bottom: 8px;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown('<div class="titulo-plan">PLAN DE INVERSIÓN DEL ANTICIPO</div>', unsafe_allow_html=True)

with st.sidebar:
    st.header("🧭 Acciones")

with st.container(border=True):
    col1, col2 = st.columns(2)
    with col1:
        datos["fecha_documento"] = st.date_input(
            "Fecha del documento",
            value=_parse_fecha(datos.get("fecha_documento")) or date.today(),
            format="DD/MM/YYYY",
            key="pia_fecha_documento",
        )
    with col2:
        datos["cantidad_meses"] = int(
            st.number_input(
                "Cantidad total de meses",
                min_value=MIN_MESES,
                value=int(datos.get("cantidad_meses", MIN_MESES) or MIN_MESES),
                step=1,
                key="pia_cantidad_meses",
            )
        )

columnas_meses = _recalcular_filas(datos, mapa_catalogo, valor_anticipo)

with st.container(border=True):
    c1, c2, c3 = st.columns(3)
    with c1:
        st.text_input("Valor del contrato", value=_formato_moneda(valor_contrato), disabled=True, key="pia_valor_contrato")
    with c2:
        st.text_input("Porcentaje del anticipo", value=f"{PORCENTAJE_ANTICIPO:.2f}%", disabled=True, key="pia_porcentaje_anticipo")
    with c3:
        st.text_input("Valor del anticipo", value=_formato_moneda(valor_anticipo), disabled=True, key="pia_valor_anticipo")

with st.container(border=True):
    c1, c2 = st.columns(2)
    with c1:
        st.text_input("Contrato de obra No.", value=_texto(contrato_obra_no), disabled=True, key="pia_contrato_obra")
        st.text_input("Contratista", value=_texto(contratista), disabled=True, key="pia_contratista")
        st.text_area("Objeto del contrato", value=_texto(objeto_contrato), disabled=True, height=110, key="pia_objeto")
    with c2:
        st.text_input("Contrato de interventoría No.", value=_texto(contrato_interventoria_no), disabled=True, key="pia_contrato_interventoria")
        st.text_input("Interventor", value=_texto(interventor), disabled=True, key="pia_interventor")

with st.container(border=True):
    st.subheader("Detalle del plan de inversión del anticipo")

    df_editor = _dataframe_para_editor(datos, columnas_meses)

    column_config = {
        "ÍTEM No.": st.column_config.TextColumn("Ítem No.", disabled=True),
        "DESCRIPCIÓN DEL ÍTEM": st.column_config.SelectboxColumn(
            "Descripción del ítem",
            options=[x["descripcion"] for x in catalogo_presupuesto],
            required=False,
            width="large",
        ),
        "VALOR": st.column_config.NumberColumn("Valor", format="$ %.2f", disabled=True),
        "VALOR PROGRAMA APROBADO": st.column_config.NumberColumn("Valor programa aprobado", format="$ %.2f", disabled=True),
        "%": st.column_config.NumberColumn("%", format="%.4f", disabled=True),
    }

    for col in columnas_meses:
        column_config[col] = st.column_config.NumberColumn(col, min_value=0.0, step=0.01, format="$ %.2f")

    columnas_editor = ["ÍTEM No.", "DESCRIPCIÓN DEL ÍTEM", "VALOR"] + columnas_meses + ["VALOR PROGRAMA APROBADO", "%"]

    df_editado = st.data_editor(
        df_editor,
        width="stretch",
        hide_index=True,
        num_rows="dynamic",
        key="plan_anticipo_editor",
        column_order=columnas_editor,
        column_config=column_config,
        disabled=["ÍTEM No.", "VALOR", "VALOR PROGRAMA APROBADO", "%"],
    )

    rows_antes = df_editado.to_dict(orient="records")
    datos["rows"] = rows_antes
    columnas_meses = _recalcular_filas(datos, mapa_catalogo, valor_anticipo)
    df_final = _dataframe_para_editor(datos, columnas_meses)

    rows_despues = df_final.to_dict(orient="records")
    if rows_despues != rows_antes:
        datos["rows"] = rows_despues
        st.rerun()

    repetidos = _duplicados_descripcion(datos["rows"])
    if repetidos:
        st.error("No se puede repetir un ítem del presupuesto en más de una fila.")

    total_programado, total_porcentaje = _sumas_totales(df_final, columnas_meses)

    st.markdown(f"**Valor del anticipo:** {_formato_moneda(valor_anticipo)}")
    st.markdown(f"**Total valor programa aprobado:** {_formato_moneda(total_programado)}")
    st.markdown(f"**Total porcentaje:** {total_porcentaje:.4f}%")

    if abs(total_programado - valor_anticipo) < 0.01:
        st.success("La suma del valor programa aprobado coincide con el valor del anticipo.")
    else:
        st.warning("La suma del valor programa aprobado debe ser igual al valor del anticipo.")

    if abs(total_porcentaje - 100.0) < 0.0001:
        st.success("La suma de porcentajes es 100%.")
    else:
        st.warning("La suma de porcentajes debe ser 100%.")

with st.container(border=True):
    st.subheader("Firmas")
    cf1, cf2 = st.columns(2)
    with cf1:
        st.text_input("Contratista", value=_texto(nombre_firma_contratista), disabled=True, key="pia_firma_contratista")
    with cf2:
        st.text_input("Interventor", value=_texto(nombre_firma_interventor), disabled=True, key="pia_firma_interventor")

with st.container(border=True):
    col_guardar, col_word = st.columns(2)

    with col_guardar:
        if st.button("💾 Guardar plan de inversión del anticipo", type="primary", width="stretch", key="pia_guardar_principal"):
            _guardar()

    with col_word:
        word_plan = _generar_word(
            datos,
            contrato_obra_no,
            contratista,
            objeto_contrato,
            contrato_interventoria_no,
            interventor,
            valor_contrato,
            valor_anticipo,
            PORCENTAJE_ANTICIPO,
            nombre_firma_contratista,
            nombre_firma_interventor,
        )
        st.download_button(
            "📥 Descargar en Word",
            data=word_plan,
            file_name="plan_inversion_anticipo.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            width="stretch",
            key="pia_descargar_word",
        )
