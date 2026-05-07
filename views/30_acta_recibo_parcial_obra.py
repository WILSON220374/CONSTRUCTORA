from io import BytesIO
from datetime import date, datetime

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt

from supabase_state import cargar_estado
from supabase_state import guardar_estado as guardar_estado_bd


CLAVE_GUARDADO = "acta_recibo_parcial_obra"


# ==========================================================
# Persistencia
# ==========================================================
def guardar_estado(clave, datos):
    def serializar(obj):
        if isinstance(obj, dict):
            return {k: serializar(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [serializar(x) for x in obj]
        if isinstance(obj, (date, datetime)):
            return obj.isoformat()
        return obj

    guardar_estado_bd(clave, serializar(datos))


def _leer_estado(clave):
    datos = cargar_estado(clave) or {}
    return datos if isinstance(datos, dict) else {}


# ==========================================================
# Helpers generales
# ==========================================================
def _texto(valor):
    if valor is None:
        return ""
    return str(valor).strip()


def _safe_float(valor, default=0.0):
    try:
        if valor is None or valor == "":
            return float(default)
        if isinstance(valor, (int, float)):
            return float(valor)
        txt = str(valor).strip().replace("$", "").replace(" ", "")
        txt = txt.replace(".", "").replace(",", ".")
        return float(txt)
    except Exception:
        return float(default)


def _parse_fecha(valor):
    if isinstance(valor, date):
        return valor
    if isinstance(valor, datetime):
        return valor.date()
    if isinstance(valor, str) and valor.strip():
        txt = valor.strip()
        try:
            return datetime.fromisoformat(txt).date()
        except Exception:
            pass
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
            try:
                return datetime.strptime(txt, fmt).date()
            except Exception:
                continue
    return None


def _fecha_input(valor):
    return _parse_fecha(valor) or date.today()


def _fecha_texto(valor):
    fecha = _parse_fecha(valor)
    return fecha.strftime("%d/%m/%Y") if fecha else ""


def _moneda(valor):
    return f"$ {_safe_float(valor, 0.0):,.2f}"


def _primero_no_vacio(*valores):
    for valor in valores:
        txt = _texto(valor)
        if txt:
            return txt
    return ""


def _key_codigo_natural(value):
    partes = []
    for chunk in _texto(value).replace("-", ".").split("."):
        if chunk.isdigit():
            partes.append(int(chunk))
        else:
            partes.append(chunk)
    return tuple(partes)


# ==========================================================
# Fuentes de datos
# ==========================================================
def _valor_contrato_obra(acta_inicio, contrato_obra):
    for valor in [
        acta_inicio.get("valor_total_contrato_obra"),
        acta_inicio.get("valor_contrato"),
        acta_inicio.get("valor"),
        contrato_obra.get("valor_total_numeros"),
        contrato_obra.get("valor_contrato"),
        contrato_obra.get("valor"),
    ]:
        numero = _safe_float(valor, 0.0)
        if numero > 0:
            return numero
    return 0.0


def _valor_anticipo_acta_inicio(acta_inicio, plan_anticipo, valor_contrato):
    for valor in [
        acta_inicio.get("valor_anticipo"),
        acta_inicio.get("valor_del_anticipo"),
        acta_inicio.get("anticipo_valor"),
        plan_anticipo.get("valor_anticipo"),
        plan_anticipo.get("anticipo"),
    ]:
        numero = _safe_float(valor, 0.0)
        if numero > 0:
            return numero

    porcentaje = _porcentaje_anticipo_acta_inicio(acta_inicio, plan_anticipo)
    if porcentaje > 0 and valor_contrato > 0:
        return round(valor_contrato * porcentaje / 100.0, 2)
    return 0.0


def _porcentaje_anticipo_acta_inicio(acta_inicio, plan_anticipo):
    for valor in [
        acta_inicio.get("porcentaje_anticipo"),
        acta_inicio.get("anticipo_porcentaje"),
        acta_inicio.get("porcentaje_del_anticipo"),
        plan_anticipo.get("porcentaje_anticipo"),
        plan_anticipo.get("porcentaje"),
    ]:
        numero = _safe_float(valor, 0.0)
        if numero > 0:
            return numero
    return 0.0


def _fecha_inicio(acta_inicio):
    return _parse_fecha(
        _primero_no_vacio(
            acta_inicio.get("fecha_inicio"),
            acta_inicio.get("fecha_presente_acta"),
        )
    )


def _datos_generales(acta_inicio, contrato_obra, contrato_interventoria):
    return {
        "contrato_no": _primero_no_vacio(acta_inicio.get("numero_contrato"), contrato_obra.get("numero_contrato")),
        "objeto": _primero_no_vacio(
            acta_inicio.get("objeto_contrato"),
            contrato_obra.get("objeto_general"),
            contrato_obra.get("objeto_contrato"),
            contrato_obra.get("objeto"),
        ),
        "contratista": _primero_no_vacio(
            acta_inicio.get("nombre_firma_contratista"),
            contrato_obra.get("nombre_contratista"),
            contrato_obra.get("contratista"),
        ),
        "contrato_interventoria_no": _primero_no_vacio(
            contrato_interventoria.get("numero_proceso_contratacion"),
            acta_inicio.get("numero_contrato_interventoria"),
        ),
        "interventor": _primero_no_vacio(
            acta_inicio.get("nombre_firma_interventor"),
            contrato_obra.get("nombre_interventor"),
            contrato_obra.get("nombre_supervisor"),
        ),
        "plazo_inicial": _primero_no_vacio(
            acta_inicio.get("plazo_ejecucion"),
            contrato_obra.get("plazo_ejecucion"),
        ),
        "fecha_inicio": _fecha_inicio(acta_inicio),
        "valor_inicial": _valor_contrato_obra(acta_inicio, contrato_obra),
    }


def _catalogo_presupuesto(presupuesto_obra):
    catalogo = []
    mapa = {}
    tablas = presupuesto_obra.get("__tablas__", {}) or {}
    grupos = tablas.get("grupos_presupuesto_obra", []) or []

    for grupo in grupos:
        if not isinstance(grupo, dict):
            continue
        capitulo = _texto(grupo.get("titulo"))
        for fila in grupo.get("rows", []) or []:
            if not isinstance(fila, dict):
                continue

            item = _texto(fila.get("ITEM"))
            if not item:
                continue

            cantidad = _safe_float(fila.get("CANT"), 0.0)
            valor_unitario = _safe_float(
                fila.get("VR AFECTADO POR FACTOR", fila.get("VR UNITARIO")),
                0.0,
            )
            valor_total = _safe_float(fila.get("VR TOTAL"), 0.0)

            registro = {
                "CAPITULO": capitulo,
                "No ITEM": item,
                "DESCRIPCIÓN ÍTEM": _texto(fila.get("DESCRIPCIÓN")),
                "FUENTE": _texto(fila.get("FUENTE")),
                "UND": _texto(fila.get("UNIDAD")),
                "CANTIDAD": cantidad,
                "VALOR UNITARIO": valor_unitario,
                "VALOR TOTAL": valor_total,
            }
            if item not in mapa:
                catalogo.append(registro)
                mapa[item] = registro

    catalogo = sorted(catalogo, key=lambda x: _key_codigo_natural(x["No ITEM"]))
    return catalogo, mapa


def _aiu(aiu_datos, presupuesto_obra):
    configuracion = presupuesto_obra.get("configuracion", {}) or {}
    return {
        "AIU %": _safe_float(aiu_datos.get("aiu_total_pct", configuracion.get("aiu_pct_global", 0.0)), 0.0),
        "UTILIDAD %": _safe_float(aiu_datos.get("utilidad_porcentaje", configuracion.get("aiu_utilidad_pct", 0.0)), 0.0),
    }


# ==========================================================
# Estado de actas
# ==========================================================
def _fila_item_vacia():
    return {
        "CAPITULO": "",
        "No ITEM": "",
        "DESCRIPCIÓN ÍTEM": "",
        "FUENTE": "",
        "UND": "",
        "CANTIDAD": 0.0,
        "VALOR UNITARIO": 0.0,
        "VALOR TOTAL": 0.0,
        "CANTIDAD ACTUALIZADA": 0.0,
        "VALOR ACTUALIZADO": 0.0,
        "CANTIDAD PRESENTE ACTA": 0.0,
        "VALOR PRESENTE ACTA": 0.0,
        "CANTIDAD ACTAS ANTERIORES": 0.0,
        "VALOR ACTAS ANTERIORES": 0.0,
        "CANTIDAD ACUMULADO": 0.0,
        "VALOR ACUMULADO": 0.0,
    }


def _normalizar_items(rows, mapa_catalogo, actas_anteriores, acta_no, tipo):
    acumulados_previos = _acumulados_acta_anterior(actas_anteriores, acta_no, tipo)
    filas = []

    for fila in rows or []:
        base = _fila_item_vacia()
        if isinstance(fila, dict):
            base.update(fila)

        item = _texto(base.get("No ITEM"))
        if item and item in mapa_catalogo:
            catalogo = mapa_catalogo[item]
            for campo in ["CAPITULO", "DESCRIPCIÓN ÍTEM", "FUENTE", "UND", "CANTIDAD"]:
                base[campo] = catalogo.get(campo, base.get(campo))
            if _safe_float(base.get("VALOR UNITARIO"), 0.0) <= 0:
                base["VALOR UNITARIO"] = catalogo.get("VALOR UNITARIO", 0.0)
            if _safe_float(base.get("VALOR TOTAL"), 0.0) <= 0:
                base["VALOR TOTAL"] = catalogo.get("VALOR TOTAL", 0.0)

        base["No ITEM"] = item
        for campo in [
            "CANTIDAD", "VALOR UNITARIO", "VALOR TOTAL", "CANTIDAD ACTUALIZADA", "VALOR ACTUALIZADO",
            "CANTIDAD PRESENTE ACTA", "VALOR PRESENTE ACTA",
        ]:
            base[campo] = _safe_float(base.get(campo), 0.0)

        anterior = acumulados_previos.get(item, {}) if item else {}
        base["CANTIDAD ACTAS ANTERIORES"] = _safe_float(anterior.get("CANTIDAD ACUMULADO"), 0.0)
        base["VALOR ACTAS ANTERIORES"] = _safe_float(anterior.get("VALOR ACUMULADO"), 0.0)
        base["CANTIDAD ACUMULADO"] = round(base["CANTIDAD PRESENTE ACTA"] + base["CANTIDAD ACTAS ANTERIORES"], 4)
        base["VALOR ACUMULADO"] = round(base["VALOR PRESENTE ACTA"] + base["VALOR ACTAS ANTERIORES"], 2)

        if item:
            filas.append(base)

    return filas


def _acumulados_acta_anterior(actas, acta_no, tipo):
    anteriores = [a for a in actas or [] if isinstance(a, dict) and int(a.get("acta_no") or 0) < int(acta_no or 1)]
    if not anteriores:
        return {}

    anterior = sorted(anteriores, key=lambda x: int(x.get("acta_no") or 0))[-1]
    rows = anterior.get("items_previstos" if tipo == "previstos" else "items_no_previstos", []) or []
    acumulados = {}
    for fila in rows:
        if isinstance(fila, dict):
            item = _texto(fila.get("No ITEM"))
            if item:
                acumulados[item] = {
                    "CANTIDAD ACUMULADO": _safe_float(fila.get("CANTIDAD ACUMULADO"), 0.0),
                    "VALOR ACUMULADO": _safe_float(fila.get("VALOR ACUMULADO"), 0.0),
                }
    return acumulados


def _acta_vacia(acta_no, generales, valor_anticipo, porcentaje_anticipo):
    return {
        "acta_no": int(acta_no),
        "fecha": date.today().isoformat(),
        "plazo_acumulado": generales.get("plazo_inicial", ""),
        "fecha_vencimiento_actual": date.today().isoformat(),
        "valor_adiciones": 0.0,
        "items_previstos": [],
        "items_no_previstos": [],
        "porcentaje_amortizado_presente_acta": 0.0,
        "valor_anticipo_otorgado": valor_anticipo,
        "porcentaje_anticipo": porcentaje_anticipo,
        "valor_en_letras": "",
        "observaciones": "",
    }


def _normalizar_acta(acta, acta_no, generales, actas, mapa_catalogo, valor_anticipo, porcentaje_anticipo):
    base = _acta_vacia(acta_no, generales, valor_anticipo, porcentaje_anticipo)
    if isinstance(acta, dict):
        base.update(acta)
    base["acta_no"] = int(base.get("acta_no") or acta_no or 1)
    base["fecha"] = _fecha_input(base.get("fecha")).isoformat()
    base["fecha_vencimiento_actual"] = _fecha_input(base.get("fecha_vencimiento_actual")).isoformat()
    base["valor_adiciones"] = _safe_float(base.get("valor_adiciones"), 0.0)
    base["items_previstos"] = _normalizar_items(base.get("items_previstos", []), mapa_catalogo, actas, base["acta_no"], "previstos")
    base["items_no_previstos"] = _normalizar_items(base.get("items_no_previstos", []), mapa_catalogo, actas, base["acta_no"], "no_previstos")
    base["porcentaje_amortizado_presente_acta"] = _safe_float(base.get("porcentaje_amortizado_presente_acta"), 0.0)
    base["valor_anticipo_otorgado"] = _safe_float(base.get("valor_anticipo_otorgado"), valor_anticipo)
    base["porcentaje_anticipo"] = _safe_float(base.get("porcentaje_anticipo"), porcentaje_anticipo)
    base["valor_en_letras"] = _texto(base.get("valor_en_letras"))
    base["observaciones"] = _texto(base.get("observaciones"))
    return base


def _normalizar_estado(cargado, generales, mapa_catalogo, valor_anticipo, porcentaje_anticipo):
    if not isinstance(cargado, dict):
        cargado = {}

    actas_originales = cargado.get("actas", []) if isinstance(cargado.get("actas"), list) else []
    if not actas_originales:
        actas_originales = [_acta_vacia(1, generales, valor_anticipo, porcentaje_anticipo)]

    actas = []
    for i, acta in enumerate(actas_originales, start=1):
        numero = int(acta.get("acta_no") or i) if isinstance(acta, dict) else i
        actas.append(_normalizar_acta(acta, numero, generales, actas_originales, mapa_catalogo, valor_anticipo, porcentaje_anticipo))

    actas = sorted(actas, key=lambda x: int(x.get("acta_no") or 0))
    activa = int(cargado.get("acta_activa") or actas[-1].get("acta_no") or 1)
    if activa not in [int(a.get("acta_no") or 0) for a in actas]:
        activa = int(actas[-1].get("acta_no") or 1)

    return {"actas": actas, "acta_activa": activa}


def _inicializar_estado(generales, mapa_catalogo, valor_anticipo, porcentaje_anticipo):
    group_id_actual = _texto(st.session_state.get("group_id"))
    cache_group = _texto(st.session_state.get("_recibo_parcial_group"))
    if cache_group != group_id_actual or "recibo_parcial_datos" not in st.session_state:
        cargado = cargar_estado(CLAVE_GUARDADO) or {}
        st.session_state["recibo_parcial_datos"] = _normalizar_estado(
            cargado,
            generales,
            mapa_catalogo,
            valor_anticipo,
            porcentaje_anticipo,
        )
        st.session_state["_recibo_parcial_group"] = group_id_actual


def _obtener_acta_activa():
    estado = st.session_state["recibo_parcial_datos"]
    activa = int(estado.get("acta_activa") or 1)
    for acta in estado.get("actas", []):
        if int(acta.get("acta_no") or 0) == activa:
            return acta
    return estado.get("actas", [])[0]


def _crear_nueva_acta(generales, valor_anticipo, porcentaje_anticipo):
    estado = st.session_state["recibo_parcial_datos"]
    actas = estado.get("actas", [])
    nuevo_numero = max([int(a.get("acta_no") or 0) for a in actas], default=0) + 1
    nueva = _acta_vacia(nuevo_numero, generales, valor_anticipo, porcentaje_anticipo)
    actas.append(nueva)
    estado["actas"] = actas
    estado["acta_activa"] = nuevo_numero
    return nuevo_numero


def _guardar():
    estado = st.session_state["recibo_parcial_datos"]

    actas_actuales = estado.get("actas", []) if isinstance(estado.get("actas"), list) else []
    guardado = cargar_estado(CLAVE_GUARDADO) or {}
    actas_guardadas = guardado.get("actas", []) if isinstance(guardado, dict) and isinstance(guardado.get("actas"), list) else []

    if actas_guardadas and len(actas_actuales) < len(actas_guardadas):
        st.error(
            "No se guardó el acta porque el estado actual tiene menos actas que el estado ya guardado. "
            "Esto evita sobrescribir actas anteriores."
        )
        return False

    guardar_estado(CLAVE_GUARDADO, estado)
    st.success("Acta de recibo parcial guardada correctamente.")
    return True


# ==========================================================
# Cálculos del acta
# ==========================================================
def _total_presente_acta(acta):
    rows = (acta.get("items_previstos", []) or []) + (acta.get("items_no_previstos", []) or [])
    return round(sum(_safe_float(f.get("VALOR PRESENTE ACTA"), 0.0) for f in rows if isinstance(f, dict)), 2)


def _amortizado_acumulado_anterior(actas, acta_no):
    anteriores = [a for a in actas or [] if isinstance(a, dict) and int(a.get("acta_no") or 0) < int(acta_no or 1)]
    if not anteriores:
        return 0.0
    anterior = sorted(anteriores, key=lambda x: int(x.get("acta_no") or 0))[-1]
    return _safe_float(anterior.get("valor_amortizado_acumulado_total"), 0.0)


def _recalcular_resumen(acta, actas):
    valor_total = _total_presente_acta(acta)
    valor_anticipo = _safe_float(acta.get("valor_anticipo_otorgado"), 0.0)
    pct_amortizado = _safe_float(acta.get("porcentaje_amortizado_presente_acta"), 0.0)
    amortizado_presente = round(valor_anticipo * pct_amortizado / 100.0, 2)
    amortizado_acumulado = _amortizado_acumulado_anterior(actas, acta.get("acta_no"))
    saldo_amortizar = round(valor_anticipo - amortizado_presente - amortizado_acumulado, 2)
    valor_con_amortizacion = round(valor_total - amortizado_presente, 2)

    acta["valor_total_obra_ejecutada"] = valor_total
    acta["valor_amortizado_presente_acta"] = amortizado_presente
    acta["valor_amortizado_acumulado"] = amortizado_acumulado
    acta["valor_amortizado_acumulado_total"] = round(amortizado_presente + amortizado_acumulado, 2)
    acta["valor_saldo_por_amortizar"] = saldo_amortizar
    acta["valor_basico_sin_amortizacion"] = valor_total
    acta["valor_basico_con_amortizacion"] = valor_con_amortizacion
    acta["valor_total_pagar_presente_acta"] = valor_con_amortizacion
    return acta


# ==========================================================
# Word helpers
# ==========================================================
def _set_doc_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(7)


def _p(doc, texto="", bold=False, align=None, size=8):
    p = doc.add_paragraph()
    run = p.add_run(str(texto))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def _cell_text(cell, texto, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(texto))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(6)
    if align:
        p.alignment = align


def _tabla_simple(doc, filas):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for etiqueta, valor in filas:
        cells = table.add_row().cells
        _cell_text(cells[0], etiqueta, bold=True)
        _cell_text(cells[1], valor)
    return table


def _tabla_dataframe(doc, df):
    if df is None or df.empty:
        _p(doc, "Sin registros.")
        return None
    df_out = df.copy()
    for col in df_out.columns:
        df_out[col] = df_out[col].apply(lambda x: _fecha_texto(x) if isinstance(x, (date, datetime, pd.Timestamp)) else x)
    table = doc.add_table(rows=1, cols=len(df_out.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(df_out.columns):
        _cell_text(table.rows[0].cells[i], col, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _, row in df_out.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df_out.columns):
            valor = row[col]
            if isinstance(valor, float) and ("VALOR" in str(col).upper() or "TOTAL" in str(col).upper()):
                valor = _moneda(valor)
            _cell_text(cells[i], valor, align=WD_ALIGN_PARAGRAPH.CENTER)
    return table


def _generar_word(acta, generales, aiu_info):
    doc = Document()
    _set_doc_defaults(doc)

    valor_acumulado_contrato = _safe_float(generales.get("valor_inicial"), 0.0) + _safe_float(acta.get("valor_adiciones"), 0.0)

    _p(doc, "ACTA DE RECIBO PARCIAL DE OBRA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    _p(doc, f"ACTA No. {int(acta.get('acta_no') or 0)}    FECHA: {_fecha_texto(acta.get('fecha'))}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "")
    _tabla_simple(
        doc,
        [
            ("CONTRATO No.", generales.get("contrato_no", "")),
            ("OBJETO", generales.get("objeto", "")),
            ("CONTRATISTA", generales.get("contratista", "")),
            ("CONTRATO DE INTERVENTORÍA No.", generales.get("contrato_interventoria_no", "")),
            ("INTERVENTOR", generales.get("interventor", "")),
            ("PLAZO INICIAL", generales.get("plazo_inicial", "")),
            ("PLAZO ACUMULADO", acta.get("plazo_acumulado", "")),
            ("FECHA DE INICIO DEL CONTRATO", _fecha_texto(generales.get("fecha_inicio"))),
            ("FECHA DE VENCIMIENTO ACTUAL", _fecha_texto(acta.get("fecha_vencimiento_actual"))),
            ("VALOR INICIAL", _moneda(generales.get("valor_inicial", 0.0))),
            ("VALOR ADICIONES", _moneda(acta.get("valor_adiciones", 0.0))),
            ("VALOR ACUMULADO", _moneda(valor_acumulado_contrato)),
        ],
    )

    columnas_items = [
        "CAPITULO", "No ITEM", "DESCRIPCIÓN ÍTEM", "FUENTE", "UND", "CANTIDAD",
        "VALOR UNITARIO", "VALOR TOTAL", "CANTIDAD ACTUALIZADA", "VALOR ACTUALIZADO",
        "CANTIDAD PRESENTE ACTA", "VALOR PRESENTE ACTA", "CANTIDAD ACTAS ANTERIORES",
        "VALOR ACTAS ANTERIORES", "CANTIDAD ACUMULADO", "VALOR ACUMULADO",
    ]

    _p(doc, "")
    _p(doc, "OBRA EJECUTADA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "ÍTEMS PREVISTOS", bold=True)
    _tabla_dataframe(doc, pd.DataFrame(acta.get("items_previstos", []), columns=columnas_items))
    _p(doc, "")
    _p(doc, "ÍTEMS NO PREVISTOS", bold=True)
    _tabla_dataframe(doc, pd.DataFrame(acta.get("items_no_previstos", []), columns=columnas_items))

    _p(doc, "")
    _tabla_simple(doc, [("VALOR TOTAL", _moneda(acta.get("valor_total_obra_ejecutada", 0.0)))])

    _p(doc, "")
    _tabla_simple(
        doc,
        [
            ("AIU", f"{_safe_float(aiu_info.get('AIU %'), 0.0):.4f} %"),
            ("Utilidad", f"{_safe_float(aiu_info.get('UTILIDAD %'), 0.0):.4f} %"),
            ("Valor Anticipo Otorgado", _moneda(acta.get("valor_anticipo_otorgado", 0.0))),
            ("Porcentaje Anticipo", f"{_safe_float(acta.get('porcentaje_anticipo'), 0.0):.4f} %"),
            ("Valor Amortizado Presente Acta", _moneda(acta.get("valor_amortizado_presente_acta", 0.0))),
            ("Porcentaje Amortizado Presente Acta", f"{_safe_float(acta.get('porcentaje_amortizado_presente_acta'), 0.0):.4f} %"),
            ("Valor Amortizado Acumulado", _moneda(acta.get("valor_amortizado_acumulado", 0.0))),
            ("Valor Saldo por Amortizar", _moneda(acta.get("valor_saldo_por_amortizar", 0.0))),
        ],
    )

    _p(doc, "")
    _p(doc, "RESUMEN PRESENTE MES PARA PAGO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _tabla_simple(
        doc,
        [
            ("VALOR BÁSICO PRESENTE MES SIN AMORTIZACIÓN DE ANTICIPO", _moneda(acta.get("valor_basico_sin_amortizacion", 0.0))),
            ("VALOR BÁSICO PRESENTE MES CON AMORTIZACIÓN DE ANTICIPO", _moneda(acta.get("valor_basico_con_amortizacion", 0.0))),
            ("VALOR TOTAL A PAGAR PRESENTE ACTA", _moneda(acta.get("valor_total_pagar_presente_acta", 0.0))),
            ("VALOR EN LETRAS", acta.get("valor_en_letras", "")),
        ],
    )

    _p(doc, "")
    _p(doc, "OBSERVACIONES", bold=True)
    _p(doc, acta.get("observaciones", ""))

    _p(doc, "")
    _p(doc, "NOTAS:", bold=True)
    notas = [
        "El Interventor certifica que revisó y verificó el pago de la Seguridad Social y Aportes de Ley por parte del Contratista, con base en los soportes presentados por éste y el cumplimiento por parte del Contratista de la normatividad vigente relacionada con la nómina electrónica cuando aplique.",
        "El Interventor y el Contratista asumen plena responsabilidad por la veracidad de la información, los valores, así como por las operaciones aritméticas contenidas en esta Acta.",
        "El pago de la presente acta se realizará con base en el principio de confiabilidad y responsabilidad con la que la Interventoría y el Contratista realizan las mediciones en campo y las revisiones efectuadas por parte de la interventoría.",
        "Con la suscripción de la presente Acta el Interventor y el Contratista certifican que las obras recibidas cumplen con los requerimientos de calidad, con las normas, especificaciones generales y particulares de construcción y demás condiciones contractuales, de acuerdo con los diseños, planos y especificaciones estipuladas para el proyecto.",
        "El valor a pagar en la presente Acta debe coincidir con el valor consignado en la factura presentada por el Contratista después de amortizar el anticipo cuando sea del caso.",
        "Con la suscripción de la presente Acta el Interventor deja constancia de que el Contratista cumplió la normatividad y gestión ambiental, calidad de los materiales y el PMT.",
    ]
    for i, nota in enumerate(notas, start=1):
        _p(doc, f"{i}. {nota}")

    _p(doc, "")
    _p(doc, "Para constancia de lo anterior firman", bold=True)
    _p(doc, "")
    tabla_firmas = doc.add_table(rows=2, cols=2)
    tabla_firmas.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cell_text(tabla_firmas.rows[0].cells[0], "Contratista", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(tabla_firmas.rows[0].cells[1], "Interventor", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(tabla_firmas.rows[1].cells[0], generales.get("contratista", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(tabla_firmas.rows[1].cells[1], generales.get("interventor", ""), align=WD_ALIGN_PARAGRAPH.CENTER)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==========================================================
# Interfaz
# ==========================================================
st.set_page_config(page_title="Acta de recibo parcial de obra", layout="wide")
st.title("30. Acta de recibo parcial de obra")

acta_inicio = _leer_estado("acta_inicio_obra")
contrato_obra = _leer_estado("contrato_obra")
contrato_interventoria = _leer_estado("contrato_interventoria")
presupuesto_obra = _leer_estado("presupuesto_obra")
aiu_datos = _leer_estado("aiu")
plan_anticipo = _leer_estado("plan_inversion_anticipo")

generales = _datos_generales(acta_inicio, contrato_obra, contrato_interventoria)
valor_anticipo = _valor_anticipo_acta_inicio(acta_inicio, plan_anticipo, generales.get("valor_inicial", 0.0))
porcentaje_anticipo = _porcentaje_anticipo_acta_inicio(acta_inicio, plan_anticipo)
catalogo, mapa_catalogo = _catalogo_presupuesto(presupuesto_obra)
aiu_info = _aiu(aiu_datos, presupuesto_obra)

_inicializar_estado(generales, mapa_catalogo, valor_anticipo, porcentaje_anticipo)
estado = st.session_state["recibo_parcial_datos"]

st.markdown("### ACTAS GENERADAS")
col_nueva, col_selector = st.columns([1, 3])
with col_nueva:
    if st.button("➕ Nueva acta", key="recibo_parcial_nueva_acta"):
        nuevo = _crear_nueva_acta(generales, valor_anticipo, porcentaje_anticipo)
        _guardar()
        st.session_state["recibo_parcial_selector"] = nuevo
        st.rerun()

actas = estado.get("actas", [])
opciones_actas = [int(a.get("acta_no") or 0) for a in actas]
activa = int(estado.get("acta_activa") or (opciones_actas[-1] if opciones_actas else 1))
with col_selector:
    seleccion = st.selectbox(
        "Seleccione acta",
        options=opciones_actas,
        index=opciones_actas.index(activa) if activa in opciones_actas else 0,
        format_func=lambda x: f"Acta No. {x}",
        key="recibo_parcial_selector",
    )
estado["acta_activa"] = int(seleccion)
acta = _obtener_acta_activa()
acta_no = int(acta.get("acta_no") or seleccion)

st.markdown("### ENCABEZADO")
col_a, col_b = st.columns(2)
with col_a:
    st.text_input("ACTA No.", value=str(acta_no), disabled=True)
with col_b:
    acta["fecha"] = st.date_input(
        "FECHA",
        value=_fecha_input(acta.get("fecha")),
        format="DD/MM/YYYY",
        key=f"recibo_fecha_{acta_no}",
    ).isoformat()

c1, c2 = st.columns(2)
with c1:
    st.text_input("CONTRATO No.", value=generales.get("contrato_no", ""), disabled=True)
    st.text_area("OBJETO", value=generales.get("objeto", ""), disabled=True, height=90)
    st.text_input("CONTRATISTA", value=generales.get("contratista", ""), disabled=True)
    st.text_input("CONTRATO DE INTERVENTORÍA No.", value=generales.get("contrato_interventoria_no", ""), disabled=True)
    st.text_input("INTERVENTOR", value=generales.get("interventor", ""), disabled=True)
with c2:
    st.text_input("PLAZO INICIAL", value=generales.get("plazo_inicial", ""), disabled=True)
    acta["plazo_acumulado"] = st.text_input("PLAZO ACUMULADO", value=_texto(acta.get("plazo_acumulado")), key=f"recibo_plazo_acumulado_{acta_no}")
    st.date_input("FECHA DE INICIO DEL CONTRATO", value=_fecha_input(generales.get("fecha_inicio")), disabled=True, format="DD/MM/YYYY")
    acta["fecha_vencimiento_actual"] = st.date_input(
        "FECHA DE VENCIMIENTO ACTUAL",
        value=_fecha_input(acta.get("fecha_vencimiento_actual")),
        format="DD/MM/YYYY",
        key=f"recibo_fecha_vencimiento_{acta_no}",
    ).isoformat()
    st.number_input("VALOR INICIAL", value=float(generales.get("valor_inicial", 0.0)), disabled=True, format="%.2f")
    acta["valor_adiciones"] = st.number_input(
        "VALOR ADICIONES",
        value=float(acta.get("valor_adiciones", 0.0)),
        format="%.2f",
        key=f"recibo_valor_adiciones_{acta_no}",
    )
    st.number_input(
        "VALOR ACUMULADO",
        value=float(generales.get("valor_inicial", 0.0)) + float(acta.get("valor_adiciones", 0.0)),
        disabled=True,
        format="%.2f",
    )

st.markdown("### OBRA EJECUTADA")
items_opciones = [""] + [x["No ITEM"] for x in catalogo]
columnas_items = [
    "CAPITULO", "No ITEM", "DESCRIPCIÓN ÍTEM", "FUENTE", "UND", "CANTIDAD",
    "VALOR UNITARIO", "VALOR TOTAL", "CANTIDAD ACTUALIZADA", "VALOR ACTUALIZADO",
    "CANTIDAD PRESENTE ACTA", "VALOR PRESENTE ACTA", "CANTIDAD ACTAS ANTERIORES",
    "VALOR ACTAS ANTERIORES", "CANTIDAD ACUMULADO", "VALOR ACUMULADO",
]


def _editor_items(titulo, key, rows, tipo):
    st.markdown(f"#### {titulo}")
    df = pd.DataFrame(rows, columns=columnas_items)
    editado = st.data_editor(
        df,
        hide_index=True,
        width="stretch",
        num_rows="dynamic",
        key=key,
        disabled=[
            "CAPITULO", "DESCRIPCIÓN ÍTEM", "FUENTE", "UND", "CANTIDAD",
            "CANTIDAD ACTAS ANTERIORES", "VALOR ACTAS ANTERIORES",
            "CANTIDAD ACUMULADO", "VALOR ACUMULADO",
        ],
        column_config={
            "No ITEM": st.column_config.SelectboxColumn("No ITEM", options=items_opciones, required=False),
            "VALOR UNITARIO": st.column_config.NumberColumn("VALOR UNITARIO", format="$ %.2f"),
            "VALOR TOTAL": st.column_config.NumberColumn("VALOR TOTAL", format="$ %.2f"),
            "CANTIDAD ACTUALIZADA": st.column_config.NumberColumn("CANTIDAD ACTUALIZADA", format="%.4f"),
            "VALOR ACTUALIZADO": st.column_config.NumberColumn("VALOR ACTUALIZADO", format="$ %.2f"),
            "CANTIDAD PRESENTE ACTA": st.column_config.NumberColumn("CANTIDAD PRESENTE ACTA", format="%.4f"),
            "VALOR PRESENTE ACTA": st.column_config.NumberColumn("VALOR PRESENTE ACTA", format="$ %.2f"),
        },
    )
    return _normalizar_items(editado.to_dict("records"), mapa_catalogo, estado.get("actas", []), acta_no, tipo)


items_previstos_actualizados = _editor_items(
    "ÍTEMS PREVISTOS",
    f"recibo_items_previstos_{acta_no}",
    acta.get("items_previstos", []),
    "previstos",
)

items_no_previstos_actualizados = _editor_items(
    "ÍTEMS NO PREVISTOS",
    f"recibo_items_no_previstos_{acta_no}",
    acta.get("items_no_previstos", []),
    "no_previstos",
)

if (
    items_previstos_actualizados != acta.get("items_previstos", [])
    or items_no_previstos_actualizados != acta.get("items_no_previstos", [])
):
    acta["items_previstos"] = items_previstos_actualizados
    acta["items_no_previstos"] = items_no_previstos_actualizados
    _recalcular_resumen(acta, estado.get("actas", []))
    st.rerun()

acta["items_previstos"] = items_previstos_actualizados
acta["items_no_previstos"] = items_no_previstos_actualizados
_recalcular_resumen(acta, estado.get("actas", []))

st.markdown("### VALOR TOTAL")
st.number_input("VALOR TOTAL", value=float(acta.get("valor_total_obra_ejecutada", 0.0)), disabled=True, format="%.2f")

st.markdown("### AIU Y ANTICIPO")
col_aiu1, col_aiu2 = st.columns(2)
with col_aiu1:
    st.number_input("AIU %", value=float(aiu_info.get("AIU %", 0.0)), disabled=True, format="%.4f")
with col_aiu2:
    st.number_input("Utilidad %", value=float(aiu_info.get("UTILIDAD %", 0.0)), disabled=True, format="%.4f")

col_ant1, col_ant2 = st.columns(2)
with col_ant2:
    st.number_input("Porcentaje Anticipo", value=float(acta.get("porcentaje_anticipo", 0.0)), disabled=True, format="%.4f")
    nuevo_pct_amortizado = st.number_input(
        "Porcentaje amortizado presente acta",
        value=float(acta.get("porcentaje_amortizado_presente_acta", 0.0)),
        format="%.4f",
        key=f"recibo_pct_amortizado_{acta_no}",
    )

if nuevo_pct_amortizado != float(acta.get("porcentaje_amortizado_presente_acta", 0.0)):
    acta["porcentaje_amortizado_presente_acta"] = nuevo_pct_amortizado
    _recalcular_resumen(acta, estado.get("actas", []))
    st.rerun()

acta["porcentaje_amortizado_presente_acta"] = nuevo_pct_amortizado
_recalcular_resumen(acta, estado.get("actas", []))

with col_ant1:
    st.number_input("Valor Anticipo Otorgado", value=float(acta.get("valor_anticipo_otorgado", 0.0)), disabled=True, format="%.2f")
    st.number_input("Valor Amortizado Presente Acta", value=float(acta.get("valor_amortizado_presente_acta", 0.0)), disabled=True, format="%.2f")
    st.number_input("Valor Amortizado Acumulado", value=float(acta.get("valor_amortizado_acumulado", 0.0)), disabled=True, format="%.2f")

with col_ant2:
    st.number_input("Valor Saldo por Amortizar", value=float(acta.get("valor_saldo_por_amortizar", 0.0)), disabled=True, format="%.2f")

st.markdown("### RESUMEN PRESENTE MES PARA PAGO")
col_r1, col_r2 = st.columns(2)
with col_r1:
    st.number_input("VALOR BÁSICO PRESENTE MES SIN AMORTIZACIÓN DE ANTICIPO", value=float(acta.get("valor_basico_sin_amortizacion", 0.0)), disabled=True, format="%.2f")
    st.number_input("VALOR BÁSICO PRESENTE MES CON AMORTIZACIÓN DE ANTICIPO", value=float(acta.get("valor_basico_con_amortizacion", 0.0)), disabled=True, format="%.2f")
with col_r2:
    st.number_input("VALOR TOTAL A PAGAR PRESENTE ACTA", value=float(acta.get("valor_total_pagar_presente_acta", 0.0)), disabled=True, format="%.2f")
    acta["valor_en_letras"] = st.text_input("VALOR EN LETRAS", value=_texto(acta.get("valor_en_letras")), key=f"recibo_valor_letras_{acta_no}")

st.markdown("### OBSERVACIONES")
acta["observaciones"] = st.text_area(
    "OBSERVACIONES",
    value=_texto(acta.get("observaciones")),
    height=140,
    label_visibility="collapsed",
    key=f"recibo_observaciones_{acta_no}",
)

st.markdown("### NOTAS")
st.write("1. El Interventor certifica que revisó y verificó el pago de la Seguridad Social y Aportes de Ley por parte del Contratista, con base en los soportes presentados por éste y el cumplimiento por parte del Contratista de la normatividad vigente relacionada con la nómina electrónica cuando aplique.")
st.write("2. El Interventor y el Contratista asumen plena responsabilidad por la veracidad de la información, los valores, así como por las operaciones aritméticas contenidas en esta Acta.")
st.write("3. El pago de la presente acta se realizará con base en el principio de confiabilidad y responsabilidad con la que la Interventoría y el Contratista realizan las mediciones en campo y las revisiones efectuadas por parte de la interventoría.")
st.write("4. Con la suscripción de la presente Acta el Interventor y el Contratista certifican que las obras recibidas cumplen con los requerimientos de calidad, con las normas, especificaciones generales y particulares de construcción y demás condiciones contractuales, de acuerdo con los diseños, planos y especificaciones estipuladas para el proyecto.")
st.write("5. El valor a pagar en la presente Acta debe coincidir con el valor consignado en la factura presentada por el Contratista después de amortizar el anticipo cuando sea del caso.")
st.write("6. Con la suscripción de la presente Acta el Interventor deja constancia de que el Contratista cumplió la normatividad y gestión ambiental, calidad de los materiales y el PMT.")

st.markdown("### FIRMAS")
df_firmas = pd.DataFrame([{"Contratista": generales.get("contratista", ""), "Interventor": generales.get("interventor", "")}])
st.dataframe(df_firmas, hide_index=True, width="stretch")

# Actualizar acta dentro del estado
for idx, existente in enumerate(estado.get("actas", [])):
    if int(existente.get("acta_no") or 0) == acta_no:
        estado["actas"][idx] = acta
        break

col_guardar, col_word = st.columns(2)
with col_guardar:
    if st.button("💾 Guardar acta", key=f"recibo_guardar_{acta_no}"):
        _guardar()

with col_word:
    word = _generar_word(acta, generales, aiu_info)
    st.download_button(
        "📄 Descargar acta en Word",
        data=word,
        file_name=f"acta_recibo_parcial_obra_{acta_no}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        width="stretch",
        key=f"recibo_word_{acta_no}",
    )
