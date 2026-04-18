import base64
from io import BytesIO
from datetime import date, datetime, timedelta

import streamlit as st
from docx import Document
from docx.shared import Inches

from supabase_state import cargar_estado
from supabase_state import guardar_estado as guardar_estado_bd


def guardar_estado(clave, datos):
    def serializar(obj):
        if isinstance(obj, dict):
            return {k: serializar(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [serializar(x) for x in obj]
        if isinstance(obj, (date, datetime)):
            return obj.isoformat()
        return obj

    guardar_estado_bd(clave, serializar(datos))


def _texto(valor):
    if valor is None:
        return ""
    return str(valor).strip()


def _parse_fecha(valor):
    if isinstance(valor, date):
        return valor
    if isinstance(valor, datetime):
        return valor.date()
    if isinstance(valor, str) and valor.strip():
        txt = valor.strip()
        try:
            return datetime.fromisoformat(txt).date()
        except Exception:
            pass
        for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%m/%d/%Y"):
            try:
                return datetime.strptime(txt, fmt).date()
            except Exception:
                continue
    return None


def _fecha_a_texto(valor):
    f = _parse_fecha(valor)
    if not f:
        return _texto(valor)
    return f.strftime("%d/%m/%Y")


def _leer_contrato_obra():
    datos = st.session_state.get("contrato_obra_datos", {})
    if not isinstance(datos, dict) or not datos:
        datos = cargar_estado("contrato_obra") or {}
    return datos if isinstance(datos, dict) else {}


def _leer_acta_inicio():
    datos = st.session_state.get("acta_inicio_obra_datos", {})
    if not isinstance(datos, dict) or not datos:
        datos = cargar_estado("acta_inicio_obra") or {}
    return datos if isinstance(datos, dict) else {}


def _leer_contrato_interventoria():
    datos = st.session_state.get("contrato_interventoria_datos", {})
    if not isinstance(datos, dict) or not datos:
        datos = cargar_estado("contrato_interventoria") or {}
    return datos if isinstance(datos, dict) else {}


def _primero_no_vacio(*valores):
    for valor in valores:
        if isinstance(valor, str):
            if valor.strip():
                return valor.strip()
        elif valor not in (None, "", []):
            return valor
    return ""


def _partes_fecha(valor):
    f = _parse_fecha(valor)
    if not f:
        return "", "", ""
    return str(f.day), str(f.month), str(f.year)


def _mostrar_imagen_guardada(info):
    if not isinstance(info, dict):
        return
    b64 = info.get("data")
    if not b64:
        return
    try:
        contenido = base64.b64decode(b64)
        st.image(contenido, use_container_width=True)
    except Exception:
        st.warning("No fue posible mostrar la imagen guardada.")


def _guardar_imagen_desde_uploader(uploaded_file):
    if uploaded_file is None:
        return None

    contenido = uploaded_file.read()
    if not contenido:
        return None

    return {
        "name": uploaded_file.name,
        "mime": uploaded_file.type or "image/png",
        "data": base64.b64encode(contenido).decode("utf-8"),
    }


def _inicializar_estado():
    if "reunion_tecnica_inicial_datos" not in st.session_state:
        st.session_state["reunion_tecnica_inicial_datos"] = cargar_estado("reunion_tecnica_inicial") or {}

    datos = st.session_state["reunion_tecnica_inicial_datos"]

    defaults = {
        "fecha_reunion": date.today(),
        "lugar_reunion": "",
        "estudios_disenos_contratista": "",
        "apu_contratista": "",
        "programa_obra": "",
        "programa_inversiones": "",
        "plan_inversion_anticipo": "",
        "flujo_caja": "",
        "requisitos_ambientales": "",
        "requisitos_sociales": "",
        "estudios_disenos_interventoria": "",
        "aprobacion_apu": "",
        "aprobacion_programa_obra": "",
        "aprobacion_programa_inversiones": "",
        "fecha_apertura_bitacora": None,
        "anexo_imagen_1": None,
        "anexo_imagen_2": None,
    }

    for k, v in defaults.items():
        if k not in datos:
            datos[k] = v

    datos["fecha_reunion"] = _parse_fecha(datos.get("fecha_reunion")) or date.today()
    datos["fecha_apertura_bitacora"] = _parse_fecha(datos.get("fecha_apertura_bitacora"))


def _guardar():
    guardar_estado("reunion_tecnica_inicial", st.session_state["reunion_tecnica_inicial_datos"])
    st.success("Acta de reunión técnica inicial guardada correctamente.")

def _agregar_imagen_docx(doc, info, titulo):
    if not isinstance(info, dict):
        return

    b64 = info.get("data")
    if not b64:
        return

    try:
        contenido = base64.b64decode(b64)
        doc.add_paragraph(titulo)
        doc.add_picture(BytesIO(contenido), width=Inches(5.5))
    except Exception:
        doc.add_paragraph(f"{titulo}: no fue posible insertar la imagen.")


def _generar_word_reunion_tecnica(
    datos,
    entidad_contratante,
    director_obra,
    director_interventoria,
    supervisor_interventoria,
    contrato_obra_no,
    fecha_contrato_obra,
    contrato_interventoria_no,
    fecha_contrato_interventoria,
    objeto_contrato_obra,
    plazo_contrato_obra,
    dia_ini,
    mes_ini,
    anio_ini,
    dia_fin,
    mes_fin,
    anio_fin,
    valor_total_contrato_obra,
    contratista_obra,
    interventor_obra,
):
    doc = Document()

    doc.add_heading("ACTA DE REUNIÓN TÉCNICA INICIAL", level=1)

    doc.add_paragraph(f"Fecha: {_fecha_a_texto(datos.get('fecha_reunion'))}")
    doc.add_paragraph(f"Entidad contratante: {_texto(entidad_contratante)}")
    doc.add_paragraph(f"Lugar de reunión: {_texto(datos.get('lugar_reunion', ''))}")
    doc.add_paragraph(f"Director de obra: {_texto(director_obra)}")
    doc.add_paragraph(f"Director de interventoría: {_texto(director_interventoria)}")
    doc.add_paragraph(f"Supervisor del contrato de interventoría: {_texto(supervisor_interventoria)}")

    doc.add_paragraph("")
    doc.add_paragraph(f"Contrato de obra No.: {_texto(contrato_obra_no)}")
    doc.add_paragraph(f"Fecha contrato de obra: {_fecha_a_texto(fecha_contrato_obra)}")
    doc.add_paragraph(f"Contrato de interventoría No.: {_texto(contrato_interventoria_no)}")
    doc.add_paragraph(f"Fecha contrato de interventoría: {_fecha_a_texto(fecha_contrato_interventoria)}")

    doc.add_paragraph("")
    doc.add_paragraph(f"Objeto del contrato de obra: {_texto(objeto_contrato_obra)}")
    doc.add_paragraph(f"Plazo del contrato de obra: {_texto(plazo_contrato_obra)}")
    doc.add_paragraph(f"Fecha de inicio del contrato de obra: {dia_ini}/{mes_ini}/{anio_ini}")
    doc.add_paragraph(f"Fecha de vencimiento del contrato de obra: {dia_fin}/{mes_fin}/{anio_fin}")
    doc.add_paragraph(f"Valor total contrato de obra: {_texto(valor_total_contrato_obra)}")
    doc.add_paragraph(f"Contratista de obra: {_texto(contratista_obra)}")
    doc.add_paragraph(f"Interventor: {_texto(interventor_obra)}")

    doc.add_paragraph("")
    doc.add_paragraph("COMPROMISOS RELACIONADOS CON EL CONTRATISTA")
    doc.add_paragraph(f"Estudios y diseños: {_texto(datos.get('estudios_disenos_contratista', ''))}")
    doc.add_paragraph(f"Análisis de precios unitarios: {_texto(datos.get('apu_contratista', ''))}")
    doc.add_paragraph(f"Programa de obra: {_texto(datos.get('programa_obra', ''))}")
    doc.add_paragraph(f"Programa de inversiones: {_texto(datos.get('programa_inversiones', ''))}")
    doc.add_paragraph(f"Plan de inversión del anticipo: {_texto(datos.get('plan_inversion_anticipo', ''))}")
    doc.add_paragraph(f"Flujo de caja: {_texto(datos.get('flujo_caja', ''))}")
    doc.add_paragraph(f"Requisitos ambientales: {_texto(datos.get('requisitos_ambientales', ''))}")
    doc.add_paragraph(f"Requisitos sociales: {_texto(datos.get('requisitos_sociales', ''))}")

    doc.add_paragraph("")
    doc.add_paragraph("COMPROMISOS RELACIONADOS CON LA INTERVENTORÍA")
    doc.add_paragraph(f"Estudios y/o diseños: {_texto(datos.get('estudios_disenos_interventoria', ''))}")
    doc.add_paragraph(f"Aprobación análisis de precios unitarios: {_texto(datos.get('aprobacion_apu', ''))}")
    doc.add_paragraph(f"Aprobación programa de obra: {_texto(datos.get('aprobacion_programa_obra', ''))}")
    doc.add_paragraph(f"Aprobación programa de inversiones: {_texto(datos.get('aprobacion_programa_inversiones', ''))}")

    doc.add_paragraph("")
    doc.add_paragraph(f"Fecha de apertura de bitácora: {_fecha_a_texto(datos.get('fecha_apertura_bitacora'))}")

    doc.add_paragraph("")
    doc.add_paragraph("ANEXOS")
    _agregar_imagen_docx(doc, datos.get("anexo_imagen_1"), "Anexo imagen 1")
    _agregar_imagen_docx(doc, datos.get("anexo_imagen_2"), "Anexo imagen 2")

    doc.add_paragraph("")
    tabla_firmas = doc.add_table(rows=2, cols=3)
    tabla_firmas.style = "Table Grid"
    tabla_firmas.cell(0, 0).text = "DIRECTOR DE OBRA"
    tabla_firmas.cell(0, 1).text = "DIRECTOR DE INTERVENTORÍA"
    tabla_firmas.cell(0, 2).text = "SUPERVISOR"
    tabla_firmas.cell(1, 0).text = _texto(director_obra)
    tabla_firmas.cell(1, 1).text = _texto(director_interventoria)
    tabla_firmas.cell(1, 2).text = _texto(supervisor_interventoria)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def _bloque_texto_etiqueta(etiqueta, valor, key, disabled=True, altura=68):
    st.markdown(f"**{etiqueta}**")
    st.text_area(
        etiqueta,
        value=_texto(valor),
        key=key,
        height=altura,
        label_visibility="collapsed",
        disabled=disabled,
    )


def _bloque_texto_editable(etiqueta, clave_estado, altura=120):
    datos = st.session_state["reunion_tecnica_inicial_datos"]
    st.markdown(f"**{etiqueta}**")
    datos[clave_estado] = st.text_area(
        etiqueta,
        value=_texto(datos.get(clave_estado, "")),
        key=clave_estado,
        height=altura,
        label_visibility="collapsed",
    )


_inicializar_estado()

datos = st.session_state["reunion_tecnica_inicial_datos"]
contrato_obra = _leer_contrato_obra()
acta_inicio = _leer_acta_inicio()
contrato_interventoria = _leer_contrato_interventoria()

plazo_dias = int(contrato_obra.get("plazo_ejecucion_dias", 0) or 0)

fecha_inicio_obra = _primero_no_vacio(
    acta_inicio.get("fecha_presente_acta"),
    contrato_obra.get("fecha_inicio"),
    contrato_obra.get("fecha_inicio_contrato"),
)

fecha_fin_obra = None
if _parse_fecha(fecha_inicio_obra) and plazo_dias > 0:
    fecha_fin_obra = _parse_fecha(fecha_inicio_obra) + timedelta(days=plazo_dias)

entidad_contratante = _primero_no_vacio(
    acta_inicio.get("entidad_contratante"),
    contrato_obra.get("nombre_entidad"),
)

director_obra = _primero_no_vacio(
    acta_inicio.get("nombre_firma_contratista"),
    contrato_obra.get("nombre_contratista"),
)

director_interventoria = _primero_no_vacio(
    acta_inicio.get("nombre_firma_interventor"),
    contrato_interventoria.get("nombre_interventor"),
    contrato_interventoria.get("firmante_interventor"),
    contrato_obra.get("nombre_interventor"),
)

supervisor_interventoria = _primero_no_vacio(
    acta_inicio.get("nombre_firma_supervisor"),
    contrato_interventoria.get("firmante_entidad"),
    contrato_obra.get("nombre_supervisor"),
)

contrato_obra_no = _primero_no_vacio(
    acta_inicio.get("numero_contrato"),
    contrato_obra.get("numero_contrato"),
)

fecha_contrato_obra = _primero_no_vacio(
    acta_inicio.get("fecha_contrato_obra"),
    contrato_obra.get("fecha_suscripcion"),
    contrato_obra.get("fecha_contrato"),
)

contrato_interventoria_no = _primero_no_vacio(
    contrato_interventoria.get("numero_proceso_contratacion"),
)

fecha_contrato_interventoria = _primero_no_vacio(
    contrato_interventoria.get("fecha_suscripcion"),
)

objeto_contrato_obra = _primero_no_vacio(
    acta_inicio.get("objeto_contrato_obra"),
    contrato_obra.get("objeto_general"),
)

plazo_contrato_obra = _primero_no_vacio(
    acta_inicio.get("plazo_contrato_obra"),
    contrato_obra.get("plazo_ejecucion"),
)

valor_total_contrato_obra = _primero_no_vacio(
    acta_inicio.get("valor_total_contrato_obra"),
    contrato_obra.get("valor_total_numeros"),
    contrato_obra.get("valor_contrato"),
)

contratista_obra = _primero_no_vacio(
    acta_inicio.get("contratista_obra"),
    contrato_obra.get("nombre_contratista"),
)

interventor_obra = _primero_no_vacio(
    acta_inicio.get("interventor"),
    contrato_interventoria.get("nombre_interventor"),
    contrato_interventoria.get("nombre_empresa_interventora"),
    contrato_obra.get("nombre_interventor"),
)

dia_ini, mes_ini, anio_ini = _partes_fecha(fecha_inicio_obra)
dia_fin, mes_fin, anio_fin = _partes_fecha(fecha_fin_obra)

st.markdown(
    """
    <style>
    .titulo-principal {
        text-align: center;
        font-size: 30px;
        font-weight: 800;
        margin-bottom: 10px;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

with st.sidebar:
    st.header("🧭 Acciones")
    if st.button("💾 Guardar acta", type="primary", key="guardar_reunion_tecnica_sidebar"):
        _guardar()

st.markdown('<div class="titulo-principal">ACTA DE REUNIÓN TÉCNICA INICIAL</div>', unsafe_allow_html=True)

with st.container(border=True):
    c1, c2 = st.columns([1, 3])
    with c1:
        st.markdown("**Fecha**")
    with c2:
        datos["fecha_reunion"] = st.date_input(
            "Fecha de reunión",
            value=datos["fecha_reunion"],
            format="DD/MM/YYYY",
            key="fecha_reunion",
            label_visibility="collapsed",
        )

with st.container(border=True):
    c1, c2 = st.columns(2)
    with c1:
        _bloque_texto_etiqueta("Entidad contratante", entidad_contratante, "rti_entidad_contratante")
    with c2:
        st.markdown("**Lugar de reunión**")
        datos["lugar_reunion"] = st.text_input(
            "Lugar de reunión",
            value=_texto(datos.get("lugar_reunion", "")),
            key="lugar_reunion",
            label_visibility="collapsed",
        )

    c3, c4, c5 = st.columns(3)
    with c3:
        _bloque_texto_etiqueta("Director de obra", director_obra, "rti_director_obra")
    with c4:
        _bloque_texto_etiqueta("Director de interventoría", director_interventoria, "rti_director_interventoria")
    with c5:
        _bloque_texto_etiqueta(
            "Supervisor del contrato de interventoría",
            supervisor_interventoria,
            "rti_supervisor_interventoria",
        )

with st.container(border=True):
    st.markdown("### Identificación de contratos")

    c1, c2 = st.columns(2)
    with c1:
        _bloque_texto_etiqueta("Contrato de obra No.", contrato_obra_no, "rti_contrato_obra_no")
    with c2:
        _bloque_texto_etiqueta("Fecha contrato de obra", _fecha_a_texto(fecha_contrato_obra), "rti_fecha_contrato_obra")

    c3, c4 = st.columns(2)
    with c3:
        _bloque_texto_etiqueta(
            "Contrato de interventoría No.",
            contrato_interventoria_no,
            "rti_contrato_interventoria_no",
        )
    with c4:
        _bloque_texto_etiqueta(
            "Fecha contrato de interventoría",
            _fecha_a_texto(fecha_contrato_interventoria),
            "rti_fecha_contrato_interventoria",
        )

with st.container(border=True):
    st.markdown("### Datos del contrato de obra")

    _bloque_texto_etiqueta("Objeto del contrato de obra", objeto_contrato_obra, "rti_objeto_obra", altura=110)

    c1, c2 = st.columns(2)
    with c1:
        _bloque_texto_etiqueta("Plazo del contrato de obra", plazo_contrato_obra, "rti_plazo_obra")
    with c2:
        _bloque_texto_etiqueta(
            "Valor total contrato de obra",
            valor_total_contrato_obra,
            "rti_valor_total_obra",
        )

    st.markdown("**Fecha de inicio del contrato de obra**")
    c3, c4, c5 = st.columns(3)
    with c3:
        st.text_input("Día inicio", value=dia_ini, disabled=True, key="rti_dia_inicio_obra")
    with c4:
        st.text_input("Mes inicio", value=mes_ini, disabled=True, key="rti_mes_inicio_obra")
    with c5:
        st.text_input("Año inicio", value=anio_ini, disabled=True, key="rti_anio_inicio_obra")

    st.markdown("**Fecha de vencimiento del contrato de obra**")
    c6, c7, c8 = st.columns(3)
    with c6:
        st.text_input("Día vencimiento", value=dia_fin, disabled=True, key="rti_dia_fin_obra")
    with c7:
        st.text_input("Mes vencimiento", value=mes_fin, disabled=True, key="rti_mes_fin_obra")
    with c8:
        st.text_input("Año vencimiento", value=anio_fin, disabled=True, key="rti_anio_fin_obra")

    c9, c10 = st.columns(2)
    with c9:
        _bloque_texto_etiqueta("Contratista de obra", contratista_obra, "rti_contratista_obra")
    with c10:
        _bloque_texto_etiqueta("Interventor", interventor_obra, "rti_interventor_obra")

with st.container(border=True):
    st.markdown("### Compromisos relacionados con el contratista")
    _bloque_texto_editable("Estudios y diseños", "estudios_disenos_contratista")
    _bloque_texto_editable("Análisis de precios unitarios", "apu_contratista")
    _bloque_texto_editable("Programa de obra", "programa_obra")
    _bloque_texto_editable("Programa de inversiones", "programa_inversiones")
    _bloque_texto_editable("Plan de inversión del anticipo", "plan_inversion_anticipo")
    _bloque_texto_editable("Flujo de caja", "flujo_caja")
    _bloque_texto_editable("Requisitos ambientales", "requisitos_ambientales")
    _bloque_texto_editable("Requisitos sociales", "requisitos_sociales")

with st.container(border=True):
    st.markdown("### Compromisos relacionados con la interventoría")
    _bloque_texto_editable("Estudios y/o diseños", "estudios_disenos_interventoria")
    _bloque_texto_editable("Aprobación análisis de precios unitarios", "aprobacion_apu")
    _bloque_texto_editable("Aprobación programa de obra", "aprobacion_programa_obra")
    _bloque_texto_editable("Aprobación programa de inversiones", "aprobacion_programa_inversiones")

with st.container(border=True):
    st.markdown("### Otros")

    st.markdown("**Anexos (fotografías, videos, etc.)**")
    col_img1, col_img2 = st.columns(2)

    with col_img1:
        upload_1 = st.file_uploader(
            "Anexo imagen 1",
            type=["png", "jpg", "jpeg", "webp"],
            key="uploader_anexo_reunion_1",
        )
        if upload_1 is not None:
            datos["anexo_imagen_1"] = _guardar_imagen_desde_uploader(upload_1)
        _mostrar_imagen_guardada(datos.get("anexo_imagen_1"))

    with col_img2:
        upload_2 = st.file_uploader(
            "Anexo imagen 2",
            type=["png", "jpg", "jpeg", "webp"],
            key="uploader_anexo_reunion_2",
        )
        if upload_2 is not None:
            datos["anexo_imagen_2"] = _guardar_imagen_desde_uploader(upload_2)
        _mostrar_imagen_guardada(datos.get("anexo_imagen_2"))

    st.markdown("**Fecha de apertura de bitácora**")
    datos["fecha_apertura_bitacora"] = st.date_input(
        "Fecha de apertura de bitácora",
        value=datos["fecha_apertura_bitacora"] or date.today(),
        format="DD/MM/YYYY",
        key="fecha_apertura_bitacora",
        label_visibility="collapsed",
    )

with st.container(border=True):
    st.markdown("### Firmas")

    c1, c2, c3 = st.columns(3)
    with c1:
        st.text_input(
            "Nombre Director de obra",
            value=_texto(director_obra),
            disabled=True,
            key="firma_director_obra",
        )
        st.markdown("**Nombre Director de obra**")

    with c2:
        st.text_input(
            "Nombre Director de interventoría",
            value=_texto(director_interventoria),
            disabled=True,
            key="firma_director_interventoria",
        )
        st.markdown("**Nombre Director de interventoría**")

    with c3:
        st.text_input(
            "Nombre Supervisor",
            value=_texto(supervisor_interventoria),
            disabled=True,
            key="firma_supervisor",
        )
        st.markdown("**Nombre Supervisor**")

col_accion_1, col_accion_2 = st.columns(2)

with col_accion_1:
    if st.button("💾 Guardar acta de reunión técnica inicial", type="primary", key="guardar_reunion_tecnica_principal"):
        _guardar()

with col_accion_2:
    word_reunion_tecnica = _generar_word_reunion_tecnica(
        datos,
        entidad_contratante,
        director_obra,
        director_interventoria,
        supervisor_interventoria,
        contrato_obra_no,
        fecha_contrato_obra,
        contrato_interventoria_no,
        fecha_contrato_interventoria,
        objeto_contrato_obra,
        plazo_contrato_obra,
        dia_ini,
        mes_ini,
        anio_ini,
        dia_fin,
        mes_fin,
        anio_fin,
        valor_total_contrato_obra,
        contratista_obra,
        interventor_obra,
    )
    st.download_button(
        "📥 Descargar acta en Word",
        data=word_reunion_tecnica,
        file_name="acta_reunion_tecnica_inicial.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        key="descargar_word_reunion_tecnica",
    )
