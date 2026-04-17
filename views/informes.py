import base64
import io
from datetime import datetime

import pandas as pd
import streamlit as st
from PIL import Image
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from streamlit_echarts import st_echarts

from supabase_state import cargar_estado, guardar_estado, get_supabase_client, cargar_apus_generados_obra


STORAGE_KEY = "informes_config"


st.title("📄 Informes")


def _safe_str(value) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _safe_bool(value, default=False) -> bool:
    if isinstance(value, bool):
        return value
    if value is None:
        return default
    if isinstance(value, str):
        return value.strip().lower() in ("true", "1", "si", "sí", "yes")
    return bool(value)


def _safe_float(value, default=0.0) -> float:
    try:
        if value is None or value == "":
            return float(default)
        return float(value)
    except Exception:
        return float(default)

def _bytes_uploader(archivo):
    if archivo is None:
        return None
    try:
        archivo.seek(0)
    except Exception:
        pass
    return base64.b64encode(archivo.getvalue()).decode("utf-8")


def _bytes_a_buffer(data_bytes):
    if not data_bytes:
        return None

    try:
        if isinstance(data_bytes, bytes):
            raw = data_bytes
        elif isinstance(data_bytes, str):
            raw = base64.b64decode(data_bytes)
        else:
            return None

        buffer = io.BytesIO(raw)
        buffer.seek(0)
        Image.open(buffer).verify()
        buffer.seek(0)
        return buffer
    except Exception:
        return None


def _normalizar_cfg(data: dict, nombres_equipo: str, fecha_sugerida: str) -> dict:
    if not isinstance(data, dict):
        data = {}

    return {
        "portada_nombre_informe": _safe_str(data.get("portada_nombre_informe")),
        "portada_Responsables": _safe_str(data.get("portada_Responsables")) or nombres_equipo,
        "portada_fecha_manual": _safe_str(data.get("portada_fecha_manual")) or fecha_sugerida,
        "logo_entidad_bytes": data.get("logo_entidad_bytes"),
        "foto_portada_bytes": data.get("foto_portada_bytes"),
        "imagen_grafico_edt_bytes": data.get("imagen_grafico_edt_bytes"),
        "imagen_gantt_bytes": data.get("imagen_gantt_bytes"),
        "incluye_portada_combinado": _safe_bool(data.get("incluye_portada_combinado", True), True),
        "incluye_alcance_combinado": _safe_bool(data.get("incluye_alcance_combinado", True), True),
        "incluye_cronograma_combinado": _safe_bool(data.get("incluye_cronograma_combinado", True), True),
        "incluye_costos_combinado": _safe_bool(data.get("incluye_costos_combinado", False), False),
        "ultima_actualizacion": data.get("ultima_actualizacion"),
    }


def _set_carta(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)


def _agregar_header_footer(section, logo_buffer, texto_footer, nombre_proyecto):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header
    footer = section.footer

    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)

    try:
        htable = header.add_table(1, 2, Inches(7.5))

        if logo_buffer is not None:
            try:
                img_logo = _bytes_a_buffer(
                    logo_buffer.getvalue() if hasattr(logo_buffer, "getvalue") else logo_buffer
                )
                if img_logo is not None:
                    img_logo.seek(0)
                    p_logo = htable.rows[0].cells[0].paragraphs[0]
                    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_logo.add_run().add_picture(img_logo, width=Inches(0.8))
            except Exception:
                pass

        h_text = htable.rows[0].cells[1].paragraphs[0]
        h_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = h_text.add_run(f"PROYECTO: {nombre_proyecto}")
        run.font.size = Pt(8)
        run.italic = True
    except Exception:
        pass

    if not footer.paragraphs:
        footer.add_paragraph()

    f_p = footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_p.text = ""
    f_run = f_p.add_run(texto_footer)
    f_run.font.size = Pt(8)

def _agregar_portada(doc, cfg, nombre_proyecto, entidad_contratante):
    logo_entidad = _bytes_a_buffer(cfg.get("logo_entidad_bytes"))
    foto_portada = _bytes_a_buffer(cfg.get("foto_portada_bytes"))

    _set_carta(doc.sections[0])

    doc.add_paragraph()

    if logo_entidad is not None:
        try:
            logo_entidad.seek(0)
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(logo_entidad, width=Inches(2.5))
            doc.add_paragraph()
        except Exception:
            pass

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_titulo = p_titulo.add_run(nombre_proyecto.upper())
    r_titulo.bold = True
    r_titulo.font.size = Pt(22)
    doc.add_paragraph()

    if foto_portada is not None:
        try:
            foto_portada.seek(0)
            p_foto = doc.add_paragraph()
            p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_foto.add_run().add_picture(foto_portada, width=Inches(5.5))
            doc.add_paragraph()
        except Exception:
            pass

    if cfg.get("portada_nombre_informe"):
        p_inf = doc.add_paragraph()
        p_inf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_inf.add_run(cfg["portada_nombre_informe"].upper()).bold = True
        p_dep = doc.add_paragraph()
        p_dep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_dep.add_run(entidad_contratante.upper()).font.size = Pt(14)

    if cfg.get("portada_resposables"):
        p_form = doc.add_paragraph()
        p_form.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for autor in [a.strip() for a in cfg["portada_Responsables"].split(",") if a.strip()]:
            p_form.add_run(autor + "\n").font.size = Pt(12)

    if cfg.get("portada_fecha_manual"):
        p_fecha = doc.add_paragraph()
        p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_fecha.add_run(cfg["portada_fecha_manual"]).font.size = Pt(12)


def _agregar_alcance(doc, cfg, datos):
    if any(p.text.strip() for p in doc.paragraphs):
        section = doc.add_section(WD_SECTION.NEW_PAGE)
    else:
        section = doc.sections[0]

    _set_carta(section)
    _agregar_header_footer(
        section,
        _bytes_a_buffer(cfg.get("logo_entidad_bytes")),
        "ALCANCE",
        datos["nombre_proyecto"],
    )

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tit.add_run("ALCANCE")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()
    doc.add_heading("1. Descripción General", level=1)
    doc.add_paragraph(datos["descripcion_proyecto"] or "Sin descripción definida.")
    doc.add_paragraph()

    doc.add_heading("2. Estructura de Desglose del Trabajo", level=1)
    doc.add_paragraph()

    imagen_grafico_edt = _bytes_a_buffer(cfg.get("imagen_grafico_edt_bytes"))
    if imagen_grafico_edt is not None:
        try:
            imagen_grafico_edt.seek(0)
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(imagen_grafico_edt, width=Inches(7.0))
            doc.add_paragraph()
        except Exception:
            pass
    else:
        p_aviso = doc.add_paragraph()
        p_aviso.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_aviso = p_aviso.add_run(
            "[ ⚠️ GRÁFICO EDT NO ADJUNTO. Cárguelo en la plataforma antes de descargar ]"
        )
        r_aviso.font.color.rgb = RGBColor(200, 0, 0)

    p_tab_tit = doc.add_paragraph()
    p_tab_tit.add_run("Lista de actividades").bold = True

    flat_table = datos["flat_table"]
    if flat_table:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Código"
        hdr_cells[1].text = "Nombre"

        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for item in flat_table:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item["Código"])
            row_cells[1].text = str(item["Nombre"])

    doc.add_paragraph()
    if datos["descripcion_edt"]:
        doc.add_paragraph(datos["descripcion_edt"])

    if datos["datos_agrupados"]:
        doc.add_heading("3. Especificaciones Técnicas", level=1)
        doc.add_paragraph()

        for grupo in datos["datos_agrupados"]:
            titulo_prod = f"{grupo['codigo_producto']} - {grupo['nombre_producto']}".upper()
            doc.add_heading(titulo_prod, level=2)

            for item in grupo["elementos"]:
                titulo_item = f"{item['codigo']} {item['nombre']}".upper()
                doc.add_heading(titulo_item, level=3)

                p_uni = doc.add_paragraph()
                p_uni.add_run("Unidad de Medida: ").bold = True
                p_uni.add_run(item["unidad"])

                specs = item["specs"]

                def agregar_bloque(titulo, clave):
                    texto = str(specs.get(clave, "")).strip()
                    if texto:
                        p_t = doc.add_paragraph()
                        p_t.add_run(f"{titulo}:").bold = True
                        doc.add_paragraph(texto)

                agregar_bloque("Descripción Detallada", "descripcion")
                agregar_bloque("Procedimiento de Ejecución", "procedimiento")
                agregar_bloque("Materiales Requeridos", "materiales")
                agregar_bloque("Herramientas", "herramientas")
                agregar_bloque("Equipos Necesarios", "equipos")
                agregar_bloque("Medición y Forma de Pago", "medicion_pago")
                agregar_bloque("Condiciones de No Conformidad", "no_conformidad")
                doc.add_paragraph()


def _agregar_cronograma(doc, cfg, datos):
    if any(p.text.strip() for p in doc.paragraphs):
        section = doc.add_section(WD_SECTION.NEW_PAGE)
    else:
        section = doc.sections[0]
        
    _set_carta(section)
    _agregar_header_footer(
        section,
        _bytes_a_buffer(cfg.get("logo_entidad_bytes")),
        "CRONOGRAMA",
        datos["nombre_proyecto"],
    )

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tit.add_run("CRONOGRAMA")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()
    doc.add_heading("1. Precedencias (Tipo y Lag)", level=1)
    doc.add_paragraph()

    df_precedencias = datos.get("tabla_precedencias_cronograma", pd.DataFrame())
    if not df_precedencias.empty:
        table = doc.add_table(rows=1, cols=len(df_precedencias.columns))
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for idx, col in enumerate(df_precedencias.columns):
            hdr_cells[idx].text = str(col)
            for paragraph in hdr_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for _, row in df_precedencias.iterrows():
            row_cells = table.add_row().cells
            for idx, col in enumerate(df_precedencias.columns):
                row_cells[idx].text = str(row[col])

        doc.add_paragraph()
    else:
        doc.add_paragraph("No hay información de precedencias disponible.")
        doc.add_paragraph()

    doc.add_heading("2. Diagrama de Gantt", level=1)
    doc.add_paragraph()

    imagen_gantt = _bytes_a_buffer(cfg.get("imagen_gantt_bytes"))
    if imagen_gantt is not None:
        try:
            imagen_gantt.seek(0)
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(imagen_gantt, width=Inches(7.0))
            doc.add_paragraph()
        except Exception:
            pass
    else:
        p_aviso = doc.add_paragraph()
        p_aviso.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_aviso = p_aviso.add_run(
            "[ ⚠️ DIAGRAMA DE GANTT NO ADJUNTO. Cárguelo en la plataforma antes de descargar ]"
        )
        r_aviso.font.color.rgb = RGBColor(200, 0, 0)

    if datos["descripcion_edt"]:
        doc.add_heading("2. Nota asociada", level=1)
        doc.add_paragraph(datos["descripcion_edt"])


def _tabla_precedencias_cronograma(alcance, cronograma_datos) -> pd.DataFrame:
    todos_los_nodos = []
    dict_id_to_cod = {}
    dict_id_to_nom = {}

    if "objetivos" in alcance and "edt_data" in alcance:
        for i, obj in enumerate(alcance["objetivos"]):
            oid = obj.get("id")
            cod_obj = f"{i+1}"
            nom_obj = obj.get("texto", "Objetivo")

            todos_los_nodos.append(
                {
                    "id": oid,
                    "codigo": cod_obj,
                    "nombre_puro": nom_obj,
                    "es_terminal": False,
                    "nivel_profundidad": 1,
                }
            )
            dict_id_to_cod[oid] = cod_obj
            dict_id_to_nom[oid] = nom_obj

            productos = alcance["edt_data"].get(oid, [])
            for j, p in enumerate(productos):
                pid = p["id"]
                cod_prod = f"{cod_obj}.{j+1}"
                nom_prod = p.get("nombre", "Producto")
                actividades = p.get("actividades", [])
                es_term_p = len(actividades) == 0

                todos_los_nodos.append(
                    {
                        "id": pid,
                        "codigo": cod_prod,
                        "nombre_puro": nom_prod,
                        "es_terminal": es_term_p,
                        "nivel_profundidad": 2,
                    }
                )
                dict_id_to_cod[pid] = cod_prod
                dict_id_to_nom[pid] = nom_prod

                if not es_term_p:
                    for k, a in enumerate(actividades):
                        aid = a["id"]
                        cod_act = f"{cod_prod}.{k+1}"
                        nom_act = a.get("nombre", "Actividad")
                        paquetes = a.get("paquetes", [])
                        es_term_a = len(paquetes) == 0

                        todos_los_nodos.append(
                            {
                                "id": aid,
                                "codigo": cod_act,
                                "nombre_puro": nom_act,
                                "es_terminal": es_term_a,
                                "nivel_profundidad": 3,
                            }
                        )
                        dict_id_to_cod[aid] = cod_act
                        dict_id_to_nom[aid] = nom_act

                        if not es_term_a:
                            for l, pq in enumerate(paquetes):
                                qid = pq["id"]
                                cod_pq = f"{cod_act}.{l+1}"
                                nom_pq = pq.get("nombre", "Paquete")

                                todos_los_nodos.append(
                                    {
                                        "id": qid,
                                        "codigo": cod_pq,
                                        "nombre_puro": nom_pq,
                                        "es_terminal": True,
                                        "nivel_profundidad": 4,
                                    }
                                )
                                dict_id_to_cod[qid] = cod_pq
                                dict_id_to_nom[qid] = nom_pq

    red_aplicada = st.session_state.get("red_dependencias", cronograma_datos.get("red_dependencias", {})) or {}
    aristas = cronograma_datos.get("aristas", {}) or {}
    pert = cronograma_datos.get("pert", {}) or {}

    def dur(nid):
        rec = pert.get(str(nid), {}) or {}
        o = rec.get("O")
        m = rec.get("M")
        p = rec.get("P")
        if o is None or m is None or p is None:
            return 1
        try:
            val = (float(o) + 4.0 * float(m) + float(p)) / 6.0
            return max(1, int(__import__("math").ceil(val)))
        except Exception:
            return 1

    filas = []
    for nodo in todos_los_nodos:
        nid = nodo["id"]

        if nodo["es_terminal"]:
            preds_ids = [orig for orig, sucs in red_aplicada.items() if nid in sucs]
            sucs_ids = red_aplicada.get(nid, [])

            preds_orden = sorted(
                preds_ids,
                key=lambda _pid: tuple(int(p) if p.isdigit() else 0 for p in str(dict_id_to_cod.get(_pid, "0")).split(".")),
            )

            tipos_list = []
            lags_list = []
            for pid in preds_orden:
                meta = aristas.get(f"{pid}::{nid}", {}) or {}
                tipos_list.append(str(meta.get("tipo", "FC")))
                try:
                    lags_list.append(str(int(meta.get("lag", 0))))
                except Exception:
                    lags_list.append("0")

            filas.append(
                {
                    "CÓDIGO": nodo["codigo"],
                    "ESTRUCTURA DE DESGLOSE (EDT)": nodo["nombre_puro"],
                    "DURACIÓN": dur(nid),
                    "PREDECESORAS": ", ".join([dict_id_to_cod.get(pid, "") for pid in preds_orden if dict_id_to_cod.get(pid, "")]),
                    "TIPO": ", ".join(tipos_list),
                    "LAG": ", ".join(lags_list),
                    "SUCESORAS": ", ".join([dict_id_to_cod.get(sid, "") for sid in sucs_ids if dict_id_to_cod.get(sid, "")]),
                }
            )
        else:
            filas.append(
                {
                    "CÓDIGO": nodo["codigo"],
                    "ESTRUCTURA DE DESGLOSE (EDT)": nodo["nombre_puro"],
                    "DURACIÓN": "",
                    "PREDECESORAS": "",
                    "TIPO": "",
                    "LAG": "",
                    "SUCESORAS": "",
                }
            )

    return pd.DataFrame(
        filas,
        columns=[
            "CÓDIGO",
            "ESTRUCTURA DE DESGLOSE (EDT)",
            "DURACIÓN",
            "PREDECESORAS",
            "TIPO",
            "LAG",
            "SUCESORAS",
        ],
    )


def _tipo_proyecto_costos() -> str:
    valor_directo = _safe_str(st.session_state.get("tipo_presupuesto_proyecto_crono"))
    if valor_directo in ("Obra", "Consultoría"):
        return valor_directo

    cronograma_datos = st.session_state.get("cronograma_datos", {}) or {}
    valor_cronograma = _safe_str(cronograma_datos.get("tipo_presupuesto_proyecto"))
    if valor_cronograma in ("Obra", "Consultoría"):
        return valor_cronograma

    return "Obra"


def _money(value):
    try:
        return f"$ {float(value):,.2f}"
    except Exception:
        return "$ 0.00"


def _cargar_catalogo_precios_informes():
    rows = []
    try:
        supabase = get_supabase_client()
        page_size = 1000
        start = 0
        while True:
            resp = (
                supabase.table("catalogo_precios_gobernacion")
                .select('"CAPITULO","SUBCAPITULO","CÓDIGO","NOMBRE DE LA ACTIVIDAD","UNIDAD","COSTO DIRECTO"')
                .eq("activo", True)
                .order('"CÓDIGO"')
                .range(start, start + page_size - 1)
                .execute()
            )
            batch = resp.data or []
            if not batch:
                break
            rows.extend(batch)
            if len(batch) < page_size:
                break
            start += page_size
    except Exception:
        rows = []

    catalogo = []
    for row in rows:
        nombre = str(row.get("NOMBRE DE LA ACTIVIDAD", "") or "").strip()
        if not nombre:
            continue
        costo_raw = str(row.get("COSTO DIRECTO", "") or "").strip().replace(".", "").replace(",", ".")
        try:
            costo = float(costo_raw) if costo_raw else 0.0
        except Exception:
            costo = 0.0
        catalogo.append(
            {
                "capitulo": str(row.get("CAPITULO", "") or "").strip(),
                "subcapitulo": str(row.get("SUBCAPITULO", "") or "").strip(),
                "codigo": str(row.get("CÓDIGO", "") or "").strip(),
                "nombre": nombre,
                "unidad": str(row.get("UNIDAD", "") or "").strip(),
                "vr_unitario": costo,
            }
        )
    st.session_state["informes_catalogo_obra"] = catalogo
    return catalogo


def _get_catalogo_precios_informes():
    items = st.session_state.get("informes_catalogo_obra", [])
    if isinstance(items, list) and items:
        return items
    return _cargar_catalogo_precios_informes()


def _normalize_text(text):
    return str(text or "").strip().upper()


def _build_catalog_index_informes():
    catalogo_items = _get_catalogo_precios_informes()
    index = {}
    for row in catalogo_items:
        codigo = _normalize_text(row.get("codigo"))
        if not codigo:
            continue
        index[codigo] = {
            "codigo": str(row.get("codigo", "") or "").strip(),
            "unidad": str(row.get("unidad", "") or "").strip(),
            "vr_unitario": _safe_float(row.get("vr_unitario", 0.0), 0.0),
            "nombre": str(row.get("nombre", "") or "").strip(),
        }
    return index


def _buscar_en_catalogo_informes(codigo, catalog_index):
    codigo_norm = _normalize_text(codigo)
    if not codigo_norm:
        return {"found": False, "unidad": "", "vr_unitario": 0.0}
    info = catalog_index.get(codigo_norm)
    if not info:
        return {"found": False, "unidad": "", "vr_unitario": 0.0}
    return {"found": True, "unidad": info.get("unidad", ""), "vr_unitario": info.get("vr_unitario", 0.0)}


def _extraer_grupos_desde_edt_obra(alcance_data):
    grupos = []
    objetivos = alcance_data.get("objetivos", []) or []
    edt_data = alcance_data.get("edt_data", {}) or {}

    for i, obj in enumerate(objetivos):
        oid = obj.get("id")
        cod_obj = f"{i + 1}"
        nom_obj = obj.get("texto", "Objetivo")
        productos = edt_data.get(oid, []) or []
        rows_obj = []

        for j, prod in enumerate(productos):
            pid = prod.get("id")
            cod_prod = f"{cod_obj}.{j + 1}"
            nom_prod = prod.get("nombre", "Producto")
            actividades = prod.get("actividades", []) or []

            if not actividades:
                rows_obj.append({"node_id": str(pid), "item": cod_prod, "descripcion": nom_prod})
                continue

            rows_prod = []
            for k, act in enumerate(actividades):
                aid = act.get("id")
                cod_act = f"{cod_prod}.{k + 1}"
                nom_act = act.get("nombre", "Actividad")
                paquetes = act.get("paquetes", []) or []

                if not paquetes:
                    rows_prod.append({"node_id": str(aid), "item": cod_act, "descripcion": nom_act})
                    continue

                grupo_act = {
                    "group_id": str(aid),
                    "group_code": cod_act,
                    "group_name": nom_act,
                    "rows": [],
                }
                for l, pq in enumerate(paquetes):
                    pqid = str(pq.get("id"))
                    cod_paq = f"{cod_act}.{l + 1}"
                    nom_paq = pq.get("nombre", "Paquete")
                    grupo_act["rows"].append({"node_id": pqid, "item": cod_paq, "descripcion": nom_paq})
                grupos.append(grupo_act)

            if rows_prod:
                grupos.append({
                    "group_id": str(pid),
                    "group_code": cod_prod,
                    "group_name": nom_prod,
                    "rows": rows_prod,
                })

        if rows_obj:
            grupos.append({
                "group_id": str(oid),
                "group_code": cod_obj,
                "group_name": nom_obj,
                "rows": rows_obj,
            })

    return grupos


def _estructura_costos_obra(alcance_data: dict) -> dict:
    datos_obra = (st.session_state.get("presupuesto_obra_datos", {}) or {}).copy()
    if not datos_obra:
        try:
            datos_obra = cargar_estado("presupuesto_obra") or {}
        except Exception:
            datos_obra = {}

    items_state = datos_obra.get("items", {}) or {}
    grupos_base = _extraer_grupos_desde_edt_obra(alcance_data)
    catalog_index = _build_catalog_index_informes()

    apus_generados = st.session_state.get("apus_generados_obra")
    if apus_generados is None:
        try:
            apus_generados = cargar_apus_generados_obra()
        except Exception:
            apus_generados = {}
    if not isinstance(apus_generados, dict):
        apus_generados = {}

    factor = _safe_float((datos_obra.get("configuracion", {}) or {}).get("factor_distancia_valor", 0.0), 0.0)

    try:
        aiu_datos = cargar_estado("aiu") or {}
    except Exception:
        aiu_datos = {}

    grupos_render = []
    costo_directo_total = 0.0

    # Primera pasada: costo directo por grupo y total
    grupos_tmp = []
    for grupo in grupos_base:
        filas_finales = []
        costo_directo_grupo = 0.0
        for row in grupo.get("rows", []):
            node_id = row["node_id"]
            descripcion = row["descripcion"]
            item_state = items_state.get(node_id, {}) or {}
            fuente = _safe_str(item_state.get("fuente"))
            item_catalogo = _safe_str(item_state.get("item_catalogo"))
            dist = _safe_str(item_state.get("dist"))
            cant = _safe_float(item_state.get("cant", 0.0), 0.0)
            item_catalogo_display = ""

            if fuente == "Precios Gobernación de Boyacá" and item_catalogo:
                catalogo = _buscar_en_catalogo_informes(item_catalogo, catalog_index)
                if catalogo["found"]:
                    unidad_display = catalogo["unidad"] or ""
                    vr_unitario = _safe_float(catalogo["vr_unitario"], 0.0)
                    item_catalogo_display = item_catalogo
                else:
                    unidad_display = ""
                    vr_unitario = 0.0
                    item_catalogo_display = item_catalogo
            elif fuente == "APU generado":
                apu_generado = apus_generados.get(str(node_id), {}) or {}
                unidad_display = _safe_str(apu_generado.get("unidad_apu")) or "GLOBAL"
                vr_unitario = _safe_float(apu_generado.get("total_apu", 0.0), 0.0)
            else:
                unidad_display = ""
                vr_unitario = 0.0

            if dist == "SI":
                factor_fila = factor
                vr_afectado = vr_unitario * (1 + factor_fila) if vr_unitario > 0 else 0.0
                vr_total = vr_afectado * cant if (vr_afectado > 0 and cant > 0) else 0.0
            else:
                factor_fila = 0.0
                vr_afectado = vr_unitario
                vr_total = vr_unitario * cant if (vr_unitario > 0 and cant > 0) else 0.0

            costo_directo_grupo += vr_total
            filas_finales.append({
                "ITEM": row["item"],
                "ITEM GOBER": item_catalogo_display,
                "DESCRIPCIÓN": descripcion,
                "FUENTE": fuente,
                "UNIDAD": unidad_display,
                "CANT": cant,
                "VR UNITARIO": vr_unitario,
                "DIST.": dist,
                "FACTOR": factor_fila,
                "VR AFECTADO POR FACTOR": vr_afectado,
                "VR TOTAL": vr_total,
            })

        costo_directo_total += costo_directo_grupo
        grupos_tmp.append({
            "group_id": grupo["group_id"],
            "titulo": f"{grupo['group_code']} {grupo['group_name']}",
            "df": pd.DataFrame(filas_finales, columns=["ITEM","ITEM GOBER","DESCRIPCIÓN","FUENTE","UNIDAD","CANT","VR UNITARIO","DIST.","FACTOR","VR AFECTADO POR FACTOR","VR TOTAL"]),
            "costo_directo_grupo": costo_directo_grupo,
        })

    aiu_total = (_safe_float(aiu_datos.get("administracion_valor", 0.0), 0.0) +
                 _safe_float(aiu_datos.get("imprevistos_valor", 0.0), 0.0) +
                 _safe_float(aiu_datos.get("utilidad_valor", 0.0), 0.0))
    aiu_pct_global = (aiu_total / costo_directo_total * 100.0) if costo_directo_total > 0 else 0.0

    for grupo in grupos_tmp:
        grupo["aiu_grupo"] = grupo["costo_directo_grupo"] * (aiu_pct_global / 100.0)
        grupos_render.append(grupo)

    subtotal_presupuesto = costo_directo_total + aiu_total

    otros_costos_indirectos = []
    alcance_costos_indirectos = alcance_data.get("otros_costos_indirectos_proyecto", []) or []
    try:
        costos_indirectos_datos = cargar_estado("costos_indirectos") or {}
    except Exception:
        costos_indirectos_datos = {}
    registros_por_oci = costos_indirectos_datos.get("registros_por_oci", {}) or {}

    for ci in alcance_costos_indirectos:
        ci_id = _safe_str(ci.get("id"))
        ci_nombre = _safe_str(ci.get("nombre"))
        if not ci_id or not ci_nombre:
            continue
        registro_oci = registros_por_oci.get(ci_id, {}) or {}
        otros_costos_indirectos.append({
            "id": ci_id,
            "nombre": ci_nombre,
            "valor": _safe_float(registro_oci.get("valor_total_final", 0.0), 0.0),
        })

    total_otros_costos = sum(_safe_float(x.get("valor", 0.0), 0.0) for x in otros_costos_indirectos)
    total_presupuesto = subtotal_presupuesto + total_otros_costos

    resumen = {
        "costo_directo_total": costo_directo_total,
        "aiu_total": aiu_total,
        "aiu_administracion_valor": _safe_float(aiu_datos.get("administracion_valor", 0.0), 0.0),
        "aiu_imprevistos_valor": _safe_float(aiu_datos.get("imprevistos_valor", 0.0), 0.0),
        "aiu_utilidad_valor": _safe_float(aiu_datos.get("utilidad_valor", 0.0), 0.0),
        "subtotal_presupuesto": subtotal_presupuesto,
        "otros_costos_indirectos": otros_costos_indirectos,
        "total_presupuesto": total_presupuesto,
    }

    return {
        "grupos": grupos_render,
        "resumen": resumen,
        "total_presupuesto": total_presupuesto,
    }


def _estructura_costos_consultoria(alcance_data: dict) -> dict:
    try:
        datos_consultoria = cargar_estado("presupuesto_consultoria") or {}
    except Exception:
        datos_consultoria = {}

    pc_items_data = st.session_state.get("pc_items_data")
    if not pc_items_data:
        pc_items_data = (st.session_state.get("presupuesto_consultoria_datos", {}) or {}).get("pc_items_data")
    if not pc_items_data:
        pc_items_data = datos_consultoria.get("pc_items_data", {}) or {}

    if not isinstance(pc_items_data, dict):
        pc_items_data = {}

    grupos_consultoria = []

    if "objetivos" in alcance_data and "edt_data" in alcance_data:
        for i, obj in enumerate(alcance_data["objetivos"]):
            oid = obj.get("id")
            cod_obj = f"{i+1}"
            nom_obj = obj.get("texto", "Objetivo")
            productos = alcance_data["edt_data"].get(oid, []) or []
            rows_obj = []

            for j, p in enumerate(productos):
                cod_prod = f"{cod_obj}.{j+1}"
                nom_prod = p.get("nombre", "Producto")
                actividades = p.get("actividades", []) or []

                if len(actividades) == 0:
                    rows_obj.append({"id": _safe_str(p.get("id")), "ITEM": cod_prod, "DESCRIPCIÓN": nom_prod})
                else:
                    rows_prod = []
                    for k, a in enumerate(actividades):
                        cod_act = f"{cod_prod}.{k+1}"
                        nom_act = a.get("nombre", "Actividad")
                        paquetes = a.get("paquetes", []) or []

                        if len(paquetes) == 0:
                            rows_prod.append({"id": _safe_str(a.get("id")), "ITEM": cod_act, "DESCRIPCIÓN": nom_act})
                        else:
                            rows_act = []
                            for l, pq in enumerate(paquetes):
                                cod_paq = f"{cod_act}.{l+1}"
                                nom_paq = pq.get("nombre", "Paquete")
                                rows_act.append({"id": _safe_str(pq.get("id")), "ITEM": cod_paq, "DESCRIPCIÓN": nom_paq})
                            if rows_act:
                                grupos_consultoria.append({"group_id": cod_act, "titulo": f"{cod_act} {nom_act}", "rows": rows_act})
                    if rows_prod:
                        grupos_consultoria.append({"group_id": cod_prod, "titulo": f"{cod_prod} {nom_prod}", "rows": rows_prod})
            if rows_obj:
                grupos_consultoria.append({"group_id": cod_obj, "titulo": f"{cod_obj} {nom_obj}", "rows": rows_obj})

    grupos_render = []
    total_presupuesto = 0.0
    for grupo in grupos_consultoria:
        gid = grupo["group_id"]
        filas_guardadas = pc_items_data.get(gid, []) or []
        filas_finales = []
        total_actividad = 0.0
        for idx, r in enumerate(grupo["rows"]):
            fila_guardada = filas_guardadas[idx] if idx < len(filas_guardadas) else {}
            item = _safe_str(fila_guardada.get("ITEM")) or _safe_str(r.get("ITEM"))
            descripcion = _safe_str(fila_guardada.get("DESCRIPCIÓN")) or _safe_str(r.get("DESCRIPCIÓN"))
            fuente = _safe_str(fila_guardada.get("FUENTE")) or "Cotización"
            unidad = _safe_str(fila_guardada.get("UNIDAD"))
            cantidad = float(pd.to_numeric(fila_guardada.get("CANTIDAD", 0.0), errors="coerce") or 0.0)
            costo_unitario = float(pd.to_numeric(fila_guardada.get("COSTO UNITARIO", 0.0), errors="coerce") or 0.0)
            subtotal = float(pd.to_numeric(fila_guardada.get("SUBTOTAL", cantidad * costo_unitario), errors="coerce") or 0.0)
            iva = float(pd.to_numeric(fila_guardada.get("IVA", subtotal * 0.19), errors="coerce") or 0.0)
            total = float(pd.to_numeric(fila_guardada.get("TOTAL", subtotal + iva), errors="coerce") or 0.0)
            total_actividad += total
            filas_finales.append({
                "ITEM": item,
                "DESCRIPCIÓN": descripcion,
                "FUENTE": fuente,
                "UNIDAD": unidad,
                "CANTIDAD": cantidad,
                "COSTO UNITARIO": costo_unitario,
                "SUBTOTAL": subtotal,
                "IVA": iva,
                "TOTAL": total,
            })
        total_presupuesto += total_actividad
        titulo_grupo = _safe_str(grupo.get("titulo"))
        partes_titulo = titulo_grupo.split(" ", 1)
        item_grupo = partes_titulo[0] if partes_titulo else ""
        nombre_grupo = partes_titulo[1] if len(partes_titulo) > 1 else titulo_grupo

        grupos_render.append({
            "group_id": gid,
            "titulo": titulo_grupo,
            "item_grupo": item_grupo,
            "nombre_grupo": nombre_grupo,
            "df": pd.DataFrame(filas_finales, columns=["ITEM","DESCRIPCIÓN","FUENTE","UNIDAD","CANTIDAD","COSTO UNITARIO","SUBTOTAL","IVA","TOTAL"]),
            "total_actividad": total_actividad,
        })

    return {"grupos": grupos_render, "total_presupuesto": total_presupuesto}


def _costos_proyecto(alcance_data: dict) -> dict:
    tipo_proyecto = _tipo_proyecto_costos()
    if tipo_proyecto == "Consultoría":
        estructura = _estructura_costos_consultoria(alcance_data)
        return {
            "tipo": "Consultoría",
            "modo": "consultoria",
            "tabla": pd.DataFrame(),
            "grupos": estructura["grupos"],
            "resumen": {},
            "total_presupuesto": estructura["total_presupuesto"],
        }

    estructura = _estructura_costos_obra(alcance_data)
    return {
        "tipo": "Obra",
        "modo": "obra",
        "tabla": pd.DataFrame(),
        "grupos": estructura["grupos"],
        "resumen": estructura["resumen"],
        "total_presupuesto": estructura["total_presupuesto"],
    }

def _flujo_fondos_obra_datos():
    try:
        flujo_data = cargar_estado("flujo_fondos") or {}
    except Exception:
        flujo_data = {}

    datos_obra = (st.session_state.get("presupuesto_obra_datos", {}) or {}).copy()
    if not datos_obra:
        try:
            datos_obra = cargar_estado("presupuesto_obra") or {}
        except Exception:
            datos_obra = {}

    try:
        costos_indirectos_datos = cargar_estado("costos_indirectos") or {}
    except Exception:
        costos_indirectos_datos = {}

    if not isinstance(flujo_data, dict) or not flujo_data:
        return {
            "df_calculado": pd.DataFrame(),
            "df_resumen": pd.DataFrame(),
            "grafico_png": None,
        }

    def _ordenar_periodos(cols):
        periodos = []
        for c in cols:
            txt = str(c).strip()
            if txt.startswith("Periodo ") and txt.endswith("%"):
                base = txt.replace(" %", "")
                periodos.append(base)
        periodos = sorted(set(periodos), key=lambda x: int(x.replace("Periodo ", "")))
        return periodos

    periodos = _ordenar_periodos(
        [k for row in flujo_data.values() if isinstance(row, dict) for k in row.keys()]
    )

    if not periodos:
        return {
            "df_calculado": pd.DataFrame(),
            "df_resumen": pd.DataFrame(),
            "grafico_png": None,
        }

    aiu_pct = _safe_float((datos_obra.get("configuracion", {}) or {}).get("aiu_pct_global", 0.0), 0.0)
    directos_guardados = datos_obra.get("flujo_fondos_directos", []) or []

    rows = []

    for rec in directos_guardados:
        node_id = _safe_str(rec.get("node_id", ""))
        descripcion = _safe_str(rec.get("DESCRIPCIÓN", rec.get("DESCRIPCION", "")))
        item = _safe_str(rec.get("ITEM", ""))
        valor_base = _safe_float(rec.get("VALOR BASE", 0.0), 0.0)

        if valor_base <= 0 or not descripcion:
            continue

        row_id = f"DIR|{node_id or item or descripcion}"
        rows.append(
            {
                "ROW_ID": row_id,
                "ITEM": item,
                "TIPO": "DIRECTO",
                "DESCRIPCIÓN": descripcion,
                "VALOR BASE": round(valor_base, 2),
                "AIU %": round(aiu_pct, 2),
                "VALOR CON AIU": round(valor_base * (1 + aiu_pct / 100.0), 2),
            }
        )

    config = datos_obra.get("configuracion") or {}
    indirectos = config.get("otros_costos_indirectos", []) or []
    registros = costos_indirectos_datos.get("registros_por_oci") or {}

    for item in indirectos:
        oci_id = _safe_str(item.get("id", ""))
        if not oci_id:
            continue

        nombre = _safe_str(item.get("nombre", "Costo indirecto"))
        reg = registros.get(oci_id, {}) or {}
        valor_base = _safe_float(reg.get("valor_total_final", item.get("valor", 0.0)), 0.0)
        if valor_base <= 0:
            continue

        row_id = f"OCI|{oci_id}"
        rows.append(
            {
                "ROW_ID": row_id,
                "ITEM": "",
                "TIPO": "INDIRECTO",
                "DESCRIPCIÓN": nombre,
                "VALOR BASE": round(valor_base, 2),
                "AIU %": 0.0,
                "VALOR CON AIU": round(valor_base, 2),
            }
        )

    base_df = pd.DataFrame(rows)
    if base_df.empty:
        return {
            "df_calculado": pd.DataFrame(),
            "df_resumen": pd.DataFrame(),
            "grafico_png": None,
        }

    filas_pct = []
    for _, row in base_df.iterrows():
        row_id = _safe_str(row["ROW_ID"])
        guardado_row = flujo_data.get(row_id, {}) if isinstance(flujo_data, dict) else {}

        rec = {
            "ROW_ID": row_id,
            "ITEM": _safe_str(row["ITEM"]),
            "TIPO": _safe_str(row["TIPO"]),
            "DESCRIPCIÓN": _safe_str(row["DESCRIPCIÓN"]),
            "VALOR BASE": round(_safe_float(row["VALOR BASE"]), 2),
            "AIU %": round(_safe_float(row["AIU %"]), 2),
            "VALOR CON AIU": round(_safe_float(row["VALOR CON AIU"]), 2),
        }

        total_pct = 0.0
        for periodo in periodos:
            col = f"{periodo} %"
            val = _safe_float(guardado_row.get(col, 0.0), 0.0)
            rec[col] = val
            total_pct += val

        rec["TOTAL %"] = round(total_pct, 2)
        filas_pct.append(rec)

    df_pct = pd.DataFrame(filas_pct)

    out = []
    out_obra = []
    for _, row in df_pct.iterrows():
        valor_con_aiu = _safe_float(row["VALOR CON AIU"], 0.0)
        cantidad_total = _safe_float(row.get("CANTIDAD TOTAL", 0.0), 0.0)

        rec = {
            "ITEM": _safe_str(row["ITEM"]),
            "TIPO": _safe_str(row["TIPO"]),
            "DESCRIPCIÓN": _safe_str(row["DESCRIPCIÓN"]),
            "VALOR CON AIU": round(valor_con_aiu, 2),
        }

        rec_obra = {
            "ITEM": _safe_str(row["ITEM"]),
            "TIPO": _safe_str(row["TIPO"]),
            "DESCRIPCIÓN": _safe_str(row["DESCRIPCIÓN"]),
            "CANTIDAD TOTAL": round(cantidad_total, 4),
        }

        total_prog = 0.0
        for periodo in periodos:
            pct = _safe_float(row.get(f"{periodo} %", 0.0), 0.0) / 100.0
            val = valor_con_aiu * pct
            rec[f"{periodo} $"] = round(val, 2)
            rec_obra[periodo] = round(cantidad_total * pct, 4)
            total_prog += val

        rec["TOTAL PROGRAMADO"] = round(total_prog, 2)
        out.append(rec)
        out_obra.append(rec_obra)

    df_calculado = pd.DataFrame(out)
    df_programa_obra = pd.DataFrame(out_obra)

    total_periodo = {}
    acumulado = {}
    pct_acum = {}
    total_general = _safe_float(df_calculado["VALOR CON AIU"].sum(), 0.0) if not df_calculado.empty else 0.0
    running = 0.0

    for periodo in periodos:
        col = f"{periodo} $"
        total = _safe_float(df_calculado[col].sum(), 0.0) if col in df_calculado.columns else 0.0
        running += total
        total_periodo[periodo] = round(total, 2)
        acumulado[periodo] = round(running, 2)
        pct_acum[periodo] = round((running / total_general) * 100.0, 2) if total_general > 0 else 0.0

    df_resumen = pd.DataFrame(
        [
            {"CONCEPTO": "TOTAL POR PERIODO", **{p: total_periodo[p] for p in periodos}},
            {"CONCEPTO": "ACUMULADO", **{p: acumulado[p] for p in periodos}},
            {"CONCEPTO": "% ACUMULADO", **{p: pct_acum[p] for p in periodos}},
        ]
    )

    grafico_png = None
    try:
        x = periodos
        y_total = [total_periodo[p] for p in periodos]
        y_acum = [acumulado[p] for p in periodos]
        y_prev = [0.0] + y_acum[:-1]

        fig, ax = plt.subplots(figsize=(10, 5))
        ax.bar(x, y_prev, label="Acumulado previo", alpha=0.35)
        ax.bar(x, y_total, bottom=y_prev, label="Ejecutado en el periodo", alpha=0.85)
        ax.plot(x, y_acum, marker="o", linewidth=2.5, label="Acumulado")

        ax.set_xlabel("Periodo")
        ax.set_ylabel("Valor")
        ax.tick_params(axis="x", rotation=0)
        ax.legend()

        buf = io.BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)
        grafico_png = buf.getvalue()
    except Exception:
        grafico_png = None

    return {
        "df_programa_obra": df_programa_obra,
        "df_calculado": df_calculado,
        "df_resumen": df_resumen,
        "grafico_png": grafico_png,
    }


def _agregar_flujo_fondos_obra(doc, cfg, datos):
    flujo = datos.get("flujo_fondos_obra", {}) or {}
    df_programa_obra = flujo.get("df_programa_obra", pd.DataFrame())
    df_calculado = flujo.get("df_calculado", pd.DataFrame())
    df_resumen = flujo.get("df_resumen", pd.DataFrame())
    grafico_png = flujo.get("grafico_png")

    if len(doc.paragraphs) > 0 and doc.paragraphs[-1].text.strip():
        section = doc.add_section(WD_SECTION.NEW_PAGE)
    else:
        section = doc.sections[0]

    _set_carta(section)
    _agregar_header_footer(
        section,
        _bytes_a_buffer(cfg.get("logo_entidad_bytes")),
        "FLUJO DE FONDOS OBRA",
        datos["nombre_proyecto"],
    )

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tit.add_run("FLUJO DE FONDOS OBRA")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()

    doc.add_paragraph("Programa de obra")
    if not df_programa_obra.empty:
        table = doc.add_table(rows=1, cols=len(df_programa_obra.columns))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for idx, col in enumerate(df_programa_obra.columns):
            hdr[idx].text = str(col)
            for paragraph in hdr[idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for _, row in df_programa_obra.iterrows():
            row_cells = table.add_row().cells
            for idx, col in enumerate(df_programa_obra.columns):
                row_cells[idx].text = str(row[col])
    else:
        doc.add_paragraph("No hay información disponible.")

    doc.add_paragraph()
    doc.add_paragraph("Progama de inversiones")

    if not df_calculado.empty:
        table = doc.add_table(rows=1, cols=len(df_calculado.columns))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for idx, col in enumerate(df_calculado.columns):
            hdr[idx].text = str(col)
            for paragraph in hdr[idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for _, row in df_calculado.iterrows():
            row_cells = table.add_row().cells
            for idx, col in enumerate(df_calculado.columns):
                row_cells[idx].text = str(row[col])
    else:
        doc.add_paragraph("No hay información disponible.")

    doc.add_paragraph()
    doc.add_paragraph("Flujo de fondos")
    if not df_resumen.empty:
        table = doc.add_table(rows=1, cols=len(df_resumen.columns))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for idx, col in enumerate(df_resumen.columns):
            hdr[idx].text = str(col)
            for paragraph in hdr[idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for _, row in df_resumen.iterrows():
            row_cells = table.add_row().cells
            for idx, col in enumerate(df_resumen.columns):
                row_cells[idx].text = str(row[col])
    else:
        doc.add_paragraph("No hay información disponible.")

    doc.add_paragraph()
    doc.add_paragraph("Gráfico de flujo de fondos")
    if grafico_png:
        img_buffer = _bytes_a_buffer(grafico_png)
        if img_buffer is not None:
            doc.add_picture(img_buffer, width=Inches(6.5))
    else:
        doc.add_paragraph("No hay gráfico disponible.")

def _flujo_fondos_consultoria_datos():
    try:
        flujo_data = cargar_estado("flujo_fondos_consultoria") or {}
    except Exception:
        flujo_data = {}

    try:
        presupuesto_cons = cargar_estado("presupuesto_consultoria") or {}
    except Exception:
        presupuesto_cons = {}

    if not isinstance(flujo_data, dict) or not flujo_data:
        return {
            "df_calculado": pd.DataFrame(),
            "df_resumen": pd.DataFrame(),
            "grafico_png": None,
        }

    def _ordenar_periodos(cols):
        periodos = []
        for c in cols:
            txt = str(c).strip()
            if txt.startswith("Periodo ") and txt.endswith("%"):
                base = txt.replace(" %", "")
                periodos.append(base)
        periodos = sorted(set(periodos), key=lambda x: int(x.replace("Periodo ", "")))
        return periodos

    periodos = _ordenar_periodos(
        [k for row in flujo_data.values() if isinstance(row, dict) for k in row.keys()]
    )

    if not periodos:
        return {
            "df_calculado": pd.DataFrame(),
            "df_resumen": pd.DataFrame(),
            "grafico_png": None,
        }

    pc_items_data = presupuesto_cons.get("pc_items_data", {}) or {}
    rows = []

    if isinstance(pc_items_data, dict):
        for group_id, filas in pc_items_data.items():
            if not isinstance(filas, list):
                continue

            for fila in filas:
                if not isinstance(fila, dict):
                    continue

                descripcion = _safe_str(fila.get("DESCRIPCIÓN"))
                item = _safe_str(fila.get("ITEM"))
                valor_base = _safe_float(fila.get("TOTAL", 0.0), 0.0)

                if valor_base <= 0 or not descripcion:
                    continue

                row_id = f"DIR|{item}"
                rows.append(
                    {
                        "ROW_ID": row_id,
                        "ITEM": item,
                        "TIPO": _safe_str(fila.get("FUENTE")),
                        "DESCRIPCIÓN": descripcion,
                        "VALOR BASE": round(valor_base, 2),
                    }
                )

    base_df = pd.DataFrame(rows)
    if base_df.empty:
        return {
            "df_calculado": pd.DataFrame(),
            "df_resumen": pd.DataFrame(),
            "grafico_png": None,
        }

    filas_pct = []
    for _, row in base_df.iterrows():
        row_id = _safe_str(row["ROW_ID"])
        guardado_row = flujo_data.get(row_id, {}) if isinstance(flujo_data, dict) else {}

        rec = {
            "ROW_ID": row_id,
            "ITEM": _safe_str(row["ITEM"]),
            "TIPO": _safe_str(row["TIPO"]),
            "DESCRIPCIÓN": _safe_str(row["DESCRIPCIÓN"]),
            "VALOR BASE": round(_safe_float(row["VALOR BASE"]), 2),
        }

        total_pct = 0.0
        for periodo in periodos:
            col = f"{periodo} %"
            val = _safe_float(guardado_row.get(col, 0.0), 0.0)
            rec[col] = val
            total_pct += val

        rec["TOTAL %"] = round(total_pct, 2)
        filas_pct.append(rec)

    df_pct = pd.DataFrame(filas_pct)

    out = []
    for _, row in df_pct.iterrows():
        valor_base = _safe_float(row["VALOR BASE"], 0.0)
        rec = {
            "ITEM": _safe_str(row["ITEM"]),
            "TIPO": _safe_str(row["TIPO"]),
            "DESCRIPCIÓN": _safe_str(row["DESCRIPCIÓN"]),
            "VALOR BASE": round(valor_base, 2),
        }
        total_prog = 0.0
        for periodo in periodos:
            pct = _safe_float(row.get(f"{periodo} %", 0.0), 0.0) / 100.0
            val = valor_base * pct
            rec[f"{periodo} $"] = round(val, 2)
            total_prog += val
        rec["TOTAL PROGRAMADO"] = round(total_prog, 2)
        out.append(rec)

    df_calculado = pd.DataFrame(out)

    total_periodo = {}
    acumulado = {}
    pct_acum = {}
    total_general = _safe_float(df_calculado["VALOR BASE"].sum(), 0.0) if not df_calculado.empty else 0.0
    running = 0.0

    for periodo in periodos:
        col = f"{periodo} $"
        total = _safe_float(df_calculado[col].sum(), 0.0) if col in df_calculado.columns else 0.0
        running += total
        total_periodo[periodo] = round(total, 2)
        acumulado[periodo] = round(running, 2)
        pct_acum[periodo] = round((running / total_general) * 100.0, 2) if total_general > 0 else 0.0

    df_resumen = pd.DataFrame(
        [
            {"CONCEPTO": "TOTAL POR PERIODO", **{p: total_periodo[p] for p in periodos}},
            {"CONCEPTO": "ACUMULADO", **{p: acumulado[p] for p in periodos}},
            {"CONCEPTO": "% ACUMULADO", **{p: pct_acum[p] for p in periodos}},
        ]
    )

    grafico_png = None
    try:
        x = periodos
        y_total = [total_periodo[p] for p in periodos]
        y_acum = [acumulado[p] for p in periodos]
        y_prev = [0.0] + y_acum[:-1]

        fig, ax = plt.subplots(figsize=(10, 5))
        ax.bar(x, y_prev, label="Acumulado previo", alpha=0.35)
        ax.bar(x, y_total, bottom=y_prev, label="Ejecutado en el periodo", alpha=0.85)
        ax.plot(x, y_acum, marker="o", linewidth=2.5, label="Acumulado")

        ax.set_xlabel("Periodo")
        ax.set_ylabel("Valor")
        ax.tick_params(axis="x", rotation=0)
        ax.legend()

        buf = io.BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)
        grafico_png = buf.getvalue()
    except Exception:
        grafico_png = None

    return {
        "df_calculado": df_calculado,
        "df_resumen": df_resumen,
        "grafico_png": grafico_png,
    }


def _agregar_flujo_fondos_consultoria(doc, cfg, datos):
    flujo = datos.get("flujo_fondos_consultoria", {}) or {}
    df_calculado = flujo.get("df_calculado", pd.DataFrame())
    df_resumen = flujo.get("df_resumen", pd.DataFrame())
    grafico_png = flujo.get("grafico_png")

    if len(doc.paragraphs) > 0 and doc.paragraphs[-1].text.strip():
        section = doc.add_section(WD_SECTION.NEW_PAGE)
    else:
        section = doc.sections[0]

    _set_carta(section)
    _agregar_header_footer(
        section,
        _bytes_a_buffer(cfg.get("logo_entidad_bytes")),
        "FLUJO DE FONDOS CONSULTORÍA",
        datos["nombre_proyecto"],
    )

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tit.add_run("FLUJO DE FONDOS CONSULTORÍA")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()

    doc.add_paragraph("Flujo de fondos calculado")
    if not df_calculado.empty:
        table = doc.add_table(rows=1, cols=len(df_calculado.columns))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for idx, col in enumerate(df_calculado.columns):
            hdr[idx].text = str(col)
            for paragraph in hdr[idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for _, row in df_calculado.iterrows():
            row_cells = table.add_row().cells
            for idx, col in enumerate(df_calculado.columns):
                row_cells[idx].text = str(row[col])
    else:
        doc.add_paragraph("No hay información disponible.")

    doc.add_paragraph()
    doc.add_paragraph("Flujo de fondos")
    if not df_resumen.empty:
        table = doc.add_table(rows=1, cols=len(df_resumen.columns))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for idx, col in enumerate(df_resumen.columns):
            hdr[idx].text = str(col)
            for paragraph in hdr[idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for _, row in df_resumen.iterrows():
            row_cells = table.add_row().cells
            for idx, col in enumerate(df_resumen.columns):
                row_cells[idx].text = str(row[col])
    else:
        doc.add_paragraph("No hay información disponible.")

    doc.add_paragraph()
    doc.add_paragraph("Gráfico de flujo de fondos")
    if grafico_png:
        img_buffer = _bytes_a_buffer(grafico_png)
        if img_buffer is not None:
            doc.add_picture(img_buffer, width=Inches(6.5))
    else:
        doc.add_paragraph("No hay gráfico disponible.")

def _agregar_costos(doc, cfg, datos):
    if len(doc.paragraphs) > 0 and doc.paragraphs[-1].text.strip():
        section = doc.add_section(WD_SECTION.NEW_PAGE)
    else:
        section = doc.sections[0]

    _set_carta(section)
    _agregar_header_footer(
        section,
        _bytes_a_buffer(cfg.get("logo_entidad_bytes")),
        "COSTOS",
        datos["nombre_proyecto"],
    )

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tit.add_run("COSTOS")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()

    estructura = datos.get("costos_proyecto", {}) or {}
    tipo_tabla = estructura.get("tipo", "Obra")
    modo = estructura.get("modo", "obra")
    grupos = estructura.get("grupos", []) or []
    resumen = estructura.get("resumen", {}) or {}
    total_presupuesto = float(pd.to_numeric(estructura.get("total_presupuesto", 0.0), errors="coerce") or 0.0)

    p_tipo = doc.add_paragraph()
    p_tipo.add_run("Tipo de proyecto: ").bold = True
    p_tipo.add_run(tipo_tabla)
    doc.add_paragraph()

    if modo == "consultoria":
        if grupos:
            for grupo in grupos:
                p_grupo = doc.add_paragraph()
                p_grupo.add_run(str(grupo.get("titulo", ""))).bold = True
                df_grupo = grupo.get("df", pd.DataFrame())
                if not df_grupo.empty:
                    table = doc.add_table(rows=1, cols=len(df_grupo.columns))
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells
                    for idx, col in enumerate(df_grupo.columns):
                        hdr_cells[idx].text = str(col)
                        for paragraph in hdr_cells[idx].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    for _, row in df_grupo.iterrows():
                        row_cells = table.add_row().cells
                        for idx, col in enumerate(df_grupo.columns):
                            row_cells[idx].text = str(row[col])
                p_total = doc.add_paragraph()
                p_total.add_run("Total actividad: ").bold = True
                p_total.add_run(_money(grupo.get("total_actividad", 0.0)))
                doc.add_paragraph()
            p_fin = doc.add_paragraph()
            p_fin.add_run("Total presupuesto: ").bold = True
            p_fin.add_run(_money(total_presupuesto))
        else:
            doc.add_paragraph("No hay información de costos disponible para este proyecto.")
        return

    # Obra
    if grupos:
        for grupo in grupos:
            p_grupo = doc.add_paragraph()
            p_grupo.add_run(str(grupo.get("titulo", ""))).bold = True
            df_grupo = grupo.get("df", pd.DataFrame())
            if not df_grupo.empty:
                table = doc.add_table(rows=1, cols=len(df_grupo.columns))
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for idx, col in enumerate(df_grupo.columns):
                    hdr_cells[idx].text = str(col)
                    for paragraph in hdr_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                for _, row in df_grupo.iterrows():
                    row_cells = table.add_row().cells
                    for idx, col in enumerate(df_grupo.columns):
                        val = row[col]
                        if isinstance(val, float):
                            row_cells[idx].text = f"{val:,.4f}" if col in ("CANT", "FACTOR") else f"{val:,.2f}"
                        else:
                            row_cells[idx].text = str(val)
            p_cd = doc.add_paragraph()
            p_cd.add_run("COSTO DIRECTO GRUPO: ").bold = True
            p_cd.add_run(_money(grupo.get("costo_directo_grupo", 0.0)))
            p_aiu = doc.add_paragraph()
            p_aiu.add_run("A.I.U.: ").bold = True
            p_aiu.add_run(_money(grupo.get("aiu_grupo", 0.0)))
            doc.add_paragraph()

    doc.add_heading("RESUMEN DEL PRESUPUESTO", level=2)
    p = doc.add_paragraph(); p.add_run("COSTO DIRECTO: ").bold = True; p.add_run(_money(resumen.get("costo_directo_total", 0.0)))
    p = doc.add_paragraph(); p.add_run("A.I.U.: ").bold = True; p.add_run(_money(resumen.get("aiu_total", 0.0)))
    p = doc.add_paragraph(); p.add_run("Administración: ").bold = True; p.add_run(_money(resumen.get("aiu_administracion_valor", 0.0)))
    p = doc.add_paragraph(); p.add_run("Imprevistos: ").bold = True; p.add_run(_money(resumen.get("aiu_imprevistos_valor", 0.0)))
    p = doc.add_paragraph(); p.add_run("Utilidad: ").bold = True; p.add_run(_money(resumen.get("aiu_utilidad_valor", 0.0)))
    p = doc.add_paragraph(); p.add_run("SUBTOTAL: ").bold = True; p.add_run(_money(resumen.get("subtotal_presupuesto", 0.0)))
    otros = resumen.get("otros_costos_indirectos", []) or []
    if otros:
        doc.add_paragraph("Otros costos indirectos:")
        for item in otros:
            p = doc.add_paragraph(style=None)
            p.add_run(f"- {_safe_str(item.get('nombre'))}: ").bold = True
            p.add_run(_money(item.get("valor", 0.0)))
    p = doc.add_paragraph(); p.add_run("TOTAL PRESUPUESTO: ").bold = True; p.add_run(_money(total_presupuesto))


def _generar_doc_portada(cfg, datos):
    doc = Document()
    _agregar_portada(doc, cfg, datos["nombre_proyecto"], datos["entidad_contratante"])
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _generar_doc_alcance(cfg, datos):
    doc = Document()
    _agregar_alcance(doc, cfg, datos)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _generar_doc_cronograma(cfg, datos):
    doc = Document()
    _agregar_cronograma(doc, cfg, datos)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _generar_doc_costos(cfg, datos):
    doc = Document()
    _agregar_costos(doc, cfg, datos)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _generar_excel_costos(costos_proyecto) -> io.BytesIO:
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        if costos_proyecto.get("modo") == "obra":
            grupos = costos_proyecto.get("grupos", []) or []
            filas_presupuesto = []

            for grupo in grupos:
                titulo_grupo = _safe_str(grupo.get("titulo"))
                filas_presupuesto.append(
                    {
                        "GRUPO": titulo_grupo,
                        "ITEM": "",
                        "ITEM GOBER": "",
                        "DESCRIPCIÓN": "",
                        "FUENTE": "",
                        "UNIDAD": "",
                        "CANT": "",
                        "VR UNITARIO": "",
                        "DIST.": "",
                        "FACTOR": "",
                        "VR AFECTADO POR FACTOR": "",
                        "VR TOTAL": "",
                    }
                )

                df_grupo = grupo.get("df", pd.DataFrame())
                if isinstance(df_grupo, pd.DataFrame) and not df_grupo.empty:
                    df_export = df_grupo.copy()

                    for col in [
                        "ITEM",
                        "ITEM GOBER",
                        "DESCRIPCIÓN",
                        "FUENTE",
                        "UNIDAD",
                        "CANT",
                        "VR UNITARIO",
                        "DIST.",
                        "FACTOR",
                        "VR AFECTADO POR FACTOR",
                        "VR TOTAL",
                    ]:
                        if col not in df_export.columns:
                            df_export[col] = ""

                    df_export = df_export[
                        [
                            "ITEM",
                            "ITEM GOBER",
                            "DESCRIPCIÓN",
                            "FUENTE",
                            "UNIDAD",
                            "CANT",
                            "VR UNITARIO",
                            "DIST.",
                            "FACTOR",
                            "VR AFECTADO POR FACTOR",
                            "VR TOTAL",
                        ]
                    ].copy()
                    df_export.insert(0, "GRUPO", "")

                    filas_presupuesto.extend(df_export.to_dict("records"))

                filas_presupuesto.append(
                    {
                        "GRUPO": "",
                        "ITEM": "",
                        "ITEM GOBER": "",
                        "DESCRIPCIÓN": "COSTO DIRECTO GRUPO",
                        "FUENTE": "",
                        "UNIDAD": "",
                        "CANT": "",
                        "VR UNITARIO": "",
                        "DIST.": "",
                        "FACTOR": "",
                        "VR AFECTADO POR FACTOR": "",
                        "VR TOTAL": grupo.get("costo_directo_grupo", ""),
                    }
                )
                filas_presupuesto.append(
                    {
                        "GRUPO": "",
                        "ITEM": "",
                        "ITEM GOBER": "",
                        "DESCRIPCIÓN": "A.I.U.",
                        "FUENTE": "",
                        "UNIDAD": "",
                        "CANT": "",
                        "VR UNITARIO": "",
                        "DIST.": "",
                        "FACTOR": "",
                        "VR AFECTADO POR FACTOR": "",
                        "VR TOTAL": grupo.get("aiu_grupo", ""),
                    }
                )
                filas_presupuesto.append(
                    {
                        "GRUPO": "",
                        "ITEM": "",
                        "ITEM GOBER": "",
                        "DESCRIPCIÓN": "SUBTOTAL GRUPO",
                        "FUENTE": "",
                        "UNIDAD": "",
                        "CANT": "",
                        "VR UNITARIO": "",
                        "DIST.": "",
                        "FACTOR": "",
                        "VR AFECTADO POR FACTOR": "",
                        "VR TOTAL": grupo.get("subtotal_grupo", ""),
                    }
                )
                filas_presupuesto.append(
                    {
                        "GRUPO": "",
                        "ITEM": "",
                        "ITEM GOBER": "",
                        "DESCRIPCIÓN": "",
                        "FUENTE": "",
                        "UNIDAD": "",
                        "CANT": "",
                        "VR UNITARIO": "",
                        "DIST.": "",
                        "FACTOR": "",
                        "VR AFECTADO POR FACTOR": "",
                        "VR TOTAL": "",
                    }
                )

            df_presupuesto = pd.DataFrame(
                filas_presupuesto,
                columns=[
                    "GRUPO",
                    "ITEM",
                    "ITEM GOBER",
                    "DESCRIPCIÓN",
                    "FUENTE",
                    "UNIDAD",
                    "CANT",
                    "VR UNITARIO",
                    "DIST.",
                    "FACTOR",
                    "VR AFECTADO POR FACTOR",
                    "VR TOTAL",
                ],
            )
            df_presupuesto.to_excel(writer, sheet_name="Presupuesto General", index=False)

            resumen = costos_proyecto.get("resumen", {}) or {}
            otros_costos_lista = resumen.get("otros_costos_indirectos", []) or []
            total_otros_costos = sum(_safe_float(x.get("valor", 0.0), 0.0) for x in otros_costos_lista if isinstance(x, dict))

            df_resumen = pd.DataFrame(
                [
                    {"CONCEPTO": "Costo directo", "VALOR": resumen.get("costo_directo_total", 0)},
                    {"CONCEPTO": "A.I.U.", "VALOR": resumen.get("aiu_total", 0)},
                    {"CONCEPTO": "Subtotal", "VALOR": resumen.get("subtotal_presupuesto", 0)},
                    {"CONCEPTO": "Otros costos indirectos", "VALOR": total_otros_costos},
                    {"CONCEPTO": "Total presupuesto", "VALOR": resumen.get("total_presupuesto", 0)},
                ]
            )

            fila_inicio_resumen = len(df_presupuesto) + 3
            df_resumen.to_excel(
                writer,
                sheet_name="Presupuesto General",
                index=False,
                startrow=fila_inicio_resumen,
            )

        else:
            grupos = costos_proyecto.get("grupos", []) or []
            filas_consultoria = []

            for grupo in grupos:
                item_grupo = _safe_str(grupo.get("item_grupo"))
                nombre_grupo = _safe_str(grupo.get("nombre_grupo"))
                rotulo_grupo = f"{item_grupo} {nombre_grupo}".strip()

                filas_consultoria.append(
                    {
                        "GRUPO": rotulo_grupo,
                        "ITEM": "",
                        "DESCRIPCIÓN": "",
                        "FUENTE": "",
                        "UNIDAD": "",
                        "CANTIDAD": "",
                        "COSTO UNITARIO": "",
                        "SUBTOTAL": "",
                        "IVA": "",
                        "TOTAL": "",
                    }
                )

                df_grupo = grupo.get("df", pd.DataFrame())
                if isinstance(df_grupo, pd.DataFrame) and not df_grupo.empty:
                    df_export = df_grupo.copy()

                    for col in [
                        "ITEM",
                        "DESCRIPCIÓN",
                        "FUENTE",
                        "UNIDAD",
                        "CANTIDAD",
                        "COSTO UNITARIO",
                        "SUBTOTAL",
                        "IVA",
                        "TOTAL",
                    ]:
                        if col not in df_export.columns:
                            df_export[col] = ""

                    df_export = df_export[
                        [
                            "ITEM",
                            "DESCRIPCIÓN",
                            "FUENTE",
                            "UNIDAD",
                            "CANTIDAD",
                            "COSTO UNITARIO",
                            "SUBTOTAL",
                            "IVA",
                            "TOTAL",
                        ]
                    ].copy()
                    df_export.insert(0, "GRUPO", "")

                    filas_consultoria.extend(df_export.to_dict("records"))

                filas_consultoria.append(
                    {
                        "GRUPO": "",
                        "ITEM": "",
                        "DESCRIPCIÓN": "TOTAL ACTIVIDAD",
                        "FUENTE": "",
                        "UNIDAD": "",
                        "CANTIDAD": "",
                        "COSTO UNITARIO": "",
                        "SUBTOTAL": "",
                        "IVA": "",
                        "TOTAL": grupo.get("total_actividad", ""),
                    }
                )
                filas_consultoria.append(
                    {
                        "GRUPO": "",
                        "ITEM": "",
                        "DESCRIPCIÓN": "",
                        "FUENTE": "",
                        "UNIDAD": "",
                        "CANTIDAD": "",
                        "COSTO UNITARIO": "",
                        "SUBTOTAL": "",
                        "IVA": "",
                        "TOTAL": "",
                    }
                )

            filas_consultoria.append(
                {
                    "GRUPO": "",
                    "ITEM": "",
                    "DESCRIPCIÓN": "TOTAL PRESUPUESTO",
                    "FUENTE": "",
                    "UNIDAD": "",
                    "CANTIDAD": "",
                    "COSTO UNITARIO": "",
                    "SUBTOTAL": "",
                    "IVA": "",
                    "TOTAL": costos_proyecto.get("total_presupuesto", 0),
                }
            )

            df_consultoria = pd.DataFrame(
                filas_consultoria,
                columns=[
                    "GRUPO",
                    "ITEM",
                    "DESCRIPCIÓN",
                    "FUENTE",
                    "UNIDAD",
                    "CANTIDAD",
                    "COSTO UNITARIO",
                    "SUBTOTAL",
                    "IVA",
                    "TOTAL",
                ],
            )
            df_consultoria.to_excel(writer, sheet_name="Presupuesto Consultoría", index=False)

            df_resumen = pd.DataFrame(
                [
                    {"CONCEPTO": "Total presupuesto", "VALOR": costos_proyecto.get("total_presupuesto", 0)}
                ]
            )
            df_resumen.to_excel(writer, sheet_name="Resumen", index=False)

    output.seek(0)
    return output


def _generar_doc_combinado(cfg, datos):
    doc = Document()
    primera = True

    if cfg.get("incluye_portada_combinado", True):
        _agregar_portada(doc, cfg, datos["nombre_proyecto"], datos["entidad_contratante"])
        primera = False

    if cfg.get("incluye_alcance_combinado", True):
        _agregar_alcance(doc, cfg, datos)
        primera = False

    if cfg.get("incluye_cronograma_combinado", True):
        _agregar_cronograma(doc, cfg, datos)
        primera = False

    if cfg.get("incluye_costos_combinado", False):
        _agregar_costos(doc, cfg, datos)
        primera = False

    if cfg.get("incluye_flujo_obra_combinado"):
        _agregar_flujo_fondos_obra(doc, cfg, datos)
        primera = False

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


alcance = st.session_state.get("alcance_datos", {}) or {}
datos_hoja_0 = st.session_state.get("integrantes", []) or []

nombres_lista = []
if isinstance(datos_hoja_0, list):
    for miembro in datos_hoja_0:
        if isinstance(miembro, dict) and "Nombre Completo" in miembro:
            nombres_lista.append(str(miembro["Nombre Completo"]).upper())

nombres_equipo = ", ".join(nombres_lista)

lugar = alcance.get("lugar_presentacion", "")
anio = alcance.get("anio_presentacion", "")
fecha_sugerida = f"{lugar}, {anio}".strip(", ")

contrato_obra = cargar_estado("contrato_obra") or {}
entidad_contratante = _safe_str(contrato_obra.get("nombre_entidad", "SIN ENTIDAD DEFINIDA"))

nombre_proyecto = alcance.get("nombre_proyecto", "SIN NOMBRE DEFINIDO")
descripcion_proyecto = alcance.get("descripcion_proyecto", "")
descripcion_edt = alcance.get("descripcion_edt", "")

if "informes_config" not in st.session_state or not isinstance(st.session_state["informes_config"], dict):
    try:
        st.session_state["informes_config"] = cargar_estado(STORAGE_KEY) or {}
    except Exception:
        st.session_state["informes_config"] = {}

cfg = _normalizar_cfg(st.session_state["informes_config"], nombres_equipo, fecha_sugerida)
st.session_state["informes_config"] = cfg

flat_table = []
echarts_data = None
altura_dinamica_str = "600px"
ancho_dinamico_str = "900px"

if "objetivos" in alcance and alcance["objetivos"]:
    nom_proy = str(nombre_proyecto).upper()
    c_l0, c_l1, c_l2, c_l3, c_l4 = "#43A047", "#9370DB", "#C2185B", "#F57C00", "#00796B"
    g_l0, g_l1, g_l2, g_l3, g_l4 = "#E0E0E0", "#EBEBEB", "#F2F2F2", "#F7F7F7", "#FAFAFA"

    echarts_data = {"name": nom_proy, "itemStyle": {"color": c_l0}, "symbolSize": [280, 85], "children": []}
    flat_table.append({"Código": "0", "Nombre": nom_proy, "ColorBG": g_l0})

    nodos_terminales = 0
    profundidad_maxima = 2

    for i, obj in enumerate(alcance["objetivos"]):
        oid = obj.get("id")
        cod1 = f"{i + 1}"
        node_l1 = {
            "name": f"{cod1}. {obj['texto']}",
            "itemStyle": {"color": c_l1},
            "symbolSize": [270, 80],
            "children": [],
        }
        flat_table.append({"Código": cod1, "Nombre": obj["texto"], "ColorBG": g_l1})

        productos = alcance.get("edt_data", {}).get(oid, [])
        if not productos:
            nodos_terminales += 1
        else:
            profundidad_maxima = max(profundidad_maxima, 3)

        for j, prod in enumerate(productos):
            cod2 = f"{cod1}.{j + 1}"
            node_l2 = {
                "name": f"{cod2}. {prod['nombre']}",
                "itemStyle": {"color": c_l2},
                "symbolSize": [260, 80],
                "children": [],
            }
            flat_table.append({"Código": cod2, "Nombre": prod["nombre"], "ColorBG": g_l2})

            actividades = prod.get("actividades", [])
            if not actividades:
                nodos_terminales += 1
            else:
                profundidad_maxima = max(profundidad_maxima, 4)

            for k, act in enumerate(actividades):
                cod3 = f"{cod2}.{k + 1}"
                node_l3 = {
                    "name": f"{cod3}. {act['nombre']}",
                    "itemStyle": {"color": c_l3},
                    "symbolSize": [250, 75],
                    "children": [],
                }
                flat_table.append({"Código": cod3, "Nombre": act["nombre"], "ColorBG": g_l3})

                paquetes = act.get("paquetes", [])
                if not paquetes:
                    nodos_terminales += 1
                else:
                    profundidad_maxima = max(profundidad_maxima, 5)
                    nodos_terminales += len(paquetes)

                for l, paq in enumerate(paquetes):
                    cod4 = f"{cod3}.{l + 1}"
                    node_l3["children"].append(
                        {
                            "name": f"{cod4}. {paq['nombre']}",
                            "itemStyle": {"color": c_l4},
                            "symbolSize": [240, 75],
                        }
                    )
                    flat_table.append({"Código": cod4, "Nombre": paq["nombre"], "ColorBG": g_l4})

                node_l2["children"].append(node_l3)

            node_l1["children"].append(node_l2)

        echarts_data["children"].append(node_l1)

    altura_calculada = max(600, int(nodos_terminales * 160))
    altura_dinamica_str = f"{altura_calculada}px"
    ancho_calculado = max(900, int(profundidad_maxima * 400))
    ancho_dinamico_str = f"{ancho_calculado}px"

datos_agrupados = []
if "edt_data" in alcance and "objetivos" in alcance:
    for i, obj in enumerate(alcance["objetivos"]):
        oid = obj.get("id")
        cod_obj = f"{i + 1}"
        productos = alcance["edt_data"].get(oid, [])

        for j, prod in enumerate(productos):
            cod_prod = f"{cod_obj}.{j + 1}"
            nombre_prod = prod.get("nombre", "PRODUCTO SIN NOMBRE")
            elementos_del_producto = []

            def evaluar_e_inyectar(item_data, item_codigo, item_nombre):
                specs = item_data.get("specs", {}) or {}
                tiene_datos = any(str(v).strip() != "" for v in specs.values() if v is not None)
                if tiene_datos:
                    elementos_del_producto.append(
                        {
                            "codigo": item_codigo,
                            "nombre": item_nombre,
                            "unidad": item_data.get("unidad", "No definida"),
                            "specs": specs,
                        }
                    )

            evaluar_e_inyectar(prod, cod_prod, nombre_prod)

            for k, act in enumerate(prod.get("actividades", [])):
                cod_act = f"{cod_prod}.{k + 1}"
                evaluar_e_inyectar(act, cod_act, act.get("nombre", "ACTIVIDAD SIN NOMBRE"))

                for l, pq in enumerate(act.get("paquetes", [])):
                    cod_paq = f"{cod_act}.{l + 1}"
                    evaluar_e_inyectar(pq, cod_paq, pq.get("nombre", "PAQUETE SIN NOMBRE"))

            if elementos_del_producto:
                datos_agrupados.append(
                    {
                        "codigo_producto": cod_prod,
                        "nombre_producto": nombre_prod,
                        "elementos": elementos_del_producto,
                    }
                )

costos_proyecto = _costos_proyecto(alcance)

flujo_fondos_obra = _flujo_fondos_obra_datos()
flujo_fondos_consultoria = _flujo_fondos_consultoria_datos()

datos_documento = {
    "nombre_proyecto": nombre_proyecto,
    "entidad_contratante": entidad_contratante,
    "descripcion_proyecto": descripcion_proyecto,
    "descripcion_edt": descripcion_edt,
    "flat_table": flat_table,
    "datos_agrupados": datos_agrupados,
    "tabla_precedencias_cronograma": _tabla_precedencias_cronograma(
        alcance,
        st.session_state.get("cronograma_datos", {}) or {},
    ),
    "costos_proyecto": costos_proyecto,
    "flujo_fondos_obra": flujo_fondos_obra,
    "flujo_fondos_consultoria": flujo_fondos_consultoria,
}

logo_entidad = _bytes_a_buffer(cfg.get("logo_entidad_bytes"))
if logo_entidad is None and cfg.get("logo_entidad_bytes"):
    cfg["logo_entidad_bytes"] = None

foto_portada = _bytes_a_buffer(cfg.get("foto_portada_bytes"))
if foto_portada is None and cfg.get("foto_portada_bytes"):
    cfg["foto_portada_bytes"] = None

imagen_grafico_edt = _bytes_a_buffer(cfg.get("imagen_grafico_edt_bytes"))
if imagen_grafico_edt is None and cfg.get("imagen_grafico_edt_bytes"):
    cfg["imagen_grafico_edt_bytes"] = None

imagen_gantt = _bytes_a_buffer(cfg.get("imagen_gantt_bytes"))
if imagen_gantt is None and cfg.get("imagen_gantt_bytes"):
    cfg["imagen_gantt_bytes"] = None

st.session_state["informes_config"] = cfg

st.info("Cada pestaña puede generar su documento independiente. La pestaña final permite construir un documento combinado.")
col_ctrl_1, col_ctrl_2 = st.columns([1, 1])

with col_ctrl_1:
    if st.button("💾 Guardar configuración e imágenes", type="primary", width="stretch"):
        cfg["ultima_actualizacion"] = datetime.now().isoformat()
        st.session_state["informes_config"] = _normalizar_cfg(cfg, nombres_equipo, fecha_sugerida)
        guardar_estado(STORAGE_KEY, st.session_state["informes_config"])
        st.success("Configuración guardada correctamente.")

with col_ctrl_2:
    if st.button("🔄 Recargar desde nube", width="stretch"):
        try:
            st.session_state["informes_config"] = _normalizar_cfg(
                cargar_estado(STORAGE_KEY) or {},
                nombres_equipo,
                fecha_sugerida,
            )
        except Exception:
            st.session_state["informes_config"] = _normalizar_cfg({}, nombres_equipo, fecha_sugerida)
        st.rerun()

st.divider()

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(
    [
        "1. Portada",
        "2. Alcance",
        "3. Cronograma",
        "4. Costos",
        "5. Flujo de fondos obra",
        "6. Documento combinado",
    ]
)

with tab1:
    st.subheader("Portada")

    with st.container(border=True):
        st.markdown("**Configuración de portada**")
        st.text_input("Nombre del proyecto", value=nombre_proyecto, disabled=True)
        st.text_input("Entidad contratante", value=entidad_contratante, disabled=True)

        cfg["portada_nombre_informe"] = st.text_input(
            "Nombre del informe",
            value=cfg.get("portada_nombre_informe", ""),
            key="portada_nombre_informe_cfg",
        )
        cfg["portada_Responsables"] = st.text_input(
            "Responsables",
            value=cfg.get("portada_Responsables", nombres_equipo),
            key="portada_Responsables_cfg",
        )
        cfg["portada_fecha_manual"] = st.text_input(
            "Fecha",
            value=cfg.get("portada_fecha_manual", fecha_sugerida),
            key="portada_fecha_manual_cfg",
        )

    with st.container(border=True):
        st.markdown("**Imágenes de portada**")
        c1, c2 = st.columns(2)

        logo_entidad_upload = c1.file_uploader(
            "Logo de la entidad",
            type=["png", "jpg", "jpeg"],
            key="logo_entidad_upload_cfg",
        )
        if logo_entidad_upload is not None:
            cfg["logo_entidad_bytes"] = _bytes_uploader(logo_entidad_upload)
            logo_entidad = _bytes_a_buffer(cfg.get("logo_entidad_bytes"))

        foto_portada_upload = c2.file_uploader(
            "Foto de portada",
            type=["png", "jpg", "jpeg"],
            key="foto_portada_upload_cfg",
        )
        if foto_portada_upload is not None:
            cfg["foto_portada_bytes"] = _bytes_uploader(foto_portada_upload)
            foto_portada = _bytes_a_buffer(cfg.get("foto_portada_bytes"))

        if logo_entidad is not None:
            c1.image(logo_entidad, width=150)

        if foto_portada is not None:
            c2.image(foto_portada, width=220)

    st.markdown("### Vista previa")
    with st.container(border=True):
        if logo_entidad is not None:
            c1, c2, c3 = st.columns([1.5, 1, 1.5])
            c2.image(logo_entidad, width="stretch")

        st.markdown(f"<h2 style='text-align:center;'>{nombre_proyecto.upper()}</h2>", unsafe_allow_html=True)

        if foto_portada is not None:
            c1, c2, c3 = st.columns([1, 2, 1])
            c2.image(foto_portada, width="stretch")

        if cfg.get("portada_nombre_informe"):
            st.markdown(
                f"<h4 style='text-align:center;'>{cfg['portada_nombre_informe'].upper()}</h4>",
                unsafe_allow_html=True,
            )

        if entidad_contratante:
            st.markdown(f"<p style='text-align:center;'>{entidad_contratante.upper()}</p>", unsafe_allow_html=True)

        if cfg.get("portada_Responsables"):
            autores_html = "<br>".join([a.strip() for a in cfg["portada_Responsables"].split(",") if a.strip()])
            st.markdown(f"<p style='text-align:center;'>{autores_html}</p>", unsafe_allow_html=True)

        if cfg.get("portada_fecha_manual"):
            st.markdown(f"<p style='text-align:center;'>{cfg['portada_fecha_manual']}</p>", unsafe_allow_html=True)

    if st.button("📥 Generar documento de Portada", key="btn_doc_portada", width="stretch"):
        st.session_state["archivo_portada"] = _generar_doc_portada(cfg, datos_documento)
        st.success("Documento de portada generado.")

    if "archivo_portada" in st.session_state:
        st.download_button(
            label="⬇️ Descargar Portada",
            data=st.session_state["archivo_portada"],
            file_name="portada.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_portada",
            width="stretch",
        )

with tab2:
    st.subheader("Alcance")

    with st.container(border=True):
        st.markdown("**Resumen del alcance**")
        st.markdown(f"**Proyecto:** {nombre_proyecto}")
        st.markdown(f"**Descripción general:** {_safe_str(descripcion_proyecto) or 'Sin descripción registrada.'}")
        st.markdown(f"**Descripción EDT:** {_safe_str(descripcion_edt) or 'Sin descripción EDT registrada.'}")

    with st.container(border=True):
        st.markdown("**Gráfico EDT**")
        if echarts_data:
            options_camara = {
                "tooltip": {"trigger": "item", "formatter": "{b}", "confine": True},
                "toolbox": {
                    "show": True,
                    "feature": {
                        "saveAsImage": {
                            "show": True,
                            "title": "📸 Capturar",
                            "name": "Grafico_EDT",
                            "pixelRatio": 2,
                            "backgroundColor": "#ffffff",
                        }
                    },
                    "iconStyle": {"borderColor": "#145A32", "borderWidth": 2},
                    "itemSize": 25,
                    "right": "5%",
                    "top": "2%",
                },
                "series": [
                    {
                        "type": "tree",
                        "data": [echarts_data],
                        "top": "5%",
                        "left": "10%",
                        "bottom": "5%",
                        "right": "15%",
                        "symbol": "rect",
                        "orient": "LR",
                        "edgeShape": "polyline",
                        "initialTreeDepth": -1,
                        "roam": False,
                        "label": {
                            "position": "inside",
                            "color": "white",
                            "fontSize": 12,
                            "overflow": "break",
                            "lineOverflow": "truncate",
                            "width": 230,
                            "height": 65,
                            "lineHeight": 16,
                        },
                        "lineStyle": {"width": 2, "curveness": 0},
                        "expandAndCollapse": True,
                        "animationDuration": 500,
                    }
                ],
            }
            st_echarts(options_camara, height=altura_dinamica_str, width=ancho_dinamico_str)
        else:
            st.warning("No hay datos EDT disponibles para graficar.")

        imagen_grafico_edt_upload = st.file_uploader(
            "Cargar imagen del gráfico EDT",
            type=["png", "jpg", "jpeg"],
            key="imagen_grafico_edt_upload_cfg",
        )
        if imagen_grafico_edt_upload is not None:
            cfg["imagen_grafico_edt_bytes"] = _bytes_uploader(imagen_grafico_edt_upload)
            imagen_grafico_edt = _bytes_a_buffer(cfg.get("imagen_grafico_edt_bytes"))

        if imagen_grafico_edt is not None:
            st.image(imagen_grafico_edt, width=280)

    with st.container(border=True):
        st.markdown("**Lista de actividades**")
        if flat_table:
            df_tabla = pd.DataFrame(flat_table)

            def aplicar_estilos_grises(row):
                return [f'background-color: {row["ColorBG"]}; color: black; font-weight: bold' for _ in row]

            st.dataframe(
                df_tabla.style.apply(aplicar_estilos_grises, axis=1),
                column_order=("Código", "Nombre"),
                width="stretch",
                hide_index=True,
            )
        else:
            st.info("No hay estructura EDT disponible.")

    with st.container(border=True):
        st.markdown("**Especificaciones técnicas**")
        if not datos_agrupados:
            st.info("No se encontraron especificaciones registradas.")
        else:
            for grupo in datos_agrupados:
                titulo_prod = f"{grupo['codigo_producto']} - {grupo['nombre_producto']}".upper()
                st.markdown(
                    f"<h4 style='color:#145A32; border-bottom:2px solid #2e7d32; padding-bottom:5px;'>{titulo_prod}</h4>",
                    unsafe_allow_html=True,
                )

                for item in grupo["elementos"]:
                    titulo_item = f"{item['codigo']} {item['nombre']}".upper()
                    st.markdown(f"**{titulo_item}**")
                    st.markdown(f"Unidad de Medida: {item['unidad']}")

                    specs = item["specs"]

                    def render_bloque(icono, titulo, clave):
                        texto = str(specs.get(clave, "")).strip().replace("\n", "<br>")
                        if texto:
                            st.markdown(
                                f"<div style='margin-top:8px;'><b>{icono} {titulo}</b></div>"
                                f"<div style='padding:8px 12px; border-left:3px solid #dbe2ea; color:#374151; text-align:justify;'>{texto}</div>",
                                unsafe_allow_html=True,
                            )

                    render_bloque("📝", "Descripción Detallada", "descripcion")
                    render_bloque("⚙️", "Procedimiento de Ejecución", "procedimiento")
                    render_bloque("🧱", "Materiales Requeridos", "materiales")
                    render_bloque("🛠️", "Herramientas", "herramientas")
                    render_bloque("🚜", "Equipos Necesarios", "equipos")
                    render_bloque("📏", "Medición y Forma de Pago", "medicion_pago")
                    render_bloque("⚠️", "Condiciones de No Conformidad", "no_conformidad")
                    st.markdown("---")

    if st.button("📥 Generar documento de Alcance", key="btn_doc_alcance", width="stretch"):
        st.session_state["archivo_alcance"] = _generar_doc_alcance(cfg, datos_documento)
        st.success("Documento de alcance generado.")

    if "archivo_alcance" in st.session_state:
        st.download_button(
            label="⬇️ Descargar Alcance",
            data=st.session_state["archivo_alcance"],
            file_name="alcance.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_alcance",
            width="stretch",
        )

with tab3:
    st.subheader("Cronograma")

    with st.container(border=True):
        st.markdown("**Precedencias (Tipo y Lag)**")
        df_precedencias = datos_documento["tabla_precedencias_cronograma"]

        if not df_precedencias.empty:
            st.dataframe(
                df_precedencias,
                width="stretch",
                hide_index=True,
            )
        else:
            st.info("No hay información de precedencias disponible.")

    with st.container(border=True):
        st.markdown("**Imagen del diagrama de Gantt**")
        imagen_gantt_upload = st.file_uploader(
            "Cargar imagen del diagrama de Gantt",
            type=["png", "jpg", "jpeg"],
            key="imagen_gantt_upload_cfg",
        )
        if imagen_gantt_upload is not None:
            cfg["imagen_gantt_bytes"] = _bytes_uploader(imagen_gantt_upload)
            imagen_gantt = _bytes_a_buffer(cfg.get("imagen_gantt_bytes"))

        if imagen_gantt is not None:
            st.image(imagen_gantt, width=320)
        else:
            st.info("Aún no se ha cargado una imagen del Gantt.")

    with st.container(border=True):
        st.markdown("**Vista previa de cronograma**")
        if imagen_gantt is not None:
            st.image(imagen_gantt, width="stretch")
        else:
            st.warning("No hay imagen cargada del diagrama de Gantt.")

    if st.button("📥 Generar documento de Cronograma", key="btn_doc_cronograma", width="stretch"):
        st.session_state["archivo_cronograma"] = _generar_doc_cronograma(cfg, datos_documento)
        st.success("Documento de cronograma generado.")

    if "archivo_cronograma" in st.session_state:
        st.download_button(
            label="⬇️ Descargar Cronograma",
            data=st.session_state["archivo_cronograma"],
            file_name="cronograma.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_cronograma",
            width="stretch",
        )

with tab4:
    st.subheader("Costos")

    estructura_costos = datos_documento.get("costos_proyecto", {}) or {}
    tipo_costos = estructura_costos.get("tipo", "Obra")
    modo_costos = estructura_costos.get("modo", "obra")
    grupos_costos = estructura_costos.get("grupos", []) or []
    resumen_costos = estructura_costos.get("resumen", {}) or {}
    total_presupuesto_costos = float(pd.to_numeric(estructura_costos.get("total_presupuesto", 0.0), errors="coerce") or 0.0)

    with st.container(border=True):
        st.markdown(f"**Tipo de proyecto detectado:** {tipo_costos}")

        if modo_costos == "consultoria":
            if grupos_costos:
                for grupo in grupos_costos:
                    st.markdown(f"### {grupo['titulo']}")
                    st.dataframe(grupo["df"], width="stretch", hide_index=True)
                    st.metric("Total actividad", _money(grupo.get("total_actividad", 0.0)))
                    st.divider()
                st.markdown("## TOTAL PRESUPUESTO")
                st.metric("Total presupuesto", _money(total_presupuesto_costos))
            else:
                st.info("No hay información de costos disponible para este proyecto.")
        else:
            if grupos_costos:
                for grupo in grupos_costos:
                    st.markdown(f"### {grupo['titulo']}")
                    st.dataframe(grupo["df"], width="stretch", hide_index=True)
                    c1, c2 = st.columns(2)
                    c1.metric("COSTO DIRECTO GRUPO", _money(grupo.get("costo_directo_grupo", 0.0)))
                    c2.metric("A.I.U.", _money(grupo.get("aiu_grupo", 0.0)))
                    st.divider()

                st.markdown("## RESUMEN DEL PRESUPUESTO")
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("COSTO DIRECTO", _money(resumen_costos.get("costo_directo_total", 0.0)))
                    st.metric("A.I.U.", _money(resumen_costos.get("aiu_total", 0.0)))
                    st.metric("Administración", _money(resumen_costos.get("aiu_administracion_valor", 0.0)))
                    st.metric("Imprevistos", _money(resumen_costos.get("aiu_imprevistos_valor", 0.0)))
                    st.metric("Utilidad", _money(resumen_costos.get("aiu_utilidad_valor", 0.0)))
                with col2:
                    st.metric("SUBTOTAL", _money(resumen_costos.get("subtotal_presupuesto", 0.0)))
                    otros = resumen_costos.get("otros_costos_indirectos", []) or []
                    if otros:
                        st.markdown("**Otros costos indirectos**")
                        for item in otros:
                            st.markdown(f"- **{_safe_str(item.get('nombre'))}**: {_money(item.get('valor', 0.0))}")
                    st.metric("TOTAL PRESUPUESTO", _money(total_presupuesto_costos))
            else:
                st.info("No hay información de costos disponible para este proyecto.")

    col_costos_1, col_costos_2 = st.columns(2)

    with col_costos_1:
        if st.button("📥 Generar documento de Costos", key="btn_doc_costos", width="stretch"):
            st.session_state["archivo_costos"] = _generar_doc_costos(cfg, datos_documento)
            st.success("Documento de costos generado.")

    with col_costos_2:
        if st.button("📊 Exportar presupuesto a Excel", key="btn_excel_costos", width="stretch"):
            st.session_state["archivo_costos_excel"] = _generar_excel_costos(costos_proyecto)
            st.success("Archivo Excel generado.")

    if "archivo_costos" in st.session_state:
        st.download_button(
            label="⬇️ Descargar Costos",
            data=st.session_state["archivo_costos"],
            file_name="costos.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_costos",
            width="stretch",
        )

    if "archivo_costos_excel" in st.session_state:
        st.download_button(
            label="⬇️ Descargar Presupuesto en Excel",
            data=st.session_state["archivo_costos_excel"],
            file_name="presupuesto.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="download_costos_excel",
            width="stretch",
        )
with tab5:
    st.subheader("Flujo de fondos obra")

    flujo = datos_documento["flujo_fondos_obra"]

    with st.container(border=True):
        st.markdown("**Programa de obra**")
        if not flujo.get("df_programa_obra", pd.DataFrame()).empty:
            st.dataframe(flujo["df_programa_obra"], width="stretch", hide_index=True)
        else:
            st.info("No hay información disponible.")

        st.markdown("**Programa de inversiones**")
        if not flujo["df_calculado"].empty:
            st.dataframe(flujo["df_calculado"], width="stretch", hide_index=True)
        else:
            st.info("No hay información disponible.")

        st.markdown("**Flujo de fondos**")
        if not flujo["df_resumen"].empty:
            st.dataframe(flujo["df_resumen"], width="stretch", hide_index=True)
        else:
            st.info("No hay información disponible.")

        st.markdown("**Gráfico**")
        if flujo["grafico_png"]:
            st.image(flujo["grafico_png"], width="stretch")
        else:
            st.info("No hay gráfico disponible.")

    if st.button("📥 Generar documento de Flujo de fondos obra", key="btn_doc_flujo_obra", width="stretch"):
        doc = Document()
        _agregar_flujo_fondos_obra(doc, cfg, datos_documento)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.session_state["archivo_flujo_obra"] = buffer
        st.success("Documento de flujo de fondos obra generado.")

    if "archivo_flujo_obra" in st.session_state:
        st.download_button(
            label="⬇️ Descargar Flujo de fondos obra",
            data=st.session_state["archivo_flujo_obra"],
            file_name="flujo_fondos_obra.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_flujo_obra",
            width="stretch",
        )

with tab6:
    st.subheader("Documento combinado")

    with st.container(border=True):
        st.markdown("**Seleccione qué módulos unir**")
        cfg["incluye_portada_combinado"] = st.checkbox(
            "Incluir portada",
            value=bool(cfg.get("incluye_portada_combinado", True)),
            key="incluye_portada_combinado_cfg",
        )
        cfg["incluye_alcance_combinado"] = st.checkbox(
            "Incluir alcance",
            value=bool(cfg.get("incluye_alcance_combinado", True)),
            key="incluye_alcance_combinado_cfg",
        )
        cfg["incluye_cronograma_combinado"] = st.checkbox(
            "Incluir cronograma",
            value=bool(cfg.get("incluye_cronograma_combinado", True)),
            key="incluye_cronograma_combinado_cfg",
        )
        cfg["incluye_costos_combinado"] = st.checkbox(
            "Incluir costos",
            value=bool(cfg.get("incluye_costos_combinado", False)),
            key="incluye_costos_combinado_cfg",
        )
        cfg["incluye_flujo_obra_combinado"] = st.checkbox(
            "Incluir flujo de fondos obra",
            value=bool(cfg.get("incluye_flujo_obra_combinado", False)),
            key="incluye_flujo_obra_combinado_cfg",
        )

    with st.container(border=True):
        st.markdown("**Resumen del documento combinado**")
        seleccionados = []
        if cfg.get("incluye_portada_combinado"):
            seleccionados.append("Portada")
        if cfg.get("incluye_alcance_combinado"):
            seleccionados.append("Alcance")
        if cfg.get("incluye_cronograma_combinado"):
            seleccionados.append("Cronograma")
        if cfg.get("incluye_costos_combinado"):
            seleccionados.append("Costos")
        if cfg.get("incluye_flujo_obra_combinado"):
            seleccionados.append("Flujo de fondos obra")

        st.markdown("Módulos seleccionados: " + (", ".join(seleccionados) if seleccionados else "Ninguno"))

    if st.button("📥 Generar documento combinado", key="btn_doc_combinado", width="stretch"):
        st.session_state["archivo_combinado"] = _generar_doc_combinado(cfg, datos_documento)
        st.success("Documento combinado generado.")

    if "archivo_combinado" in st.session_state:
        st.download_button(
            label="⬇️ Descargar Documento Combinado",
            data=st.session_state["archivo_combinado"],
            file_name="reporte_proyecto.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_combinado",
            width="stretch",
        )
st.session_state["informes_config"] = _normalizar_cfg(cfg, nombres_equipo, fecha_sugerida)
