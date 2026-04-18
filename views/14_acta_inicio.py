import re
from io import BytesIO
from datetime import date, datetime, timedelta

import pandas as pd
import streamlit as st
from docx import Document

from supabase_state import cargar_estado
from supabase_state import guardar_estado as guardar_estado_bd


def guardar_estado(clave, datos):
    def serializar(obj):
        if isinstance(obj, dict):
            return {k: serializar(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [serializar(x) for x in obj]
        if isinstance(obj, (date, datetime)):
            return obj.isoformat()
        return obj

    guardar_estado_bd(clave, serializar(datos))


def _leer_contrato_obra():
    datos = st.session_state.get("contrato_obra_datos", {})
    if not isinstance(datos, dict) or not datos:
        datos = cargar_estado("contrato_obra") or {}
    return datos if isinstance(datos, dict) else {}


def _parse_fecha(valor):
    if isinstance(valor, date):
        return valor
    if isinstance(valor, datetime):
        return valor.date()
    if isinstance(valor, str) and valor.strip():
        txt = valor.strip()
        try:
            return datetime.fromisoformat(txt).date()
        except Exception:
            pass
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
            try:
                return datetime.strptime(txt, fmt).date()
            except Exception:
                continue
    return None


def _fecha_texto(valor):
    f = _parse_fecha(valor)
    if not f:
        return ""
    return f.strftime("%d/%m/%Y")


def _extraer_dias_plazo(valor):
    if isinstance(valor, (int, float)):
        return int(valor)

    txt = str(valor or "").strip()
    if not txt:
        return 0

    match = re.search(r"\d+", txt.replace(".", "").replace(",", ""))
    if not match:
        return 0

    try:
        return int(match.group())
    except Exception:
        return 0


def _texto_seguro(valor):
    if valor is None:
        return ""
    return str(valor)


def _inicializar_estado():
    if "acta_inicio_obra_datos" not in st.session_state:
        st.session_state["acta_inicio_obra_datos"] = cargar_estado("acta_inicio_obra") or {}

    datos = st.session_state["acta_inicio_obra_datos"]

    valores_defecto = {
        "fecha_presente_acta": None,
        "requisitos": [
            {
                "REQUISITOS": "Garantías y su aprobación",
                "ESTADO": "N.A.",
            }
        ],
        "certificaciones": [
            {
                "CERTIFICACION": "Afiliacion a Riesgos Laborales",
                "ESTADO": "N.A.",
                "FECHA DE INICIO DE COBERTURA": None,
            }
        ],
        "nombre_firma_interventor": "",
        "nombre_firma_contratista": "",
        "nombre_firma_supervisor": "",
    }

    for k, v in valores_defecto.items():
        if k not in datos:
            datos[k] = v

    datos["fecha_presente_acta"] = _parse_fecha(datos.get("fecha_presente_acta"))

    if not isinstance(datos.get("requisitos"), list) or len(datos["requisitos"]) == 0:
        datos["requisitos"] = [
            {
                "REQUISITOS": "Garantías y su aprobación",
                "ESTADO": "N.A.",
            }
        ]

    if not isinstance(datos.get("certificaciones"), list) or len(datos["certificaciones"]) == 0:
        datos["certificaciones"] = [
            {
                "CERTIFICACION": "Afiliacion a Riesgos Laborales",
                "ESTADO": "N.A.",
                "FECHA DE INICIO DE COBERTURA": None,
            }
        ]

    for fila in datos["certificaciones"]:
        if isinstance(fila, dict):
            fila["FECHA DE INICIO DE COBERTURA"] = _parse_fecha(fila.get("FECHA DE INICIO DE COBERTURA"))


def _guardar():
    guardar_estado("acta_inicio_obra", st.session_state["acta_inicio_obra_datos"])
    st.success("Acta de inicio guardada correctamente.")


def _generar_word_acta_inicio(contrato, datos, fecha_terminacion):
    doc = Document()

    doc.add_heading("ACTA DE INICIACIÓN CONTRATO", level=1)

    doc.add_paragraph(f"FECHA PRESENTE ACTA: {_fecha_texto(datos.get('fecha_presente_acta'))}")
    doc.add_paragraph(f"CONTRATO No: {_texto_seguro(contrato.get('numero_contrato', ''))}")
    doc.add_paragraph(f"CONTRATANTE: {_texto_seguro(contrato.get('nombre_entidad', ''))}")
    doc.add_paragraph(f"NIT. C.C. CONTRATANTE: {_texto_seguro(contrato.get('nit_entidad', ''))}")
    doc.add_paragraph(f"CONTRATISTA: {_texto_seguro(contrato.get('nombre_contratista', ''))}")
    doc.add_paragraph(f"NIT. C.C. CONTRATISTA: {_texto_seguro(contrato.get('nit_contratista', ''))}")
    doc.add_paragraph(
        f"INTERVENTOR: {_texto_seguro(contrato.get('nombre_interventor') or contrato.get('nombre_supervisor') or '')}"
    )
    doc.add_paragraph(f"OBJETO: {_texto_seguro(contrato.get('objeto_general', ''))}")
    doc.add_paragraph(f"VALOR: {_texto_seguro(contrato.get('valor_total_numeros', ''))}")
    doc.add_paragraph(f"PLAZO DE EJECUCIÓN: {_texto_seguro(contrato.get('plazo_ejecucion', ''))}")
    doc.add_paragraph(f"FECHA DE TERMINACIÓN: {_fecha_texto(fecha_terminacion)}")

    doc.add_paragraph("")
    doc.add_paragraph("REQUISITOS")
    for fila in datos.get("requisitos", []):
        doc.add_paragraph(
            f"{_texto_seguro(fila.get('REQUISITOS', ''))}: {_texto_seguro(fila.get('ESTADO', ''))}"
        )

    doc.add_paragraph("")
    doc.add_paragraph("CERTIFICACIÓN")
    for fila in datos.get("certificaciones", []):
        doc.add_paragraph(
            f"{_texto_seguro(fila.get('CERTIFICACION', ''))}: "
            f"{_texto_seguro(fila.get('ESTADO', ''))} | "
            f"FECHA DE INICIO DE COBERTURA: {_fecha_texto(fila.get('FECHA DE INICIO DE COBERTURA'))}"
        )

    doc.add_paragraph("")
    doc.add_paragraph(
        "En constancia se firma por los que en ésta intervinieron, dejando constancia que se han reunido "
        "todos y cada uno de los requisitos necesarios tanto para la legalización del contrato como para su ejecución."
    )

    doc.add_paragraph("")

    tabla_firmas = doc.add_table(rows=2, cols=3)
    tabla_firmas.style = "Table Grid"

    tabla_firmas.cell(0, 0).text = "INTERVENTOR Y/O SUPERVISOR"
    tabla_firmas.cell(0, 1).text = "CONTRATISTA"
    tabla_firmas.cell(0, 2).text = "SUPERVISOR"

    tabla_firmas.cell(1, 0).text = _texto_seguro(datos.get("nombre_firma_interventor", ""))
    tabla_firmas.cell(1, 1).text = _texto_seguro(datos.get("nombre_firma_contratista", ""))
    tabla_firmas.cell(1, 2).text = _texto_seguro(datos.get("nombre_firma_supervisor", ""))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

_inicializar_estado()
datos = st.session_state["acta_inicio_obra_datos"]
contrato = _leer_contrato_obra()

plazo_dias = int(contrato.get("plazo_ejecucion_dias", 0) or 0)
if plazo_dias <= 0:
    plazo_dias = _extraer_dias_plazo(contrato.get("plazo_ejecucion", ""))

if not datos.get("fecha_presente_acta"):
    datos["fecha_presente_acta"] = date.today()

fecha_terminacion = None
if isinstance(datos.get("fecha_presente_acta"), date):
    fecha_terminacion = datos["fecha_presente_acta"] + timedelta(days=plazo_dias)

if not datos.get("nombre_firma_interventor"):
    datos["nombre_firma_interventor"] = _texto_seguro(
        contrato.get("nombre_interventor") or contrato.get("nombre_supervisor") or ""
    )

if not datos.get("nombre_firma_contratista"):
    datos["nombre_firma_contratista"] = _texto_seguro(contrato.get("nombre_contratista", ""))

if not datos.get("nombre_firma_supervisor"):
    datos["nombre_firma_supervisor"] = _texto_seguro(contrato.get("nombre_supervisor", ""))

st.markdown(
    """
    <style>
    .acta-titulo {
        text-align: center;
        font-size: 32px;
        font-weight: 800;
        margin-bottom: 14px;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

with st.sidebar:
    st.header("🧭 Acciones")
    if st.button("💾 Guardar acta", type="primary", key="guardar_acta_inicio_sidebar"):
        _guardar()

st.markdown('<div class="acta-titulo">ACTA DE INICIACIÓN CONTRATO</div>', unsafe_allow_html=True)

with st.container(border=True):
    st.markdown("### FECHA PRESENTE ACTA:")
    datos["fecha_presente_acta"] = st.date_input(
        "FECHA PRESENTE ACTA",
        value=datos["fecha_presente_acta"],
        format="DD/MM/YYYY",
        label_visibility="collapsed",
        key="fecha_presente_acta",
    )

with st.container(border=True):
    st.markdown("### CONTRATO No:")
    st.text_input(
        "CONTRATO No",
        value=_texto_seguro(contrato.get("numero_contrato", "")),
        label_visibility="collapsed",
        disabled=True,
        key="acta_numero_contrato",
    )

with st.container(border=True):
    c1, c2 = st.columns([4, 2])
    with c1:
        st.markdown("### CONTRATANTE:")
        st.text_input(
            "CONTRATANTE",
            value=_texto_seguro(contrato.get("nombre_entidad", "")),
            label_visibility="collapsed",
            disabled=True,
            key="acta_contratante",
        )
    with c2:
        st.markdown("### NIT. C.C.")
        st.text_input(
            "NIT. C.C. CONTRATANTE",
            value=_texto_seguro(contrato.get("nit_entidad", "")),
            label_visibility="collapsed",
            disabled=True,
            key="acta_nit_contratante",
        )

with st.container(border=True):
    c1, c2 = st.columns([4, 2])
    with c1:
        st.markdown("### CONTRATISTA:")
        st.text_input(
            "CONTRATISTA",
            value=_texto_seguro(contrato.get("nombre_contratista", "")),
            label_visibility="collapsed",
            disabled=True,
            key="acta_contratista",
        )
    with c2:
        st.markdown("### NIT. C.C.")
        st.text_input(
            "NIT. C.C. CONTRATISTA",
            value=_texto_seguro(contrato.get("nit_contratista", "")),
            label_visibility="collapsed",
            disabled=True,
            key="acta_nit_contratista",
        )

with st.container(border=True):
    st.markdown("### INTERVENTOR:")
    st.text_input(
        "INTERVENTOR",
        value=_texto_seguro(contrato.get("nombre_interventor") or contrato.get("nombre_supervisor") or ""),
        label_visibility="collapsed",
        disabled=True,
        key="acta_interventor",
    )

with st.container(border=True):
    st.markdown("### OBJETO:")
    st.text_area(
        "OBJETO",
        value=_texto_seguro(contrato.get("objeto_general", "")),
        label_visibility="collapsed",
        disabled=True,
        height=120,
        key="acta_objeto",
    )

with st.container(border=True):
    st.markdown("### VALOR:")
    st.text_input(
        "VALOR",
        value=_texto_seguro(contrato.get("valor_total_numeros", "")),
        label_visibility="collapsed",
        disabled=True,
        key="acta_valor",
    )

with st.container(border=True):
    st.markdown("### PLAZO DE EJECUCIÓN:")
    st.text_input(
        "PLAZO DE EJECUCIÓN",
        value=_texto_seguro(contrato.get("plazo_ejecucion", "")),
        label_visibility="collapsed",
        disabled=True,
        key="acta_plazo",
    )

with st.container(border=True):
    st.markdown("### FECHA DE TERMINACIÓN:")
    st.date_input(
        "FECHA DE TERMINACIÓN",
        value=fecha_terminacion or date.today(),
        format="DD/MM/YYYY",
        label_visibility="collapsed",
        disabled=True,
        key="acta_fecha_terminacion",
    )

st.markdown("Marque con una X el requisito que aplique para la legalización y ejecución del contrato:")

df_requisitos = pd.DataFrame(datos["requisitos"])
df_requisitos_edit = st.data_editor(
    df_requisitos,
    width="stretch",
    hide_index=True,
    num_rows="fixed",
    key="acta_requisitos_editor",
    column_order=["REQUISITOS", "ESTADO"],
    column_config={
        "REQUISITOS": st.column_config.TextColumn("REQUISITOS", disabled=True),
        "ESTADO": st.column_config.SelectboxColumn(
            "MARCAR",
            options=["SI", "NO", "N.A."],
            required=True,
        ),
    },
    disabled=["REQUISITOS"],
)
datos["requisitos"] = df_requisitos_edit.to_dict(orient="records")

st.markdown("")

df_cert = pd.DataFrame(datos["certificaciones"])

if "FECHA DE INICIO DE COBERTURA" in df_cert.columns:
    df_cert["FECHA DE INICIO DE COBERTURA"] = df_cert["FECHA DE INICIO DE COBERTURA"].apply(_parse_fecha)

df_cert_edit = st.data_editor(
    df_cert,
    width="stretch",
    hide_index=True,
    num_rows="fixed",
    key="acta_certificaciones_editor",
    column_order=["CERTIFICACION", "ESTADO", "FECHA DE INICIO DE COBERTURA"],
    column_config={
        "CERTIFICACION": st.column_config.TextColumn("CERTIFICACION", disabled=True),
        "ESTADO": st.column_config.SelectboxColumn(
            "MARCAR",
            options=["SI", "N.A."],
            required=True,
        ),
        "FECHA DE INICIO DE COBERTURA": st.column_config.DateColumn(
            "FECHA DE INICIO DE COBERTURA",
            format="DD/MM/YYYY",
        ),
    },
    disabled=["CERTIFICACION"],
)

certificaciones = []
for fila in df_cert_edit.to_dict(orient="records"):
    if isinstance(fila, dict):
        fila["FECHA DE INICIO DE COBERTURA"] = _parse_fecha(fila.get("FECHA DE INICIO DE COBERTURA"))
        certificaciones.append(fila)

datos["certificaciones"] = certificaciones

st.markdown(
    """
    En constancia se firma por los que en ésta intervinieron, dejando constancia que se han reunido todos
    y cada uno de los requisitos necesarios tanto para la legalización del contrato como para su ejecución.
    """
)

col_firma_1, col_firma_2 = st.columns(2)

with col_firma_1:
    datos["nombre_firma_interventor"] = st.text_input(
        "Nombre interventor y/o supervisor",
        value=_texto_seguro(datos.get("nombre_firma_interventor", "")),
        key="nombre_firma_interventor",
    )
    st.markdown("**INTERVENTOR Y/O SUPERVISOR**")

with col_firma_2:
    datos["nombre_firma_contratista"] = st.text_input(
        "Nombre contratista",
        value=_texto_seguro(datos.get("nombre_firma_contratista", "")),
        key="nombre_firma_contratista",
    )
    st.markdown("**CONTRATISTA**")

col_firma_3, col_firma_4 = st.columns(2)

with col_firma_3:
    datos["nombre_firma_supervisor"] = st.text_input(
        "Nombre supervisor",
        value=_texto_seguro(datos.get("nombre_firma_supervisor", "")),
        key="nombre_firma_supervisor",
    )
    st.markdown("**SUPERVISOR**")

with col_firma_4:
    st.markdown("")

st.divider()

col_accion_1, col_accion_2 = st.columns(2)

with col_accion_1:
    if st.button("💾 Guardar acta de inicio", type="primary", key="guardar_acta_inicio_principal"):
        _guardar()

with col_accion_2:
    word_acta_inicio = _generar_word_acta_inicio(contrato, datos, fecha_terminacion)
    st.download_button(
        "📥 Descargar acta de inicio en Word",
        data=word_acta_inicio,
        file_name="acta_inicio_obra.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        key="descargar_word_acta_inicio",
    )
